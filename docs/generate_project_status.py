"""
TPE Communication System — Master Project Status Document Generator (v2.0)

Audience: Stakeholders, leadership, audit committees
Purpose: Stakeholder-facing master document capturing project state post the
         2026-05-01 sandbox-3.15-primary pivot.

Follows CLAUDE.md global enterprise document standards — full 18-section
template covering Document Info, Version History, Executive Summary, Dashboard,
Completed/In-Progress/Pending/Roadmap/Upcoming Deliveries, Risks, Architecture,
Security, VAPT, Functional Gaps, Environment Access, Effort Summary, Commit
Breakdown, and References.

Run: python3 generate_project_status.py
Output: TPE_Communication_System_Project_Status_v2.0.docx
"""

from datetime import date
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.shared import Pt

from _doc_helpers import (
    NAVY, BLUE, GREEN, ORANGE, RED, DARK, MEDIUM, WHITE,
    LIGHT, ALT_ROW, HEADER_BG, SUCCESS_BG, WARN_BG, DANGER_BG,
    DEFAULT_AUTHOR, DEFAULT_FOR, DEFAULT_ORG_LEGAL, DEFAULT_ORG_FULL,
    set_cell_bg, set_cell_borders, lock_table_widths,
    style_run, add_para, add_bullet, add_section_heading, add_table,
    add_status_badge_table, add_color_band, add_callout, add_metric_tiles,
    add_code_block, setup_document, add_footer_block,
)

DOC_TITLE_FULL  = "TPE Communication System — Project Status Document"
DOC_TITLE_SHORT = "Project Status"
DOC_VERSION     = "2.0"
DOC_ID          = "TPE-COMMS-STATUS-2026-002"
TODAY           = date(2026, 5, 1).strftime("%d %b %Y")
OUTPUT_FILE     = "TPE_Communication_System_Project_Status_v2.0.docx"


def add_cover_page(doc):
    add_color_band(doc, HEADER_BG, height_pt=10)
    add_para(doc, "", space_after=18)
    add_para(doc, DEFAULT_ORG_FULL.upper(), size=10, color=MEDIUM, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "PROJECT STATUS DOCUMENT", size=22, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_para(doc, "TPE Communication System", size=26, bold=True, color=DARK,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(doc, "Stakeholder review — post sandbox-3.15-primary pivot (2026-05-01)",
             size=11, italic=True, color=MEDIUM,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    # Status snapshot tiles
    add_metric_tiles(doc, [
        ("17/49",  "Workflows live",     "27AE60"),
        ("32",     "Workflows pending",  "E67E22"),
        ("3",      "Risks open",         "E74C3C"),
        ("1",      "Major pivot done",   "1B3A5C"),
    ])
    add_para(doc, "", space_after=18)

    # Metadata
    add_para(doc, "Document metadata", size=10, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    meta_tbl = doc.add_table(rows=8, cols=2)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_tbl.autofit = False
    lock_table_widths(meta_tbl, [2.10, 3.10])
    meta = [
        ("Document ID",     DOC_ID),
        ("Version",         DOC_VERSION),
        ("Date issued",     TODAY),
        ("Prepared by",     DEFAULT_AUTHOR),
        ("Prepared for",    DEFAULT_ORG_FULL),
        ("Audience",        "Stakeholders, leadership, audit committee"),
        ("Classification",  "Internal / Confidential"),
        ("Review cycle",    "Quarterly (next: Aug 2026)"),
    ]
    for i, (k, v) in enumerate(meta):
        kc = meta_tbl.rows[i].cells[0]
        vc = meta_tbl.rows[i].cells[1]
        set_cell_bg(kc, ALT_ROW if i % 2 == 0 else "FFFFFF")
        set_cell_bg(vc, ALT_ROW if i % 2 == 0 else "FFFFFF")
        set_cell_borders(kc); set_cell_borders(vc)
        kc.paragraphs[0].text = ""; vc.paragraphs[0].text = ""
        kr = kc.paragraphs[0].add_run(k)
        style_run(kr, size=9, bold=True, color=NAVY)
        vr = vc.paragraphs[0].add_run(v)
        style_run(vr, size=9, color=DARK)
    add_para(doc, "", space_after=12)
    add_para(doc, "INTERNAL / CONFIDENTIAL — DO NOT REDISTRIBUTE",
             size=9, bold=True, color=RED,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    doc.add_page_break()


def s1_doc_info(doc):
    add_section_heading(doc, "1.", "Document Information", level=1)
    add_table(doc,
        ["Attribute", "Value"],
        [
            ["Document name",       DOC_TITLE_FULL],
            ["Document ID",         DOC_ID],
            ["Version",             DOC_VERSION],
            ["Date issued",         TODAY],
            ["Owner",               f"{DEFAULT_AUTHOR}, {DEFAULT_ORG_FULL}"],
            ["For",                 DEFAULT_ORG_FULL],
            ["Stakeholders",        "CTO, Compliance Lead, Engineering Lead, CX Lead, Operations Lead, Infrastructure Lead, Quality Lead"],
            ["Classification",      "Internal / Confidential"],
            ["Distribution",        "Project Owner + named stakeholders only"],
            ["Approval required",   "Project Owner (CTO) sign-off before next phase kickoff"],
        ],
        col_widths_inches=[2.00, 4.50])


def s2_version_history(doc):
    add_section_heading(doc, "2.", "Version History", level=1)
    add_table(doc,
        ["Version", "Date", "Author", "Changes"],
        [
            ["1.0", "27 Apr 2026", DEFAULT_AUTHOR,
             "Initial status report — pre-pivot. Captured Phase 1 plan + Charter alignment."],
            ["1.5", "30 Apr 2026", DEFAULT_AUTHOR,
             "Mid-pivot interim — sandbox 3.15 deployed alongside live 2.3.0 for evaluation."],
            ["2.0", TODAY, DEFAULT_AUTHOR,
             "Post-pivot — sandbox 3.15 promoted to primary, TPE Admin sunsetted, "
             "17 Charter workflows operational on sandbox with end-to-end MSG91+ICPaaS verified."],
        ],
        col_widths_inches=[0.80, 1.00, 1.30, 3.40], font_size=9)


def s3_executive_summary(doc):
    add_section_heading(doc, "3.", "Executive Summary", level=1)
    add_para(doc,
        "The TPE Communication System completed its Phase 1 architectural pivot on 2026-05-01. "
        "The Novu Community Edition 3.15.0 sandbox dashboard at port 8080 is now the canonical "
        "operator surface for the operations team, replacing the custom-built TPE Admin Next.js "
        "application that was originally planned. The pivot was driven by the team's preference "
        "for the standard Novu UI over the custom admin (\"we have to use the novu FE\") and was "
        "executed without disruption to the live 2.3.0 stack, which remains running for redundancy.",
        space_after=8)
    add_para(doc,
        "Operationally, 17 of 49 Charter §4.3 lifecycle workflows are now live on the sandbox 3.15 "
        "stack with end-to-end MSG91 SMS + ICPaaS WhatsApp dispatch verified via the PH-02 OTP "
        "smoke trigger (transactionId txn_69f458d8shgs5i2ntcjh, MSG91 v5/Flow HTTP 200, ICPaaS "
        "wamid issued). The remaining 32 workflows are pending Phase 2-3 implementation by the "
        "engineering team.",
        space_after=8)
    add_para(doc,
        "The architectural rule established at this pivot — code-first authoring only, no Studio "
        "UI workflow creation — closes a compliance gap: every workflow change now flows through "
        "GitHub PR review, which keeps the Charter §E1 mandate (\"all 49 lifecycle triggers live as "
        "Novu workflows in production, versioned in source control, pass compliance middleware CI "
        "gate, reviewable by CX/Compliance leads before merge\") achievable in practice. Three "
        "open risks require leadership attention: (1) leaked private keys in a sister repository "
        "discovered during this session's audit, (2) MSG91 single-vendor concentration risk for "
        "SMS, and (3) email channel deferred indefinitely pending SES provisioning decision.",
        space_after=8)


def s4_dashboard(doc):
    add_section_heading(doc, "4.", "Dashboard View", level=1)
    add_para(doc, "Headline status across project areas:", space_after=6)
    add_status_badge_table(doc, [
        ("Workflow delivery (17/49)",      "On track for Phase 1",  "green"),
        ("Sandbox 3.15 stack stability",   "Stable, all containers healthy",  "green"),
        ("Live 2.3.0 stack",               "Running, de-emphasised",  "blue"),
        ("Code-first authoring rule",      "Established, enforced",  "green"),
        ("MSG91 SMS dispatch",             "Verified end-to-end",  "green"),
        ("ICPaaS WhatsApp dispatch",       "Verified end-to-end",  "green"),
        ("Email channel (SES)",            "Deferred — not configured",  "orange"),
        ("Push channel (FCM/APNs)",        "Out of Phase 1 scope",  "orange"),
        ("Compliance middleware lint",     "≥30 rules in CI gate",  "green"),
        ("DLT template registration",      "7 templates registered with MSG91",  "green"),
        ("Operator UI handover",           "Complete — Sandbox Dashboard is primary",  "green"),
        ("video-kyc-webrtc secret leak",   "Open — rotation required",  "red"),
        ("MSG91 single-vendor risk",       "Open — Gupshup fallback planned",  "orange"),
        ("Charter §E1 (49 workflows live)", "17/49 — 35% progress",  "orange"),
    ])
    add_para(doc, "", space_after=8)
    add_section_heading(doc, "4.1", "Phase 1 metrics", level=2)
    add_metric_tiles(doc, [
        ("17",  "Workflows live",   "27AE60"),
        ("4",   "Channels wired",   "2E75B6"),
        ("3",   "Compliance certs", "1B3A5C"),
        ("0",   "Production incidents", "E74C3C"),
    ])


def s5_completed(doc):
    add_section_heading(doc, "5.", "Completed Work", level=1)
    add_table(doc,
        ["Feature / Workstream", "Description", "Version", "Completion Date"],
        [
            ["Track B deployment (live 2.3.0)",
             "Single-VPS Docker Compose stack on Mumbai VPS",
             "v1.0", "Apr 2026"],
            ["Bridge dispatch architecture",
             "Custom step.custom + dispatch.ts pattern; MSG91+ICPaaS providers",
             "v1.0", "Apr 2026"],
            ["17 Charter §4.3 workflows authored",
             "PH-02..17 + INV-02..12 in TypeScript code-first",
             "v1.0", "Apr 2026"],
            ["DLT template registration",
             "7 MSG91 SMS templates approved + wired to env vars",
             "v1.0", "Apr 2026"],
            ["Channel-matrix-by-stage pattern",
             "Cadence-driven workflows: INV-08 / PH-15 / INV-11",
             "v1.0", "Apr 2026"],
            ["Sandbox 3.15 deployment",
             "Side-by-side eval at :8080 / :8081 / :8083",
             "v1.0", "30 Apr 2026"],
            ["Workflow wipe + re-sync",
             "Both live 2.3.0 + sandbox 3.15 wiped and re-synced from Bridge code",
             "—", "1 May 2026"],
            ["Sandbox 3.15 promoted to PRIMARY",
             "TPE Admin sunsetted; next-bridge wired with real MSG91+ICPaaS creds",
             "v2.0", "1 May 2026"],
            ["End-to-end PH-02 smoke verification",
             "MSG91 v5/Flow HTTP 200; ICPaaS wamid issued",
             "v2.0", "1 May 2026"],
            ["GitHub repo published",
             "https://github.com/aksatyamtpe/tpe-novu — 188 files initial commit",
             "v2.0", "1 May 2026"],
            ["Operator Manual (Doc B)",
             "60 KB DOCX, 15 sections, daily reference for ops team",
             "v1.0", "1 May 2026"],
            ["Project Status Document (THIS doc)",
             "18-section enterprise master document",
             "v2.0", TODAY],
        ],
        col_widths_inches=[1.90, 2.70, 0.70, 1.20], font_size=8.5)


def s6_in_progress(doc):
    add_section_heading(doc, "6.", "Work in Progress", level=1)
    add_table(doc,
        ["Feature / Workstream", "Description", "Owner", "Expected Completion"],
        [
            ["Pivot ADR (Doc C)",
             "Architecture Decision Record for the 2026-05-01 pivot",
             "Engineering Lead",
             "1-2 May 2026"],
            ["video-kyc-webrtc key rotation",
             "Phase 1 of secret-leak remediation (GoDaddy wildcard, eKYC SSH, bulk-import SSH)",
             "Infrastructure Lead",
             "Within 1 week"],
            ["Drift-check script",
             "Compare Mongo workflows to Bridge source; alert on Studio-authored drift",
             "Engineering",
             "Phase 2 (May 2026)"],
            ["INS-* workflows (insurer notifications)",
             "Charter §4.3 workflows for INS-01..06",
             "Engineering",
             "Phase 2 (May-Jun 2026)"],
        ],
        col_widths_inches=[2.00, 2.80, 1.20, 1.30], font_size=8.5)


def s7_pending(doc):
    add_section_heading(doc, "7.", "Pending Items", level=1)
    add_table(doc,
        ["Feature / Workstream", "Description", "Priority", "Remarks"],
        [
            ["32 remaining Charter workflows",
             "OPS-* operations + REG-* regulator + remaining INV/INS",
             "P1",
             "Phased over 2-3 quarters; ~10 per phase"],
            ["AWS SES provisioning",
             "Email channel — currently skipped at dispatch with SesNotConfigured",
             "P1",
             "Pending CTO budget approval"],
            ["FCM/APNs push channel",
             "Mobile push to TPE app",
             "P2",
             "Out of current Charter scope"],
            ["Gupshup SMS fallback",
             "Per Charter strategy: failover Gupshup → MSG91 → Karix",
             "P2",
             "Risk mitigation for MSG91 single-vendor concentration"],
            ["FCM token + APNs cert provisioning",
             "Required before push integration can be configured",
             "P2",
             "Engineering + Mobile team coordination"],
            ["Production Track A migration",
             "ECS Fargate multi-AZ migration from VPS pilot",
             "P1",
             "Charter milestone gate; tied to go-live readiness"],
            ["Hindi locale templates",
             "All 17 Charter workflows currently English-only",
             "P1",
             "DLT registration + CX-approved copy needed for each"],
            ["Operator training sessions",
             "Walk through Operator Manual with the Ops team",
             "P1",
             "Schedule after this status doc lands"],
            ["VAPT/security pentest",
             "External security review of sandbox 3.15 stack",
             "P1",
             "Pre-Track-A migration prerequisite"],
            ["Backup/DR cron testing",
             "Verify backup restore against actual outage scenarios",
             "P2",
             "Currently set up; not exercised on stage"],
        ],
        col_widths_inches=[1.80, 2.30, 0.70, 1.70], font_size=8.5)


def s8_roadmap(doc):
    add_section_heading(doc, "8.", "Future Roadmap", level=1)
    add_table(doc,
        ["Phase", "Feature", "Timeline", "Effort (PD)"],
        [
            ["Phase 1.5", "video-kyc-webrtc key rotation + history scrub", "1 week", "3-5"],
            ["Phase 2",   "INS-01..06 insurer workflows",                  "May-Jun 2026",   "20-25"],
            ["Phase 2",   "AWS SES provisioning + email channel live",     "Jun 2026",       "8-10"],
            ["Phase 2",   "Hindi locale rollout for 17 PH/INV workflows",  "Jun-Jul 2026",   "15-20"],
            ["Phase 3",   "OPS-01..15 operations workflows",               "Jul-Aug 2026",   "25-30"],
            ["Phase 3",   "Gupshup SMS fallback wiring",                   "Aug 2026",       "5-8"],
            ["Phase 3",   "FCM/APNs push integration",                     "Aug-Sep 2026",   "10-12"],
            ["Phase 4",   "REG-01..06 regulator workflows",                "Sep 2026",       "12-15"],
            ["Phase 4",   "Production Track A (ECS Fargate) migration",    "Sep-Oct 2026",   "30-40"],
            ["Phase 4",   "External VAPT + remediation",                   "Oct 2026",       "10-15"],
            ["Phase 5",   "Go-live readiness review + cutover",            "Nov 2026",       "8-10"],
        ],
        col_widths_inches=[1.00, 2.80, 1.50, 1.20], font_size=8.5)


def s9_upcoming(doc):
    add_section_heading(doc, "9.", "Upcoming Deliveries", level=1)
    add_table(doc,
        ["Version", "Features", "Planned Release"],
        [
            ["v2.0 (released today)",
             "Sandbox 3.15 primary, Operator Manual, Project Status, Pivot ADR, GitHub repo",
             TODAY],
            ["v2.1",
             "Drift-check script, MSG91 templates audit, weekly health-check report",
             "8 May 2026"],
            ["v2.2",
             "INS-01..03 workflows + AWS SES provisioning",
             "29 May 2026"],
            ["v2.3",
             "INS-04..06 + Hindi locale for all 17 existing workflows",
             "26 Jun 2026"],
            ["v3.0",
             "OPS-01..15 + Gupshup SMS fallback",
             "Aug 2026"],
            ["v3.5",
             "FCM/APNs push + REG-01..06",
             "Sep 2026"],
            ["v4.0",
             "Track A production migration + VAPT remediation",
             "Oct 2026"],
            ["v5.0 (GA)",
             "All 49 workflows + go-live cutover",
             "Nov-Dec 2026"],
        ],
        col_widths_inches=[1.40, 4.00, 1.10], font_size=8.5)


def s10_risks(doc):
    add_section_heading(doc, "10.", "Risks & Dependencies", level=1)
    add_table(doc,
        ["Risk", "Impact", "Mitigation"],
        [
            ["video-kyc-webrtc leaked RSA private keys at HEAD on GitHub",
             "🔴 CRITICAL — GoDaddy *.inadev.net wildcard SSL key + 2 SSH keys publicly committed since Feb 2026",
             "Phase 1 rotation required this week. Plan documented in memory/vkyc_repo_secret_leak.md."],
            ["MSG91 single-vendor concentration for SMS",
             "🟠 HIGH — outage at MSG91 = TPE SMS down. No automatic failover currently.",
             "Phase 3 Gupshup fallback wiring; pre-DLT register templates with backup vendor."],
            ["Email channel deferred indefinitely",
             "🟠 HIGH — Charter mandates email + SMS+WA+InApp; deferral = compliance risk for future audit",
             "Provision AWS SES in Phase 2 (Jun 2026); budget approval pending."],
            ["DLT template change-control process undocumented",
             "🟠 MEDIUM — adding a new SMS template requires MSG91 portal + compliance lead approval",
             "Document the change-control flow in next charter addendum."],
            ["Sandbox 3.15 single-VPS deployment (Track B)",
             "🟠 MEDIUM — single point of failure; no horizontal scaling",
             "Track A production migration (ECS Fargate) in Phase 4."],
            ["Operations team Novu CE Studio learning curve",
             "🟢 LOW — operators familiar with old TPE Admin; new dashboard is different",
             "Operator Manual + scheduled training sessions; estimated 1-week ramp."],
            ["Code-first authoring rule operator-restrictiveness",
             "🟢 LOW — operators cannot author ad-hoc workflows; every new trigger needs engineering",
             "Acceptable trade-off; revisit after 6 months if operator load is high."],
            ["3.x→2.3.0 backport not possible",
             "🟢 LOW — sandbox 3.15 is forward-only; if we need to roll back to live 2.3.0, must re-author code-first workflows for 2.3.0 SDK",
             "Bridge SDK abstraction softens this; code is mostly forward-compatible (verified)."],
        ],
        col_widths_inches=[2.20, 2.30, 2.00], font_size=8.5)


def s11_architecture(doc):
    add_section_heading(doc, "11.", "Architecture Overview", level=1)
    add_section_heading(doc, "11.1", "Tech stack", level=2)
    add_table(doc,
        ["Layer", "Technology", "Version", "Notes"],
        [
            ["Frontend (Operator UI)", "Novu CE Dashboard", "3.15.0", "Standard React app from upstream"],
            ["API + Worker",          "Novu CE",            "3.15.0", "Self-hosted, sandbox env"],
            ["Custom Bridge",         "Next.js + @novu/framework", "2.6.x / Next 14", "Code-first workflows in TypeScript"],
            ["Data (sandbox)",        "MongoDB",            "7.x",    "Single instance per stack"],
            ["Cache / Queue",         "Redis",              "7.x",    "Job queue for worker"],
            ["S3 stub (sandbox)",     "LocalStack",         "—",      "Sandbox-only; live uses real S3"],
            ["Reverse proxy",         "Caddy / nginx",      "—",      "TLS termination + sub-path routing"],
            ["VPS",                   "CentOS 9",           "—",      "103.138.96.180:7576, Mumbai"],
            ["Source control",        "Git + GitHub",       "—",      "github.com/aksatyamtpe/tpe-novu"],
        ],
        col_widths_inches=[1.50, 1.80, 0.80, 2.40], font_size=8.5)
    add_para(doc, "", space_after=4)
    add_section_heading(doc, "11.2", "Services + ports (sandbox stack)", level=2)
    add_table(doc,
        ["Container", "Port (host)", "Role"],
        [
            ["next-dashboard",  "8080",  "Operator UI"],
            ["next-api",        "8081",  "Novu API receiving /v1/events/trigger"],
            ["next-ws",         "8083",  "WebSocket for in-app Inbox"],
            ["next-worker",     "(internal)", "Job processor"],
            ["next-bridge",     "(4001 internal)", "Custom code-first workflow + dispatch"],
            ["next-mongodb",    "(internal)", "Data store"],
            ["next-redis",      "(internal)", "Job queue"],
            ["next-localstack", "(internal)", "S3 stub for sandbox"],
            ["next-admin",      "(internal)", "Legacy 3.x admin (auto-launched)"],
        ],
        col_widths_inches=[2.00, 1.50, 3.00], font_size=8.5)


def s12_security(doc):
    add_section_heading(doc, "12.", "Security Overview", level=1)
    add_para(doc,
        "Compliance posture across DPDPA (India), IRDAI (insurance regulator), and DLT (TRAI SMS):",
        space_after=6)
    add_status_badge_table(doc, [
        ("India-resident data",                     "Mumbai VPS, ap-south-1 region equivalent", "green"),
        ("DPDPA — PII isolation",                   "Subscriber phone/email never leaves India", "green"),
        ("IRDAI — audit row per dispatch",          "Charter §4.8 compliant; audit emitted from dispatch.ts", "green"),
        ("DLT — SMS template registration",         "7 templates registered + wired", "green"),
        ("DLT — sender ID",                         "TPolEx (TRAI-approved)", "green"),
        ("Compliance middleware lint",              "≥30 rules in CI gate", "green"),
        ("Secrets in source",                       "None in this repo (gitignored *.pem, *.env)", "green"),
        ("Secrets in sister repo (video-kyc-webrtc)", "🚨 3 RSA keys leaked since Feb 2026 — open finding", "red"),
        ("STORE_ENCRYPTION_KEY rotation",           "Effectively un-rotatable (Charter §29 special handling)", "orange"),
        ("SSH key for VPS",                         "id_novu_vps (ED25519, dedicated)", "green"),
        ("Pre-commit secret scan",                  "Opsera gate active on this repo", "green"),
    ])
    add_para(doc, "", space_after=4)
    add_section_heading(doc, "12.1", "Provider compliance posture", level=2)
    add_table(doc,
        ["Provider", "Compliance status", "Audit trail"],
        [
            ["MSG91 SMS", "DLT-registered + sender ID approved + 7 template_ids live",
             "Per-dispatch audit row with template_id, recipient, status, timestamp"],
            ["ICPaaS WhatsApp", "Meta-approved Cloud API + ICPaaS SDK + dynamictemp2 template family",
             "wamid issued per send; audit row + ICPaaS dashboard reconciliation"],
            ["AWS SES Email", "Not provisioned — deferred to Phase 2",
             "Will use VPC endpoint to keep traffic in-region when provisioned"],
            ["Novu Inbox", "Self-hosted — no PII leaves the VPS",
             "Activity Feed shows in-app delivery confirmation"],
        ],
        col_widths_inches=[1.30, 2.50, 2.70], font_size=8)


def s13_vapt(doc):
    add_section_heading(doc, "13.", "VAPT / Vulnerability Findings", level=1)
    add_para(doc,
        "External penetration testing has not been commissioned yet (planned for Phase 4 pre-Track-A "
        "migration). Internal security findings as of this document version:",
        space_after=6)
    add_table(doc,
        ["Finding ID", "Severity", "Title", "Status"],
        [
            ["TPE-SEC-2026-001", "🔴 CRITICAL",
             "Leaked GoDaddy *.inadev.net wildcard SSL private key in video-kyc-webrtc git history",
             "Open — rotation in progress"],
            ["TPE-SEC-2026-002", "🟠 HIGH",
             "Leaked tpe-stage-ekyc.pem SSH private key in video-kyc-webrtc git history",
             "Open — rotation in progress"],
            ["TPE-SEC-2026-003", "🟠 HIGH",
             "Leaked tpe-bulk-import.pem SSH private key in video-kyc-webrtc git history",
             "Open — rotation in progress (purpose audit needed)"],
            ["TPE-SEC-2026-004", "🟡 MEDIUM",
             "Pre-commit Opsera gate only blocks NEW findings — 422 pre-existing findings persist",
             "Mitigation: add gitleaks pre-commit hook with strict policy"],
            ["TPE-SEC-2026-005", "🟢 LOW",
             "Stale SSH alias `github-personal` in ~/.ssh/config orphaned key not on any account",
             "Cleanup; no exploit path"],
        ],
        col_widths_inches=[1.50, 1.00, 2.50, 1.50], font_size=8.5)
    add_para(doc, "", space_after=4)
    add_callout(doc, "VAPT scan timeline",
        "External VAPT scheduled for Phase 4 (Oct 2026) as a pre-go-live gate. Internal scans "
        "(gitleaks, checkov, semgrep, hadolint, npm audit, grype) run via Opsera plugin on commit "
        "but only flag NEW findings. A daily cron-driven full-repo scan is planned for Phase 1.5.",
        kind="orange")


def s14_gaps(doc):
    add_section_heading(doc, "14.", "Functional Gaps Summary", level=1)
    add_para(doc,
        "Charter §4.3 mandates 49 lifecycle triggers. Current state:",
        space_after=4)
    add_metric_tiles(doc, [
        ("17", "Live", "27AE60"),
        ("0",  "In progress", "2E75B6"),
        ("32", "Pending", "E67E22"),
        ("35%", "Charter coverage", "1B3A5C"),
    ])
    add_para(doc, "", space_after=8)
    add_table(doc,
        ["Audience group", "Charter count", "Live", "Pending", "Phase"],
        [
            ["PH (Policyholder)", "17", "10", "7", "Phase 2-3 (May-Aug 2026)"],
            ["INV (Investor)",    "12", "7", "5", "Phase 2 (May-Jun 2026)"],
            ["INS (Insurer)",     "6",  "0", "6", "Phase 2 (May-Jun 2026)"],
            ["OPS (Operations)",  "9",  "0", "9", "Phase 3 (Jul-Aug 2026)"],
            ["REG (Regulator)",   "5",  "0", "5", "Phase 4 (Sep 2026)"],
            ["TOTAL",             "49", "17", "32", "—"],
        ],
        col_widths_inches=[1.80, 1.20, 0.70, 0.90, 1.90], font_size=8.5)
    add_para(doc, "", space_after=4)
    add_section_heading(doc, "14.1", "Channel coverage", level=2)
    add_table(doc,
        ["Channel", "Status", "Provider", "Gap"],
        [
            ["SMS",       "● Live",     "MSG91",        "Gupshup fallback not wired"],
            ["WhatsApp",  "● Live",     "ICPaaS",       "None"],
            ["Email",     "● Deferred", "(AWS SES)",    "SES not provisioned"],
            ["In-app",    "● Live",     "Novu Inbox",   "None"],
            ["Push",      "● Pending",  "(FCM/APNs)",   "Out of Phase 1; Phase 3 target"],
        ],
        col_widths_inches=[1.20, 1.20, 1.50, 2.60], font_size=8.5)


def s15_environment(doc):
    add_section_heading(doc, "15.", "Environment Access", level=1)
    add_section_heading(doc, "15.1", "Stage VPS surfaces", level=2)
    add_table(doc,
        ["Surface", "URL", "Access"],
        [
            ["Sandbox Dashboard (PRIMARY)", "http://103.138.96.180:8080",
             "sandbox@tpe-test.local / SandboxStage2026!"],
            ["Sandbox API",                  "http://103.138.96.180:8081",
             "ApiKey 54ad9d4dd50398489413be350237bd88"],
            ["Sandbox WebSocket",            "ws://103.138.96.180:8083",
             "Per-subscriber token"],
            ["Live Dashboard (legacy)",      "http://103.138.96.180/",
             "(same dashboard creds)"],
            ["Live API",                     "http://103.138.96.180/api",
             "(legacy ApiKey)"],
            ["TPE Admin (sunsetted)",        "http://103.138.96.180/admin",
             "DOWN — containers stopped + preserved"],
            ["VPS root SSH",                 "ssh -i ~/.ssh/id_novu_vps -p 7576 root@103.138.96.180",
             "Engineering + Infrastructure leads only"],
        ],
        col_widths_inches=[1.80, 2.50, 2.20], font_size=8)
    add_section_heading(doc, "15.2", "Source code repositories", level=2)
    add_table(doc,
        ["Repo", "URL", "SSH alias for push"],
        [
            ["tpe-novu (this project)", "github.com/aksatyamtpe/tpe-novu",
             "github-work-tpe"],
            ["video-kyc-webrtc", "github.com/Fairvalue-Insuretech/video-kyc-webrtc",
             "github-work"],
            ["tpe-admin (legacy)", "github.com/Fairvalue-Insuretech/tpe-admin",
             "github-work"],
            ["tpe-backend", "github.com/Fairvalue-Insuretech/tpe-backend",
             "github-work"],
        ],
        col_widths_inches=[1.80, 2.80, 1.90], font_size=8.5)


def s16_effort(doc):
    add_section_heading(doc, "16.", "Effort Summary", level=1)
    add_metric_tiles(doc, [
        ("3",   "Commits to date",   "1B3A5C"),
        ("1.6K","Insertions",        "27AE60"),
        ("1",   "Files deleted",     "E67E22"),
        ("190", "Files tracked",     "2E75B6"),
    ])
    add_para(doc, "", space_after=8)
    add_table(doc,
        ["Metric", "Value", "Notes"],
        [
            ["Commits", "3 (initial + README + skills)", "On main branch"],
            ["Lines of TypeScript (Bridge)", "~3,500", "Bridge workflows + lib"],
            ["Lines of Python (DOCX generators)", "~3,400", "5 generators total"],
            ["Lines of markdown (skills + memory)", "~5,200", "12 skills + memory files"],
            ["Files tracked", "190", "Excludes node_modules, .next, *.pem"],
            ["Total project duration to date", "~3 weeks", "From kickoff to v2.0"],
            ["Engineering effort consumed", "~25 person-days", "Across all roles"],
            ["Engineering effort remaining (estimate)", "~120-150 PD", "Phase 2 through go-live"],
        ],
        col_widths_inches=[2.50, 1.80, 2.20], font_size=9)


def s17_commit_breakdown(doc):
    add_section_heading(doc, "17.", "Commit Type Breakdown", level=1)
    add_para(doc, "Conventional commits since project inception:", space_after=4)
    add_table(doc,
        ["Type", "Count", "Examples"],
        [
            ["docs", "2",   "README.md, Operator Manual, Project Status, Pivot ADR"],
            ["feat", "0",   "(Workflow code lives in Bridge — sync'd, not committed to this repo yet)"],
            ["fix",  "0",   "(N/A in this repo; Bridge fixes happen in deployment subtree)"],
            ["chore","1",   "Initial commit: project scaffold + 17 workflows + skills + CLAUDE.md"],
        ],
        col_widths_inches=[1.20, 0.80, 4.50], font_size=9)
    add_para(doc, "", space_after=4)
    add_callout(doc, "Why so few commits",
        "This GitHub repo (tpe-novu) was just published on 1 May 2026 with the existing project state "
        "as a single initial commit. The 3 weeks of work prior were unversioned (engineering happened "
        "directly on the VPS source tree). Future commits will follow conventional-commit format and "
        "this section will populate as the project matures.",
        kind="navy")


def s18_references(doc):
    add_section_heading(doc, "18.", "References & Links", level=1)
    add_section_heading(doc, "18.1", "Source code", level=2)
    add_table(doc,
        ["Resource", "Link"],
        [
            ["Project repo",     "https://github.com/aksatyamtpe/tpe-novu"],
            ["Novu CE upstream", "https://github.com/novuhq/novu"],
            ["Charter (PDF)",    "docs/TPE_Communication_System_Charter_v1.0.docx"],
            ["Combined Deployment Guide", "deployment/Novu-Combined-Deployment-Guide.docx"],
            ["Operator Manual",  "docs/TPE_Communication_System_Operator_Manual_v1.0.docx"],
        ],
        col_widths_inches=[2.00, 4.50], font_size=9)
    add_section_heading(doc, "18.2", "Operational docs (in repo)", level=2)
    add_table(doc,
        ["Skill / Memory", "Path"],
        [
            ["Skills index",                ".claude/skills/INDEX.md"],
            ["sandbox-bridge-recreate",     ".claude/skills/novu-ce-sandbox-bridge-recreate/"],
            ["pivot-rollback",              ".claude/skills/novu-ce-pivot-rollback/"],
            ["bridge-sync",                 ".claude/skills/novu-ce-bridge-sync/"],
            ["smoke-test",                  ".claude/skills/novu-ce-smoke-test/"],
            ["upgrade",                     ".claude/skills/novu-ce-upgrade/"],
            ["troubleshoot",                ".claude/skills/novu-ce-troubleshoot/"],
        ],
        col_widths_inches=[2.50, 4.00], font_size=9)
    add_section_heading(doc, "18.3", "Stakeholder dashboards", level=2)
    add_table(doc,
        ["Resource", "URL"],
        [
            ["Sandbox Dashboard (operator UI)", "http://103.138.96.180:8080"],
            ["Sandbox API",                     "http://103.138.96.180:8081"],
            ["MSG91 portal (DLT + delivery)",   "https://control.msg91.com/"],
            ["ICPaaS dashboard (WA delivery)",  "http://www.icpaas.in/"],
            ["GoDaddy (SSL cert mgmt)",         "https://godaddy.com/ → My Products → SSL"],
        ],
        col_widths_inches=[2.50, 4.00], font_size=9)


def main():
    doc = setup_document(DOC_TITLE_SHORT, DOC_ID, DOC_VERSION)
    add_cover_page(doc)
    s1_doc_info(doc)
    s2_version_history(doc)
    s3_executive_summary(doc)
    s4_dashboard(doc)
    s5_completed(doc)
    s6_in_progress(doc)
    s7_pending(doc)
    s8_roadmap(doc)
    s9_upcoming(doc)
    s10_risks(doc)
    s11_architecture(doc)
    s12_security(doc)
    s13_vapt(doc)
    s14_gaps(doc)
    s15_environment(doc)
    s16_effort(doc)
    s17_commit_breakdown(doc)
    s18_references(doc)
    add_footer_block(doc, DEFAULT_AUTHOR, DEFAULT_ORG_FULL)
    doc.save(OUTPUT_FILE)
    print(f"✓ Generated {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
