"""
TPE Communication System — Charter Addendum Generator (v1.1)

Companion to TPE_Communication_System_Charter_v1.0.docx.

Adds the Tier 1 and Tier 2 gap-fill sections identified in the gap analysis:
  Tier 1 (block sign-off):
    A1. Effort Estimation (person-days)
    A2. Per-deliverable Acceptance Criteria
    A3. Service Level Agreements (operational + delivery)
    A4. Quality Assurance & Testing Strategy
    A5. Security & Privacy Framework (DPDPA-aligned)
    A6. Cost & Budget Estimate
    A7. Glossary
  Tier 2 (recommended):
    A8. Document Control & Approval Matrix
    A9. Communication Management Plan
    A10. Operational Readiness & Go-Live Checklist

Run: python3 generate_charter_addendum.py
Output: TPE_Communication_System_Charter_Addendum_v1.1.docx
"""

from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────────────────────
# Brand palette
# ─────────────────────────────────────────────────────────────────────────────
NAVY    = RGBColor(0x1B, 0x3A, 0x5C)
BLUE    = RGBColor(0x2E, 0x75, 0xB6)
GREEN   = RGBColor(0x27, 0xAE, 0x60)
ORANGE  = RGBColor(0xE6, 0x7E, 0x22)
RED     = RGBColor(0xE7, 0x4C, 0x3C)
DARK    = RGBColor(0x2C, 0x3E, 0x50)
MEDIUM  = RGBColor(0x7F, 0x8C, 0x8D)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT   = "F0F4F8"
ALT_ROW = "F8F9FA"
HEADER_BG = "1B3A5C"
SUCCESS_BG = "E8F5E9"
WARN_BG    = "FFF8E1"
DANGER_BG  = "FFEBEE"

DOC_TITLE   = "TPE Communication System — Charter Addendum v1.1"
DOC_VERSION = "1.1"
DOC_ID      = "TPE-COMMS-CHARTER-ADDENDUM-2026-001"
TODAY       = date(2026, 4, 27).strftime("%d %b %Y")
AUTHOR      = "Ashish Kumar Satyam"
ORG_NAME    = "The Policy Exchange"
ORG_LEGAL   = "Fairvalue Insuretech Pvt. Ltd."
ORG_FULL    = f"{ORG_NAME} ({ORG_LEGAL})"
PARENT_DOC  = "TPE_Communication_System_Charter_v1.0.docx"

# ─────────────────────────────────────────────────────────────────────────────
# XML-level cell-width helpers — identical pattern to archived generators
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_width_xml(cell, width_inches):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(int(width_inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')

def set_table_grid_xml(table, widths_inches):
    tbl = table._tbl
    grid = tbl.find(qn('w:tblGrid'))
    if grid is None:
        grid = OxmlElement('w:tblGrid')
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths_inches:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(int(w * 1440)))
        grid.append(col)

def set_table_width_xml(table, total_inches):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(int(total_inches * 1440)))
    tblW.set(qn('w:type'), 'dxa')

def lock_table_widths(table, widths_inches):
    total = sum(widths_inches)
    set_table_grid_xml(table, widths_inches)
    set_table_width_xml(table, total)
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            if ci < len(widths_inches):
                set_cell_width_xml(cell, widths_inches[ci])

def set_cell_borders(cell, color="BFBFBF", size=4):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(size))
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def style_run(run, *, font="Arial", size=10, bold=False, italic=False, color=DARK):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rFonts.set(qn(attr), font)

def add_para(doc, text="", *, font="Arial", size=10, bold=False, italic=False,
             color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4, indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if text:
        r = p.add_run(text)
        style_run(r, font=font, size=size, bold=bold, italic=italic, color=color)
    return p

def add_bullet(doc, text, *, indent=0.5, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    style_run(r, size=size)
    return p

def add_section_heading(doc, num, title, level=1):
    """All heading bars / colored pointers removed — plain typographic
    hierarchy only. H1 is large bold dark text with a thin underline rule;
    H2 is medium bold navy text; H3 is italic blue."""
    if level == 1:
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
        # Plain typographic heading — no navy bar
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{num}  {title}")
        style_run(run, font="Arial", size=14, bold=True, color=DARK)
        # Thin underline rule using a 1-cell zero-height table
        from docx.enum.table import WD_ROW_HEIGHT_RULE
        rule = doc.add_table(rows=1, cols=1)
        rule.autofit = False
        lock_table_widths(rule, [6.50])
        rc = rule.rows[0].cells[0]
        set_cell_bg(rc, "2C3E50")
        rc.paragraphs[0].text = ""
        rc.paragraphs[0].paragraph_format.space_before = Pt(0)
        rc.paragraphs[0].paragraph_format.space_after = Pt(0)
        rule.rows[0].height = Pt(1)
        rule.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
    elif level == 2:
        # Plain bold navy heading — no left bar
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{num}  {title}")
        style_run(r, size=12, bold=True, color=NAVY)
    else:
        add_para(doc, f"{num}  {title}", size=11, bold=True, italic=True,
                 color=BLUE, space_after=2)

def add_table(doc, headers, rows, *, col_widths_inches=None, header_bg=None,
              header_fg=None, alt_rows=True, font_size=9, code_cols=None):
    """Header row is now plain — no navy fill, no white text. Bold dark text on
    white background, with a slightly heavier bottom border to separate header
    from data. (Per request: 'remove all header pointers from table format'.)"""
    code_cols = code_cols or set()
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = False
    if col_widths_inches:
        lock_table_widths(tbl, col_widths_inches)
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        # No background fill on header row
        set_cell_bg(c, "FFFFFF")
        # Heavier bottom border to keep visual separation from data rows
        set_cell_borders(c, color="2C3E50", size=8)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.paragraphs[0].text = ""
        r = c.paragraphs[0].add_run(h)
        # Bold dark text instead of white-on-navy
        style_run(r, font="Arial", size=font_size, bold=True, color=DARK)
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]
        bg = ALT_ROW if (alt_rows and ri % 2 == 0) else "FFFFFF"
        for ci, val in enumerate(row):
            c = tr.cells[ci]
            set_cell_bg(c, bg)
            set_cell_borders(c)
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            c.paragraphs[0].text = ""
            font = "Consolas" if ci in code_cols else "Arial"
            r = c.paragraphs[0].add_run(str(val))
            style_run(r, font=font, size=font_size, color=DARK)
            c.paragraphs[0].paragraph_format.space_after = Pt(0)
    return tbl

def add_color_band(doc, hex_color, height_pt=4):
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    lock_table_widths(tbl, [6.50])
    c = tbl.rows[0].cells[0]
    set_cell_bg(c, hex_color)
    c.paragraphs[0].text = ""
    c.paragraphs[0].paragraph_format.space_before = Pt(0)
    c.paragraphs[0].paragraph_format.space_after = Pt(0)
    tbl.rows[0].height = Pt(height_pt)
    tbl.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

def add_callout(doc, title, body, kind="blue"):
    color_map = {
        'blue':   (BLUE,   LIGHT),
        'green':  (GREEN,  SUCCESS_BG),
        'orange': (ORANGE, WARN_BG),
        'red':    (RED,    DANGER_BG),
        'navy':   (NAVY,   LIGHT),
    }
    fg, bg = color_map[kind]
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    lock_table_widths(tbl, [6.50])
    c = tbl.rows[0].cells[0]
    set_cell_bg(c, bg)
    set_cell_borders(c, color="BFBFBF")
    c.paragraphs[0].text = ""
    r = c.paragraphs[0].add_run(title)
    style_run(r, size=10, bold=True, color=fg)
    p = c.add_paragraph()
    r2 = p.add_run(body)
    style_run(r2, size=9.5, color=DARK)
    add_para(doc, "", space_after=2)

# ─────────────────────────────────────────────────────────────────────────────
# Page setup
# ─────────────────────────────────────────────────────────────────────────────

def setup_document():
    doc = Document()
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.4)
        section.footer_distance = Inches(0.4)
        h = section.header.paragraphs[0]
        h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rh = h.add_run(f"Charter Addendum v{DOC_VERSION}  |  Companion to {PARENT_DOC}")
        style_run(rh, size=8, color=MEDIUM, italic=True)
        f = section.footer.paragraphs[0]
        f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rf = f.add_run(f"{DOC_ID}  |  Internal & Confidential  |  Page ")
        style_run(rf, size=8, color=MEDIUM)
        run = f.add_run()
        fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText'); instr.text = "PAGE"
        fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_end)
        style_run(run, size=8, color=MEDIUM, bold=True)
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.font.color.rgb = DARK
    return doc

# ─────────────────────────────────────────────────────────────────────────────
# Cover page
# ─────────────────────────────────────────────────────────────────────────────

def add_cover_page(doc):
    add_color_band(doc, HEADER_BG, height_pt=10)
    add_para(doc, "", space_after=18)
    add_para(doc, ORG_FULL.upper(), size=10, color=MEDIUM, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "CHARTER ADDENDUM", size=22, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_para(doc, "TPE Communication System", size=26, bold=True, color=DARK,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(doc, "Tier 1 + Tier 2 gap-fill addendum to the v1.0 Internal Project Charter",
             size=11, italic=True, color=MEDIUM,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    # Snapshot tiles
    tbl = doc.add_table(rows=1, cols=3)
    tbl.autofit = False
    lock_table_widths(tbl, [2.166, 2.166, 2.168])
    tiles = [
        ("10",  "New sections",       "1B3A5C"),
        ("7",   "Tier 1 (critical)",  "E74C3C"),
        ("3",   "Tier 2 (important)", "E67E22"),
    ]
    for i, (n, lbl, col) in enumerate(tiles):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, "FAFBFC")
        set_cell_borders(c, color=col, size=18)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.paragraphs[0].text = ""
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.paragraphs[0].add_run(n)
        style_run(r, size=22, bold=True, color=RGBColor.from_string(col))
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(lbl)
        style_run(r2, size=8.5, color=DARK)
    add_para(doc, "", space_after=24)

    # Metadata table
    add_para(doc, "Document metadata", size=10, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    meta_tbl = doc.add_table(rows=8, cols=2)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_tbl.autofit = False
    lock_table_widths(meta_tbl, [2.10, 3.10])
    meta = [
        ("Document ID",     DOC_ID),
        ("Version",         DOC_VERSION),
        ("Date",            TODAY),
        ("Prepared by",     AUTHOR),
        ("Parent document", PARENT_DOC),
        ("Status",          "Pending Sponsor approval"),
        ("Distribution",    "Same as parent Charter"),
        ("Classification",  "Internal & Confidential"),
    ]
    for ri, (k, v) in enumerate(meta):
        bg = ALT_ROW if ri % 2 == 0 else "FFFFFF"
        c0 = meta_tbl.rows[ri].cells[0]
        c1 = meta_tbl.rows[ri].cells[1]
        for c in (c0, c1):
            set_cell_bg(c, bg)
            set_cell_borders(c, color="DDDDDD")
            c.paragraphs[0].text = ""
        r0 = c0.paragraphs[0].add_run(k)
        style_run(r0, size=9.5, bold=True, color=NAVY)
        r1 = c1.paragraphs[0].add_run(v)
        style_run(r1, size=9.5, color=DARK)
    add_para(doc, "", space_after=18)
    add_color_band(doc, "2E75B6", height_pt=4)
    add_para(doc, "This addendum supplements — does not replace — the v1.0 Charter. "
                  "Once Sponsor-approved, its sections become effective immediately and "
                  "are folded into v1.1 of the Charter at the next regenerated revision.",
             size=9, italic=True, color=MEDIUM,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    doc.add_page_break()

def add_intro(doc):
    add_section_heading(doc, "0", "Purpose & Reading Notes")
    add_para(doc,
        "This addendum closes ten gaps identified during a structured review "
        "of the v1.0 Charter on " + TODAY + ". The Charter is a strong scope "
        "and ownership document, but several enterprise-grade sections were "
        "absent. Rather than reissue the full Charter and disturb its "
        "structure, this addendum is published as a controlled companion: "
        "Sponsor approval here brings each section into force immediately.",
        space_after=8)

    add_para(doc, "How to read this document", size=11, bold=True, color=NAVY,
             space_after=4)
    add_table(doc, ["Section", "Tier", "Why it was missing matters"], [
        ["A1  Effort Estimation (person-days)",       "Critical",  "Charter has 6 calendar weeks but no person-day budget — cannot resource-plan"],
        ["A2  Per-deliverable Acceptance Criteria",   "Critical",  "Phase exit criteria exist but no per-deliverable 'done' definition"],
        ["A3  Service Level Agreements",              "Critical",  "No operational or delivery SLAs — both teams need them"],
        ["A4  Quality Assurance & Testing Strategy",  "Critical",  "Test pack is a deliverable but the strategy is not specified"],
        ["A5  Security & Privacy Framework",          "Critical",  "DPDPA referenced but no PII classification or controls matrix"],
        ["A6  Cost & Budget Estimate",                "Critical",  "No budget envelope — Sponsor cannot approve spend"],
        ["A7  Glossary",                               "Critical",  "DLT, BSP, IRDAI, DPDPA, FATCA undefined for non-technical readers"],
        ["A8  Document Control & Approval Matrix",    "Important", "Sign-off block exists but no formal state machine"],
        ["A9  Communication Management Plan",         "Important", "How does the project itself communicate"],
        ["A10 Operational Readiness & Go-Live Checklist","Important","Cutover mentioned but no go/no-go criteria"],
    ], col_widths_inches=[2.85, 0.95, 2.70], font_size=9)

    add_callout(doc, "Source of truth precedence",
        "Where this addendum and the v1.0 Charter both speak to a topic, the "
        "addendum prevails for the items it explicitly covers. The Charter "
        "remains authoritative for Sections 1–10 of the original. On approval "
        "of this addendum and subsequent regeneration into Charter v1.1, this "
        "precedence rule is dissolved.", kind="navy")
    doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# Tier 1 Sections
# ─────────────────────────────────────────────────────────────────────────────

def section_a1_effort(doc):
    add_section_heading(doc, "A1", "Effort Estimation (Person-Days)")
    add_para(doc,
        "The Charter's six calendar weeks are made up of the following "
        "person-day allocations. Engineering is expressed as full-time "
        "equivalents (FTE × days); reviewer time is calendar-distributed.",
        space_after=6)

    add_section_heading(doc, "A1.1", "Engineering effort by phase", level=2)
    add_table(doc, ["Phase", "Calendar weeks", "Backend Eng (×2 FTE)", "Infra/DevOps (×0.5)", "QA (×0.5)", "Total person-days"], [
        ["P1  Foundation",                                 "Week 1",   "10",  "2.5",  "2.5",  "15.0"],
        ["P2  Policyholder Workflows",                      "Week 2",   "10",  "2.5",  "2.5",  "15.0"],
        ["P3  Investor / Insurer / Ops / Reg + HI + Multi-tenant", "Weeks 3–4", "20",  "5.0",  "5.0",  "30.0"],
        ["P4  Compliance Middleware + Observability",      "Week 5",   "10",  "2.5",  "2.5",  "15.0"],
        ["P5  Custom Providers + Production Hardening",    "Week 6",   "10",  "2.5",  "2.5",  "15.0"],
        ["Hypercare",                                       "Week 7",   "5",   "2.5",  "1.0",  "8.5"],
        ["TOTAL",                                           "7 weeks",   "65",  "17.5", "16.0", "98.5"],
    ], col_widths_inches=[2.20, 0.90, 1.05, 1.00, 0.65, 0.70], font_size=8.5)

    add_section_heading(doc, "A1.2", "Reviewer effort (CX, Compliance, Brand, Operations)", level=2)
    add_table(doc, ["Reviewer", "Allocation", "Total person-days", "When"], [
        ["CX Lead",           "~4h/week × 7 weeks",  "3.5",  "Continuous content review"],
        ["Compliance Lead",   "~4h/week × 7 weeks",  "3.5",  "Continuous lint review + DLT/IRDAI/BSP sign-off"],
        ["Brand Reviewer",    "~2h/week × 5 weeks",  "1.25", "P2–P3 voice + tone review"],
        ["Operations Lead",   "~2 days total",        "2.0",  "Insurer registry inputs + cadence rules"],
        ["Infrastructure Lead","~3 days total",       "3.0",  "AWS account, IAM, observability hookup"],
        ["Sponsor/Owner",     "Phase-gate × 5 + escalations", "2.5", "Phase-gate sign-offs"],
        ["TOTAL reviewer effort", "—",                 "15.75", "Distributed across 7 weeks"],
    ], col_widths_inches=[1.55, 1.85, 1.20, 1.90], font_size=9)

    add_section_heading(doc, "A1.3", "Headline numbers", level=2)
    add_table(doc, ["Metric", "Value"], [
        ["Total engineering person-days (P1–P5)",  "90.0"],
        ["Hypercare person-days",                   "8.5"],
        ["Reviewer person-days",                    "15.75"],
        ["GRAND TOTAL person-days",                 "114.25"],
        ["Calendar duration (build + hypercare)",   "7 weeks"],
        ["FTE-equivalent for build phase",          "≈ 3.0 FTE for 6 weeks"],
    ], col_widths_inches=[3.50, 3.00], font_size=9.5)

    add_callout(doc, "How to read these numbers",
        "Engineering person-days assume two backend engineers full-time, one "
        "infra engineer at 50%, and one QA engineer at 50%. If the team "
        "composition differs, multiply or divide accordingly. Reviewer time "
        "is non-substitutable — Compliance review cannot be back-filled by "
        "anyone outside the Compliance function.", kind="navy")

def section_a2_acceptance(doc):
    add_section_heading(doc, "A2", "Per-deliverable Acceptance Criteria")
    add_para(doc,
        "Each of the ten deliverables in Charter §5 is accepted only when "
        "the criteria below are demonstrated to the Project Owner. The "
        "criteria are written as observable behaviour, not as project "
        "milestones — so they survive any later schedule changes.",
        space_after=6)

    add_table(doc, ["#", "Deliverable", "Acceptance criteria"], [
        ["1", "Self-hosted Novu deployment (Multi-AZ, ap-south-1)",
         "Novu Web reachable via SSO; data tier in Multi-AZ; documented RPO ≤ 24h / RTO ≤ 4h; first PoC workflow execution succeeds end-to-end; IaC checked into repo."],
        ["2", "49 workflow definitions",
         "All 49 IDs deployed in production environment; each has at least one successful execution against a representative subscriber; trigger ID stable; per-trigger detail (5 axes) documented."],
        ["3", "Custom provider plugins (MSG91, Karix) + Gupshup config",
         "All 3 providers reachable; failover Gupshup → MSG91 → Karix tested end-to-end; sandbox dry-run produces provider-valid payload; receipt ingestion verified."],
        ["4", "Compliance middleware (≥ 30 rules across 7 regimes)",
         "≥ 30 lint rules active; CI gate blocks PR on rule violation; pre-execution gate blocks send on rule violation; rule definitions externalised as YAML; override path logged."],
        ["5", "Multi-tenant insurer registry (11 carriers)",
         "Each carrier has sender identity, MIS template, contact, escalation, brand-voice; tenant resolution at workflow execution; INS-02 reconciles 100% in first month."],
        ["6", "Observability stack: logs + metrics + audit trail",
         "Structured logs queryable; SLA metrics dashboard live; per-message audit row written with PII tokenised; subject-access export under 60 seconds for any subscriber."],
        ["7", "One-time subscriber data sync",
         "All in-scope PH and INV subscribers loaded; diff report produced and reconciled; non-conforming rows surfaced with action; idempotent re-run supported."],
        ["8", "Documentation set (5 documents)",
         "Workflow author guide, infra runbook, compliance reference card, provider plugin docs, operations runbook — all in repo, reviewed, and tagged to the release."],
        ["9", "Test pack: schema, lint, render, execution, E2E",
         "≥ 30 schema tests; ≥ 30 lint rules covered; render goldens for all (audience × channel × language) pairs; 49 workflow execution tests; 5 E2E tests covering the highest-volume triggers."],
        ["10", "Production cutover plan + 1 week hypercare",
         "Cutover plan signed by Project Owner; staging cutover successful before production; 1 week of named on-call; zero open Sev-1 at hypercare exit."],
    ], col_widths_inches=[0.35, 2.20, 3.95], font_size=8.5)

    add_callout(doc, "Acceptance is binary",
        "A deliverable either meets all of its acceptance criteria or it "
        "does not. There is no 'partial accept' — partial work returns to "
        "the build team with a documented gap list. The Project Owner has "
        "five business days to accept or return after demonstration; "
        "silence past five days is implicit acceptance.", kind="green")

def section_a3_sla(doc):
    add_section_heading(doc, "A3", "Service Level Agreements")
    add_para(doc,
        "Two SLA sets: operational SLAs the delivered system commits to "
        "post-launch, and delivery SLAs the build team commits to during "
        "the project. Both are tracked from go-live and during build "
        "respectively; breach handling is in §A3.3.",
        space_after=6)

    add_section_heading(doc, "A3.1", "Operational SLAs (post go-live)", level=2)
    add_table(doc, ["Service", "Metric", "Target", "Measurement"], [
        ["Workflow approval queue",   "Time from PR open to merge",          "≤ 24h business day",            "GitHub PR metrics"],
        ["Lint engine availability",  "Uptime (rolling 30 days)",             "99.5%",                         "CI metrics"],
        ["Render correctness",        "Sev-1 incidents per quarter",          "0",                              "Production observation"],
        ["OTP delivery",              "P95 delivery latency",                  "≤ 10 seconds",                   "Provider receipts"],
        ["Premium-due dispatch",      "Send-window adherence",                 "100% within DLT 9–21 IST window","DLT compliance report"],
        ["Audit-trail export",        "Time to produce per-template audit",    "≤ 60 seconds",                   "Operational test"],
        ["Subject-access response",   "DPDPA SAR turnaround",                  "≤ 7 calendar days",              "Incident log"],
        ["Failover (provider)",       "Time to switch provider on 5xx",        "≤ 30 seconds",                   "Workflow execution log"],
    ], col_widths_inches=[2.00, 2.00, 1.50, 1.00], font_size=9)

    add_section_heading(doc, "A3.2", "Delivery SLAs (during build)", level=2)
    add_table(doc, ["Activity", "Target SLA", "Owner"], [
        ["Phase-gate review pack ready",         "T-2 business days before review",  "Engineering Lead"],
        ["Compliance review of workflow PR",      "≤ 24h business day",                "Compliance Lead"],
        ["Brand review of workflow content",      "≤ 24h business day",                "CX Lead"],
        ["Reviewer-blocked CR triage",            "Same business day",                  "Engineering Lead"],
        ["Sev-1 defect response (in build)",      "1 hour",                             "Engineering"],
        ["Weekly status report publication",      "Friday 16:00 IST",                   "Engineering Lead"],
        ["Change Request decision (Material+)",   "Within 5 business days of raised",   "Project Owner"],
    ], col_widths_inches=[2.85, 2.05, 1.60], font_size=9)

    add_section_heading(doc, "A3.3", "Breach handling", level=2)
    for b in [
        "First breach in a rolling 30-day window: documented in weekly status, root-cause noted, no further escalation.",
        "Second breach in 30 days: Engineering Lead escalates to Project Owner; mitigation plan due within 3 business days.",
        "Three or more breaches in 30 days: phase plan re-baselined if material; surfaced to Sponsor.",
        "Breaches are visible — no breach is concealed by re-defining the metric.",
        "SLAs are subject to formal review at each phase gate; revisions follow the Change Management process (Charter §8).",
    ]:
        add_bullet(doc, b)

def section_a4_qa_testing(doc):
    add_section_heading(doc, "A4", "Quality Assurance & Testing Strategy")
    add_para(doc,
        "Quality is engineered, not reviewed in. Five test levels each "
        "catch a different defect class. The test pack itself (Charter "
        "Deliverable #9) is the substrate; the strategy below is how the "
        "pack is grown and operated.",
        space_after=6)

    add_section_heading(doc, "A4.1", "The five test levels", level=2)
    add_table(doc, ["Level", "Catches", "Example", "Run when"], [
        ["L1  Schema",
         "Malformed workflow definitions, missing required fields",
         "Workflow with no audience tag fails parse",
         "Pre-commit + CI"],
        ["L2  Lint",
         "Compliance violations — DLT, IRDAI, WA-BSP, DPDPA, brand-voice",
         "WhatsApp Marketing template without opt-out",
         "Pre-commit + workflow pre-execution"],
        ["L3  Render",
         "Channel-renderer regressions: char limits, button schema, encoding",
         "SMS body that exceeds 160 chars after variable substitution",
         "CI on every push"],
        ["L4  Workflow execution",
         "Workflow logic regressions — branches, delays, fan-out",
         "Cadence step skips a tier under specific subscriber data",
         "CI nightly + before phase-gate"],
        ["L5  End-to-end",
         "Full execution: trigger → workflow → lint → render → provider sandbox",
         "PH-03 KYC week-2 produces provider-valid Gupshup payload",
         "Phase-gate + before any production release"],
    ], col_widths_inches=[1.05, 2.10, 2.05, 1.30], font_size=8.5, code_cols={0})

    add_section_heading(doc, "A4.2", "Test pack composition (per phase)", level=2)
    add_table(doc, ["Phase", "Pack composition target"], [
        ["P1", "≥ 20 schema tests + 5 lint rules + 2 render goldens + 1 E2E for the PoC trigger"],
        ["P2", "+ 1 workflow execution test per PH trigger; render goldens for Email + WA × all PH active statuses"],
        ["P3", "+ INV/INS/OPS/REG workflow execution; Hindi-language render goldens; 11-tenant routing test"],
        ["P4", "+ ≥ 30 lint rules across 7 regimes; pre-execution gate enforcement test; audit-trail export test"],
        ["P5", "+ MSG91 / Karix sandbox E2E goldens; failover test; receipt ingestion test"],
    ], col_widths_inches=[0.55, 5.95], font_size=9, code_cols={0})

    add_section_heading(doc, "A4.3", "Defect severity & SLA", level=2)
    add_table(doc, ["Severity", "Definition", "Response SLA", "Resolution SLA"], [
        ["Sev-1", "Compliance violation in active workflow OR data exposure",
                  "1 hour",                  "Same business day; rollback if needed"],
        ["Sev-2", "Active workflow broken (cannot execute) OR active trigger silent",
                  "4 hours",                 "1 business day"],
        ["Sev-3", "Active workflow flawed but still producing output (typo, mild wording)",
                  "1 business day",          "1 sprint"],
        ["Sev-4", "Cosmetic / docs / non-active workflow",
                  "1 sprint",                "Backlog"],
    ], col_widths_inches=[0.85, 3.00, 1.30, 1.35], font_size=9)

def section_a5_security(doc):
    add_section_heading(doc, "A5", "Security & Privacy Framework")
    add_para(doc,
        "The Charter calls out DPDPA in passing. This section provides the "
        "data classification, controls matrix and DPDPA-principle alignment "
        "the Compliance Lead needs for sign-off.",
        space_after=6)

    add_section_heading(doc, "A5.1", "Data classification", level=2)
    add_table(doc, ["Class", "Examples in this system", "Handling rule"], [
        ["PII — High",
         "ph_name, ph_phone, ph_email, inv_name, PAN, bank_last4",
         "Tokenise in workflow inputs; never log raw; mask in audit; HTTPS only in transit"],
        ["PII — Medium",
         "policy_no, premium_amount, due_date, surrender_value",
         "Encrypted at rest; access on least-privilege; redact in screenshots"],
        ["Operational",
         "trigger_id, workflow_id, status, cadence",
         "Public within company; safe to log"],
        ["Brand confidential",
         "Disclaimer text, brand voice rules, insurer-specific cards",
         "Internal-only; not for external comms"],
    ], col_widths_inches=[1.30, 3.20, 2.00], font_size=9)

    add_section_heading(doc, "A5.2", "Controls matrix", level=2)
    add_table(doc, ["Control area", "Control"], [
        ["Access control",       "Role-based: author / reviewer / operator / admin; least-privilege; reviewed quarterly"],
        ["Authentication",        "SSO via TPE identity provider; MFA mandatory for reviewer/admin/infra"],
        ["Authorisation gates",   "Workflow PR merge requires Compliance + Brand reviewer; no self-merge for active workflows"],
        ["Encryption — transit",  "TLS 1.2+ for all network calls; certificate-pinned for BSP / SMS gateway"],
        ["Encryption — rest",     "AES-256 for repository storage; key rotation per TPE security policy"],
        ["Secret management",     "Provider API keys in vault; never in repo; rotated per provider policy"],
        ["Audit log",              "Append-only; per-workflow execution + state transitions + sends; PII masked"],
        ["PII minimisation",      "Subscriber data referenced by ID at workflow runtime, not stored in audit log"],
        ["Backup & recovery",     "Daily snapshot; documented RPO ≤ 24h / RTO ≤ 4h within ap-south-1"],
        ["Vulnerability mgmt",    "Quarterly dependency scan; security review at every major version"],
        ["Network segmentation",  "Novu in private subnets; no direct public ingress; egress allow-listed"],
    ], col_widths_inches=[1.65, 4.85], font_size=9)

    add_section_heading(doc, "A5.3", "DPDPA principle alignment", level=2)
    add_table(doc, ["DPDPA principle", "How this system aligns"], [
        ["Lawful processing",         "Workflows carry compliance_class; lint blocks promotional sends without consent reference"],
        ["Purpose limitation",        "Trigger taxonomy makes intent explicit; workflows linked to triggers; reuse is auditable"],
        ["Data minimisation",          "Variables declare what is needed; lint flags unused PII variables"],
        ["Storage limitation",         "Audit retains state transitions, not raw PII; backups retained per data-retention policy"],
        ["Accuracy",                    "Variables source from system of record; fallbacks prevent broken renders"],
        ["Notice & consent",            "Opt-out wiring on every promotional template; opt-out evidence stored at BSP"],
        ["Rights of data principals",  "Subject-access response supported via audit-trail export; deletion on request"],
        ["Breach notification",         "Incident runbook covers detection, escalation, regulator notification within statutory window"],
    ], col_widths_inches=[1.85, 4.65], font_size=9)

def section_a6_budget(doc):
    add_section_heading(doc, "A6", "Cost & Budget Estimate")
    add_para(doc,
        "Indicative budget envelope built up from §A1 effort estimation "
        "plus infrastructure plus channel-provider pricing. Numbers are "
        "ranges, not commitments; they exist so the Sponsor can validate "
        "the engagement is the size that was assumed.",
        space_after=6)

    add_section_heading(doc, "A6.1", "Internal effort cost (build + hypercare)", level=2)
    add_table(doc, ["Cost driver", "Quantum", "Indicative range (₹)"], [
        ["Engineering — backend (×2 FTE × 6 weeks)",  "60 person-days",  "₹6L – ₹12L (range depends on band)"],
        ["Engineering — infra/DevOps (×0.5 × 6 weeks)","15 person-days",  "₹1.5L – ₹3L"],
        ["Engineering — QA (×0.5 × 6 weeks)",          "15 person-days",  "₹1.2L – ₹2.5L"],
        ["Hypercare (×1 FTE × 1 week)",                "8.5 person-days", "₹0.8L – ₹1.7L"],
        ["Reviewer time (CX + Compliance + Brand + Ops)", "~13 person-days","Internal — embedded cost"],
        ["Sponsor / Owner phase-gate time",             "~2.5 person-days","Internal — embedded cost"],
        ["INTERNAL TOTAL",                                "112 person-days","≈ ₹9.5L – ₹19.2L"],
    ], col_widths_inches=[3.10, 1.45, 1.95], font_size=9)

    add_section_heading(doc, "A6.2", "Infrastructure cost (annualised)", level=2)
    add_table(doc, ["Infrastructure item", "Indicative monthly", "Annualised"], [
        ["AWS compute — Novu services (multi-AZ ap-south-1)", "₹35K – ₹60K",   "₹4.2L – ₹7.2L"],
        ["AWS managed services — RDS, Redis, S3, CloudWatch","₹25K – ₹45K",   "₹3.0L – ₹5.4L"],
        ["Backup + cross-AZ + observability tooling",         "₹10K – ₹15K",   "₹1.2L – ₹1.8L"],
        ["INFRA SUBTOTAL (steady state)",                      "₹70K – ₹120K",  "₹8.4L – ₹14.4L"],
    ], col_widths_inches=[3.30, 1.55, 1.65], font_size=9)

    add_section_heading(doc, "A6.3", "Channel provider cost (volume-driven)", level=2)
    add_table(doc, ["Provider", "Channel", "Pricing model", "Indicative monthly (steady)"], [
        ["Gupshup",            "WhatsApp + SMS + RCS",  "Per-message + monthly minimums",   "₹1.0L – ₹2.0L"],
        ["MSG91",              "SMS + WhatsApp + Email", "Per-message",                       "₹0.5L – ₹1.0L (failover)"],
        ["Karix",              "SMS + WhatsApp",         "Per-message",                       "₹0.3L – ₹0.7L (failover)"],
        ["AWS SES (Email)",    "Email transactional",    "Per-1000 emails",                   "₹0.05L – ₹0.15L"],
        ["PROVIDER SUBTOTAL",  "—",                       "—",                                  "₹1.85L – ₹3.85L / month"],
    ], col_widths_inches=[1.30, 1.60, 1.85, 1.75], font_size=9)

    add_section_heading(doc, "A6.4", "Total cost summary", level=2)
    add_table(doc, ["Bucket", "One-time", "Annualised steady state"], [
        ["Internal effort (build + hypercare)",  "₹9.5L – ₹19.2L",     "—"],
        ["Infrastructure (AWS)",                  "—",                    "₹8.4L – ₹14.4L"],
        ["Channel providers",                      "—",                    "₹22.2L – ₹46.2L"],
        ["TOTAL",                                  "₹9.5L – ₹19.2L (one-time)", "₹30.6L – ₹60.6L /year"],
    ], col_widths_inches=[2.85, 1.85, 1.80], font_size=9.5)

    add_callout(doc, "What these numbers do and do not include",
        "Included: engineering effort, AWS steady-state, channel-provider "
        "transactional pricing at projected steady-state volume. Excluded: "
        "vendor procurement legal time, DLT registration fees (one-time, "
        "small), WA BSP onboarding fees (vendor-specific), penetration "
        "testing or formal security audit (out of scope per Charter §4.12). "
        "Volume-sensitive providers should be re-baselined six months "
        "post-launch against actual send volume.", kind="orange")

def section_a7_glossary(doc):
    add_section_heading(doc, "A7", "Glossary")
    add_para(doc,
        "Terms that appear in the Charter without definition. Listed "
        "alphabetically. Charter Section §1 onwards uses these terms freely "
        "— this glossary is the reference any non-technical reader will "
        "need.",
        space_after=6)
    add_table(doc, ["Term", "Definition"], [
        ["Assignment",          "Legal transfer of policy ownership from policyholder (assignor) to investor (assignee)"],
        ["BSP",                  "Business Solution Provider — Meta-approved partner for WhatsApp Business Cloud API (Gupshup, Karix, etc.)"],
        ["Charter",              "This document — internal scope and ownership document for the project"],
        ["CKYC",                 "Central KYC Registry — India-wide KYC repository to avoid repeating KYC across financial institutions"],
        ["DLT",                  "Distributed Ledger Technology — TRAI-mandated registry for SMS sender headers and templates in India"],
        ["DPDPA",                "Digital Personal Data Protection Act, 2023 — India's data-protection law"],
        ["DPDPA SAR",            "Subject-Access Request under DPDPA — a data principal's right to retrieve their data"],
        ["FATCA",                "Foreign Account Tax Compliance Act — US-driven; applies to NRI / foreign-asset declarations on investor accounts"],
        ["FTE",                  "Full-Time Equivalent — one person-day of one full-time worker"],
        ["Hypercare",            "Period of intensified support immediately after go-live, typically 1 week, with named on-call"],
        ["IRDAI",                "Insurance Regulatory and Development Authority of India"],
        ["MIS",                  "Management Information System — in this Charter: the per-insurer monthly Excel file mapping premium remittances to policies"],
        ["MIS Pipeline",         "Existing TPE project (TPE-MIS-PLAN-2026-001) for ingesting insurer allocation files"],
        ["NOC",                  "No-Objection Certificate from the insurer for an assignment to proceed"],
        ["Novu",                 "Open-source notification framework (https://novu.co); used self-hosted as the workflow runtime"],
        ["Novu Framework",        "Novu's TypeScript SDK for defining workflows-as-code"],
        ["OTP",                  "One-Time Passcode — short-lived authentication credential delivered via SMS / WA / Email"],
        ["PE ID",                "Principal Entity ID — DLT registration identifier for TPE on TRAI's registry"],
        ["PII",                  "Personally Identifiable Information — protected under DPDPA"],
        ["RCS",                  "Rich Communication Services — successor to SMS supporting buttons, carousels, branded sender"],
        ["RPO / RTO",            "Recovery Point Objective / Recovery Time Objective — data-loss tolerance / restoration time targets"],
        ["SAR",                  "Subject-Access Request (see DPDPA SAR)"],
        ["Surrender value",      "Cash value paid by the insurer if the policyholder discontinues the policy mid-term"],
        ["TDS",                  "Tax Deducted at Source — Indian income-tax withholding mechanism"],
        ["TRAI",                 "Telecom Regulatory Authority of India — owns the DLT regime"],
        ["Trigger",              "A defined lifecycle event that initiates one or more communications; one trigger maps to one Novu workflow"],
        ["Workflow",             "A Novu workflow definition — TypeScript expression of statuses, channels, cadence and content"],
    ], col_widths_inches=[1.65, 4.85], font_size=9)

# ─────────────────────────────────────────────────────────────────────────────
# Tier 2 Sections
# ─────────────────────────────────────────────────────────────────────────────

def section_a8_doc_control(doc):
    add_section_heading(doc, "A8", "Document Control & Approval Matrix")
    add_para(doc,
        "Formal control of the Charter (and of this addendum) as enterprise "
        "documents. Closes the gap of 'sign-off block at the end with no "
        "front-of-document approval matrix'.",
        space_after=6)

    add_section_heading(doc, "A8.1", "Approval matrix", level=2)
    add_table(doc, ["Role", "Approval scope", "Approval mechanism"], [
        ["Executive Sponsor",   "Commercial intent, budget envelope, sign-off",
                                 "Wet/electronic signature on Charter §10"],
        ["Project Owner (CTO)", "Architecture, scope, exit criteria, day-to-day",
                                 "Electronic signature; phase-gate review"],
        ["Compliance Lead",     "Regulatory scope (DLT, IRDAI, DPDPA, FATCA, TDS)",
                                 "Electronic signature; lint-rule sign-off"],
        ["CX Lead",              "Audience voice, brand sign-off on customer-facing content",
                                 "Electronic signature; review queue"],
        ["Engineering Lead",    "Plan integrity, weekly status, technical hires",
                                 "Owns this document; cuts new versions"],
    ], col_widths_inches=[1.85, 2.95, 1.70], font_size=9)

    add_section_heading(doc, "A8.2", "Document state machine", level=2)
    add_table(doc, ["State", "Meaning", "Next states"], [
        ["draft",         "Author iterating; not yet circulated",                 "in_review, archived"],
        ["in_review",     "Circulated to approvers; comments open",                "approved, draft (returned)"],
        ["approved",      "All approvers signed; document is the project bible",   "amended (new version), archived"],
        ["amended",       "Substantive change in flight; previous version retained","approved, archived"],
        ["archived",      "Superseded; read-only; retained for audit",              "(terminal)"],
    ], col_widths_inches=[1.10, 3.50, 1.90], font_size=9)

    add_section_heading(doc, "A8.3", "Change control on the Charter", level=2)
    for b in [
        "MAJOR version (x.0) — change in scope, deliverables, effort, budget. Requires re-approval by all roles in the matrix.",
        "MINOR version (1.x) — clarification, additional detail, error correction. Approval by Project Owner and one other approver.",
        "PATCH version (1.1.x) — typographical or formatting only. Engineering Lead authority alone.",
        "On approval of this addendum, the Charter regenerates as v1.1 and this addendum moves to archived.",
    ]:
        add_bullet(doc, b)

def section_a9_comms(doc):
    add_section_heading(doc, "A9", "Communication Management Plan")
    add_para(doc,
        "How the project itself communicates — distinct from the "
        "communication system being built. Defined cadence, owners and "
        "channels prevent the project from generating its own coordination "
        "drag.",
        space_after=6)
    add_table(doc, ["Forum", "Frequency", "Owner", "Audience", "Output"], [
        ["Daily standup",          "Daily 10:30 IST",       "Engineering Lead", "Engineering",        "Daily update + blockers"],
        ["Weekly status",           "Friday 16:00 IST",      "Engineering Lead", "All approvers",      "1-pager + Slack"],
        ["Phase-gate review",       "End of P1, P2, P3, P4, P5", "Engineering Lead", "All approvers",  "Live demo + sign-off"],
        ["Compliance review queue", "Continuous",            "Compliance Lead",   "Authors",            "PR approvals"],
        ["Brand review queue",      "Continuous",            "CX Lead",           "Authors",            "PR approvals"],
        ["Risk review",             "Bi-weekly",             "Engineering Lead", "Project Owner + Sponsor", "Updated risk log"],
        ["Stakeholder digest",      "Monthly",               "Engineering Lead", "All stakeholders",   "Email"],
        ["Post-mortem",             "Per Sev-1/Sev-2",       "Engineering Lead", "All approvers",      "Markdown doc"],
    ], col_widths_inches=[1.55, 1.50, 1.10, 1.20, 1.15], font_size=8.5)

    add_callout(doc, "Information radiator, not push",
        "Status is published to one place (the Slack channel + the weekly "
        "1-pager). Approvers pull when they need to; the Engineering Lead "
        "does not chase or duplicate. Push is reserved for blockers and "
        "phase-gate decisions.", kind="navy")

def section_a10_ops_readiness(doc):
    add_section_heading(doc, "A10", "Operational Readiness & Go-Live Checklist")
    add_para(doc,
        "Production cutover is mentioned in Charter §4.11. This section "
        "defines the go/no-go criteria. A phase that does not meet its "
        "criteria does not transition; the cause is documented and reset.",
        space_after=6)

    add_section_heading(doc, "A10.1", "Per-phase readiness checklist", level=2)
    add_table(doc, ["Dimension", "P1", "P2", "P3", "P4", "P5"], [
        ["Schema validation passes",            "✓","✓","✓","✓","✓"],
        ["Lint coverage adequate",              "—","Partial","Partial","✓","✓"],
        ["Render goldens up-to-date",           "✓","✓","✓","✓","✓"],
        ["Workflow execution tests pass",       "—","✓","✓","✓","✓"],
        ["E2E for representative trigger",       "✓","✓","✓","✓","✓"],
        ["Audit-trail export tested",            "—","—","—","✓","✓"],
        ["Compliance Lead signed",                "—","✓","✓","✓","✓"],
        ["CX Lead signed",                        "—","✓","✓","—","—"],
        ["Project Owner signed phase gate",      "✓","✓","✓","✓","✓"],
        ["Provider sandbox tests (P5 only)",      "—","—","—","—","Required"],
    ], col_widths_inches=[3.50, 0.60, 0.60, 0.60, 0.60, 0.60], font_size=8.5)

    add_section_heading(doc, "A10.2", "System-wide go-live criteria", level=2)
    for b in [
        "All P1–P5 acceptance criteria (§A2) signed by named approvers.",
        "Risk register shows zero open Sev-1; Sev-2 risks have documented mitigations.",
        "Incident runbook tested in tabletop with engineering + compliance.",
        "Rollback procedure documented and tested for the most-recently-changed workflow family.",
        "Subject-access request (DPDPA SAR) end-to-end tested for at least one investor and one policyholder.",
        "DLT registration, WhatsApp BSP template approval and channel-provider sandbox all green.",
        "Communication Management Plan transitions to BAU operations cadence.",
        "On-call rota agreed for the 1-week hypercare period; named primary + secondary.",
    ]:
        add_bullet(doc, b)

    add_callout(doc, "What 'go-live' means here",
        "The Charter's value is realised at production cutover — the moment "
        "the team stops drafting in spreadsheets and starts authoring in "
        "Novu. Hypercare is one week of named on-call; after hypercare exit "
        "with zero open Sev-1, the project closes and ownership transitions "
        "to BAU.", kind="green")

# ─────────────────────────────────────────────────────────────────────────────
# Sign-off page
# ─────────────────────────────────────────────────────────────────────────────

def section_addendum_signoff(doc):
    add_section_heading(doc, "A11", "Addendum Sign-off")
    add_para(doc,
        "Signature below brings sections A1–A10 into force as supplements "
        "to the v1.0 Charter. On Sponsor approval, the Charter is "
        "regenerated as v1.1 with these sections folded in, and this "
        "addendum moves to archived.",
        space_after=8)
    add_table(doc, ["Role", "Name", "Signature", "Date"], [
        ["Executive Sponsor",        "", "", ""],
        ["Project Owner (CTO)",       "", "", ""],
        ["Compliance Lead",          "", "", ""],
        ["CX Lead",                   "", "", ""],
        ["Engineering Lead (Author)", AUTHOR, "", TODAY],
    ], col_widths_inches=[1.90, 2.10, 1.45, 1.05], font_size=9.5)

# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

def main():
    doc = setup_document()

    add_cover_page(doc)
    add_intro(doc)

    # Tier 1
    section_a1_effort(doc)
    doc.add_page_break()
    section_a2_acceptance(doc)
    doc.add_page_break()
    section_a3_sla(doc)
    doc.add_page_break()
    section_a4_qa_testing(doc)
    doc.add_page_break()
    section_a5_security(doc)
    doc.add_page_break()
    section_a6_budget(doc)
    doc.add_page_break()
    section_a7_glossary(doc)
    doc.add_page_break()

    # Tier 2
    section_a8_doc_control(doc)
    doc.add_page_break()
    section_a9_comms(doc)
    doc.add_page_break()
    section_a10_ops_readiness(doc)
    doc.add_page_break()

    # Sign-off
    section_addendum_signoff(doc)

    out = ("/Users/aksatyam/TPE_DIRECTORY/TPE - Communication System/docs/"
           "TPE_Communication_System_Charter_Addendum_v1.1.docx")
    doc.save(out)
    print(f"Saved: {out}  (v{DOC_VERSION} — 10 gap-fill sections)")

if __name__ == "__main__":
    main()
