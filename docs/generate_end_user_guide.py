#!/usr/bin/env python3
"""
Enterprise End-User Guide generator — TPE Communication System.

Output:  docs/TPE_Communication_System_End_User_Guide_v1.0.docx

Follows the CLAUDE.md enterprise standards:
  - Arial body + Consolas for code; US Letter, 1" margins
  - Navy (#1B3A5C) header rows + alternating row shading
  - Status badges using `●` (Word-compatible, NEVER emoji)
  - Classification "Internal / Confidential" in red
  - Cover page + footer attribution to Fairvalue Insuretech (The Policy Exchange)
"""
from __future__ import annotations
import datetime as _dt
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Brand palette + page setup
# ---------------------------------------------------------------------------
NAVY       = RGBColor(0x1B, 0x3A, 0x5C)
BLUE       = RGBColor(0x2E, 0x75, 0xB6)
GREEN      = RGBColor(0x27, 0xAE, 0x60)
ORANGE     = RGBColor(0xE6, 0x7E, 0x22)
RED        = RGBColor(0xE7, 0x4C, 0x3C)
DARK       = RGBColor(0x2C, 0x3E, 0x50)
MEDIUM     = RGBColor(0x7F, 0x8C, 0x8D)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG   = "F0F4F8"
ALT_ROW    = "F8F9FA"
SUCCESS_BG = "E8F5E9"
WARN_BG    = "FFF8E1"
DANGER_BG  = "FFEBEE"

DOC_TITLE     = "End User Guide"
DOC_SUBTITLE  = "TPE Communication System — Operator Manual"
DOC_ID        = "TPE-COMMS-EUG-2026-001"
DOC_VERSION   = "1.0"
DOC_DATE      = _dt.date.today().strftime("%d %b %Y")
CLASSIF       = "Internal / Confidential"
ORG_FOR       = "The Policy Exchange"
ORG_LEGAL     = "Fairvalue Insuretech Pvt. Ltd."
ORG_FULL      = f"{ORG_FOR} ({ORG_LEGAL})"
AUTHOR        = "Ashish Kumar Satyam"

OUTPUT = Path(__file__).resolve().parent / \
    f"TPE_Communication_System_End_User_Guide_v{DOC_VERSION}.docx"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, hex_color: str) -> None:
    """Apply solid background color to a table cell. ShadingType.CLEAR per CLAUDE.md."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color: str = "B0B0B0", size: str = "4") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def style_run(run, font="Arial", size=11, bold=False, italic=False,
              color: RGBColor | None = None):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    # Force eastAsia + cs to same font for Word's odd font-fallback behavior
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)
    rFonts.set(qn("w:eastAsia"), font)


def add_p(doc, text: str = "", *, font="Arial", size=11, bold=False,
          italic=False, color=None, align=None, space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if text:
        r = p.add_run(text)
        style_run(r, font=font, size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading_styled(doc, text: str, level: int = 1):
    sizes  = {1: 18, 2: 14, 3: 12}
    afters = {1: 8,  2: 6,  3: 4}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(afters.get(level, 4))
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    style_run(r, font="Arial", size=sizes.get(level, 11), bold=True, color=NAVY)
    return p


def add_code_block(doc, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.15)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "C0C0C0")
        pBdr.append(b)
    pPr.append(pBdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)
    for i, line in enumerate(code.split("\n")):
        if i:
            p.add_run().add_break(WD_BREAK.LINE)
        r = p.add_run(line)
        style_run(r, font="Consolas", size=9, color=DARK)


def add_callout(doc, label: str, body: str, *, tone: str = "info"):
    """Coloured note/warning/tip box."""
    palette = {
        "info":    (BLUE,   "EBF3FB"),
        "warn":    (ORANGE, WARN_BG),
        "danger":  (RED,    DANGER_BG),
        "success": (GREEN,  SUCCESS_BG),
    }[tone]
    border_color, fill_hex = palette
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.5)
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(6.5)
    set_cell_shading(cell, fill_hex)
    set_cell_borders(cell, color="C0C0C0", size="6")
    cell.paragraphs[0].paragraph_format.space_before = Pt(4)
    cell.paragraphs[0].paragraph_format.space_after = Pt(4)
    r = cell.paragraphs[0].add_run(label + " ")
    style_run(r, font="Arial", size=10, bold=True, color=border_color)
    r2 = cell.paragraphs[0].add_run(body)
    style_run(r2, font="Arial", size=10, color=DARK)
    add_p(doc, "", space_after=6)


def add_status_badge(paragraph, text: str, tone: str = "ok"):
    color = {"ok": GREEN, "warn": ORANGE, "danger": RED, "muted": MEDIUM}[tone]
    r = paragraph.add_run("● ")
    style_run(r, font="Arial", size=10, bold=True, color=color)
    r2 = paragraph.add_run(text)
    style_run(r2, font="Arial", size=10, color=color)


def add_table_styled(doc, headers, rows, *, col_widths=None, header_color=NAVY,
                     alt_row=True):
    """Header row navy + white text; alternating body row shading."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            tbl.columns[i].width = Inches(w)
    # Header
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        if col_widths:
            cell.width = Inches(col_widths[i])
        set_cell_shading(cell, f"{header_color.rgb if hasattr(header_color, 'rgb') else '%02X%02X%02X' % tuple(header_color)}")
        set_cell_borders(cell, color="FFFFFF", size="4")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(h)
        style_run(r, font="Arial", size=10, bold=True, color=WHITE)
    # Body rows
    for ridx, row in enumerate(rows):
        is_alt = alt_row and (ridx % 2 == 1)
        for cidx, val in enumerate(row):
            cell = tbl.rows[ridx + 1].cells[cidx]
            if col_widths:
                cell.width = Inches(col_widths[cidx])
            if is_alt:
                set_cell_shading(cell, ALT_ROW)
            set_cell_borders(cell, color="DDDDDD", size="4")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(val))
            style_run(r, font="Arial", size=10, color=DARK)
    return tbl


def add_bullet(doc, text: str, *, indent=0.25):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("•  ")
    style_run(r, font="Arial", size=11, bold=True, color=NAVY)
    r2 = p.add_run(text)
    style_run(r2, font="Arial", size=11, color=DARK)


def add_numbered(doc, items, *, indent=0.25):
    for i, t in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(indent)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{i}.  ")
        style_run(r, font="Arial", size=11, bold=True, color=NAVY)
        r2 = p.add_run(t)
        style_run(r2, font="Arial", size=11, color=DARK)


# ---------------------------------------------------------------------------
# COVER PAGE
# ---------------------------------------------------------------------------

def cover_page(doc):
    # Top navy banner with classification
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.5)
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(6.5)
    set_cell_shading(cell, "1B3A5C")
    cell.paragraphs[0].paragraph_format.space_before = Pt(8)
    cell.paragraphs[0].paragraph_format.space_after = Pt(8)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cell.paragraphs[0].add_run(ORG_FOR.upper())
    style_run(r, font="Arial", size=14, bold=True, color=WHITE)

    # Classification banner
    p = add_p(doc, CLASSIFICATION_LINE, bold=True, color=RED,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    p.paragraph_format.space_before = Pt(12)

    # Big title block
    add_p(doc, DOC_TITLE.upper(), font="Arial", size=36, bold=True, color=NAVY,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_p(doc, DOC_SUBTITLE, font="Arial", size=14, italic=True, color=DARK,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    # A thin navy separator
    sep = doc.add_table(rows=1, cols=1)
    sep.columns[0].width = Inches(2.5)
    sep.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = sep.rows[0].cells[0]
    cell.width = Inches(2.5)
    set_cell_shading(cell, "1B3A5C")
    cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    r = cell.paragraphs[0].add_run(" ")
    style_run(r, font="Arial", size=2)

    add_p(doc, "", space_after=18)

    # Metadata table
    meta_rows = [
        ("Document ID",       DOC_ID),
        ("Version",           DOC_VERSION),
        ("Date",              DOC_DATE),
        ("Classification",    CLASSIFICATION_LINE),
        ("Author",            AUTHOR),
        ("Prepared For",      ORG_FULL),
        ("Audience",          "Operations / Customer-Care / Marketing — operators of the TPE Comms platform"),
    ]
    add_table_styled(doc, ["Field", "Value"], meta_rows,
                     col_widths=[1.6, 4.9], alt_row=True)

    # Footer block (Navy)
    add_p(doc, "", space_after=24)
    foot = doc.add_table(rows=2, cols=1)
    foot.autofit = False
    foot.columns[0].width = Inches(6.5)
    for r_idx in range(2):
        cell = foot.rows[r_idx].cells[0]
        cell.width = Inches(6.5)
        set_cell_shading(cell, "1B3A5C")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].paragraph_format.space_before = Pt(4 if r_idx == 0 else 2)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2 if r_idx == 0 else 4)
    r = foot.rows[0].cells[0].paragraphs[0].add_run(f"Prepared by {AUTHOR}")
    style_run(r, font="Arial", size=12, bold=True, color=WHITE)
    r = foot.rows[1].cells[0].paragraphs[0].add_run(
        f"For {ORG_FULL}  •  www.thepolicyexchange.com  •  {DOC_DATE}")
    style_run(r, font="Arial", size=10, color=WHITE)

    doc.add_page_break()


CLASSIFICATION_LINE = CLASSIF


# ---------------------------------------------------------------------------
# DOCUMENT BODY
# ---------------------------------------------------------------------------

def section_doc_control(doc):
    add_heading_styled(doc, "1.  Document Control", level=1)

    add_p(doc, "Version history", bold=True, color=NAVY, space_after=4)
    add_table_styled(doc, ["Version", "Date", "Author", "Summary of changes"],
        [
            ("1.0", DOC_DATE, AUTHOR,
             "Initial publication. Covers TPE Admin Console (Phase 1A–1C) + "
             "Novu CE 2.3.0 Studio dashboard authoring."),
        ],
        col_widths=[0.7, 1.1, 1.6, 3.1])

    add_p(doc, "", space_after=6)
    add_p(doc, "Document owner & distribution", bold=True, color=NAVY, space_after=4)
    add_table_styled(doc, ["Role", "Name / team"],
        [
            ("Document Owner",          "Engineering Lead, The Policy Exchange"),
            ("Reviewer (Compliance)",   "Compliance Lead, The Policy Exchange"),
            ("Reviewer (CX)",           "Customer Experience Lead, The Policy Exchange"),
            ("Approver",                "Project Owner / CTO, The Policy Exchange"),
            ("Distribution",            "Operations, Customer-Care, Marketing operators "
                                         "of the TPE Comms platform; copy on file with the "
                                         "Project Owner"),
            ("Classification",          CLASSIFICATION_LINE),
            ("Retention",               "Living document — superseded versions archived "
                                         "for 7 years per Charter §4.8"),
        ],
        col_widths=[1.8, 4.7])
    doc.add_page_break()


def section_intro(doc):
    add_heading_styled(doc, "2.  Introduction", level=1)

    add_heading_styled(doc, "2.1  What this guide is", level=2)
    add_p(doc,
          "The TPE Communication System is a self-hosted notification platform that "
          "lets operators send transactional and operational notifications across "
          "in-app, SMS, WhatsApp, and email — driven either by individual triggers "
          "or by data-selected campaigns from the analytics database. This guide is "
          "the operator-facing manual: it walks through every screen and action in "
          "the TPE Admin Console and the Novu CE dashboard, end to end.")

    add_heading_styled(doc, "2.2  Audience", level=2)
    add_p(doc,
          "This guide is written for non-developer operators in Operations, "
          "Customer-Care, and Marketing who will run the platform day-to-day. It "
          "assumes you can use a web browser and read a table; it does not assume "
          "any programming knowledge. Any task that requires code changes is "
          "explicitly called out and routed to Engineering.")

    add_heading_styled(doc, "2.3  Prerequisites", level=2)
    add_bullet(doc, "A laptop or desktop with a modern browser (Chrome, Edge, Firefox, Safari).")
    add_bullet(doc, "Network access to the staging hostname (currently the IP address 103.138.96.180; DNS will land later).")
    add_bullet(doc, "Login credentials for the TPE Admin Console (operator password) and the Novu CE Dashboard (operator account).")
    add_bullet(doc, "For email-touching workflows: confirmation that AWS SES credentials have been provisioned by Engineering. Until then, email steps gracefully skip.")

    add_heading_styled(doc, "2.4  What this guide does NOT cover", level=2)
    add_bullet(doc, "Server administration, container restarts, or any task requiring SSH — those live in the Combined Deployment Guide.")
    add_bullet(doc, "Authoring Bridge (TypeScript) workflows — that is engineering-owned and documented separately in bridge/workflows/AUTHORING.md.")
    add_bullet(doc, "Integration setup (creating MSG91 / ICPaaS / SES accounts, registering DLT templates) — those one-time setups live in the Solution Document and the Provider Strategy memo.")
    doc.add_page_break()


def section_overview(doc):
    add_heading_styled(doc, "3.  System Overview", level=1)
    add_p(doc,
          "The platform has two operator-facing surfaces. Most day-to-day work is "
          "in the TPE Admin Console; advanced workflow authoring and integration "
          "configuration is in the Novu CE Dashboard. Both run on the same server.")

    add_heading_styled(doc, "3.1  The two surfaces", level=2)
    add_table_styled(doc,
        ["Surface", "URL", "What you do here", "Login"],
        [
            ("TPE Admin Console", "http://103.138.96.180/admin",
             "Channel allowlist, operator email/in-app templates, audience "
             "campaigns from analytics_db, scheduled triggers, activity history.",
             "Single shared password (issued by Engineering)."),
            ("Novu CE Dashboard", "http://103.138.96.180/",
             "Workflow authoring (Studio), integration providers, subscribers "
             "list, environment settings, API keys.",
             "Operator account (issued by Engineering)."),
        ],
        col_widths=[1.5, 2.0, 2.0, 1.0])

    add_heading_styled(doc, "3.2  How they fit together", level=2)
    add_code_block(doc,
        "(operator)\n"
        "   │\n"
        "   ├─ Novu CE Dashboard ─── authors a Workflow → registers it on the API\n"
        "   │                        wires Integrations (e.g. Email = SES)\n"
        "   │                        manages Subscribers (recipients)\n"
        "   │\n"
        "   └─ TPE Admin Console ─── builds a Campaign that targets a Workflow\n"
        "                            picks a data slice (analytics_db tables)\n"
        "                            schedules / fires it\n"
        "                            audits past runs\n"
        "\n"
        "Behind the scenes:\n"
        "   campaign fire → /v1/events/trigger on Novu API\n"
        "                 → Novu Worker picks up\n"
        "                 → executes workflow steps (in-app / email / SMS / WhatsApp)\n"
        "                 → audit row written to bridge stdout + Mongo")

    add_heading_styled(doc, "3.3  Roles & responsibilities", level=2)
    add_table_styled(doc,
        ["Role", "Owns these screens / actions"],
        [
            ("Operations",            "/admin/channels (allowlist), /admin/schedules (one-off triggers), Novu Subscribers"),
            ("Customer-Care",         "Replies in /admin/history (audit), retry-failed-rows on campaigns"),
            ("Marketing",             "/admin/campaigns (audience-driven sends), /admin/templates (email + in-app HTML)"),
            ("Engineering",           "Workflow authoring on Novu CE Dashboard, integration setup, server administration"),
            ("Compliance Lead",       "Approves new templates, audits the activity feed, owns Charter §4.8 retention"),
        ],
        col_widths=[1.4, 5.1])
    doc.add_page_break()


def section_access(doc):
    add_heading_styled(doc, "4.  Access & Login", level=1)

    add_heading_styled(doc, "4.1  Logging into the TPE Admin Console", level=2)
    add_numbered(doc, [
        "Open http://103.138.96.180/admin in a browser. The page shows a single "
        "password field on a card titled “TPE Admin”.",
        "Type the operator password (issued by Engineering — never circulate over "
        "chat). Click Sign in.",
        "On success the browser lands on /admin (the Operator Dashboard). A "
        "session cookie keeps you logged in for 12 hours; after that, sign in "
        "again.",
        "Sign out via the operator avatar (top-right) → Sign out.",
    ])

    add_callout(doc, "Reminder.",
        "The admin password is shared across all operators in this MVP. Do not commit it to chat, "
        "tickets, or screenshots. If you suspect compromise, ask Engineering to rotate the "
        "TPE_ADMIN_PASSWORD environment variable — every active session is invalidated when the "
        "password changes (the session token is derived from the password).",
        tone="warn")

    add_heading_styled(doc, "4.2  Logging into the Novu CE Dashboard", level=2)
    add_numbered(doc, [
        "Open http://103.138.96.180/ — the Novu sign-in page loads.",
        "Enter the operator email + password issued to you. There is no public "
        "sign-up on this instance — accounts are created by Engineering on "
        "request.",
        "On success you land on the Workflows tab of the Development environment. "
        "The environment switcher (top-left, below the org name) lets you flip "
        "between Development and Production (when Production is provisioned).",
    ])

    add_heading_styled(doc, "4.3  Session timeouts & multi-tab use", level=2)
    add_bullet(doc, "Both consoles are safe to open in multiple browser tabs at once. State is shared via cookies.")
    add_bullet(doc, "If a Save or Fire button suddenly stops working with a 401 error, your session has expired. Refresh the page to be redirected to the sign-in screen.")
    add_bullet(doc, "Always sign out when leaving a shared workstation.")
    doc.add_page_break()


def section_workflow_authoring(doc):
    add_heading_styled(doc, "5.  Authoring a Notification Workflow", level=1)
    add_p(doc,
          "A workflow is the recipe for a notification — it defines which channels "
          "fire (in-app / email / SMS / WhatsApp), in what order, and with what "
          "content. Workflows are authored in the Novu CE Dashboard and triggered "
          "from anywhere (manually via the API, via a TPE Admin campaign, or via "
          "an integration).")

    add_heading_styled(doc, "5.1  Step-by-step: your first workflow", level=2)
    add_numbered(doc, [
        "Open http://103.138.96.180/ and sign in.",
        "Click Workflows in the sidebar → click + New Workflow (top-right). "
        "A dialog asks for a name.",
        "Enter a short slug (lowercase, hyphenated): e.g. welcome-test-2026. "
        "Click Create. You land in the workflow editor.",
        "Drag In-App from the Channels panel onto the canvas. Set: "
        "Subject = “Welcome {{firstName}}”, Body = “Hi {{firstName}}, your "
        "account is live. Visit the portal at https://thepolicyexchange.com/.”. "
        "Save the step.",
        "Drag Email from the Channels panel below the In-App step. Set: "
        "Subject = “Welcome to TPE”, Body = a short HTML message. Save.",
        "Click the Settings tab on the workflow header — confirm the Notification "
        "Group (defaults to “General”) and the trigger key.",
        "Click Save (top-right).",
        "From the Workflows list, toggle the workflow to Active (green dot). "
        "Inactive workflows return 400 when triggered.",
    ])

    add_callout(doc, "Variables in workflow content.",
        "Anywhere you can type, you can use {{var}} placeholders. The values come from the trigger's "
        "payload field. Example: {{firstName}} pulls from payload.firstName. The TPE Admin Campaign "
        "builder fills these from your Postgres data via the Field Map step.",
        tone="info")

    add_heading_styled(doc, "5.2  Channel selection guidance", level=2)
    add_table_styled(doc,
        ["Channel", "Status on stage", "Use when..."],
        [
            ("In-App",   "● Working out of the box",
             "Always — confirms the trigger fired even when external channels are unconfigured."),
            ("Email",    "● Available; SES creds pending",
             "Long-form content (statements, NOCs, decision letters). Until SES creds land, the step gracefully skips."),
            ("SMS",      "● Working via MSG91 (Bridge workflows only)",
             "Short transactional notices (OTPs, payment due). Studio-authored workflows do NOT have an MSG91 provider — coordinate with Engineering."),
            ("WhatsApp", "● Working via ICPaaS (Bridge workflows only)",
             "Engagement / nudges with rich formatting. Same caveat — Studio workflows can't reach ICPaaS."),
        ],
        col_widths=[1.0, 2.4, 3.1])

    add_heading_styled(doc, "5.3  Verifying the workflow appears in TPE Admin", level=2)
    add_numbered(doc, [
        "Open http://103.138.96.180/admin in another tab.",
        "On the Operator Dashboard, the “Workflows deployed” metric should be one "
        "higher than before. If it isn't, the workflow either (a) isn't Active or "
        "(b) the API key is misconfigured — open a ticket.",
        "Open /admin/campaigns/new → in the Target workflow dropdown, your new "
        "workflow should appear. This proves it is reachable from the Admin "
        "Console for campaign-driven sends.",
    ])
    doc.add_page_break()


def section_channels(doc):
    add_heading_styled(doc, "6.  Channel Allowlist (Safe-Mode Switch)", level=1)
    add_p(doc,
          "The /admin/channels page controls whether each of the four channels "
          "actually fires. It is the single biggest safety net on the platform — "
          "a power switch that lets you disable channels per environment without "
          "touching workflow code.")

    add_heading_styled(doc, "6.1  Why this exists", level=2)
    add_p(doc,
          "Sometimes you want to test a workflow end-to-end without messaging real "
          "customers. Sometimes a provider has an outage and you want all SMS sends "
          "to skip cleanly until the issue is fixed. Sometimes a regulator asks you "
          "to demonstrate that no messages will be sent during a specific window. "
          "All three use the same control.")

    add_heading_styled(doc, "6.2  How it works", level=2)
    add_numbered(doc, [
        "Open /admin/channels — four checkboxes (SMS, WhatsApp, Email, In-app) "
        "with a status badge next to each.",
        "Untick any channel → click Save. The change persists to the "
        "tpe_channel_gating Mongo collection within milliseconds.",
        "The Bridge container picks up the change within ~10 seconds (cache TTL). "
        "From that moment, every dispatch on a disabled channel returns "
        "status=skipped with reason “channel ... disabled in /admin/channels” "
        "and the audit row records it.",
        "Workflows continue to run normally; only the channel-step that's "
        "disabled is skipped. Other channel-steps in the same workflow still "
        "fire if their channel is enabled.",
    ])

    add_heading_styled(doc, "6.3  Recommended defaults by environment", level=2)
    add_table_styled(doc,
        ["Channel", "Stage / Dev", "Production (when launched)"],
        [
            ("SMS",      "● On — MSG91 v5 Flow with DLT-registered templates",
                         "● On"),
            ("WhatsApp", "● On — ICPaaS with Meta-approved templates",
                         "● On"),
            ("Email",    "● Off — SES creds not yet provisioned",
                         "● On (once SES production access is granted)"),
            ("In-app",   "● Off on Stage Dev — keeps the inbox clean during testing",
                         "● On"),
        ],
        col_widths=[1.0, 3.0, 2.5])

    add_callout(doc, "Tip.",
        "When testing a new campaign for the first time, untick ALL four channels, fire-now once, "
        "inspect the audit history (every row should show status=skipped), then re-enable the "
        "channels you want and fire for real.",
        tone="info")
    doc.add_page_break()


def section_templates(doc):
    add_heading_styled(doc, "7.  Operator-Managed Templates", level=1)
    add_p(doc,
          "The /admin/templates page stores email + in-app HTML snippets that "
          "Bridge workflows pull at dispatch time. This is independent of the "
          "templates you author inside the Novu Dashboard's workflow editor; it "
          "exists for cases where Marketing wants to update copy without touching "
          "the workflow code.")

    add_heading_styled(doc, "7.1  When to use admin-managed templates", level=2)
    add_bullet(doc, "Marketing updates email body copy without an Engineering deployment.")
    add_bullet(doc, "Same template needs to be reused across multiple workflows.")
    add_bullet(doc, "A/B testing — flip a single template's content and the change reaches every workflow that pulls it.")
    add_bullet(doc, "When NOT to use: SMS bodies (DLT requires registered templates) and WhatsApp bodies (Meta requires approved templates) are managed in MSG91 + Meta — not here.")

    add_heading_styled(doc, "7.2  Step-by-step: create a template", level=2)
    add_numbered(doc, [
        "Open /admin/templates → click + New template.",
        "Fill in: Name (lowercase-hyphenated, the workflow looks it up by this), "
        "Channel (Email or In-app), Language (English or Hindi — register two "
        "templates if you need both), Subject, HTML body, optional Text body.",
        "Click Save. The template is now available to any Bridge workflow that "
        "queries it by name.",
        "Edit by clicking the pencil icon. Delete by clicking the trash icon — "
        "operators are prompted to confirm; deleted templates are NOT recoverable.",
    ])
    doc.add_page_break()


def section_campaigns(doc):
    add_heading_styled(doc, "8.  Audience-Driven Campaigns", level=1)
    add_p(doc,
          "A campaign is a saved instruction to fire a workflow against a slice "
          "of the analytics database. It encapsulates: the source table, a filter, "
          "the target workflow, and a mapping from row columns to the workflow's "
          "subscriber identity + payload. This is the central feature of the TPE "
          "Admin Console.")

    add_heading_styled(doc, "8.1  Anatomy of a campaign", level=2)
    add_table_styled(doc,
        ["Field", "Meaning"],
        [
            ("Source table",
             "Which Postgres table to read. Curated allowlist: tbl_allocation "
             "(60,946 rows of customer + contact), data_allocation, "
             "tbl_lapsedata_apr (lapsed policies), final_table, policy_data."),
            ("Filter",
             "Optional WHERE clauses. Each filter is column + operator + value. "
             "Operators allowed: = != > >= < <= ILIKE IS NULL IS NOT NULL."),
            ("Target workflow",
             "The Novu workflow to fire once per matching row."),
            ("Subscriber ID column",
             "Which row column becomes the Novu subscriberId. Convention is "
             "POLICY_NUMBER for tbl_allocation. The actual subscriberId is "
             "namespaced as <table>_<value> (e.g. tbl_allocation_20356001) so "
             "rows from different tables don't collide."),
            ("Email / Phone / Name columns",
             "Optional. When mapped, the trigger will create or update the Novu "
             "subscriber with these values. Empty values skip that field."),
            ("Payload mapping",
             "Map workflow payload fields to source columns. e.g., "
             "policyNumber → POLICY_NUMBER. Anywhere your workflow uses "
             "{{policyNumber}} will pull from the row at fire time."),
            ("Schedule",
             "Manual fire (default; you press Fire when ready) OR Schedule for "
             "(operator-picked datetime; the admin-worker auto-fires within "
             "30 seconds of due time)."),
        ],
        col_widths=[2.0, 4.5])

    add_heading_styled(doc, "8.2  Step-by-step: build & fire a campaign", level=2)
    add_numbered(doc, [
        "Open /admin/campaigns → click + New campaign. The builder page opens "
        "with a left panel (configuration) and a right panel (preview).",
        "Section 1 — Source table: pick a table from the dropdown. Suggested "
        "starting point is tbl_allocation. Approximate row counts are shown.",
        "Section 2 — Filter audience: click + Add filter. Pick a column, an "
        "operator, and a value. Add multiple filters to narrow further. Click "
        "Refresh preview when ready.",
        "Watch the right panel: it shows the matching row count and the first "
        "five rows. The SQL Generated query card (collapsible) shows the literal "
        "SQL that will run — useful for sanity-checking complex filters.",
        "Section 3 — Target workflow: pick from the dropdown (populated from "
        "Novu's /v1/workflows). If your workflow isn't there, it isn't Active.",
        "Section 4 — Map row → trigger: confirm the Subscriber ID column "
        "(suggestions pre-fill per table). Map Email / Phone / First name "
        "columns when present. Add Payload mapping rows: each one binds a "
        "workflow payload field to a row column.",
        "Section 5 — Save campaign: type a Name, optional Description, choose "
        "Manual fire OR Schedule for, and click Save campaign.",
        "You're back on /admin/campaigns. To fire immediately, click the ▶ "
        "Play button on your campaign's row. To wait for a scheduled time, "
        "the admin-worker will fire it automatically.",
    ])

    add_callout(doc, "Safety note.",
        "Before firing for the first time, set the channel allowlist (/admin/channels) to all-OFF, "
        "fire-now, then inspect the audit. Every row should appear with status=skipped and reason "
        "“channel ... disabled in /admin/channels”. Once verified, re-enable the channels you want "
        "and fire for real.",
        tone="warn")

    add_heading_styled(doc, "8.3  Reading the fire result", level=2)
    add_bullet(doc, "Toast banner: “Fired N of M” — N is the number of rows that successfully triggered, M is the total matching rows.")
    add_bullet(doc, "Status badge on the campaign row: Fired (all sent), Partial (some errored), Failed (all errored), Empty audience (filter matched 0 rows), or Fired (all skipped).")
    add_bullet(doc, "Click the campaign name → opens /admin/campaigns/[id] — the Run history page with a row per fire, drill-down per subscriberId, retry-failed action.")

    add_heading_styled(doc, "8.4  Throttling & caps", level=2)
    add_p(doc,
          "Each fire is throttled to about 10 rows per second and capped at 200 "
          "rows per single run. This protects the providers (MSG91 / Meta) from "
          "sudden bursts. If your filter matches more than 200 rows, only the "
          "first 200 fire — narrow the filter further or wait for Phase 2 (which "
          "will move the runner into a background queue).")
    doc.add_page_break()


def section_schedules(doc):
    add_heading_styled(doc, "9.  Scheduling One-Off Triggers", level=1)
    add_p(doc,
          "When you need to fire a single workflow at a specific future time for a "
          "single subscriber — without going through the full campaign flow — use "
          "/admin/schedules. This is the simpler primitive: one workflow, one "
          "subscriber, one time.")

    add_heading_styled(doc, "9.1  When to use a schedule vs a campaign", level=2)
    add_table_styled(doc,
        ["Use a Schedule", "Use a Campaign"],
        [
            ("One subscriber, one workflow, one fire time",
             "Many rows from a Postgres table, all firing the same workflow"),
            ("You know the subscriberId already",
             "Subscriber identity is derived from a data row's column"),
            ("Quick ad-hoc tests",
             "Recurring or scheduled marketing sends"),
            ("No data lookup needed",
             "Filter criteria + per-row content from analytics_db"),
        ],
        col_widths=[3.25, 3.25])

    add_heading_styled(doc, "9.2  Step-by-step", level=2)
    add_numbered(doc, [
        "Open /admin/schedules → click + Schedule trigger.",
        "Select Workflow from the dropdown.",
        "Type the Subscriber ID — must already exist in Novu OR will auto-create "
        "on first fire.",
        "Pick Fire at — datetime-local input, your local timezone.",
        "Optional: edit the Payload JSON — defaults to a stub with a fresh "
        "triggerInstanceId.",
        "Click Schedule. The row appears in the table with status pending.",
        "The admin-worker polls the tpe_scheduled_triggers Mongo collection "
        "every 30 seconds. Once your fireAt time has passed, the worker claims "
        "the row, fires the trigger, and updates status to fired (or failed).",
        "To cancel before fire: click the X icon on a pending row — confirmation "
        "dialog appears.",
    ])
    doc.add_page_break()


def section_history(doc):
    add_heading_styled(doc, "10.  Activity History (Audit)", level=1)
    add_p(doc,
          "Every trigger fired through the platform produces audit rows in two "
          "places: Novu's own activity feed (visible at /admin/history) and the "
          "Bridge's structured stdout log (consumed by the Engineering log "
          "aggregator). Operators read /admin/history; Compliance reviews the "
          "structured logs for the formal audit trail.")

    add_heading_styled(doc, "10.1  Reading the /admin/history page", level=2)
    add_p(doc,
          "Open /admin/history. The table shows the most recent triggers (capped "
          "at the latest batch returned by Novu's API). Each row shows:")
    add_bullet(doc, "Time — when Novu received the trigger (with relative “3m ago” hint).")
    add_bullet(doc, "Workflow — the workflowId.")
    add_bullet(doc, "Subscriber — the subscriberId.")
    add_bullet(doc, "transactionId — Novu's correlation key, used to look up the trigger in any other system.")
    add_bullet(doc, "Channels — colored chips for the channels the workflow used.")

    add_heading_styled(doc, "10.2  Filtering", level=2)
    add_p(doc,
          "The search box at the top filters by workflowId, subscriberId, or "
          "transactionId. Type any substring; matches highlight in real time. "
          "Click the Refresh button to fetch the latest batch.")

    add_heading_styled(doc, "10.3  Per-channel skip reasons", level=2)
    add_p(doc,
          "When a channel-step is skipped (rather than sent), the audit row in "
          "the structured log carries a `reason` field. Common values:")
    add_table_styled(doc,
        ["Reason", "What it means", "Action"],
        [
            ("channel '<x>' disabled in /admin/channels",
             "The operator allowlist on /admin/channels has the channel unticked.",
             "If unintended: re-enable on /admin/channels and re-fire."),
            ("SesNotConfigured: missing env ...",
             "AWS SES credentials are not yet provisioned for this environment.",
             "Engineering must add the SES env vars; once added, re-fire."),
            ("MSG91_OTP_TEMPLATE_ID not configured",
             "The DLT template ID for this trigger is not yet wired.",
             "Engineering must register the template with MSG91 and add the env var."),
            ("ICPaaS template not configured",
             "Same as above for WhatsApp via ICPaaS.",
             "Operations must approve the Meta template, then Engineering wires it."),
            ("no mobile number on subscriber",
             "The subscriber doesn't have a phone field set.",
             "Update the subscriber via Novu Dashboard or the campaign's field map."),
            ("subscriberIdColumn empty for this row",
             "The campaign row's subscriber-ID column was null/empty.",
             "Tighten the campaign filter to exclude rows with null IDs."),
        ],
        col_widths=[2.2, 2.4, 1.9])
    doc.add_page_break()


def section_retry(doc):
    add_heading_styled(doc, "11.  Retrying Failed Sends", level=1)
    add_p(doc,
          "When a campaign fan-out has errored rows — most commonly transient "
          "provider errors (MSG91 5xx, Meta rate limits) — operators can retry "
          "JUST those rows without re-firing the entire audience.")

    add_heading_styled(doc, "11.1  When to retry vs investigate", level=2)
    add_bullet(doc, "Retry when: errored rows are a small fraction (< 10%) AND the reason looks transient (HTTP 5xx, network errors, rate limits).")
    add_bullet(doc, "Investigate first when: every row errored (the workflow itself has a problem), the same subscriberId errors repeatedly, or the reason is a known config issue (missing template, missing creds).")

    add_heading_styled(doc, "11.2  Step-by-step", level=2)
    add_numbered(doc, [
        "From /admin/campaigns, click the campaign name to open its Run history "
        "page.",
        "Find the run that has an Errored count > 0. Expand it (click the row).",
        "The per-row drill-down shows each row's status. Click the Errored "
        "filter to show only failed entries.",
        "Click the Retry button on that run (top-right of the row). The platform "
        "extracts only the errored subscriberIds and fires them again — same "
        "workflow, same payloads, only those rows.",
        "A new run appears in Run history with triggeredBy = retry-of-<runId>. "
        "Outcome is reported as a toast and reflected in the run's row counts.",
    ])

    add_callout(doc, "Cap.",
        "Retry caps at 200 errored rows per attempt. If you have more than 200 errors, retry once "
        "(handles the first 200), then retry again (handles the next 200), etc. Each retry is logged "
        "as its own run.",
        tone="info")
    doc.add_page_break()


def section_provider_notes(doc):
    add_heading_styled(doc, "12.  Provider-Specific Notes", level=1)

    add_heading_styled(doc, "12.1  SMS via MSG91", level=2)
    add_bullet(doc, "Sender ID: TPolEx (DLT-registered, Promotional + Transactional categories).")
    add_bullet(doc, "Endpoint: MSG91 v5 Flow API (https://control.msg91.com/api/v5/flow/).")
    add_bullet(doc, "Templates: every workflow needs its own DLT-approved template (registered with the DLT registry, approved by the operator). The template ID is stored in an env var per trigger; Engineering wires this once.")
    add_bullet(doc, "Variables: the template uses {#var#} numeric slots; the workflow's vars: { var1, var2, ... } object maps to them. Mismatches cause silent drops at MSG91 — always verify the first send.")
    add_bullet(doc, "DLT: registration takes 2–24 hours per template; expect rejections for ambiguous wording.")

    add_heading_styled(doc, "12.2  WhatsApp via ICPaaS", level=2)
    add_bullet(doc, "Phone-number ID: 919953061191037 (TPE's Meta-verified WhatsApp Business number).")
    add_bullet(doc, "Endpoint: http://www.icpaas.in/v23.0/.../messages with single Bearer auth.")
    add_bullet(doc, "Templates: every workflow needs a Meta-approved template (UTILITY category for transactional). Operations submits via Meta Business Manager → ICPaaS BSP → Engineering wires the templateName.")
    add_bullet(doc, "Approval lead time: 24–72 hours. Common rejection reasons: marketing language in UTILITY, missing opt-out, ambiguous variables.")

    add_heading_styled(doc, "12.3  Email via AWS SES", level=2)
    add_bullet(doc, "Region: ap-south-1 (Mumbai) per Charter §4.1 India-residency rule.")
    add_bullet(doc, "Sender: no-reply@thepolicyexchange.com (verified SES identity).")
    add_bullet(doc, "Status: production credentials are not yet provisioned — every email step gracefully skips with reason “SesNotConfigured”. Once Engineering adds the EMAIL_SES_REGION + access keys to the environment, no operator action is required; emails start firing on the next workflow run.")
    add_bullet(doc, "Ramp-up: SES sandbox limits 200 emails per day until production access is requested.")

    add_heading_styled(doc, "12.4  In-App via Novu Inbox", level=2)
    add_bullet(doc, "Always works. No external provider; messages are stored in Novu's own Mongo and rendered by the <Inbox/> widget on the customer-facing portal.")
    add_bullet(doc, "Application Identifier (X0pp0lKPr8sb on stage) ties the widget to the dashboard.")
    add_bullet(doc, "Ideal first channel for any new workflow — confirms the trigger fired even when external channels are unconfigured.")
    doc.add_page_break()


def section_compliance(doc):
    add_heading_styled(doc, "13.  Compliance & Operational Constraints", level=1)
    add_p(doc,
          "This section summarises the compliance frame the platform operates "
          "under. The full controls and audit retention policies are in the "
          "Charter §4.8 and the Compliance Constraints memo.")

    add_heading_styled(doc, "13.1  Data residency (DPDPA)", level=2)
    add_bullet(doc, "All primary data stores are in India (Mumbai region — Hosting World Pvt Ltd VPS, AWS ap-south-1).")
    add_bullet(doc, "No Novu Cloud, no US/EU managed services for primary stores. This is a hard rule — never override.")
    add_bullet(doc, "Cross-border data movement requires Compliance Lead sign-off + a Data Processing Agreement.")

    add_heading_styled(doc, "13.2  Quiet hours", level=2)
    add_bullet(doc, "Default: 21:00–09:00 IST quiet hours suppress non-urgent reminders.")
    add_bullet(doc, "OTP, payment-due, and similar urgent triggers bypass quiet hours by design.")
    add_bullet(doc, "Quiet-hours enforcement is currently a Phase 2 deliverable — workflows are coded to respect the policy but the runtime gate lands when the quiet-hours middleware ships.")

    add_heading_styled(doc, "13.3  Audit retention (IRDAI / Income Tax / Insurance Act)", level=2)
    add_bullet(doc, "Per-message audit row retained 12 months hot, 5 years warm, 8 years cold.")
    add_bullet(doc, "Subject-access export (a customer asks for their data) returns within 60 seconds — see /admin/history for the operator-facing entry point.")
    add_bullet(doc, "PII is tokenised in audit rows; no raw email / phone numbers in the audit log.")

    add_heading_styled(doc, "13.4  IRDAI message-content policy", level=2)
    add_bullet(doc, "Insurance-related communications must include the regulatory disclaimer required by IRDAI for the relevant product category. Marketing must verify each template against the current IRDAI guidance before submission.")
    add_bullet(doc, "Compliance Lead sign-off is required on every regulatory (REG-*) trigger template body.")

    add_heading_styled(doc, "13.5  Opt-outs", level=2)
    add_bullet(doc, "Every transactional SMS / WhatsApp / email surface includes the platform-standard opt-out language per provider rules (DLT for SMS, Meta for WhatsApp, SES for email).")
    add_bullet(doc, "Subscriber-level opt-outs are honoured in dispatch — when a subscriber has opt-out flags set, the channel-step skips with reason “opt-out”.")
    doc.add_page_break()


def section_troubleshoot(doc):
    add_heading_styled(doc, "14.  Troubleshooting & FAQ", level=1)

    add_heading_styled(doc, "14.1  Common questions", level=2)
    faqs = [
        ("My campaign fired but I see status=skipped on every row.",
         "Check /admin/channels — at least one channel is unticked. The skip "
         "reason in the audit will say “channel ... disabled in /admin/channels”. "
         "Re-enable the channel(s) and fire again."),
        ("Email steps say SesNotConfigured.",
         "AWS SES credentials are not yet provisioned for this environment. "
         "Open a ticket with Engineering. No operator action needed once Engineering "
         "adds the env vars — the next fire will pick up the new credentials."),
        ("My campaign matched 0 rows.",
         "Open the SQL preview card on /admin/campaigns/new and inspect the WHERE "
         "clause. Common pitfalls: case-sensitivity (POLICY_NUMBER ≠ policy_number), "
         "leading/trailing spaces in IDs, NULL values that ILIKE doesn't match. "
         "Use IS NULL / IS NOT NULL filters explicitly."),
        ("The dashboard shows “Workflows deployed: 0”.",
         "There are no Active workflows. Either author one in the Novu Dashboard "
         "or activate an existing one. The TPE Admin Console only sees Active "
         "workflows."),
        ("How do I add another operator?",
         "TPE Admin: there is no per-user account in this MVP — everyone shares "
         "the operator password. To rotate, ask Engineering to update the "
         "TPE_ADMIN_PASSWORD environment variable. Novu Dashboard: ask Engineering "
         "to create the user via the Novu API or the dashboard's user-invite flow "
         "if enabled."),
        ("My session keeps expiring.",
         "TPE Admin sessions are 12 hours. After that, sign in again. Note: "
         "rotating the operator password invalidates every active session."),
        ("Can I edit a workflow after publishing?",
         "Yes — open the workflow in the Novu Dashboard, edit, save. The change "
         "takes effect on the next trigger. (Novu CE 3.x adds an explicit "
         "Publish-changes step; on the current 2.3.0 stage edits go live "
         "immediately.)"),
        ("How do I send a test message to myself?",
         "Add yourself as a subscriber in the Novu Dashboard (Subscribers → "
         "+ Add Subscriber). Then on /admin/schedules schedule a trigger to your "
         "subscriberId for ~2 minutes in the future. Or fire via the API "
         "directly using the curl snippet in this guide."),
        ("Why are some workflows using SMS without an MSG91 template?",
         "Stage Dev workflows fall back to MSG91_OTP_TEMPLATE_ID when no per-trigger "
         "template is registered. This is intentional for testing; in Production "
         "every trigger has its own DLT-registered template."),
    ]
    for q, a in faqs:
        add_p(doc, q, bold=True, color=NAVY, space_after=2)
        add_p(doc, a, color=DARK, space_after=8)

    add_heading_styled(doc, "14.2  Escalation", level=2)
    add_table_styled(doc,
        ["Severity", "When", "Who to call", "How fast"],
        [
            ("● Sev-1", "Live customer-facing outage (no notifications fire at all)",
             "Engineering on-call (PagerDuty or Slack #tpe-comms-oncall)", "Immediate"),
            ("● Sev-2", "Single channel down (e.g. SMS via MSG91 broken; others fine)",
             "Engineering Lead via Slack", "Within 30 minutes"),
            ("● Sev-3", "Bug in admin console UI, missing data in /admin/history",
             "File a ticket; tag Engineering", "Next business day"),
            ("● Sev-4", "Documentation question, feature request",
             "Slack #tpe-comms-help", "Best-effort"),
        ],
        col_widths=[0.7, 2.2, 2.2, 1.4])
    doc.add_page_break()


def section_glossary(doc):
    add_heading_styled(doc, "15.  Glossary", level=1)
    add_table_styled(doc,
        ["Term", "Meaning"],
        [
            ("Audit row",
             "A structured log entry written for every channel-step. Carries "
             "transactionId, workflowId, triggerId, audienceGroup, "
             "subscriberToken (PII-tokenised), channel, status, reason."),
            ("Bridge",
             "The TypeScript service that hosts code-first workflows. Engineering-owned. "
             "Operators don't interact with it directly; they interact with "
             "Studio (the dashboard authoring UI) and the TPE Admin Console."),
            ("Campaign",
             "A saved record on /admin/campaigns describing source data + filter "
             "+ workflow + field map + schedule. The atomic unit of operator-driven "
             "data sends."),
            ("Channel",
             "A delivery surface: SMS, WhatsApp, Email, In-app."),
            ("Channel allowlist",
             "The toggle on /admin/channels that gates dispatch per channel. "
             "Operator-managed kill switch."),
            ("DLT",
             "Distributed Ledger Technology — TRAI mandate for India SMS. Every "
             "Sender ID + template body must be registered before a single SMS "
             "can fire."),
            ("Dispatch",
             "The act of sending one channel's content for one subscriber. Each "
             "trigger produces one dispatch per channel-step."),
            ("Environment",
             "Novu's notion of a logical scope. Stage has Development + Production. "
             "Workflows authored in Development are visible there only until "
             "promoted."),
            ("ICPaaS",
             "TPE's WhatsApp BSP (Business Solution Provider) — Indian-hosted, "
             "DPDPA-compliant. Routes to Meta's WhatsApp Cloud API."),
            ("MSG91",
             "TPE's primary SMS provider via DLT-registered templates."),
            ("Novu Cloud",
             "The hosted offering at novu.co. Forbidden by Charter — TPE "
             "self-hosts the Community Edition."),
            ("Operator",
             "A non-developer human who uses the TPE Admin Console or the Novu "
             "Dashboard."),
            ("Quiet hours",
             "21:00–09:00 IST window during which non-urgent triggers are "
             "suppressed."),
            ("Run",
             "A single execution of a campaign's fan-out. Multiple runs per "
             "campaign — each fires the same campaign config but produces a new "
             "row in tpe_campaign_runs."),
            ("Schedule (one-off)",
             "A single deferred trigger on /admin/schedules. The admin-worker "
             "picks it up at fireAt and fires once."),
            ("SES",
             "AWS Simple Email Service — TPE's email provider, region "
             "ap-south-1."),
            ("Studio",
             "Novu Dashboard's workflow authoring UI."),
            ("Subscriber",
             "A notification recipient, identified by subscriberId. Created "
             "automatically on first trigger or explicitly via the dashboard."),
            ("transactionId",
             "Novu's unique key per trigger. Used to correlate dispatches "
             "across channels in audit rows."),
            ("Workflow",
             "A series of channel-steps + delays + decisions. The recipe for a "
             "notification."),
        ],
        col_widths=[1.7, 4.8])
    doc.add_page_break()


def section_references(doc):
    add_heading_styled(doc, "16.  References & Links", level=1)

    add_heading_styled(doc, "16.1  Internal documents", level=2)
    add_table_styled(doc,
        ["Document", "Path / location"],
        [
            ("Charter v1.0 + Addendum v1.1",   "docs/TPE_Communication_System_Charter_*.docx"),
            ("Solution Document v1.0",          "docs/TPE_Communication_System_Novu_Solution_v1.0_turn1_review.docx"),
            ("Trigger Conditions v2.0",         "docs/TPE_Communication_System_Trigger_Conditions_v2.0.docx"),
            ("Combined Deployment Guide",       "deployment/Novu-Combined-Deployment-Guide.docx"),
            ("Bridge AUTHORING.md",             "bridge/workflows/AUTHORING.md  (engineering-only)"),
        ],
        col_widths=[2.4, 4.1])

    add_heading_styled(doc, "16.2  External references", level=2)
    add_table_styled(doc,
        ["Resource", "URL"],
        [
            ("Novu CE Documentation",  "https://docs.novu.co/"),
            ("MSG91 v5 Flow API",      "https://docs.msg91.com/"),
            ("Meta WhatsApp Cloud API", "https://developers.facebook.com/docs/whatsapp/cloud-api"),
            ("AWS SES Developer Guide", "https://docs.aws.amazon.com/ses/"),
            ("DPDPA 2023 (text)",      "https://www.meity.gov.in/data-protection-framework"),
        ],
        col_widths=[2.4, 4.1])

    add_heading_styled(doc, "16.3  Operator URLs (Stage)", level=2)
    add_table_styled(doc,
        ["Surface", "URL"],
        [
            ("TPE Admin Console",      "http://103.138.96.180/admin"),
            ("Novu CE Dashboard 2.3.0", "http://103.138.96.180/"),
            ("Novu CE Sandbox 3.15.0",  "http://103.138.96.180:8080/  (eval only)"),
        ],
        col_widths=[2.4, 4.1])

    add_heading_styled(doc, "16.4  Support contacts", level=2)
    add_table_styled(doc,
        ["Role / channel", "Contact"],
        [
            ("Engineering Lead",          "The Policy Exchange — Slack #tpe-comms-help"),
            ("On-call (Sev-1)",           "PagerDuty / Slack #tpe-comms-oncall"),
            ("Compliance Lead",           "compliance@thepolicyexchange.com"),
            ("CX Lead",                   "cx@thepolicyexchange.com"),
        ],
        col_widths=[2.4, 4.1])

    add_p(doc, "", space_after=24)
    add_p(doc, "— End of document —", italic=True, color=MEDIUM,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)


# ---------------------------------------------------------------------------
# Page setup
# ---------------------------------------------------------------------------

def setup_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Header (right-aligned, doc title + version)
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"{DOC_TITLE} — {DOC_SUBTITLE}  •  v{DOC_VERSION}")
    style_run(r, font="Arial", size=9, color=MEDIUM)

    # Footer (page number centred)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(CLASSIFICATION_LINE + "  •  ")
    style_run(r1, font="Arial", size=9, color=RED, bold=True)
    r2 = p.add_run("Page ")
    style_run(r2, font="Arial", size=9, color=MEDIUM)
    # Add PAGE field
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " PAGE ")
    rChild = OxmlElement("w:r")
    tChild = OxmlElement("w:t")
    tChild.text = "1"
    rPrChild = OxmlElement("w:rPr")
    rChild.append(rPrChild)
    rChild.append(tChild)
    fld.append(rChild)
    p._p.append(fld)


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------

def main():
    doc = Document()
    setup_page(doc)

    # Body styles fallback
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    cover_page(doc)
    section_doc_control(doc)
    section_intro(doc)
    section_overview(doc)
    section_access(doc)
    section_workflow_authoring(doc)
    section_channels(doc)
    section_templates(doc)
    section_campaigns(doc)
    section_schedules(doc)
    section_history(doc)
    section_retry(doc)
    section_provider_notes(doc)
    section_compliance(doc)
    section_troubleshoot(doc)
    section_glossary(doc)
    section_references(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"✓ Wrote {OUTPUT}")
    print(f"  size: {OUTPUT.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
