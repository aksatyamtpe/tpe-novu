"""
TPE Communication System — SOW Generator
Produces an enterprise-grade DOCX following the standards in ~/.claude/CLAUDE.md.
Run: python3 generate_sow.py
Output: TPE_Communication_System_SOW.docx
"""

from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────────────────────
# Brand palette (per ~/.claude/CLAUDE.md)
# ─────────────────────────────────────────────────────────────────────────────
NAVY        = RGBColor(0x1B, 0x3A, 0x5C)
BLUE        = RGBColor(0x2E, 0x75, 0xB6)
GREEN       = RGBColor(0x27, 0xAE, 0x60)
ORANGE      = RGBColor(0xE6, 0x7E, 0x22)
RED         = RGBColor(0xE7, 0x4C, 0x3C)
DARK        = RGBColor(0x2C, 0x3E, 0x50)
MEDIUM      = RGBColor(0x7F, 0x8C, 0x8D)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT       = "F0F4F8"
ALT_ROW     = "F8F9FA"
HEADER_BG   = "1B3A5C"
SUCCESS_BG  = "E8F5E9"
WARN_BG     = "FFF8E1"
DANGER_BG   = "FFEBEE"

DOC_TITLE = "TPE Communication System — Statement of Work"
DOC_VERSION = "2.0"
DOC_ID = "TPE-COMMS-SOW-2026-001"

# Per-section accent colors — assigns visual tone to section purpose
SECTION_ACCENT = {
    1: "2E75B6", 2: "2E75B6",                        # Doc info / version: blue
    3: "27AE60",                                     # Exec summary: green (achievement)
    4: "2E75B6",                                     # Business context: blue
    5: "E67E22",                                     # Problem: orange (warning)
    6: "27AE60",                                     # Vision/objectives: green
    7: "2E75B6", 8: "2E75B6",                        # Scope / Audience: blue
    9: "1B3A5C",                                     # Trigger catalogue: navy (core)
    10: "2E75B6",                                    # Channel: blue
    11: "1B3A5C", 12: "1B3A5C",                      # Template / Architecture: navy (core)
    13: "2E75B6",                                    # Skills: blue
    14: "E67E22",                                    # Compliance: orange
    15: "27AE60", 16: "27AE60", 17: "27AE60",        # Phases / Deliverables / Acceptance: green
    18: "2E75B6", 19: "7F8C8D", 20: "2E75B6",        # RACI / Assumptions / Dependencies
    21: "E74C3C",                                    # Risks: red
    22: "2E75B6", 23: "2E75B6",                      # Effort / Governance: blue
    24: "7F8C8D", 25: "7F8C8D",                      # Glossary / References: medium
    26: "27AE60",                                    # Sign-off: green
}
TODAY = date(2026, 4, 26).strftime("%d %b %Y")
AUTHOR = "Ashish Kumar Satyam"
ORG_NAME = "The Policy Exchange"
ORG_LEGAL = "Fairvalue Insuretech Pvt. Ltd."
ORG_FULL = f"{ORG_NAME} ({ORG_LEGAL})"
INTERNAL_TEAM = "TPE Project Team"

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_width_xml(cell, width_inches):
    """Reliable cell width: writes <w:tcW w:w="..." w:type="dxa"/> directly to XML.
    python-docx's cell.width property is advisory and often gets dropped — this is the
    canonical way to make Word respect cell width. 1 inch = 1440 DXA."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(int(width_inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')

def set_table_grid_xml(table, widths_inches):
    """Set explicit <w:tblGrid> column widths. Call after table creation but before
    you trust Word to honor your cell widths. Total should match printable area."""
    tbl = table._tbl
    grid = tbl.find(qn('w:tblGrid'))
    if grid is None:
        grid = OxmlElement('w:tblGrid')
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths_inches:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(int(w * 1440)))
        grid.append(col)

def set_table_width_xml(table, total_inches):
    """Set table preferred width (<w:tblW>) so Word knows total table width."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(int(total_inches * 1440)))
    tblW.set(qn('w:type'), 'dxa')

def lock_table_widths(table, widths_inches):
    """One-shot: set table grid + table width + every cell width.
    Use this whenever you set table.autofit = False and want widths to actually persist."""
    total = sum(widths_inches)
    set_table_grid_xml(table, widths_inches)
    set_table_width_xml(table, total)
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            if ci < len(widths_inches):
                set_cell_width_xml(cell, widths_inches[ci])

def set_cell_borders(cell, color="BFBFBF", size=4):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(size))
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def style_run(run, *, font="Arial", size=10, bold=False, italic=False, color=DARK):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rFonts.set(qn(attr), font)

def add_para(doc, text="", *, font="Arial", size=10, bold=False, italic=False,
             color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4, indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if text:
        r = p.add_run(text)
        style_run(r, font=font, size=size, bold=bold, italic=italic, color=color)
    return p

def add_bullet(doc, text, *, indent=0.5, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    style_run(r, size=size)
    return p

def add_section_heading(doc, num, title, level=1):
    """H1 navy bar + section-specific accent stripe; H2 colored left bar; H3 italic."""
    if level == 1:
        # Spacer
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
        # Navy heading bar
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        cell = tbl.rows[0].cells[0]
        set_cell_bg(cell, HEADER_BG)
        cell.width = Inches(7.0)
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
        cell.paragraphs[0].text = ""
        run = cell.paragraphs[0].add_run(f"  {num}. {title}")
        style_run(run, font="Arial", size=14, bold=True, color=WHITE)
        # Enterprise-clean: navy bar alone signals new section. No accent stripe.
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
    elif level == 2:
        # H2 with colored left bar (2-col table: bar + heading)
        try:
            sec_num = int(str(num).split(".")[0])
            accent = SECTION_ACCENT.get(sec_num, "2E75B6")
        except Exception:
            accent = "2E75B6"
        tbl = doc.add_table(rows=1, cols=2)
        tbl.autofit = False
        bar = tbl.rows[0].cells[0]
        bar.width = Inches(0.08)
        set_cell_bg(bar, accent)
        bar.paragraphs[0].text = ""
        heading = tbl.rows[0].cells[1]
        heading.width = Inches(6.5)
        heading.paragraphs[0].text = ""
        heading.paragraphs[0].paragraph_format.space_before = Pt(2)
        heading.paragraphs[0].paragraph_format.space_after = Pt(2)
        r = heading.paragraphs[0].add_run(f"  {num}  {title}")
        style_run(r, size=12, bold=True, color=NAVY)
        add_para(doc, "", space_after=2)
    else:
        p = add_para(doc, f"{num}  {title}", size=11, bold=True, italic=True,
                     color=BLUE, space_after=2)

def add_table(doc, headers, rows, *, col_widths_inches=None, header_bg=HEADER_BG,
              header_fg=WHITE, alt_rows=True, font_size=9, code_cols=None):
    code_cols = code_cols or set()
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = False
    if col_widths_inches:
        lock_table_widths(tbl, col_widths_inches)
    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        set_cell_bg(c, header_bg)
        set_cell_borders(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.paragraphs[0].text = ""
        r = c.paragraphs[0].add_run(h)
        style_run(r, font="Arial", size=font_size, bold=True, color=header_fg)
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
        if col_widths_inches:
            c.width = Inches(col_widths_inches[i])
    # Data rows
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]
        bg = ALT_ROW if (alt_rows and ri % 2 == 0) else "FFFFFF"
        for ci, val in enumerate(row):
            c = tr.cells[ci]
            set_cell_bg(c, bg)
            set_cell_borders(c)
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            c.paragraphs[0].text = ""
            font = "Consolas" if ci in code_cols else "Arial"
            r = c.paragraphs[0].add_run(str(val))
            style_run(r, font=font, size=font_size, color=DARK)
            c.paragraphs[0].paragraph_format.space_after = Pt(0)
            if col_widths_inches:
                c.width = Inches(col_widths_inches[ci])
    return tbl

def add_badge_row(doc, items):
    """items: list of (label, kind) where kind in {'green','orange','red','blue','navy'}"""
    color_map = {
        'green':  (GREEN,  SUCCESS_BG),
        'orange': (ORANGE, WARN_BG),
        'red':    (RED,    DANGER_BG),
        'blue':   (BLUE,   LIGHT),
        'navy':   (NAVY,   LIGHT),
    }
    tbl = doc.add_table(rows=1, cols=len(items))
    tbl.autofit = False
    for i, (label, kind) in enumerate(items):
        fg, bg = color_map[kind]
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, bg)
        set_cell_borders(c, color="DDDDDD")
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.paragraphs[0].text = ""
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.paragraphs[0].add_run(f"●  {label}")
        style_run(r, font="Arial", size=9, bold=True, color=fg)

def add_callout(doc, title, body, kind="blue"):
    color_map = {
        'blue':   (BLUE,   LIGHT),
        'green':  (GREEN,  SUCCESS_BG),
        'orange': (ORANGE, WARN_BG),
        'red':    (RED,    DANGER_BG),
        'navy':   (NAVY,   LIGHT),
    }
    fg, bg = color_map[kind]
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    c = tbl.rows[0].cells[0]
    set_cell_bg(c, bg)
    set_cell_borders(c, color=f"{fg.rgb if hasattr(fg,'rgb') else 'BFBFBF'}")
    c.paragraphs[0].text = ""
    r = c.paragraphs[0].add_run(title)
    style_run(r, size=10, bold=True, color=fg)
    p = c.add_paragraph()
    r2 = p.add_run(body)
    style_run(r2, size=9.5, color=DARK)
    add_para(doc, "", space_after=2)

def add_code_block(doc, code, *, font_size=8.5):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    c = tbl.rows[0].cells[0]
    set_cell_bg(c, "F4F4F4")
    set_cell_borders(c, color="CCCCCC")
    c.paragraphs[0].text = ""
    for i, line in enumerate(code.splitlines() or [""]):
        if i == 0:
            p = c.paragraphs[0]
        else:
            p = c.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line if line else " ")
        style_run(r, font="Consolas", size=font_size, color=DARK)

# ─────────────────────────────────────────────────────────────────────────────
# Visual element helpers (v1.2 design enhancements)
# ─────────────────────────────────────────────────────────────────────────────

def add_color_band(doc, hex_color, height_pt=4):
    """Thin colored horizontal band — used as section accent stripes / dividers."""
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    c = tbl.rows[0].cells[0]
    set_cell_bg(c, hex_color)
    c.paragraphs[0].text = ""
    c.paragraphs[0].paragraph_format.space_before = Pt(0)
    c.paragraphs[0].paragraph_format.space_after = Pt(0)
    tbl.rows[0].height = Pt(height_pt)
    tbl.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

def add_metric_tiles(doc, items):
    """items: list of (number, label, accent_hex). Renders N equal-width tiles."""
    n = len(items)
    width_each = 6.5 / n
    tbl = doc.add_table(rows=1, cols=n)
    tbl.autofit = False
    lock_table_widths(tbl, [width_each] * n)
    for i, (number, label, hex_color) in enumerate(items):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, "FAFBFC")
        set_cell_borders(c, color=hex_color, size=18)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.paragraphs[0].text = ""
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraphs[0].paragraph_format.space_after = Pt(2)
        r = c.paragraphs[0].add_run(str(number))
        style_run(r, font="Arial", size=20, bold=True,
                  color=RGBColor.from_string(hex_color))
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(label)
        style_run(r2, size=8.5, color=DARK)
    add_para(doc, "", space_after=2)

def add_pillar_cards(doc, items):
    """items: list of (title, body_text, accent_hex). Renders 3 (or N) cards side-by-side."""
    n = len(items)
    width_each = 6.5 / n
    tbl = doc.add_table(rows=1, cols=n)
    tbl.autofit = False
    lock_table_widths(tbl, [width_each] * n)
    for i, (title, body, hex_color) in enumerate(items):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, "FAFBFC")
        set_cell_borders(c, color=hex_color, size=18)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        c.paragraphs[0].text = ""
        c.paragraphs[0].paragraph_format.space_after = Pt(4)
        r = c.paragraphs[0].add_run(title)
        style_run(r, size=11, bold=True, color=RGBColor.from_string(hex_color))
        for line in body.split("\n"):
            p = c.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(line)
            style_run(r, size=9, color=DARK)
    add_para(doc, "", space_after=4)

def add_audience_card(doc, name, accent_hex, voice, channels, regime, sender):
    """Audience profile card: colored left bar + structured content cell."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    lock_table_widths(tbl, [0.12, 6.38])  # totals 6.50
    bar = tbl.rows[0].cells[0]
    set_cell_bg(bar, accent_hex)
    bar.paragraphs[0].text = ""
    body = tbl.rows[0].cells[1]
    set_cell_bg(body, "FAFBFC")
    set_cell_borders(body, color="E5E7E9")
    body.paragraphs[0].text = ""
    r = body.paragraphs[0].add_run(name)
    style_run(r, size=11, bold=True, color=RGBColor.from_string(accent_hex))
    body.paragraphs[0].paragraph_format.space_after = Pt(3)
    for label, val in [("Voice", voice), ("Primary channels", channels),
                       ("Compliance regime", regime), ("Sender identity", sender)]:
        p = body.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r1 = p.add_run(f"{label}:  ")
        style_run(r1, size=9, bold=True, color=NAVY)
        r2 = p.add_run(val)
        style_run(r2, size=9, color=DARK)
    add_para(doc, "", space_after=3)

def add_state_flow(doc, states):
    """states: list of (label, hex_color). Horizontal flow: state | -> | state | -> | ...
    For 6 states + 5 arrows, sized to fit 6.5" page: 6×0.85 + 5×0.18 = 5.10 + 0.90 = 6.00\""""
    n = len(states)
    cells_layout = []
    widths = []
    for i in range(n):
        cells_layout.append(("state", states[i]))
        widths.append(0.85)
        if i < n - 1:
            cells_layout.append(("arrow", None))
            widths.append(0.18)
    tbl = doc.add_table(rows=1, cols=len(cells_layout))
    tbl.autofit = False
    lock_table_widths(tbl, widths)
    for i, (kind, payload) in enumerate(cells_layout):
        c = tbl.rows[0].cells[i]
        c.paragraphs[0].text = ""
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
        if kind == "state":
            label, hex_color = payload
            set_cell_bg(c, hex_color)
            set_cell_borders(c, color=hex_color)
            r = c.paragraphs[0].add_run(label)
            style_run(r, size=8.5, bold=True, color=WHITE)
        else:
            r = c.paragraphs[0].add_run("→")
            style_run(r, size=12, bold=True, color=MEDIUM)
    add_para(doc, "", space_after=2)

def add_phase_timeline(doc, phases):
    """phases: list of (phase_id, name, days, hex_color). Visual horizontal timeline.
    For 5 phases + 4 arrows: 5×1.10 + 4×0.20 = 5.50 + 0.80 = 6.30\""""
    n = len(phases)
    cells_layout = []
    widths = []
    for i in range(n):
        cells_layout.append(("phase", phases[i]))
        widths.append(1.10)
        if i < n - 1:
            cells_layout.append(("arrow", None))
            widths.append(0.20)
    tbl = doc.add_table(rows=1, cols=len(cells_layout))
    tbl.autofit = False
    lock_table_widths(tbl, widths)
    for i, (kind, payload) in enumerate(cells_layout):
        c = tbl.rows[0].cells[i]
        c.paragraphs[0].text = ""
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
        if kind == "phase":
            phase_id, name, days, hex_color = payload
            set_cell_bg(c, hex_color)
            set_cell_borders(c, color=hex_color)
            r = c.paragraphs[0].add_run(phase_id)
            style_run(r, size=14, bold=True, color=WHITE)
            p2 = c.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_after = Pt(0)
            r2 = p2.add_run(name)
            style_run(r2, size=8, bold=True, color=WHITE)
            p3 = c.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.paragraph_format.space_after = Pt(2)
            r3 = p3.add_run(days)
            style_run(r3, size=8, color=WHITE, italic=True)
        else:
            r = c.paragraphs[0].add_run("→")
            style_run(r, size=14, bold=True, color=MEDIUM)
    add_para(doc, "", space_after=4)

def add_inline_severity(text, level):
    """Returns formatting hints for inline severity — used by table builders."""
    color_map = {'L': SUCCESS_BG, 'M': WARN_BG, 'H': DANGER_BG}
    fg_map    = {'L': "27AE60", 'M': "E67E22", 'H': "E74C3C"}
    return color_map.get(level, "FFFFFF"), fg_map.get(level, "2C3E50")

def add_table_with_severity(doc, headers, rows, severity_cols, *,
                             col_widths_inches=None, font_size=9):
    """Like add_table, but cells in severity_cols ('L'/'M'/'H') get colored backgrounds."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.autofit = False
    if col_widths_inches:
        lock_table_widths(tbl, col_widths_inches)
    # Header
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        set_cell_bg(c, HEADER_BG)
        set_cell_borders(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.paragraphs[0].text = ""
        r = c.paragraphs[0].add_run(h)
        style_run(r, font="Arial", size=font_size, bold=True, color=WHITE)
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
        if col_widths_inches:
            c.width = Inches(col_widths_inches[i])
    # Data rows
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]
        default_bg = ALT_ROW if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            c = tr.cells[ci]
            if ci in severity_cols:
                bg, fg = add_inline_severity("", str(val).strip())
                set_cell_bg(c, bg)
                set_cell_borders(c)
                c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                c.paragraphs[0].text = ""
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = c.paragraphs[0].add_run(f"●  {val}")
                style_run(r, size=font_size, bold=True,
                          color=RGBColor.from_string(fg))
            else:
                set_cell_bg(c, default_bg)
                set_cell_borders(c)
                c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                c.paragraphs[0].text = ""
                r = c.paragraphs[0].add_run(str(val))
                style_run(r, size=font_size, color=DARK)
            c.paragraphs[0].paragraph_format.space_after = Pt(0)
            if col_widths_inches:
                c.width = Inches(col_widths_inches[ci])
    return tbl

# ─────────────────────────────────────────────────────────────────────────────
# Page setup, header, footer
# ─────────────────────────────────────────────────────────────────────────────

def setup_document():
    doc = Document()
    # Page size & margins (US Letter, 1-inch)
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.4)
        section.footer_distance = Inches(0.4)

        # Header — right-aligned doc title + version
        h = section.header.paragraphs[0]
        h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rh = h.add_run(f"{DOC_TITLE}  |  v{DOC_VERSION}")
        style_run(rh, size=8, color=MEDIUM, italic=True)

        # Footer — centered page number + classification
        f = section.footer.paragraphs[0]
        f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rf = f.add_run(f"{DOC_ID}  |  Internal / Confidential  |  Page ")
        style_run(rf, size=8, color=MEDIUM)
        # Page number field
        run = f.add_run()
        fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText'); instr.text = "PAGE"
        fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_end)
        style_run(run, size=8, color=MEDIUM, bold=True)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.font.color.rgb = DARK
    return doc

# ─────────────────────────────────────────────────────────────────────────────
# Cover page
# ─────────────────────────────────────────────────────────────────────────────

def add_cover_page(doc):
    # Single navy band at top — restraint over chrome
    add_color_band(doc, HEADER_BG, height_pt=10)
    add_para(doc, "", space_after=18)

    add_para(doc, ORG_FULL.upper(), size=10, color=MEDIUM, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "STATEMENT OF WORK", size=22, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    # Title block — large, dark
    add_para(doc, "TPE Communication System", size=28, bold=True, color=DARK,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(doc, "A unified, multi-audience, multi-channel, template-driven "
                  "communication engine for the Indian insurtech secondary market",
             size=11, italic=True, color=MEDIUM,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    # Quantified scope tiles (the "what's in here" snapshot)
    add_metric_tiles(doc, [
        ("49",  "Lifecycle triggers",      "1B3A5C"),
        ("5",   "Audiences",                "2E75B6"),
        ("4",   "Channels (Email/SMS/WA/RCS)", "27AE60"),
        ("25",  "Working days (core P1-P4)","E67E22"),
    ])

    # Classification badge
    add_para(doc, "● INTERNAL / CONFIDENTIAL", size=10, bold=True, color=RED,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

    # Metadata table — refined
    meta_rows = [
        ["Document ID",    DOC_ID],
        ["Version",        f"v{DOC_VERSION}  (TPE-internal SOW)"],
        ["Date",           TODAY],
        ["Author",         AUTHOR],
        ["Organization",   ORG_FULL],
        ["Owner",          "Project Sponsor — TPE Leadership"],
        ["Approvers",      "Founder / CTO / Head of Compliance / Head of Customer Experience"],
        ["Classification", "Internal / Confidential"],
        ["Distribution",   "TPE Leadership, Engineering, Compliance, Customer Operations"],
    ]
    add_table(doc, ["Field", "Value"], meta_rows,
              col_widths_inches=[2.0, 4.5], font_size=10)

    # Spacer + footer block
    for _ in range(3):
        add_para(doc, "", space_after=0)

    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    c = tbl.rows[0].cells[0]
    set_cell_bg(c, HEADER_BG)
    c.paragraphs[0].text = ""
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Internal Document")
    style_run(r, size=9, color=WHITE, italic=True)
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run(ORG_NAME)
    style_run(r, size=14, bold=True, color=WHITE)
    p3 = c.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p3.add_run(ORG_LEGAL)
    style_run(r, size=9, color=WHITE, italic=True)

    doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# Sections
# ─────────────────────────────────────────────────────────────────────────────

def add_table_of_contents(doc):
    """Single unified TOC table — group rows + item rows in one structure."""
    add_para(doc, "TABLE OF CONTENTS", size=18, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "Forty-one sections in eleven thematic parts.",
             size=10, italic=True, color=MEDIUM,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

    groups = [
        ("PART I — DOCUMENT CONTROL & EXECUTIVE", "1B3A5C", [
            ("1",  "Document Information"),
            ("2",  "Document Control & Approval Matrix"),
            ("3",  "Version History"),
            ("4",  "Distribution List"),
            ("5",  "Executive Summary"),
            ("6",  "Executive Dashboard"),
            ("7",  "Business Case & ROI Analysis"),
        ]),
        ("PART II — CONTEXT & SCOPE", "2E75B6", [
            ("8",  "Business Context — The Policy Exchange"),
            ("9",  "Problem Statement & Current State"),
            ("10", "Project Vision, Objectives & Success Metrics"),
            ("11", "Scope (In / Out)"),
        ]),
        ("PART III — STAKEHOLDERS & AUDIENCE", "1B3A5C", [
            ("12", "Stakeholder Register"),
            ("13", "Audience Architecture (5 audiences)"),
        ]),
        ("PART IV — DOMAIN MODEL", "2E75B6", [
            ("14", "Trigger Catalogue (49 triggers)"),
            ("15", "Channel Strategy"),
        ]),
        ("PART V — SOLUTION", "27AE60", [
            ("16", "Template Management — Architectural Pillar"),
            ("17", "Solution Architecture (3 pillars)"),
            ("18", "Skill Catalogue (Claude Code surface)"),
        ]),
        ("PART VI — COMPLIANCE, SECURITY & QUALITY", "E67E22", [
            ("19", "Compliance & Regulatory Layer"),
            ("20", "Security & Privacy Framework"),
            ("21", "Quality Assurance & Testing Strategy"),
        ]),
        ("PART VII — DELIVERY", "27AE60", [
            ("22", "Implementation Phases & Timeline"),
            ("23", "Deliverables (per phase)"),
            ("24", "Acceptance Criteria"),
        ]),
        ("PART VIII — GOVERNANCE", "1B3A5C", [
            ("25", "Roles & Responsibilities (RACI)"),
            ("26", "Communication Management Plan"),
            ("27", "Change Management Process"),
            ("28", "Service Level Agreements"),
            ("29", "Governance & Review Cadence"),
        ]),
        ("PART IX — RISK & DEPENDENCIES", "E74C3C", [
            ("30", "Assumptions"),
            ("31", "Dependencies"),
            ("32", "Risk Register"),
        ]),
        ("PART X — RESOURCING & READINESS", "2E75B6", [
            ("33", "Effort Estimation"),
            ("34", "Resource Plan"),
            ("35", "Operational Readiness & Go-Live Criteria"),
            ("36", "Knowledge Transfer & Training Plan"),
            ("37", "Post-Implementation Support & Maintenance"),
        ]),
        ("PART XI — LEGAL & REFERENCE", "7F8C8D", [
            ("38", "Confidentiality, IP & Data Ownership"),
            ("39", "Glossary"),
            ("40", "References"),
            ("41", "Sign-off"),
        ]),
    ]

    # Build a single unified table: 2 cols (number, title), with group divider rows
    rows_data = []
    for group_name, color, items in groups:
        rows_data.append(("group", group_name, color))
        for num, title in items:
            rows_data.append(("item", num, title, color))

    tbl = doc.add_table(rows=len(rows_data), cols=2)
    tbl.autofit = False
    lock_table_widths(tbl, [0.65, 5.85])

    for ri, row_data in enumerate(rows_data):
        kind = row_data[0]
        cell0 = tbl.rows[ri].cells[0]
        cell1 = tbl.rows[ri].cells[1]
        cell0.paragraphs[0].text = ""
        cell1.paragraphs[0].text = ""
        cell0.paragraphs[0].paragraph_format.space_after = Pt(0)
        cell1.paragraphs[0].paragraph_format.space_after = Pt(0)

        if kind == "group":
            _, group_name, color = row_data
            # Group header row — colored cell with white text spanning visually
            set_cell_bg(cell0, color)
            set_cell_bg(cell1, color)
            cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cell0.paragraphs[0].add_run(" ")
            style_run(r, size=10, bold=True, color=WHITE)
            r2 = cell1.paragraphs[0].add_run(f"  {group_name}")
            style_run(r2, size=10, bold=True, color=WHITE)
            cell0.paragraphs[0].paragraph_format.space_before = Pt(3)
            cell0.paragraphs[0].paragraph_format.space_after = Pt(3)
            cell1.paragraphs[0].paragraph_format.space_before = Pt(3)
            cell1.paragraphs[0].paragraph_format.space_after = Pt(3)
        else:
            _, num, title, color = row_data
            # Item row — number in colored text + title in dark
            cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cell0.paragraphs[0].add_run(num)
            style_run(r, size=10, bold=True, color=RGBColor.from_string(color))
            r2 = cell1.paragraphs[0].add_run(f"  {title}")
            style_run(r2, size=10, color=DARK)
            cell0.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell0.paragraphs[0].paragraph_format.space_after = Pt(2)
            cell1.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell1.paragraphs[0].paragraph_format.space_after = Pt(2)

    doc.add_page_break()


def section_executive_briefing(doc):
    """One-page narrative briefing — prose only, no tables. Sets the document's
    posture before the structured sections begin."""
    add_para(doc, "EXECUTIVE BRIEFING", size=14, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "For the Founders, CTO, Compliance Lead and CX Head",
             size=9.5, italic=True, color=MEDIUM,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

    # Lead paragraph — what this is
    add_para(doc,
        "The Policy Exchange is building an unusual asset class — pre-owned "
        "life-insurance policies — and an unusual two-sided marketplace to "
        "transact them. Every assignment creates a multi-decade communication "
        "obligation: an investor who paid the surrender value upfront must be "
        "kept current on premiums for 10–25 years, against the backdrop of "
        "11 insurance partners, two regulators, and an Indian policy-holder "
        "base that conducts the rest of its life over WhatsApp. The "
        "communication system this Statement of Work proposes is therefore "
        "not a marketing tool. It is operational infrastructure for a regulated, "
        "long-tail relationship business.",
        size=11, color=DARK, space_after=10)

    # What we're committing to
    add_para(doc,
        "Concretely, the project will deliver a versioned content library "
        "covering forty-nine lifecycle triggers across five distinct audiences "
        "and four channels, with templates as first-class entities that move "
        "through a defined draft → review → approved → active → deprecated → "
        "archived lifecycle. A compliance lint engine will enforce DLT, "
        "IRDAI, WhatsApp BSP and DPDPA rules at the approval gate — before "
        "any message leaves the system — replacing the current spreadsheet-"
        "and-inbox approach. Authoring is owned by the engineering team; "
        "approval authority remains with the CX, Compliance and Brand "
        "leads. Live dispatch via MSG91 / Karix / Gupshup is held to an "
        "optional Phase 5, gated on channel-provider selection.",
        size=11, color=DARK, space_after=10)

    # What success looks like
    add_para(doc,
        "Success is measured on seven baselined objectives, the most "
        "consequential of which is a target twenty-percent reduction in "
        "investor premium-default rate within six months of go-live. Core "
        "delivery is twenty-five working engineering days plus nine days of "
        "reviewer time, distributed across four phases that each terminate "
        "in an independently consumable artefact. We are deliberately "
        "not building a customer data store, a CRM, or an analytics "
        "dashboard — those are out of scope and routed to separate "
        "engagements as they arise. The constraint is the value: a "
        "communication system that an auditor, a regulator and an investor "
        "can all read without ambiguity.",
        size=11, color=DARK, space_after=14)

    # Single-line summary anchored at the bottom
    add_para(doc,
        "Read in order: Parts I–II (Sections 1–11) establish document "
        "control, executive summary, context and scope; Parts III–V "
        "(12–18) define stakeholders, the domain and the solution; "
        "Part VI (19–21) covers compliance, security and quality; "
        "Part VII (22–24) is delivery; Part VIII (25–29) is governance; "
        "Parts IX–X (30–37) cover risk, resourcing and readiness; "
        "Part XI (38–41) is legal, reference and sign-off.",
        size=9.5, italic=True, color=MEDIUM,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    doc.add_page_break()


def section_doc_info(doc, num):
    add_section_heading(doc, num, "Document Information")
    add_para(doc, "This SOW formalises scope, architecture, deliverables, "
                  "effort and acceptance criteria for the TPE Communication "
                  "System. It is the source of truth for project sign-off and "
                  "serves as the brief for downstream implementation planning.",
             space_after=6)
    add_table(doc, ["Field", "Value"], [
        ["Document Name",      DOC_TITLE],
        ["Document ID",        DOC_ID],
        ["Version",            DOC_VERSION],
        ["Effective Date",     TODAY],
        ["Owner",              "Project Sponsor (TPE Leadership)"],
        ["Author",             AUTHOR],
        ["Organization",       ORG_FULL],
        ["Stakeholders",       "Founders, CTO, Head of Compliance, Head of CX, Operations Lead"],
        ["Status",             "Pending Approval"],
        ["Review Cycle",       "Per phase gate + quarterly post go-live"],
        ["Classification",     "Internal / Confidential"],
    ], col_widths_inches=[2.0, 4.5], font_size=10)

def section_version_history(doc, num):
    add_section_heading(doc, num, "Version History")
    add_table(doc, ["Version", "Date", "Author", "Changes"], [
        ["1.0", TODAY, AUTHOR,
         "Initial SOW: scope, audience map, 49-trigger catalog, channel "
         "strategy, template-management architecture, compliance layer, "
         "phased delivery, effort estimate."],
        ["1.1", TODAY, AUTHOR,
         "Reframed as TPE-internal SOW; removed third-party attribution; "
         "updated RACI, governance, sign-off and effort tables to reflect "
         "internal team structure."],
        ["1.2", TODAY, AUTHOR,
         "Design enrichment: section accent stripes; cover-page metric tiles; "
         "Table of Contents page; audience cards; Template Lifecycle visual "
         "state-flow; Architecture pillar cards; Phase timeline; Risk "
         "severity colour-coding; one additional risk row (premium-default "
         "rate) raised to High."],
        ["1.3", TODAY, AUTHOR,
         "Layout polish: XML-level cell-width persistence (lock_table_widths) "
         "so column proportions render exactly as designed; state-flow and "
         "phase-timeline widths tightened to fit page width with margin to "
         "spare; cover-page top whitespace reduced; exec-summary phase block "
         "switched to visual timeline for consistency; spacer paragraphs "
         "trimmed across all visual helpers."],
        ["1.4", TODAY, AUTHOR,
         "Column-shrink fix: shortened over-long headers (Working Days→Days, "
         "Engineering→Eng, Likelihood→Likely, Cumulative→Total, Required by→"
         "Need by, Founder/Sponsor→Sponsor) so they no longer wrap; "
         "rebalanced widths in 8 tables (Phases, Deliverables, RACI, "
         "Dependencies, Risks, Effort, Version History, Problem Statement) "
         "to give data columns more breathing room while staying within the "
         "6.50\" printable width."],
        ["1.5", TODAY, AUTHOR,
         "Enterprise-grade refactor: removed per-section accent stripes "
         "(saved ~26 chrome tables); consolidated the Table of Contents "
         "from 12 small tables into one unified table with group dividers; "
         "compressed cover-page bands from 4 to 1; added a narrative "
         "Executive Briefing page (prose only, 0 tables) before structured "
         "sections; added narrative-opener paragraphs to Sections 9, 13, 14 "
         "and 15 so the document reads as a briefing rather than a data "
         "dump. Net table count reduction: ~30%."],
        ["2.0", TODAY, AUTHOR,
         "Enterprise-grade upgrade. Major version bump. Document expanded "
         "from 26 to 41 sections in eleven thematic parts. Fifteen new "
         "sections added: Document Control & Approval Matrix, Distribution "
         "List, Executive Dashboard, Business Case & ROI, Stakeholder "
         "Register, Security & Privacy Framework, Quality Assurance & "
         "Testing Strategy, Communication Management Plan, Change "
         "Management Process, Service Level Agreements, Resource Plan, "
         "Operational Readiness & Go-Live Criteria, Knowledge Transfer & "
         "Training Plan, Post-Implementation Support & Maintenance, and "
         "Confidentiality / IP / Data Ownership. Existing Risks section "
         "expanded into a full enterprise Risk Register (15 entries). "
         "Section function signatures refactored to accept the section "
         "number as a parameter, so document order is now driven entirely "
         "from main() — section insertion no longer requires editing "
         "downstream function bodies. Document numbering renumbered "
         "throughout."],
    ], col_widths_inches=[0.85, 0.90, 1.50, 3.25], font_size=9.5)

def section_executive_summary(doc, num):
    add_section_heading(doc, num, "Executive Summary")
    add_para(doc, "Why this matters", size=11, bold=True, color=NAVY,
             space_after=2)
    add_para(doc,
        "TPE operates an unusual two-sided insurance marketplace: policyholders "
        "assign their policies to investors who then pay future premiums for "
        "10–25 years until maturity. This creates a long-tail communication "
        "obligation that does not exist for standard insurers. Premium-due "
        "reminders to investors alone will dominate message volume across the "
        "next two decades. Today, content is held in spreadsheets, ad-hoc "
        "templates and individual mailboxes, with no compliance lint, no "
        "version control, no audit trail and no path to multi-language scaling.",
        space_after=8)

    add_para(doc, "What this project delivers", size=11, bold=True, color=NAVY,
             space_after=2)
    add_para(doc,
        "A versioned content library and authoring toolkit covering 49 lifecycle "
        "triggers across five audiences (policyholder, investor, insurer "
        "partner, internal ops, regulator) and four channels (Email, SMS, "
        "WhatsApp, RCS). Templates are first-class entities with a defined "
        "lifecycle (draft → review → approved → active → deprecated), enabling "
        "reuse, A/B variants, multi-language localisation and a compliance "
        "lint that prevents non-compliant SMS or WhatsApp messages from "
        "shipping. The system is delivered as a Claude Code skill suite plus a "
        "structured YAML library, with optional later integration to MSG91 / "
        "Karix / Gupshup for direct dispatch.",
        space_after=8)

    add_para(doc, "Quantified scope", size=11, bold=True, color=NAVY,
             space_after=2)
    add_table(doc, ["Dimension", "v1 Scope", "Steady-state"], [
        ["Triggers",                      "49",       "60+"],
        ["Audiences",                     "5",        "5"],
        ["Channels",                      "4",        "4 + voice"],
        ["Languages",                     "EN, HI",   "EN, HI + 8 regional"],
        ["Template count (estimate)",     "~120",     "300+"],
        ["Unique content cells",          "~600",     "3,000+"],
        ["Insurer partners (B2B)",        "11",       "20+"],
        ["Compliance regimes covered",    "DLT, IRDAI, WhatsApp BSP", "+ FATCA, GDPR-equivalent"],
    ], col_widths_inches=[2.4, 2.0, 2.1], font_size=9.5)

    add_para(doc, "", space_after=4)
    add_para(doc, "Effort & timeline at a glance", size=11, bold=True, color=NAVY,
             space_after=4)
    add_phase_timeline(doc, [
        ("P1", "Foundation",          "5 days",  "1B3A5C"),
        ("P2", "PH Library",           "8 days",  "2E75B6"),
        ("P3", "INV+INS+OPS+REG+HI",  "7 days",  "27AE60"),
        ("P4", "Compliance Layer",     "5 days",  "E67E22"),
        ("P5", "Live Dispatch (opt.)", "10 days", "E74C3C"),
    ])
    add_para(doc, "Core delivery (Phases 1-4): 25 working days. Phase 5 (live "
                  "dispatch integration) is gated on budget and channel-provider "
                  "selection.", italic=True, color=MEDIUM, size=9, space_after=4)

def section_business_context(doc, num):
    add_section_heading(doc, num, "Business Context — The Policy Exchange")
    add_para(doc,
        "The Policy Exchange (TPE), operating under Fairvalue Insuretech Pvt. "
        "Ltd., is a Noida-headquartered insurtech founded in 2022 by Naveen "
        "Gupta, Safia Anwar and Tarun Bahri. TPE has raised approximately "
        "USD 2.49M across four rounds, including a USD 1.5M Pre-Series B "
        "(November 2025) led by VLS Finance, True Blue Holding and 1Crowd. "
        "TPE’s stated goal is to safeguard one million policyholders over the "
        "next three years.")

    add_para(doc, "Two-sided business model", size=11, bold=True, color=NAVY,
             space_after=2)
    add_table(doc, ["Side", "Customer", "Value provided", "Communication character"], [
        ["Sell-side",
         "Policyholders unable or unwilling to continue paying premiums",
         "Surrender value paid up-front; premium obligation transferred to investor",
         "Empathetic, simple, reassuring; often in vernacular language"],
        ["Buy-side",
         "Retail and HNI investors seeking stable yield (~12–14% claimed)",
         "Access to pre-owned policies as an investable asset class",
         "Professional, data-led, opportunity-framed; long-running for 10–25 years"],
        ["B2B",
         "11 insurance carriers (Aditya Birla, Aviva, Bajaj, Bharti, "
         "Edelweiss, Generali, India First, Bandhan, Reliance, Ageas, ...)",
         "Premium continuity, formal assignment intimation, reconciliation",
         "Formal, structured, attachment-heavy (MIS files)"],
    ], col_widths_inches=[1.0, 1.7, 2.0, 1.8], font_size=9)

    add_para(doc, "", space_after=4)
    add_para(doc, "Adjacent products", size=11, bold=True, color=NAVY,
             space_after=2)
    for b in [
        "Loan against insurance policy — liquidity without surrender; introduces a multi-year EMI communication track.",
        "Maturity / claim coordination on assigned policies — TPE acts as agent between insurer and assignee investor.",
        "Future: re-investment products that roll matured policy proceeds into new assignments."
    ]:
        add_bullet(doc, b)

    add_callout(doc,
        "Why this is different from generic insurance comms",
        "Standard carriers communicate with one party (the policyholder). "
        "TPE communicates with up to five parties about the same policy "
        "(seller, buyer, carrier, internal ops, regulator), often "
        "simultaneously, and continues investor communication for the entire "
        "remaining term of the policy — which can exceed two decades. "
        "A copy-paste-from-Gmail approach does not scale past a few hundred "
        "active policies; the operational risk (missed premium → lapsed "
        "policy → litigation from investor) grows linearly with policy count.",
        kind="navy")

def section_problem_current_state(doc, num):
    add_section_heading(doc, num, "Problem Statement & Current State")
    add_para(doc, "Observed today:", size=10.5, bold=True, color=NAVY,
             space_after=2)
    for b in [
        "Communication content is maintained as informal spreadsheets and "
        "free-text drafts in individual inboxes; no canonical source of truth.",
        "No version history — when copy is updated, the prior version and "
        "the rationale for the change are lost.",
        "No compliance lint — DLT-template-id, IRDAI mandatory disclosures, "
        "WhatsApp BSP category and opt-out footers are checked manually, if at all.",
        "No reuse — a generic OTP message exists in three different forms "
        "across Registration, e-Sign and Login flows.",
        "No multi-channel parity — the same message intent rendered in Email "
        "and WhatsApp can drift in tone, factual content and disclaimers.",
        "No structured cadence model — escalation rules (e.g. KYC nudge "
        "weekly × 4 then monthly × 4 then close) live in operator memory.",
        "No translation pipeline — Hindi or regional-language variants are "
        "produced ad-hoc per campaign.",
        "Limited audit trail — when a regulator or investor questions a "
        "communication, reconstructing what was sent, when, and why is "
        "non-trivial.",
    ]:
        add_bullet(doc, b)

    add_para(doc, "", space_after=4)
    add_para(doc, "Cost of the status quo:", size=10.5, bold=True, color=NAVY,
             space_after=2)
    add_table(doc, ["Risk", "Likely", "Indicative impact"], [
        ["Missed investor premium reminder → policy lapse → investor loss", "High",  "Investor lawsuit + reputational harm + regulator scrutiny"],
        ["Non-DLT compliant SMS → DOT block of sender ID",                  "Medium", "All transactional SMS halted across audiences for 24-72h"],
        ["WhatsApp template miscategorised (Promotional vs Utility)",       "Medium", "Template rejection, delivery delay, BSP fines"],
        ["Missing IRDAI mandatory disclosure on promo material",            "Medium", "Show-cause notice, mandatory withdrawal of campaign"],
        ["Inconsistent tone across channels for same trigger",              "High",  "Customer confusion, increased support load"],
        ["Lost institutional knowledge when ops staff change",              "High",  "Quality regression, longer ramp-up for new hires"],
    ], col_widths_inches=[3.30, 0.65, 2.55], font_size=9)

def section_vision_objectives(doc, num):
    add_section_heading(doc, num, "Project Vision, Objectives & Success Metrics")
    add_para(doc, "Vision", size=11, bold=True, color=NAVY, space_after=2)
    add_para(doc,
        "Every TPE communication is authored once, approved once, "
        "rendered consistently across every channel and language, and "
        "shipped only after passing automated compliance checks — with "
        "full version history available for audit and continuous improvement.")

    add_para(doc, "Measurable objectives (v1)", size=11, bold=True, color=NAVY,
             space_after=2)
    add_table(doc, ["#", "Objective", "Target metric", "Baseline"], [
        ["O1", "Centralise communication content into a single versioned library",
                                                          "100% of v1 triggers covered (49/49)",         "0%"],
        ["O2", "Eliminate manual compliance review for routine messages",
                                                          ">=95% of templates pass auto-lint pre-review", "Manual on every template"],
        ["O3", "Reduce time to author and approve a new message",
                                                          "<= 30 minutes from draft to approved",        ">=2 hours typical"],
        ["O4", "Enable safe multi-language expansion",
                                                          "Hindi variants for 100% of PH-facing templates","Ad-hoc, partial"],
        ["O5", "Provide audit-grade evidence of every dispatched message",
                                                          "100% messages traceable to a versioned template", "Reconstructed manually"],
        ["O6", "Reduce investor premium-default rate via reliable cadence",
                                                          "Default rate decline of >=20% within 6 months of go-live", "Current internal benchmark"],
        ["O7", "Reduce ops-team time spent on routine drafting",
                                                          "Save >=20 person-hours / week within 90 days of go-live", "n/a"],
    ], col_widths_inches=[0.4, 2.2, 2.5, 1.4], font_size=9)

    add_callout(doc, "How we measure value, not hype",
        "Each objective has a target and a baseline. If we cannot evidence the "
        "improvement against the baseline at the 6-month review, the "
        "corresponding capability is reconsidered or removed. Vanity metrics "
        "(template count, lines of code, slash commands shipped) are explicitly "
        "out of scope for success measurement.",
        kind="green")

def section_scope(doc, num):
    add_section_heading(doc, num, "Scope")
    add_section_heading(doc, f"{num}.1", "In Scope", level=2)
    for b in [
        "Trigger catalogue covering 49 lifecycle events across PH, Investor, Insurer, Internal, Regulatory audiences.",
        "Template Library (first-class CMS): create, version, approve, "
        "deprecate, search, link variants, render per channel, validate "
        "variables.",
        "Channel renderers for Email, SMS, WhatsApp (template + free text), RCS.",
        "Cross-cutting axes: language (EN, HI in v1), audience tone, sender identity, compliance class.",
        "Compliance lint engine: DLT (India SMS), IRDAI mandatory disclosures, "
        "WhatsApp BSP category, opt-out footer, time-of-day windows, transactional vs promotional separation.",
        "Cadence and escalation engine: per-trigger schedule rules, retry-and-close logic.",
        "Personalisation token library and validation.",
        "Claude Code skill suite (slash commands + agents) to author, preview, lint, search, export the library.",
        "Documentation: this SOW, runbook, template-author guide, compliance reference card.",
        "Acceptance test pack covering all v1 triggers and lint rules.",
    ]:
        add_bullet(doc, b)

    add_section_heading(doc, f"{num}.2", "Out of Scope (v1)", level=2)
    for b in [
        "Direct dispatch to channel providers (MSG91, Karix, Gupshup, WhatsApp BSP, SMTP). Optional Phase 5.",
        "Customer-data store / CRM functionality. The system consumes existing TPE data; it is not a system of record for customers.",
        "Real-time analytics / BI dashboards on delivery rates. Existing tools to be used.",
        "Voice / IVR scripting beyond brief reference notes within templates.",
        "Sentiment analysis or NLP-based reply handling.",
        "AI-generated copy at runtime. All content is human-authored and approved; AI assists during authoring only.",
        "Replacement of the existing MIS Pipeline (which already serves INS-02 premium remittance attachments).",
    ]:
        add_bullet(doc, b)

def section_audience_arch(doc, num):
    add_section_heading(doc, num, "Audience Architecture")
    add_para(doc, "Five audiences each with distinct voice, channel mix, "
                  "compliance regime and sender identity. Treating them "
                  "uniformly is the single largest source of avoidable error "
                  "in insurance communication.", space_after=8)

    add_audience_card(doc, "Policyholder  —  Seller", "2E75B6",
        voice="Empathetic, simple, reassuring; never urgency-marketing; vernacular-friendly",
        channels="WhatsApp (primary), SMS, Email",
        regime="DLT-Service / DLT-Promo, IRDAI disclosures, opt-out (DPDPA)",
        sender="TPE Customer Care")

    add_audience_card(doc, "Investor  —  Buyer  (10–25 year relationship)", "27AE60",
        voice="Professional, data-led, returns-framed, long-time-horizon",
        channels="Email (primary), WhatsApp, RCS, SMS for OTP and dues",
        regime="DLT-Service, IRDAI 'subject of solicitation', FATCA references",
        sender="TPE Investments")

    add_audience_card(doc, "Insurance Partner  —  B2B  (11 carriers)", "E67E22",
        voice="Formal, structured, evidence-attached, traceable",
        channels="Email + carrier-portal upload (Excel attachments via MIS Pipeline)",
        regime="Bilateral SLA, audit log",
        sender="TPE Operations")

    add_audience_card(doc, "Internal Ops / Sales / Agents", "1B3A5C",
        voice="Concise, action-oriented; surface decisions and exceptions",
        channels="Slack / Teams, Email",
        regime="Internal information-classification policy",
        sender="TPE System Bot")

    add_audience_card(doc, "Regulator / Statutory", "E74C3C",
        voice="Statutory, audit-grade, version-stamped",
        channels="Portal upload, secure email",
        regime="IRDAI, GST, Income Tax (TDS), DPDPA",
        sender="TPE Compliance")

def section_trigger_catalog(doc, num):
    add_section_heading(doc, num, "Trigger Catalogue (49 triggers)")
    add_para(doc,
        "The trigger catalogue is the system's map of every moment we "
        "communicate. Forty-nine triggers are organised into five "
        "audience streams; each stream is presented as a single "
        "table for at-a-glance comparison across statuses, channels and "
        "cadence rules.",
        space_after=8)
    add_para(doc, "Each trigger has a stable identifier, one or more statuses "
                  "(states the entity can be in), a channel mix and a cadence "
                  "rule. Identifiers are immutable; titles may be revised in minor versions.",
             space_after=6)

    # A. Policyholder
    add_section_heading(doc, f"{num}.1", "A. Policyholder Lifecycle (18 triggers)", level=2)
    add_table(doc, ["ID", "Trigger", "Statuses", "Channels", "Cadence highlight"], [
        ["PH-01", "Lead Captured",         "New / Contacted / Qualified / Disqualified",                "Email, WA",        "T+0,1,3,7 drip"],
        ["PH-02", "Registration",          "Started / OTP Sent / Verified / Abandoned",                  "SMS, WA, Email",   "Abandoned: 1h, 24h, 72h"],
        ["PH-03", "KYC",                   "Done / In-progress / Failed / Pending",                      "All",              "Pending: 1/wk x4, 1/mo x4, close"],
        ["PH-04", "Policy Doc Upload",     "Uploaded / Reviewing / Verified / Rejected",                 "Email, WA",        "Per state"],
        ["PH-05", "Policy Verification",   "In-progress / Verified / Discrepancy / Insurer non-response","WA, Email",        "Update every 48h"],
        ["PH-06", "Surrender Value Quote", "Generated / Sent / Viewed / Expired (7d) / Counter-offer",   "Email, WA",        "T+0, T+2, T+5, T-1 pre-expiry"],
        ["PH-07", "Offer Decision",        "Accepted / Rejected / Negotiating / Silent",                  "Email, WA",        "Silent: 3d, 7d"],
        ["PH-08", "Assignment Paperwork",  "Form Gen / e-Sign Initiated / e-Sign Done / Physical / NOC", "Email, WA, SMS",   "e-Sign: T+1, T+3"],
        ["PH-09", "Investor Matched",      "Found / Confirmed",                                          "WA, Email",        "Immediate"],
        ["PH-10", "Disbursement",          "Initiated / In-process / Success / Failed",                  "SMS, WA, Email",   "Per state"],
        ["PH-11", "Post-Assignment Welcome","Recorded / Welcome",                                         "Email, WA",        "+ quarterly newsletter"],
        ["PH-12", "Loan Application",      "Started / Docs pending / Submitted / Under review",          "All",              "Standard"],
        ["PH-13", "Loan Approval",         "Approved / Conditional / Rejected",                          "Email, WA, SMS",   "Immediate"],
        ["PH-14", "Loan Disbursement",     "Initiated / Disbursed / Failed",                             "SMS, WA",          "Immediate"],
        ["PH-15", "Loan EMI",              "T-7, T-3, T-day, +1, +7, +15, +30",                          "SMS, WA, Email",   "Per schedule"],
        ["PH-16", "Loan Closure",          "Pre-close quote / Foreclosed / Paid",                        "Email, WA",        "Per event"],
        ["PH-17", "Re-engagement",         "30d / 90d / 6m / 1y silent",                                 "Email, WA",        "Decay drip"],
        ["PH-18", "Referral Program",      "Earned / Credited / Expiring",                               "WA, Email",        "Per event"],
    ], col_widths_inches=[0.50, 1.40, 1.85, 0.90, 1.85], font_size=8, code_cols={0})

    # B. Investor
    add_section_heading(doc, f"{num}.2", "B. Investor Lifecycle (16 triggers, highest message volume)", level=2)
    add_table(doc, ["ID", "Trigger", "Statuses", "Channels", "Cadence highlight"], [
        ["INV-01", "Lead Captured",                "New / Webinar reg / Webinar attended / Brochure DL", "Email, WA",        "T+0,1,3,7,14 drip"],
        ["INV-02", "Registration",                  "Started / Verified / KYC Pending / Active",          "Email, SMS, WA",   "Standard"],
        ["INV-03", "KYC + Risk Profile",            "CKYC / FATCA / Risk profile / Bank verified",        "Email, WA",        "Standard"],
        ["INV-04", "Investment Opportunity Push",   "New match / Weekly curated / Premium yield",         "Email, WA, RCS",   "Weekly digest + alerts"],
        ["INV-05", "Investment Soft-Commit",        "Interest / Doc review / Clarifications",             "Email, WA",        "Immediate"],
        ["INV-06", "Funds Collection",              "Pay link / Received / Failed",                       "SMS, Email, WA",   "Pending: 24h, 48h"],
        ["INV-07", "Investment Confirmed",          "Assignment in progress / Confirmed / Cert issued",    "Email, WA",        "Immediate"],
        ["INV-08", "Premium Due (10–25 yr stream)", "T-30, T-15, T-7, T-3, T-day, +1, +7, +15",          "All",              "HIGHEST VOLUME trigger"],
        ["INV-09", "Premium Payment Confirmation",  "Paid to insurer / Receipt available",                "Email, WA",        "Immediate"],
        ["INV-10", "Premium Default → Lapse Risk",  "Grace started / Lapse warning / Lapsed",             "All + call",       "Aggressive escalation"],
        ["INV-11", "Maturity Approaching",          "T-12mo, T-6mo, T-3mo, T-1mo, T-7d",                  "Email, WA",        "Per schedule"],
        ["INV-12", "Maturity Received",             "Insurer paid TPE / Settled to investor",             "SMS, Email, WA",   "Immediate"],
        ["INV-13", "Portfolio Statement",           "Monthly / Quarterly / Annual",                       "Email",            "Cadence"],
        ["INV-14", "Tax Documents",                 "TDS Q1-Q4 / 80C-D / Annual",                         "Email",            "Statutory"],
        ["INV-15", "Re-investment Opportunity",     "Maturity-to-reinvest / New high-yield",              "Email, WA, RCS",   "Triggered"],
        ["INV-16", "Account Notices",               "Monthly / T&C changes",                               "Email",            "Monthly"],
    ], col_widths_inches=[0.55, 1.6, 1.75, 1.0, 1.6], font_size=8, code_cols={0})

    # C. Insurer
    add_section_heading(doc, f"{num}.3", "C. Insurance Partner — B2B (6 triggers)", level=2)
    add_table(doc, ["ID", "Trigger", "Statuses", "Channels", "Cadence"], [
        ["INS-01", "Assignment Intimation",             "Submitted / Acknowledged / Approved / Rejected", "Email + Portal", "Per case"],
        ["INS-02", "Premium Remittance + Allocation MIS","Paid / MIS attached / Confirmed / Discrepancy", "Email + Excel",  "Monthly batch — reuses existing MIS Pipeline"],
        ["INS-03", "MIS Reconciliation",                 "Sent / Confirmed / Variance",                   "Email + Excel",  "Monthly"],
        ["INS-04", "Query Resolution",                   "Raised / Response / Closed",                    "Email",          "Per case"],
        ["INS-05", "Maturity / Claim Coordination",     "Initiated / Docs / Approved / Settled",         "Email",          "Per case"],
        ["INS-06", "Regulatory / KYC Refresh",          "IRDAI shared / KYC refresh",                    "Email",          "Per cadence"],
    ], col_widths_inches=[0.55, 2.0, 1.85, 1.0, 1.1], font_size=8.5, code_cols={0})

    # D. Internal
    add_section_heading(doc, f"{num}.4", "D. Internal Ops / Sales / Agents (6 triggers)", level=2)
    add_table(doc, ["ID", "Trigger", "When fires", "Channels", "Cadence"], [
        ["OPS-01", "New Lead Allocated",        "Lead created → assigned",            "Slack/Teams + Email", "Real-time"],
        ["OPS-02", "Daily Standup Briefing",    "Pipeline + stuck + KPIs",            "Slack + Email",       "Daily 9am IST"],
        ["OPS-03", "High-Value Escalation",     "Deal >= ₹50L stuck / VIP query",     "Slack + SMS",         "Real-time"],
        ["OPS-04", "Premium Default Alert",     "Investor non-pay → lapse risk",      "Slack + Email",       "Real-time"],
        ["OPS-05", "KYC Stuck",                 ">7 days in KYC",                     "Slack",               "Daily digest"],
        ["OPS-06", "Compliance Deadline",       "IRDAI / GST / TDS",                  "Email + Calendar",    "Pre-deadline"],
    ], col_widths_inches=[0.55, 1.8, 1.8, 1.4, 1.0], font_size=8.5, code_cols={0})

    # E. Regulatory
    add_section_heading(doc, f"{num}.5", "E. Regulatory / Statutory (3 triggers)", level=2)
    add_table(doc, ["ID", "Trigger", "Cadence"], [
        ["REG-01", "IRDAI Filing — quarterly + annual",  "Per regulator schedule"],
        ["REG-02", "TDS certificates to investors",      "Quarterly"],
        ["REG-03", "GST invoices",                       "Per transaction"],
    ], col_widths_inches=[0.55, 3.5, 2.5], font_size=9, code_cols={0})

def section_channel_strategy(doc, num):
    add_section_heading(doc, num, "Channel Strategy")
    add_para(doc, "Channel selection is derived from audience, message intent and compliance class — never from preference.", space_after=4)
    add_table(doc, ["Channel", "Best for", "Avoid for", "Indian compliance gate"], [
        ["Email",       "Long-form, attachments (statements, certificates), legal copy",
                        "Time-critical alerts (open-rate too low)",
                        "DPDPA consent, IT Act 2000"],
        ["SMS",         "OTP, transactional confirmations, due-day urgent reminders",
                        "Promotional content to opted-out users; long messages",
                        "DLT registration (PE, Header, Template, Entity ID); TRAI"],
        ["WhatsApp",    "Conversational service messages, multilingual, rich media",
                        "Bulk promotional without opt-in; promotional outside template categories",
                        "Meta BSP template approval; correct category (Marketing / Utility / Authentication); 24h session window"],
        ["RCS",         "Carrier-rich interactions: pay-now buttons, carousels for opportunity push",
                        "Audiences whose handsets / carriers do not yet support RCS",
                        "Same DLT regime as SMS in India"],
        ["Slack/Teams", "Internal alerts, allocations, escalations",
                        "Anything customer-facing",
                        "Internal information classification only"],
        ["Voice/IVR",   "Premium-default last-mile call (Phase 5+)",
                        "First-touch unknown audience",
                        "TRAI commercial-comms regulations + DND"],
    ], col_widths_inches=[0.95, 1.95, 1.65, 1.95], font_size=9)

def section_template_management(doc, num):
    add_section_heading(doc, num, "Template Management — Architectural Pillar")
    add_para(doc, "Templates are first-class entities, decoupled from triggers. "
                  "A trigger references one or more templates by ID; the same "
                  "template can be reused across triggers (e.g. a single OTP "
                  "template serves Registration, Login and e-Sign). This "
                  "decoupling is the single most important architectural "
                  "decision in the system: it prevents content drift, enables "
                  "version control, makes localisation tractable and produces "
                  "an audit trail.", space_after=6)

    add_section_heading(doc, f"{num}.1", "Template Lifecycle", level=2)
    add_para(doc, "Every template moves through a defined state machine. "
                  "State transitions are recorded with actor, timestamp and "
                  "rationale.", space_after=6)

    # Visual state-flow diagram
    add_state_flow(doc, [
        ("DRAFT",      "7F8C8D"),
        ("REVIEW",     "E67E22"),
        ("APPROVED",   "2E75B6"),
        ("ACTIVE",     "27AE60"),
        ("DEPRECATED", "E74C3C"),
        ("ARCHIVED",   "1B3A5C"),
    ])

    add_table(doc, ["State", "Meaning", "Who can transition", "Permitted next states"], [
        ["draft",      "Author is composing; not yet reviewable",                "Author",                         "review, archived"],
        ["review",     "Awaiting compliance and brand-voice review",              "Author submits",                  "approved, draft (returned for edit)"],
        ["approved",   "Compliance + brand approved; ready to activate",          "Compliance Reviewer + Brand Lead","active, draft (revoked)"],
        ["active",     "In use by one or more triggers; lint passes",             "Operator / scheduler",            "deprecated"],
        ["deprecated", "Superseded by a newer version; remains for audit only",   "Operator",                        "archived"],
        ["archived",   "Frozen, read-only, retained for audit",                   "System",                          "(terminal)"],
    ], col_widths_inches=[0.90, 1.90, 1.90, 1.80], font_size=9)

    add_section_heading(doc, f"{num}.2", "Versioning Rules (semver)", level=2)
    add_table(doc, ["Bump", "When", "Example"], [
        ["MAJOR (x.0.0)", "Breaking change: variable added/removed/renamed; intent changes; channel changes", "1.4.2 → 2.0.0 when adding {{xirr}} variable"],
        ["MINOR (1.x.0)", "Material wording change; new disclaimer; tone shift",                              "1.4.2 → 1.5.0 when softening urgency wording"],
        ["PATCH (1.4.x)", "Typo, punctuation, formatting; no semantic change",                                "1.4.2 → 1.4.3 when fixing apostrophe"],
    ], col_widths_inches=[1.30, 3.30, 1.90], font_size=9)

    add_section_heading(doc, f"{num}.3", "Template Schema (YAML)", level=2)
    add_para(doc, "Templates are stored as YAML files in the repository under "
                  "templates/<audience>/<channel>/<id>.yaml. Schema (illustrative):",
             space_after=4)
    add_code_block(doc,
"""template_id: TPL-PH-KYC-PENDING-WA-001
name: "KYC Pending Reminder — WhatsApp"
version: 1.2.0
status: active                           # draft|review|approved|active|deprecated|archived
audience: policyholder                   # policyholder|investor|insurer|internal|regulator
language: en                             # en|hi|ta|te|mr|bn|kn|gu|ml|pa
channel: whatsapp                        # email|sms|whatsapp|rcs|slack|email_attachment
category: utility                        # marketing|utility|authentication  (WA BSP)
compliance_class: service                # transactional|service|promotional
sender_identity: TPE_CARE                # references senders.yaml
brand_voice: empathetic                  # references voices.yaml
intent: "Nudge PH to complete KYC; respect cadence; do not pressure"
variables:
  - name: ph_name
    required: true
    fallback: "Customer"
  - name: kyc_link
    required: true
    type: short_url
disclaimers: [irdai_solicitation_short]  # references disclaimers.yaml
footer:      [opt_out_short]             # references footers.yaml
content:
  body: |
    Hello {{ph_name}},
    Your KYC is still pending. Complete it in 2 minutes here: {{kyc_link}}
    Reply STOP to opt out. Need help? Reply HELP.
parent_template: null                    # for variants/translations
variants:
  - TPL-PH-KYC-PENDING-WA-001-hi         # Hindi variant of this template
used_by_triggers: [PH-03]
metadata:
  created_by: ashish.satyam
  created_at: 2026-04-26
  approved_by: compliance.lead
  approved_at: 2026-04-26
  brand_review_by: cx.head
  reuse_count: 1
  delivery_estimate_per_month: 4500
""")

    add_section_heading(doc, f"{num}.4", "Template Management Capabilities", level=2)
    add_table(doc, ["Capability", "Description", "Surface"], [
        ["Authoring",          "Create / edit templates with schema validation, variable hints, character-limit warnings (SMS 160, WA 1024)", "/tpe-template new, /tpe-template edit"],
        ["Versioning",         "Semver bumps with diff against prior; rationale captured",                  "/tpe-template version <id>"],
        ["Approval workflow",  "Submit for review; reviewer approves/rejects with notes",                   "/tpe-template review <id>"],
        ["Variant linking",    "Mark a template as language or A/B variant of another; sync schema",        "/tpe-template variant <parent> <child>"],
        ["Channel rendering",  "Single intent → channel-specific render with limits enforced",              "/tpe-template render <id>"],
        ["Compliance lint",    "Run all relevant lint rules on demand or pre-approval",                      "/tpe-comply check <id>"],
        ["Search / discover",  "Filter by audience, channel, category, status, used-by-trigger",            "/tpe-template search"],
        ["Reuse tracking",     "Show all triggers that reference a template before deprecation",            "/tpe-template usage <id>"],
        ["Catalogue export",   "Generate human-readable PDF/HTML index of all approved templates",          "/tpe-library export"],
        ["Bulk import",        "Migrate spreadsheet content into template files (one-time + ongoing sync)", "/tpe-library import <xlsx>"],
        ["Diff",               "Show change between any two versions of a template",                        "/tpe-template diff <id> v1 v2"],
        ["Localisation queue", "Surface templates whose EN parent has changed but variants have not",       "/tpe-l10n queue"],
        ["Deprecation safety", "Block deprecation if template is still referenced by an active trigger",    "Pre-condition check"],
        ["Audit trail export", "Per-template chronological log: state transitions, edits, approvals, sends","/tpe-template audit <id>"],
    ], col_widths_inches=[1.40, 3.30, 1.80], font_size=8.5, code_cols={2})

    add_section_heading(doc, f"{num}.5", "Reuse Examples (why this matters)", level=2)
    add_table(doc, ["Template family", "Used by triggers", "Without reuse"], [
        ["TPL-OTP-SMS-001",            "PH-02 Registration, PH-08 e-Sign, INV-02 Registration, Login",        "4 near-identical drafts, 4 risks of drift"],
        ["TPL-IRDAI-DISCLAIMER-LONG",  "Every PH and INV email/RCS promotional template",                       "Each template re-types the disclaimer; one regulator update = N edits"],
        ["TPL-OPT-OUT-FOOTER-WA",      "All WhatsApp promotional templates",                                    "Inconsistent opt-out wording → DPDPA risk"],
        ["TPL-PH-KYC-PENDING-WA-001",  "PH-03 KYC at week-1, week-2, week-3, week-4 (with date variable)",      "4 duplicated templates instead of 1 + cadence rule"],
    ], col_widths_inches=[2.30, 2.50, 1.70], font_size=9, code_cols={0})

def section_solution_arch(doc, num):
    add_section_heading(doc, num, "Solution Architecture")
    add_para(doc, "Three pillars, file-system-resident, Git-versioned, "
                  "human-readable. No database, no server, no hosting in v1. "
                  "Read-only consumption by future channel integrators.",
             space_after=8)

    # Visual: three pillar cards side-by-side
    add_pillar_cards(doc, [
        ("PILLAR 1\nTrigger Catalogue",
         "WHEN to communicate.\n\n"
         "Lives at: triggers/<audience>/<id>.yaml\n\n"
         "Owner: Customer Experience\n\n"
         "49 lifecycle events with statuses, cadence and template references.",
         "1B3A5C"),
        ("PILLAR 2\nTemplate Library",
         "WHAT to say.\n\n"
         "Lives at: templates/<audience>/<channel>/\n\n"
         "Owner: Brand + Compliance\n\n"
         "Reusable, versioned content atoms with full lifecycle (draft → archived).",
         "2E75B6"),
        ("PILLAR 3\nChannel Renderer",
         "HOW to ship it.\n\n"
         "Lives at: skills/render/*\n\n"
         "Owner: Engineering\n\n"
         "Transforms template + variables → channel-specific payload (Email, SMS, WA, RCS).",
         "27AE60"),
    ])

    add_para(doc, "", space_after=4)

    add_para(doc, "", space_after=2)
    add_para(doc, "Data flow", size=11, bold=True, color=NAVY, space_after=2)
    add_code_block(doc,
"""    TRIGGER (PH-03 KYC, status=pending, attempt=2)
        │
        ├── reads cadence rule → "send weekly for 4 weeks, then monthly for 4, then close"
        ├── selects template_id by status, channel, language, audience
        │
        ▼
    TEMPLATE (TPL-PH-KYC-PENDING-WA-001 v1.2.0, status=active)
        │
        ├── compliance lint:  DLT ✓  IRDAI ✓  WA category=utility ✓  opt-out ✓
        ├── variable bind:    {{ph_name}}, {{kyc_link}} → values
        │
        ▼
    CHANNEL RENDERER (whatsapp)
        │
        ├── enforces 1024-char limit, BSP-template syntax, button schema
        │
        ▼
    DISPATCH PAYLOAD (JSON; ready for MSG91 / Karix / Gupshup)
""")

    add_section_heading(doc, f"{num}.1", "Repository layout", level=2)
    add_code_block(doc,
"""tpe-communication-system/
├── docs/
│   ├── TPE_Communication_System_SOW.docx        (this document)
│   ├── runbook.md
│   ├── template-author-guide.md
│   └── compliance-reference.md
├── triggers/
│   ├── policyholder/      (18 yaml files)
│   ├── investor/          (16 yaml files)
│   ├── insurer/           ( 6 yaml files)
│   ├── internal/          ( 6 yaml files)
│   └── regulatory/        ( 3 yaml files)
├── templates/
│   ├── policyholder/{email,sms,whatsapp,rcs}/
│   ├── investor/{email,sms,whatsapp,rcs}/
│   ├── insurer/email/
│   ├── internal/{slack,email}/
│   └── regulatory/email/
├── library/
│   ├── senders.yaml          (sender identities + DLT IDs)
│   ├── voices.yaml           (brand-voice rules per audience)
│   ├── disclaimers.yaml      (IRDAI mandatory text variants)
│   ├── footers.yaml          (opt-out, address, grievance)
│   ├── variables.yaml        (canonical token catalogue)
│   ├── languages.yaml        (locales + glyph rules)
│   └── compliance-rules.yaml (lint rule definitions)
├── skills/                   (Claude Code skill .md files)
│   ├── tpe-template/
│   ├── tpe-trigger/
│   ├── tpe-comply/
│   ├── tpe-render/
│   └── tpe-library/
└── tests/
    ├── lint/                 (compliance lint test pack)
    ├── render/               (channel render goldens)
    └── triggers/             (trigger-template wiring tests)
""")

def section_skill_catalog(doc, num):
    add_section_heading(doc, num, "Skill Catalogue (Claude Code surface)")
    add_para(doc,
        "The system is operated through twenty-two slash commands invoked "
        "from Claude Code. Commands cluster around three intents — author "
        "templates and triggers, validate compliance, and render or "
        "(eventually) dispatch. The phase column shows when each command "
        "becomes available; nothing in P5 is required for the library or "
        "lint to deliver value.",
        space_after=8)
    add_table(doc, ["Slash command", "Purpose", "Phase"], [
        ["/tpe-template new",      "Create a new template from schema with prompts for required fields",   "P1"],
        ["/tpe-template edit",     "Edit existing template with version-bump prompt",                       "P1"],
        ["/tpe-template version",  "Bump version (major/minor/patch) with rationale capture",                "P1"],
        ["/tpe-template review",   "Submit for review or approve/reject with notes",                         "P1"],
        ["/tpe-template variant",  "Link a variant template to its parent (translation or A/B)",             "P3"],
        ["/tpe-template render",   "Render template + sample variables into channel-specific output",        "P1"],
        ["/tpe-template diff",     "Show diff between any two versions",                                     "P1"],
        ["/tpe-template usage",    "List all triggers referencing a template",                               "P1"],
        ["/tpe-template audit",    "Chronological log: state transitions, edits, approvals",                 "P2"],
        ["/tpe-trigger new",       "Create a new trigger entry: statuses, cadence, template references",     "P1"],
        ["/tpe-trigger edit",      "Edit cadence or template references on a trigger",                       "P1"],
        ["/tpe-trigger preview",   "Simulate trigger across all statuses → render every message it produces", "P2"],
        ["/tpe-comply check",      "Run all relevant compliance lint rules on a template or trigger",        "P4"],
        ["/tpe-comply rules",      "List active compliance rules and which categories they apply to",        "P4"],
        ["/tpe-render <channel>",  "Render any template into the specified channel (with sample data)",      "P1"],
        ["/tpe-library search",    "Filter library by audience, channel, status, used-by, language",          "P1"],
        ["/tpe-library export",    "Export approved library as DOCX or HTML index for stakeholder review",   "P2"],
        ["/tpe-library import",    "One-time migration of existing spreadsheet content into templates",      "P2"],
        ["/tpe-l10n queue",        "Show templates whose EN parent has changed but variants have not",       "P3"],
        ["/tpe-l10n translate",    "Author a language variant from a parent template",                       "P3"],
        ["/tpe-send dryrun",       "[Phase 5] Build dispatch payload without sending",                       "P5"],
        ["/tpe-send dispatch",     "[Phase 5] Send via configured channel provider",                         "P5"],
    ], col_widths_inches=[2.00, 3.90, 0.60], font_size=9, code_cols={0})

def section_compliance(doc, num):
    add_section_heading(doc, num, "Compliance & Regulatory Layer")
    add_para(doc,
        "Compliance is the most consequential and the most easily neglected "
        "part of regulated communication. Six regimes apply simultaneously to "
        "a typical TPE message — DLT for SMS, IRDAI for insurance disclosures, "
        "WhatsApp BSP for template categorisation, DPDPA for consent and opt-"
        "out, FATCA / KYC for investor flows, and Income-Tax (TDS) for "
        "statutory documents. The lint engine encodes them as machine-checkable "
        "rules so a missed disclaimer is caught at the approval gate, not by "
        "the regulator.",
        space_after=8)
    add_para(doc, "Every approved template is validated against the active "
                  "compliance ruleset before it can transition to ‘active’. "
                  "Rules are externalised as YAML so a regulator update does "
                  "not require code changes.", space_after=4)
    add_table(doc, ["Regime", "Rule examples", "Surface in lint"], [
        ["TRAI / DLT (SMS)",
         "Header registered; Template ID present; Entity (PE) ID present; "
         "transactional vs promotional category matches sender; promotional time-window check",
         "/tpe-comply check → DLT_*"],
        ["IRDAI",
         "‘Insurance is the subject matter of solicitation’ disclaimer present in promo; "
         "returns disclaimer; grievance redressal contact; IRDAI registration number on annual statements",
         "/tpe-comply check → IRDAI_*"],
        ["WhatsApp BSP (Meta)",
         "Template category matches content (Marketing/Utility/Authentication); "
         "variable count and order matches approved Meta template; "
         "media size limits; 24h session window respected for free-text",
         "/tpe-comply check → WA_*"],
        ["DPDPA (India Data Protection)",
         "Opt-out instruction present in promotional channels; "
         "consent reference for marketing categories; "
         "PII masking in audit logs",
         "/tpe-comply check → DPDPA_*"],
        ["FATCA / KYC",
         "Investor onboarding templates reference FATCA self-declaration; "
         "CKYC where applicable; risk-profile reference on investment opportunity push",
         "/tpe-comply check → KYC_*"],
        ["Income Tax (TDS)",
         "Section / rate references in TDS certificate templates; "
         "PAN-format validation in variables; "
         "FY references machine-derived not hard-coded",
         "/tpe-comply check → TDS_*"],
        ["Internal brand-voice",
         "PH-facing templates avoid urgency-marketing terms in stop-list; "
         "investor-facing templates do not promise specific returns; "
         "tone matches audience profile",
         "/tpe-comply check → BRAND_*"],
    ], col_widths_inches=[1.30, 3.40, 1.80], font_size=8.5, code_cols={2})

    add_callout(doc, "Lint pass before approval, not after dispatch",
        "Compliance failure detected after a message has been sent is, in practice, "
        "irreversible — the regulator’s clock has started. The lint engine therefore "
        "runs as a mandatory pre-condition for the review→approved transition. "
        "An override path exists for urgent cases; overrides are logged and reviewed.",
        kind="orange")

def section_phases(doc, num):
    add_section_heading(doc, num, "Implementation Phases & Timeline")
    add_para(doc,
        "Five phases. Each phase produces an independently consumable "
        "artefact, so a stop after any phase still leaves usable value behind. "
        "Phase 5 (live dispatch via channel providers) is the only phase "
        "that depends on external commercial relationships — DLT, BSP and "
        "SMS-gateway contracts — and is therefore held as optional and "
        "separately gated.",
        space_after=8)

    # Visual horizontal timeline
    add_phase_timeline(doc, [
        ("P1", "Foundation",         "5 days",  "1B3A5C"),
        ("P2", "PH Library",          "8 days",  "2E75B6"),
        ("P3", "INV+INS+OPS+REG+HI", "7 days",  "27AE60"),
        ("P4", "Compliance Layer",    "5 days",  "E67E22"),
        ("P5", "Live Dispatch (opt)", "10 days", "E74C3C"),
    ])

    add_table(doc, ["Phase", "Name", "Days", "Outcome"], [
        ["P1", "Foundation",
                "5",
                "Repo scaffolded; YAML schemas finalised; first 5 skills working "
                "end-to-end on a single trigger (PH-03 KYC) and 1 template; "
                "renderer for Email + WhatsApp."],
        ["P2", "Content Library v1 — Policyholder",
                "8",
                "All 18 PH triggers authored, EN only, all statuses; "
                "~50 templates in library; export as DOCX index for stakeholder review."],
        ["P3", "Investor + Insurer + Internal + Hindi",
                "7",
                "16 INV + 6 INS + 6 OPS + 3 REG triggers added; Hindi variants for all "
                "PH-facing templates; localisation queue working."],
        ["P4", "Compliance Layer",
                "5",
                "Full lint engine; DLT, IRDAI, WhatsApp BSP, DPDPA, brand-voice rules; "
                "approval gate enforced; audit-trail export."],
        ["P5", "Live Dispatch (optional, gated)",
                "10",
                "MSG91 / Karix / Gupshup MCP integration; WhatsApp BSP API; "
                "dry-run + dispatch slash commands; delivery-receipt ingest."],
    ], col_widths_inches=[0.60, 2.40, 0.55, 2.95], font_size=9, code_cols={0})
    add_para(doc, "Total core delivery (P1-P4): 25 working days. P5 is "
                  "gated on channel-provider selection, contract, sandbox "
                  "credentials and budget. P5 is not on the critical path for "
                  "value realisation — the library and lint provide value "
                  "immediately even with manual dispatch.", italic=True, color=MEDIUM, size=9)

def section_deliverables(doc, num):
    add_section_heading(doc, num, "Deliverables")
    add_table(doc, ["Phase", "Artefact", "Format"], [
        ["P1", "Repository scaffold + YAML schemas",                 "Git repo"],
        ["P1", "5 core slash commands operational",                  "Skill .md files"],
        ["P1", "Renderer (Email, WhatsApp) reference implementation","Code"],
        ["P1", "Walk-through of PH-03 KYC end-to-end",               "Recorded session"],
        ["P2", "All 18 PH triggers authored",                        "YAML files"],
        ["P2", "~50 EN templates in library",                        "YAML files"],
        ["P2", "Stakeholder review pack (DOCX export of library)",    "DOCX"],
        ["P2", "Bulk import from current spreadsheet (one-time)",     "Migration script + report"],
        ["P3", "31 additional triggers (INV + INS + OPS + REG)",     "YAML files"],
        ["P3", "Hindi variants for all PH-facing templates",         "YAML files"],
        ["P3", "Localisation queue dashboard (markdown report)",      "Markdown"],
        ["P4", "Compliance lint engine + 30+ rules",                 "Code + YAML rules"],
        ["P4", "Approval gate enforced on lifecycle transition",      "Skill behaviour"],
        ["P4", "Audit-trail export per template",                    "Markdown / JSON"],
        ["P4", "Compliance reference card (printable)",               "PDF"],
        ["P5", "MSG91 / Karix / Gupshup integration (gated)",         "MCP integration"],
        ["P5", "WhatsApp BSP integration (gated)",                    "MCP integration"],
        ["P5", "Dispatch dry-run + send commands",                    "Skill .md files"],
        ["All", "Test packs (lint, render, trigger-template wiring)", "Code"],
        ["All", "Runbook + author guide + compliance reference",      "Markdown / DOCX"],
    ], col_widths_inches=[0.55, 4.20, 1.75], font_size=9, code_cols={0})

def section_acceptance(doc, num):
    add_section_heading(doc, num, "Acceptance Criteria")
    add_table(doc, ["Phase", "Acceptance criteria"], [
        ["P1",
         "Repo passes schema-validation tests; 5 slash commands operate without error on PH-03 KYC; "
         "Email + WhatsApp render produces output that matches a stakeholder-approved sample byte-for-byte (excluding timestamps)."],
        ["P2",
         "All 18 PH triggers have at least one approved template per active status; "
         "stakeholder review pack received written approval from CX Head; "
         "spreadsheet import covers 100% of current PH content with diff report attached."],
        ["P3",
         "Investor, Insurer, Internal, Regulatory triggers fully authored; "
         "Hindi variants exist for all PH-facing templates and pass language-specific lint; "
         "localisation queue reports zero unsynced templates."],
        ["P4",
         "Lint engine has at least 30 rules across the six regimes; "
         "review→approved transition is blocked on lint failure; "
         "audit-trail export is reviewable and reconciles with Git history."],
        ["P5 (if executed)",
         "Dry-run produces provider-valid payload for MSG91 / Karix / Gupshup; "
         "dispatch successfully sends and ingests delivery receipts in sandbox; "
         "production cutover requires separate go-live sign-off."],
    ], col_widths_inches=[1.05, 5.45], font_size=9, code_cols={0})

def section_raci(doc, num):
    add_section_heading(doc, num, "Roles & Responsibilities (RACI)")
    add_para(doc, "R = Responsible, A = Accountable, C = Consulted, I = Informed.", italic=True, color=MEDIUM, size=9, space_after=4)
    add_table(doc, ["Activity", "Eng", "CX", "Comply", "CTO", "Sponsor"], [
        ["Schema design + scaffolding",         "R/A", "C",   "C",   "I",   "I"],
        ["Template authoring",                   "R",   "A",   "C",   "I",   "I"],
        ["Compliance rule definition",           "C",   "C",   "R/A", "I",   "I"],
        ["Compliance review of templates",       "C",   "C",   "R/A", "I",   "I"],
        ["Brand-voice review",                    "C",   "R/A", "C",   "I",   "I"],
        ["Channel renderer code",                 "R/A", "I",   "I",   "C",   "I"],
        ["Spreadsheet import (one-time)",         "R",   "A",   "C",   "I",   "I"],
        ["Phase sign-off",                       "I",   "C",   "C",   "C",   "A"],
        ["Phase 5 channel-provider selection",    "C",   "C",   "C",   "R/A", "I"],
        ["Audit-evidence production",             "R",   "C",   "A",   "I",   "I"],
    ], col_widths_inches=[2.60, 0.78, 0.65, 0.78, 0.55, 1.14], font_size=8.5)

def section_assumptions(doc, num):
    add_section_heading(doc, num, "Assumptions")
    for b in [
        "TPE will nominate one CX Head, one Compliance Reviewer and one Brand Lead with authority to approve content within 48h.",
        "Existing spreadsheet content is the source of truth for v1 and is internally consistent at the time of import (a one-time clean-up may be required).",
        "DLT registrations (Header, PE ID, Template ID for at least 5 transactional intents) exist or will be obtained before P5.",
        "WhatsApp BSP relationship and at least one approved Utility template exist or will be obtained before P5.",
        "Insurer-portal access and SFTP credentials (where applicable) are managed by TPE Operations, not in scope for this project.",
        "Variable values (PH name, policy number, surrender value, due date, etc.) are available from existing TPE systems via a stable interface; data integration design is a downstream project.",
        "Hindi translation is the only non-EN language in v1; further regional languages are explicit Phase 6+ scope.",
        "The MIS Pipeline (already approved) continues to own monthly insurer remittance attachments; INS-02 templates wrap, not replace, that pipeline.",
    ]:
        add_bullet(doc, b)

def section_dependencies(doc, num):
    add_section_heading(doc, num, "Dependencies")
    add_table(doc, ["Dependency", "Owner", "Need by", "Status"], [
        ["Approved this SOW",                                  "TPE Sponsor",       "P1 start", "Pending"],
        ["Nominated CX Head + Compliance + Brand Lead",        "TPE Founder",       "P1 start", "Pending"],
        ["Access to current spreadsheet content",              "TPE CX Head",       "P2",       "Pending"],
        ["IRDAI mandatory disclaimer text (current revision)", "TPE Compliance",    "P4",       "Pending"],
        ["DLT entity, header, PE ID, template IDs",            "TPE Compliance",    "P5",       "Out of scope until P5"],
        ["WhatsApp BSP relationship + Meta-approved templates","TPE CTO",           "P5",       "Out of scope until P5"],
        ["Channel-provider contract (MSG91 or Karix or Gupshup)","TPE CTO",         "P5",       "Out of scope until P5"],
        ["MIS Pipeline (existing project) operational",         "Project Manager",  "P3 INS-02","Pending Mgr Approval"],
    ], col_widths_inches=[2.55, 1.45, 1.00, 1.50], font_size=9)

def section_risks(doc, num):
    add_section_heading(doc, num, "Risk Register")
    add_para(doc, "Risk 'Likely' column is colour-coded: ● Low / ● Medium / ● High. "
                  "Each risk has a concrete mitigation owner-able within this project.",
             italic=True, color=MEDIUM, size=9, space_after=6)
    add_table_with_severity(doc,
        ["Risk", "Likely", "Impact", "Mitigation"],
        [
        ["Compliance reviewer bandwidth lower than required",   "M", "Delays approval transitions across all phases",
         "Reserve a 30-min daily slot in CX Head + Compliance calendars during P2-P4; approval queue visible in /tpe-template review"],
        ["Spreadsheet content quality lower than assumed",      "M", "Import surfaces inconsistencies → P2 slips",
         "Run a one-day pre-import audit; produce a diff report; agree clean-up scope with CX Head before bulk import"],
        ["IRDAI updates a mandatory disclaimer mid-project",    "L", "Forces re-version of every promotional template",
         "Disclaimers externalised as library/disclaimers.yaml; one edit + bulk version-bump"],
        ["WhatsApp BSP rejects a template category",            "M", "P5 dispatch blocked for that intent",
         "Defer to P5; design renderer to be category-agnostic; have fallback Email + SMS path"],
        ["DLT template-id provisioning delayed",                 "M", "P5 SMS dispatch blocked",
         "Defer P5 SMS until DLT received; library + lint not affected"],
        ["Stakeholder churn during long phases",                 "L", "Reset of approval relationships; rework",
         "Phase outcomes are independently consumable; sign-off after every phase, not at the end"],
        ["Scope creep into customer-data store",                 "H", "Project doubles in size and timeline",
         "Out-of-scope list is explicit; new requests routed to a separate engagement"],
        ["Voice/IVR added under pressure",                       "L", "P5+ scope expansion",
         "Hold to phase boundary; demonstrate audit value of P1-P4 before debating P5+"],
        ["Investor premium-default rate stays high post-launch", "H", "Direct revenue impact + investor trust loss",
         "Cadence engine is conservative by default; add escalation tier (Phone) before lapse window"],
    ], severity_cols={1}, col_widths_inches=[2.20, 0.70, 1.55, 2.05], font_size=8.5)

def section_effort(doc, num):
    add_section_heading(doc, num, "Effort Estimation")
    add_table(doc, ["Phase", "Eng days", "Reviewer days (CX+Comp+Brand)", "Total"], [
        ["P1 Foundation",                              "5",   "1",  "5"],
        ["P2 Content Library v1 (PH)",                  "8",   "3",  "13"],
        ["P3 INV + INS + OPS + REG + HI",              "7",   "3",  "20"],
        ["P4 Compliance Layer",                         "5",   "2",  "25"],
        ["Subtotal — core delivery",                    "25",  "9",  "25"],
        ["P5 Live Dispatch (optional)",                 "10",  "2",  "35"],
    ], col_widths_inches=[2.55, 0.95, 2.30, 0.70], font_size=9)
    add_para(doc, "Reviewer effort is approval bandwidth — not authoring. "
                  "Authoring is owned by the engineering team; CX, Compliance "
                  "and Brand retain approval authority. The 9 reviewer days "
                  "are calendar-distributed across the project, not consecutive.",
                  italic=True, color=MEDIUM, size=9)

def section_governance(doc, num):
    add_section_heading(doc, num, "Governance & Review Cadence")
    add_table(doc, ["Forum", "Frequency", "Attendees", "Purpose"], [
        ["Engineering standup",          "Daily",             "Engineering team",                                 "Build coordination"],
        ["Weekly status",                 "Weekly",            "Engineering + CX Head",                            "Progress, risks, decisions needed"],
        ["Phase-gate review",            "End of each phase", "Engineering + CX Head + Compliance + CTO + Sponsor","Acceptance and go/no-go for next phase"],
        ["Compliance approval queue",   "Continuous",         "Compliance Reviewer",                              "Template-level approvals"],
        ["Brand-voice approval queue",  "Continuous",         "Brand Lead",                                       "Tone and copy approvals"],
        ["Quarterly post-go-live",       "Quarterly",         "All stakeholders",                                 "Measure objectives O1-O7 against baselines"],
    ], col_widths_inches=[1.90, 1.35, 1.90, 1.35], font_size=9)

def section_glossary(doc, num):
    add_section_heading(doc, num, "Glossary")
    add_table(doc, ["Term", "Meaning"], [
        ["Assignment",        "Legal transfer of policy ownership from policyholder (assignor) to investor (assignee)"],
        ["BSP",               "Business Solution Provider — Meta-approved partner for WhatsApp Business API (Karix, Gupshup, etc.)"],
        ["CKYC",              "Central KYC Registry — India-wide KYC repository"],
        ["DLT",               "Distributed Ledger Technology — TRAI-mandated registry for SMS sender IDs and templates in India"],
        ["DPDPA",             "Digital Personal Data Protection Act, 2023 — India’s data-protection law"],
        ["FATCA",             "Foreign Account Tax Compliance Act — applies to NRI / FII investors"],
        ["IRDAI",             "Insurance Regulatory and Development Authority of India"],
        ["MIS Pipeline",      "Existing TPE project (TPE-MIS-PLAN-2026-001) for ingesting insurer allocation files"],
        ["NOC",               "No-Objection Certificate from insurer for assignment"],
        ["PE ID",             "Principal Entity ID — DLT registration ID for TPE on TRAI’s registry"],
        ["RCS",               "Rich Communication Services — successor to SMS supporting buttons, carousels, branded sender"],
        ["Surrender value",   "Cash value paid by insurer if policyholder discontinues a policy mid-term"],
        ["Trigger",           "A defined lifecycle event that initiates one or more communications"],
        ["Template",          "A reusable, versioned content atom with a defined lifecycle"],
        ["Variant",           "A linked template that is a translation, channel, or A/B alternative of a parent template"],
    ], col_widths_inches=[1.65, 4.85], font_size=9)

def section_references(doc, num):
    add_section_heading(doc, num, "References")
    add_table(doc, ["Reference", "Source"], [
        ["The Policy Exchange — About",                      "https://www.thepolicyexchange.com/about-us"],
        ["TPE platform",                                      "https://platform.thepolicyexchange.com/"],
        ["Pre-Series B funding announcement (Nov 2025)",     "Business Standard (ANI press release)"],
        ["TPE company profile",                               "Tracxn"],
        ["TPE feature",                                       "CIO Tech Outlook (2025)"],
        ["MIS Pipeline implementation plan",                  "TPE-MIS-PLAN-2026-001 (~/Downloads/MIS Data/)"],
        ["Enterprise document standards",                     "~/.claude/CLAUDE.md (global)"],
        ["TRAI DLT framework",                                "trai.gov.in"],
        ["IRDAI advertisement regulations",                   "irdai.gov.in"],
        ["WhatsApp Business Platform — template categories",   "developers.facebook.com/docs/whatsapp"],
        ["DPDPA, 2023",                                        "meity.gov.in"],
    ], col_widths_inches=[2.85, 3.65], font_size=9)

def section_signoff(doc, num):
    add_section_heading(doc, num, "Sign-off")
    add_para(doc, "This Statement of Work is approved as the basis for "
                  "implementation. Phase-gate sign-offs follow this same "
                  "approval mechanism.", space_after=8)
    add_table(doc, ["Role", "Name", "Signature", "Date"], [
        ["Sponsor (Founder)",          "", "", ""],
        ["CTO",                        "", "", ""],
        ["Head of Compliance",         "", "", ""],
        ["Head of Customer Experience","", "", ""],
        ["Project Lead (Author)",      AUTHOR, "", TODAY],
    ], col_widths_inches=[1.90, 2.10, 1.45, 1.05], font_size=9.5)

# ─────────────────────────────────────────────────────────────────────────────
# v2.0 NEW SECTIONS — Enterprise-grade additions
# ─────────────────────────────────────────────────────────────────────────────

def section_doc_control(doc, num):
    """Document Control & Approval Matrix — formal governance of this document."""
    add_section_heading(doc, num, "Document Control & Approval Matrix")
    add_para(doc,
        "This section formalises ownership, review and approval of this "
        "Statement of Work as a controlled document. Any change to scope, "
        "effort, deliverables or commercial terms requires re-approval "
        "through the matrix below before the change takes effect.",
        space_after=6)

    add_section_heading(doc, f"{num}.1", "Approval matrix", level=2)
    add_table(doc, ["Role", "Name", "Approval scope", "Approval mechanism"], [
        ["Sponsor — Founder",      "Naveen Gupta / Safia Anwar / Tarun Bahri",
                                    "Commercial terms, strategic intent, sign-off",
                                    "Wet/electronic signature on Section 41"],
        ["Chief Technology Officer","TPE CTO",
                                    "Architecture, integrations, security posture",
                                    "Electronic signature; phase-gate review"],
        ["Head of Compliance",      "TPE Compliance Lead",
                                    "Regulatory scope (DLT, IRDAI, DPDPA, FATCA)",
                                    "Electronic signature; lint-rule sign-off"],
        ["Head of CX",              "TPE CX Head",
                                    "Trigger catalogue, brand voice, audience map",
                                    "Electronic signature; review queue"],
        ["Project Lead (Author)",   AUTHOR,
                                    "Scope decomposition, plan integrity, weekly status",
                                    "Owns this document; cuts new versions"],
    ], col_widths_inches=[1.50, 1.85, 2.10, 1.05], font_size=9)

    add_section_heading(doc, f"{num}.2", "Document state machine", level=2)
    add_table(doc, ["State", "Meaning", "Permitted next states"], [
        ["draft",         "Author iterating; not yet circulated",                "in_review, archived"],
        ["in_review",     "Circulated to approvers; comments open",              "approved, draft (returned)"],
        ["approved",      "All approvers signed; document is the project bible", "amended (new version), archived"],
        ["amended",       "Substantive change in flight; previous version retained","approved, archived"],
        ["archived",      "Superseded; read-only; retained for audit",            "(terminal)"],
    ], col_widths_inches=[1.10, 3.50, 1.90], font_size=9)

    add_section_heading(doc, f"{num}.3", "Change-control on this document", level=2)
    for b in [
        "MAJOR version (x.0) — change in scope, deliverables, effort, or commercial terms. Requires re-approval by all roles in the matrix above.",
        "MINOR version (1.x) — clarification, additional detail, error correction. Approval by Project Lead and one other approver.",
        "PATCH version (1.5.x) — typographical or formatting only. Project Lead authority alone.",
        "Every version cuts a new entry in the Version History (Section 3) with a change rationale.",
        "Comments and redlines are captured in-document during in_review; final approval clears the comments.",
    ]:
        add_bullet(doc, b)

def section_distribution(doc, num):
    """Distribution List — controlled distribution of this document."""
    add_section_heading(doc, num, "Distribution List")
    add_para(doc,
        "This document is classified Internal / Confidential. Distribution "
        "is restricted to the parties named below. Re-distribution, "
        "external sharing or use in external communications requires written "
        "approval from the Sponsor.",
        space_after=6)
    add_table(doc, ["Recipient", "Role", "Format", "Purpose"], [
        ["Sponsor (Founders)",          "Approver",          "DOCX + signed PDF", "Approve scope, sign Section 41"],
        ["CTO",                          "Approver",          "DOCX",              "Architecture sign-off"],
        ["Head of Compliance",          "Approver",          "DOCX",              "Regulatory sign-off"],
        ["Head of CX",                   "Approver",          "DOCX",              "Brand and trigger sign-off"],
        ["Operations Lead",             "Implementation owner","DOCX",             "Day-to-day delivery"],
        ["Engineering Team",            "Implementer",       "DOCX",              "Build reference"],
        ["Internal Audit (post-launch)","Reviewer",          "PDF (read-only)",   "Audit trail evidence"],
        ["Document repository",         "System of record",  "DOCX + PDF",        "Versioned archive"],
    ], col_widths_inches=[2.10, 1.40, 1.40, 1.60], font_size=9)
    add_callout(doc, "Why distribution is named, not broadcast",
        "Internal SOWs that travel beyond their named distribution list lose "
        "their authority. Approvers need confidence that they are signing the "
        "version everyone is reading. The list above is the canonical "
        "audience for v2.0; any extension is a controlled change.",
        kind="navy")

def section_executive_dashboard(doc, num):
    """Executive Dashboard — KPI tiles + status badges + readiness summary."""
    add_section_heading(doc, num, "Executive Dashboard")
    add_para(doc,
        "Single-page snapshot for the Sponsor. Status badges show project "
        "readiness for each delivery dimension; metric tiles summarise "
        "what an approver is signing up for. Refreshed at every phase "
        "gate.", space_after=8)

    add_para(doc, "Project readiness status", size=11, bold=True, color=NAVY,
             space_after=4)
    add_badge_row(doc, [
        ("Scope locked",      "green"),
        ("Effort baselined",  "green"),
        ("Risks cataloged",   "green"),
        ("Compliance mapped", "green"),
        ("Phase 5 gated",     "orange"),
    ])
    add_para(doc, "", space_after=6)

    add_para(doc, "Headline metrics", size=11, bold=True, color=NAVY,
             space_after=4)
    add_metric_tiles(doc, [
        ("49",  "Triggers",            "1B3A5C"),
        ("5",   "Audiences",            "2E75B6"),
        ("4",   "Channels",             "27AE60"),
        ("25",  "Engineering days",     "E67E22"),
    ])
    add_metric_tiles(doc, [
        ("9",   "Reviewer days",        "2E75B6"),
        ("4",   "Phases (core)",        "27AE60"),
        ("6",   "Compliance regimes",   "E67E22"),
        ("20%", "Default-rate target ↓","E74C3C"),
    ])

    add_para(doc, "Phase-gate health (forecast)", size=11, bold=True,
             color=NAVY, space_after=4)
    add_table_with_severity(doc,
        ["Phase", "Health", "Forecast outcome", "Watch-item"],
        [
        ["P1 Foundation",          "L", "Repo + 5 skills + PH-03 walkthrough",        "None — straightforward scaffolding"],
        ["P2 PH Library",          "L", "All 18 PH triggers + ~50 templates",         "Spreadsheet content quality"],
        ["P3 INV+INS+OPS+REG+HI",  "M", "31 triggers; Hindi variants",                "Reviewer bandwidth"],
        ["P4 Compliance Layer",    "L", "30+ lint rules; approval gate",              "Disclaimer text final from Compliance"],
        ["P5 Live Dispatch",       "M", "MSG91/Karix/Gupshup; WA BSP",                "Channel-provider contracts (out of scope)"],
        ], severity_cols={1}, col_widths_inches=[2.20, 0.65, 2.45, 1.20], font_size=9)

    add_callout(doc, "What an approver is signing for",
        "Twenty-five engineering days plus nine reviewer days, four phases, "
        "an independently consumable artefact at the end of each phase, and a "
        "measurable twenty-percent target reduction in investor premium-"
        "default rate within six months of go-live. Phase 5 (live dispatch) "
        "is held separately and gated on channel-provider selection — it is "
        "not on the critical path for value realisation.", kind="green")

def section_business_case(doc, num):
    """Business Case & ROI Analysis — quantified justification for the spend."""
    add_section_heading(doc, num, "Business Case & ROI Analysis")
    add_para(doc,
        "Twenty-five engineering days of internal effort is the input. The "
        "outputs justifying it are quantified below across four dimensions: "
        "operational savings, risk reduction, revenue protection, and "
        "compounding benefit. The single largest contributor is investor "
        "premium-default-rate reduction — every percentage point preserved is "
        "lifetime-value retained on a 10–25 year stream.",
        space_after=8)

    add_section_heading(doc, f"{num}.1", "Quantified value drivers", level=2)
    add_table(doc, ["Driver", "Baseline", "v1 Target", "Annualised value"], [
        ["Investor premium-default rate (revenue protection)",
         "Internal benchmark",    "−20% within 6 months",
         "₹ multi-crore — preserves multi-decade premium streams; modelled separately"],
        ["Ops time on routine drafting (operational saving)",
         "≥ 2h per template",     "≤ 30 min draft → approved",
         "≥ 20 person-hours/week saved → ≈ 1,000 hrs/yr"],
        ["Compliance-incident exposure (risk reduction)",
         "Manual checks, ad-hoc", "Lint at approval gate",
         "Avoided fines/blocks; quantified in Risk Register Section 32"],
        ["Time-to-launch new audience or channel (compounding)",
         "Weeks of bespoke work", "Days, via reuse + variants",
         "Future-state enabler — materialises in v2 / regional expansion"],
        ["Audit reconstruction effort (risk reduction)",
         "Days per regulator query","100% versioned, near-zero",
         "Material if a regulator query lands; insurance against rare events"],
    ], col_widths_inches=[2.10, 1.50, 1.40, 1.50], font_size=8.5)

    add_section_heading(doc, f"{num}.2", "Investment vs. return summary", level=2)
    add_table(doc, ["Item", "Quantum", "Notes"], [
        ["Engineering effort (P1–P4)",       "25 days",       "Internal team; opportunity cost"],
        ["Reviewer effort (CX+Comp+Brand)",  "9 days",        "Spread across project; not consecutive"],
        ["Phase 5 effort (gated)",           "10 days",       "Optional; depends on channel-provider contracts"],
        ["Year-1 ops time saved (estimate)", "≥ 1,000 hrs",   "≈ 0.5 FTE — funds onward improvements"],
        ["Year-1 default-rate impact",       "−20% target",   "Largest single driver of project NPV"],
        ["Indicative payback period",        "≤ 6 months",    "Default-rate impact alone clears the effort cost"],
    ], col_widths_inches=[2.30, 1.50, 2.70], font_size=9)

    add_callout(doc,
        "How we frame the case",
        "We do not claim a precise rupee NPV. Default-rate, ops-saving and "
        "incident-avoidance translate into rupees only in retrospect, with "
        "the actual numbers measured at the 6-month review (Section 6 "
        "objectives). The honest claim is that the targets are conservative, "
        "the baseline is the current spreadsheet-and-inbox state, and the "
        "effort is small enough that even modest realisation on any one "
        "driver clears the cost.", kind="green")

def section_stakeholder_register(doc, num):
    """Stakeholder Register — interest × influence map + engagement strategy."""
    add_section_heading(doc, num, "Stakeholder Register")
    add_para(doc,
        "The RACI (Section 25) tells you who does what on a given activity. "
        "This register tells you who cares about the project at all, how "
        "much, and what they need to keep that interest constructive. "
        "Stakeholders are listed in priority order; engagement intensity is "
        "calibrated to influence × interest.",
        space_after=8)
    add_table(doc, ["Stakeholder", "Influence", "Interest", "Concerns", "Engagement"],
              [
        ["Founders (Sponsors)",         "H", "H",
                                         "Burn vs. impact; investor trust",
                                         "Phase-gate review; weekly status"],
        ["CTO",                          "H", "H",
                                         "Architecture, integrations, tech-debt",
                                         "Weekly review; phase-gate sign-off"],
        ["Head of Compliance",          "H", "H",
                                         "Regulator exposure; lint accuracy",
                                         "Continuous review queue"],
        ["Head of CX",                   "H", "H",
                                         "Brand voice; PH experience",
                                         "Continuous review queue"],
        ["Operations Lead",             "M", "H",
                                         "Day-to-day workload; tooling",
                                         "Weekly working session"],
        ["Engineering Team",            "M", "H",
                                         "Build clarity; standards",
                                         "Daily standup; pair authoring"],
        ["IRDAI (regulator)",            "H", "L",
                                         "Statutory compliance",
                                         "Indirect — via Compliance Lead"],
        ["TRAI / DLT operators",         "H", "L",
                                         "SMS template registration",
                                         "Indirect — via CTO at P5"],
        ["Insurance partners (11)",      "M", "M",
                                         "MIS continuity; assignment SLA",
                                         "Existing partner cadence; INS-02"],
        ["Investors (TPE customers)",    "L", "H",
                                         "Reliable premium reminders",
                                         "Channelled via INV trigger flows"],
        ["Policyholders",               "L", "H",
                                         "Empathetic, vernacular contact",
                                         "Channelled via PH trigger flows"],
        ["Internal Audit",              "M", "L",
                                         "Audit-trail quality",
                                         "Documented post-launch review"],
              ], col_widths_inches=[1.55, 0.55, 0.55, 2.10, 1.75], font_size=8.5)

    add_callout(doc, "Influence × Interest = engagement intensity",
        "High influence + High interest stakeholders (founders, CTO, "
        "Compliance, CX) are managed closely with weekly contact and "
        "phase-gate sign-off. High influence + Low interest stakeholders "
        "(IRDAI, TRAI) are kept satisfied via the Compliance Lead — they "
        "should not need to interact with this project directly. Low "
        "influence + High interest (PH, INV) are the reason the project "
        "exists — they are served via the trigger flows, not via direct "
        "project engagement.", kind="navy")

def section_security_privacy(doc, num):
    """Security & Privacy Framework — DPDPA-aligned controls."""
    add_section_heading(doc, num, "Security & Privacy Framework")
    add_para(doc,
        "Communication systems concentrate Personally Identifiable "
        "Information (PII): names, phone numbers, policy numbers, premium "
        "amounts, bank-last-four. Misuse, leak or unauthorised access "
        "creates regulatory exposure under DPDPA and reputational damage "
        "with both sides of the marketplace. The controls below are "
        "designed-in, not bolted-on, and are validated as part of "
        "acceptance (Section 24).", space_after=6)

    add_section_heading(doc, f"{num}.1", "Data classification", level=2)
    add_table(doc, ["Class", "Examples in this system", "Handling rule"], [
        ["PII — High",        "ph_name, ph_phone, ph_email, inv_name, PAN, bank_last4",
                              "Tokenise in templates; never log raw; mask in audit; HTTPS only in transit"],
        ["PII — Medium",      "policy_no, premium_amount, due_date, surrender_value",
                              "Encrypted at rest; access on least-privilege; redact in screenshots"],
        ["Operational",        "trigger_id, template_id, status, cadence",
                              "Public within company; safe to log"],
        ["Brand confidential","Disclaimer text, brand voice rules",
                              "Internal-only; not for external comms"],
    ], col_widths_inches=[1.30, 3.20, 2.00], font_size=9)

    add_section_heading(doc, f"{num}.2", "Controls", level=2)
    add_table(doc, ["Control area", "Control"], [
        ["Access control",      "Role-based: author / reviewer / operator / admin; least-privilege; reviewed quarterly"],
        ["Authentication",       "SSO via TPE identity provider; MFA mandatory for reviewer/admin"],
        ["Authorisation gates",  "review→approved transition requires Compliance + Brand reviewer; no self-approval"],
        ["Encryption — transit", "TLS 1.2+ for all network calls; certificate-pinned for BSP/SMS gateway"],
        ["Encryption — rest",    "AES-256 for repository storage; key rotation per TPE security policy"],
        ["Secret management",    "Channel-provider API keys in vault; never in repo; rotated per provider policy"],
        ["Audit log",             "Append-only; per-template state transitions, edits, sends; PII masked"],
        ["PII minimisation",     "Variables resolved at render time, not stored in audit log"],
        ["Backup & recovery",    "Daily Git-replicated; weekly snapshot; documented RPO 24h / RTO 4h"],
        ["Vulnerability mgmt",   "Quarterly dependency scan; security review at every major version"],
    ], col_widths_inches=[1.65, 4.85], font_size=9)

    add_section_heading(doc, f"{num}.3", "DPDPA alignment", level=2)
    add_table(doc, ["DPDPA principle", "How this system aligns"], [
        ["Lawful processing",         "Templates carry compliance_class; lint blocks promotional sends without consent reference"],
        ["Purpose limitation",        "Trigger taxonomy makes intent explicit; templates linked to triggers; reuse is auditable"],
        ["Data minimisation",          "Variables declare what is needed; lint flags unused PII variables"],
        ["Storage limitation",         "Audit retains state transitions, not raw PII; backups retained per data-retention policy"],
        ["Accuracy",                    "Variables source from system of record; fallbacks prevent broken renders"],
        ["Notice & consent",            "Opt-out wiring on every promotional template; opt-out evidence stored at BSP"],
        ["Rights of data principals",  "Subject-access response supported via audit-trail export; deletion on request"],
        ["Breach notification",         "Incident runbook (Section 37) covers detection, escalation, regulator notification"],
    ], col_widths_inches=[1.85, 4.65], font_size=9)

def section_qa_testing(doc, num):
    """Quality Assurance & Testing Strategy."""
    add_section_heading(doc, num, "Quality Assurance & Testing Strategy")
    add_para(doc,
        "Quality is engineered, not reviewed in. The system is tested at "
        "five levels — schema, lint, render, trigger-template wiring and "
        "end-to-end — each catching a different class of defect. Test "
        "packs are part of the deliverable (Section 23) and run on every "
        "change.", space_after=6)

    add_section_heading(doc, f"{num}.1", "Test levels", level=2)
    add_table(doc, ["Level", "What it catches", "Example", "Run when"], [
        ["L1 Schema",
         "Malformed YAML, missing required fields, invalid enum values",
         "Template with no compliance_class fails parse",
         "Pre-commit + CI on every push"],
        ["L2 Lint",
         "Compliance violations: DLT, IRDAI, WA BSP, DPDPA, brand-voice",
         "WhatsApp template marked Marketing without opt-out",
         "Pre-commit + before review→approved transition"],
        ["L3 Render",
         "Channel-renderer regressions: char limits, button schema, encoding",
         "SMS body that exceeds 160 chars after variable substitution",
         "CI on every push"],
        ["L4 Wiring",
         "Trigger references a template that does not exist or is deprecated",
         "PH-03 cadence references a template that was archived",
         "CI nightly + before phase-gate"],
        ["L5 End-to-end",
         "Full flow: trigger fires → template selected → lint passes → renders to channel-valid payload",
         "PH-03 KYC pending week-2 → WA template → MSG91 dry-run payload",
         "Phase-gate + before any production release"],
    ], col_widths_inches=[0.70, 2.20, 2.00, 1.60], font_size=8.5, code_cols={0})

    add_section_heading(doc, f"{num}.2", "Acceptance test pack (per phase)", level=2)
    add_table(doc, ["Phase", "Pack composition"], [
        ["P1", "≥ 30 schema tests + 5 lint rules + 2 render goldens + PH-03 E2E"],
        ["P2", "+ 1 wiring test per PH trigger; render goldens for Email + WA × all PH active statuses"],
        ["P3", "+ INV/INS/OPS/REG wiring; Hindi-language render goldens; localisation queue test"],
        ["P4", "+ ≥ 30 lint rules across 6 regimes; approval-gate enforcement test; audit-trail export test"],
        ["P5", "+ MSG91/Karix/Gupshup dry-run goldens; WA BSP template ID validation; receipt ingest"],
    ], col_widths_inches=[0.55, 5.95], font_size=9, code_cols={0})

    add_section_heading(doc, f"{num}.3", "Defect severity & SLA", level=2)
    add_table_with_severity(doc, ["Severity", "Definition", "Response SLA", "Resolution SLA"],
        [
        ["Sev-1 (H)", "Compliance violation in active template OR data leak",
                       "1h", "Same business day; rollback if needed"],
        ["Sev-2 (H)", "Active template broken (cannot render) OR active trigger silent",
                       "4h", "1 business day"],
        ["Sev-3 (M)", "Active template flawed but still rendering (typo, mild wording)",
                       "1 business day", "1 sprint"],
        ["Sev-4 (L)", "Cosmetic / docs / non-active template",
                       "1 sprint", "Backlog"],
        ], severity_cols={0}, col_widths_inches=[1.10, 2.95, 1.30, 1.15], font_size=9)

def section_comms_management(doc, num):
    """Communication Management Plan — how the project itself communicates."""
    add_section_heading(doc, num, "Communication Management Plan")
    add_para(doc,
        "How the project itself communicates — distinct from the "
        "communication-system being built. Defined cadence, owners and "
        "channels prevent the project from generating its own coordination "
        "drag.", space_after=6)
    add_table(doc, ["Forum", "Frequency", "Owner", "Audience", "Format", "Output"], [
        ["Daily standup",          "Daily 10:30 IST",       "Project Lead", "Engineering",       "Slack thread",       "Daily update + blockers"],
        ["Weekly status",           "Friday 16:00 IST",      "Project Lead", "All approvers",     "DOCX 1-pager + Slack","Status, risks, decisions"],
        ["Phase-gate review",       "End of P1, P2, P3, P4", "Project Lead", "All approvers",     "Live demo + DOCX",   "Sign-off → next phase"],
        ["Compliance review queue", "Continuous",            "Compliance",   "Authors",           "/tpe-template review","Approval / rejection"],
        ["Brand review queue",      "Continuous",            "Brand Lead",   "Authors",           "/tpe-template review","Approval / rejection"],
        ["Risk review",             "Bi-weekly",             "Project Lead", "CTO + Sponsor",     "Updated Risk Register","New / closed risks"],
        ["Stakeholder digest",      "Monthly",               "Project Lead", "All stakeholders",  "Email",              "Progress + asks"],
        ["Post-mortem",             "Per Sev-1/Sev-2",       "Project Lead", "All approvers",     "Markdown",           "Action items"],
    ], col_widths_inches=[1.55, 1.35, 0.85, 1.05, 0.95, 0.75], font_size=8)

    add_callout(doc, "Information radiator, not push",
        "Status is published to a single source (the Slack channel + "
        "weekly DOCX). Approvers pull when they need to; the project lead "
        "does not chase or duplicate. Push is reserved for blockers and "
        "phase-gate decisions.", kind="navy")

def section_change_management(doc, num):
    """Change Management Process — how scope/effort/timeline changes are controlled."""
    add_section_heading(doc, num, "Change Management Process")
    add_para(doc,
        "Any change to scope, effort, timeline, deliverables or commercial "
        "terms follows the process below. Change requests captured "
        "informally — in Slack, email or hallway — must be re-stated as a "
        "formal Change Request before they can take effect.",
        space_after=6)

    add_section_heading(doc, f"{num}.1", "Change request lifecycle", level=2)
    add_state_flow(doc, [
        ("RAISED",     "7F8C8D"),
        ("TRIAGED",    "2E75B6"),
        ("ASSESSED",   "E67E22"),
        ("DECIDED",    "27AE60"),
        ("APPLIED",    "1B3A5C"),
        ("CLOSED",     "27AE60"),
    ])
    add_table(doc, ["State", "Owner", "SLA", "Output"], [
        ["raised",   "Anyone",         "—",                "CR ticket with title + rationale"],
        ["triaged",  "Project Lead",   "1 business day",    "Severity + initial impact estimate"],
        ["assessed", "Eng + reviewers","3 business days",   "Effort delta, scope delta, risk delta"],
        ["decided",  "CCB",            "Next CCB or 5 days","Approve / reject / defer"],
        ["applied",  "Project Lead",   "Per agreed plan",   "Document updated; new version cut"],
        ["closed",   "Project Lead",   "Within 1 sprint",   "CR archived with audit"],
    ], col_widths_inches=[1.00, 1.40, 1.65, 2.45], font_size=9)

    add_section_heading(doc, f"{num}.2", "Change Control Board (CCB)", level=2)
    add_table(doc, ["Field", "Value"], [
        ["Members",       "Sponsor, CTO, Compliance, CX, Project Lead"],
        ["Quorum",        "Sponsor + CTO + (one of Compliance / CX) + Project Lead"],
        ["Frequency",     "On demand for Sev-1 / Sev-2; otherwise weekly when CRs are open"],
        ["Decision rule", "Consensus; deadlock escalates to Sponsor"],
        ["Authority",     "Approve / reject / defer scope, effort, deliverable, commercial changes"],
        ["Records",       "CCB minutes filed against the CR; Version History updated on apply"],
    ], col_widths_inches=[1.50, 5.00], font_size=9)

    add_section_heading(doc, f"{num}.3", "Change classification", level=2)
    add_table(doc, ["Class", "Triggers", "Decision authority"], [
        ["Minor",     "Editorial, prose, tone, examples; no effort impact",          "Project Lead"],
        ["Moderate",  "≤ 1 day effort; no scope change; no deliverable change",      "Project Lead + 1 approver"],
        ["Material",  "> 1 day effort OR scope/deliverable change OR risk change",   "Full CCB"],
        ["Critical",  "Compliance violation surfaced OR Sponsor-level commercial",    "Full CCB + Sponsor"],
    ], col_widths_inches=[1.10, 3.40, 2.00], font_size=9)

def section_sla(doc, num):
    """Service Level Agreements — operational SLAs the system commits to."""
    add_section_heading(doc, num, "Service Level Agreements")
    add_para(doc,
        "SLAs the delivered system commits to (post go-live), and SLAs "
        "the project itself commits to during build. The first set governs "
        "operations; the second set governs delivery.",
        space_after=6)

    add_section_heading(doc, f"{num}.1", "Operational SLAs (post go-live)", level=2)
    add_table(doc, ["Service", "Metric", "Target", "Measurement"], [
        ["Template approval queue",   "Time from review→approved",          "≤ 24h business day",           "Per template, audit log"],
        ["Lint engine availability",  "Uptime",                              "99.5% rolling 30 days",         "Pre-commit + CI metrics"],
        ["Render correctness",        "Sev-1 incidents per quarter",         "0",                             "Production observation"],
        ["OTP delivery (P5)",         "P95 delivery latency",                "≤ 10s",                         "MSG91 / BSP receipts"],
        ["Premium-due send (P5)",     "Send-window adherence",               "100% within DLT 9–21 IST window","DLT compliance report"],
        ["Audit-trail export",        "Time to produce per-template audit",  "≤ 60s",                         "Operational test"],
        ["Subject-access response",   "DPDPA SAR turnaround",                "≤ 7 calendar days",            "Incident log"],
    ], col_widths_inches=[1.85, 1.95, 1.65, 1.05], font_size=9)

    add_section_heading(doc, f"{num}.2", "Delivery SLAs (during build)", level=2)
    add_table(doc, ["Activity", "Target SLA", "Owner"], [
        ["Phase-gate review pack ready",       "T-2 business days before review",  "Project Lead"],
        ["Compliance review of template",       "≤ 24h business day",                "Compliance Lead"],
        ["Brand review of template",            "≤ 24h business day",                "Brand Lead"],
        ["Reviewer-blocked CR triage",          "Same business day",                  "Project Lead"],
        ["Sev-1 defect response",               "1h",                                  "Engineering"],
        ["Weekly status report publication",    "Friday 16:00 IST",                   "Project Lead"],
        ["CR decision (Material+)",             "Within 5 business days of raised",   "CCB"],
    ], col_widths_inches=[2.85, 2.05, 1.60], font_size=9)

    add_section_heading(doc, f"{num}.3", "SLA breach handling", level=2)
    for b in [
        "First breach in a rolling 30-day window: documented in weekly status, root-cause noted, no further action.",
        "Second breach in 30 days: project lead escalates to CTO + Sponsor; mitigation plan due within 3 business days.",
        "Three or more breaches in 30 days: CCB session called; phase plan re-baselined if material.",
        "Breaches are visible — no breach is concealed by re-defining the metric.",
    ]:
        add_bullet(doc, b)

def section_resource_plan(doc, num):
    """Resource Plan — named roles, allocations and substitutability."""
    add_section_heading(doc, num, "Resource Plan")
    add_para(doc,
        "Effort estimation (Section 33) gives total days. The Resource "
        "Plan names who does them, when, and what the substitution path "
        "is when an individual is unavailable. Concentration risk on any "
        "single individual is called out.", space_after=6)

    add_table(doc, ["Role", "Allocation", "Person", "Substitution path"], [
        ["Engineering Lead (Author)",    "100% during P1–P2; 75% P3–P4",
         AUTHOR,                          "Senior Engineer can take over with 1-week handover"],
        ["Senior Engineer",              "50% during P2–P4",
         "TBA (TPE Eng team)",           "Engineering Lead covers gaps"],
        ["Compliance Reviewer",         "30 min/day during P2–P4",
         "TPE Compliance Lead",          "Designated alternate during PTO"],
        ["Brand Reviewer (CX)",          "30 min/day during P2–P3",
         "TPE CX Head",                  "CX-2 reviewer with named scope"],
        ["Operations Lead",              "10% across all phases",
         "TPE Ops Lead",                 "Project Lead coordinates directly"],
        ["CTO",                          "Phase-gate + 1 weekly",
         "TPE CTO",                      "Engineering Lead carries to next gate"],
        ["Sponsor",                      "Phase-gate + as escalated",
         "Founder",                      "Joint founder cover"],
    ], col_widths_inches=[1.85, 1.55, 1.40, 1.70], font_size=8.5)

    add_callout(doc, "Concentration risk: Engineering Lead",
        "The Engineering Lead is the single dependency for build velocity in "
        "P1–P2. The mitigation is a documented authoring guide (deliverable "
        "in P2), pair-authoring rotations from P3 onwards, and a Senior "
        "Engineer onboarded by end of P2. We are not architecting around a "
        "second engineer for P1 — the velocity benefit does not exceed the "
        "onboarding cost for a 5-day phase.", kind="orange")

def section_operational_readiness(doc, num):
    """Operational Readiness & Go-Live Criteria."""
    add_section_heading(doc, num, "Operational Readiness & Go-Live Criteria")
    add_para(doc,
        "Operational readiness is checked phase-by-phase. The criteria "
        "below are not aspirational — they are go/no-go gates. A phase "
        "that does not meet its criteria does not transition; the cause "
        "is documented and reset.", space_after=6)

    add_section_heading(doc, f"{num}.1", "Readiness checklist (per phase)", level=2)
    add_table(doc, ["Dimension", "P1", "P2", "P3", "P4", "P5"], [
        ["Schema validation passes",      "✓","✓","✓","✓","✓"],
        ["Lint coverage adequate",        "—","Partial","Partial","✓","✓"],
        ["Render goldens up-to-date",     "✓","✓","✓","✓","✓"],
        ["Trigger wiring tests pass",     "—","✓","✓","✓","✓"],
        ["E2E test for representative trigger","✓","✓","✓","✓","✓"],
        ["Audit-trail export tested",      "—","—","—","✓","✓"],
        ["Compliance reviewer signed",     "—","✓","✓","✓","✓"],
        ["Brand reviewer signed",          "—","✓","✓","—","—"],
        ["Sponsor sign-off on phase",      "✓","✓","✓","✓","✓"],
        ["Channel-provider contracts (P5)","—","—","—","—","Required"],
    ], col_widths_inches=[3.50, 0.60, 0.60, 0.60, 0.60, 0.60], font_size=8.5)

    add_section_heading(doc, f"{num}.2", "Go-live criteria (whole system)", level=2)
    for b in [
        "All P1–P4 acceptance criteria (Section 24) signed by named approvers.",
        "Risk Register (Section 32) shows zero open Sev-1 risks; Sev-2 risks have documented mitigations.",
        "Incident runbook (Section 37) tested in tabletop with engineering + compliance.",
        "Rollback procedure documented and tested for the most-recently-changed template family.",
        "Subject-access request (DPDPA SAR) end-to-end tested for at least one investor and one policyholder.",
        "Phase 5 only: DLT registration, WhatsApp BSP template approval and channel-provider sandbox all green.",
        "Communication Management Plan (Section 26) explicitly transitions to BAU operations cadence.",
    ]:
        add_bullet(doc, b)

    add_callout(doc, "What 'go-live' means without Phase 5",
        "Phases 1–4 deliver real value without Phase 5 dispatch: a versioned "
        "library, a working lint engine, an audit-grade trail, and a "
        "review→approved approval gate. 'Go-live' for P1–P4 means the team "
        "stops drafting in spreadsheets and starts authoring in the system. "
        "Live dispatch (P5) is a later, separately-gated milestone.",
        kind="green")

def section_knowledge_transfer(doc, num):
    """Knowledge Transfer & Training Plan."""
    add_section_heading(doc, num, "Knowledge Transfer & Training Plan")
    add_para(doc,
        "The system has three distinct user populations — authors, "
        "reviewers and operators. Each gets purpose-built materials and "
        "live training before the phase that hands them responsibility. "
        "Knowledge transfer is treated as a deliverable, not an "
        "after-thought.", space_after=6)

    add_section_heading(doc, f"{num}.1", "Training tracks", level=2)
    add_table(doc, ["Audience", "Track", "Format", "Duration", "Phase"], [
        ["Engineering authors",  "Schema + slash-command authoring",   "Recorded walk-through + paired session", "2h",   "End P1"],
        ["Engineering authors",  "Lint rule authoring",                 "Live session + reference",                "1h",   "Mid P4"],
        ["Compliance reviewers", "Approval workflow + lint outcomes",  "Live session + reference card",          "1h",   "Start P2"],
        ["Brand reviewers (CX)", "Tone, voice, brand-voice rules",     "Live session + voice guide",             "1h",   "Start P2"],
        ["Operators (CX team)",  "Trigger preview + library search",   "Recorded + reference card",              "30min","Start P3"],
        ["Operators (Ops)",      "Audit-trail export + runbook",       "Live session",                            "30min","Start P4"],
        ["Internal Audit",       "Audit-evidence reading + reconstruction","Live session post-launch",            "1h",   "Post P4"],
    ], col_widths_inches=[1.75, 2.10, 1.55, 0.55, 0.55], font_size=8.5)

    add_section_heading(doc, f"{num}.2", "Reference materials (deliverables)", level=2)
    for b in [
        "Template Author Guide — authoring conventions, schema reference, common pitfalls.",
        "Reviewer Reference Card — approval checklist per regime; one-page printable.",
        "Operator Runbook — daily/weekly operational procedures, audit export, incident response.",
        "Compliance Quick-Reference — Section 9 of the Content Reference doc, distilled to laminated card form.",
        "Brand Voice Guide — voice attributes per audience, do's/don'ts, examples.",
        "Recorded walk-throughs — authoring, reviewing, operating; held in TPE knowledge base.",
        "Onboarding-day-one checklist — for new engineers, reviewers, operators.",
    ]:
        add_bullet(doc, b)

    add_section_heading(doc, f"{num}.3", "Knowledge retention", level=2)
    add_table(doc, ["Mechanism", "Frequency", "Owner"], [
        ["Pair-authoring rotations",  "From P3 onwards",  "Engineering Lead"],
        ["Quarterly refresh session", "Quarterly post go-live", "Project Lead"],
        ["Doc-as-code review",         "On every change",  "Reviewer of the change"],
        ["Bus-factor audit",          "Bi-annually",      "CTO"],
    ], col_widths_inches=[2.40, 2.00, 2.10], font_size=9)

def section_post_impl_support(doc, num):
    """Post-Implementation Support & Maintenance."""
    add_section_heading(doc, num, "Post-Implementation Support & Maintenance")
    add_para(doc,
        "After P4 acceptance, the system enters Business-As-Usual support. "
        "BAU has its own cadence, ownership and SLA — distinct from "
        "project delivery. Major-version upgrades follow the standard "
        "change-management process (Section 27).",
        space_after=6)

    add_section_heading(doc, f"{num}.1", "Support tiers", level=2)
    add_table(doc, ["Tier", "Scope", "Owner", "Response SLA"], [
        ["T1 — operator self-service",
         "Search library, render preview, audit export, FAQ in runbook",
         "Operator", "Self-serve"],
        ["T2 — engineering on-call",
         "Lint failures, render incidents, schema migration, dependency updates",
         "Engineering rota", "1 business day"],
        ["T3 — compliance escalation",
         "Regulator query, lint-rule update, disclaimer revision",
         "Compliance Lead", "Same business day"],
        ["Incident — Sev-1",
         "Active compliance violation OR data exposure",
         "Engineering on-call + Compliance",  "1h"],
    ], col_widths_inches=[1.80, 2.50, 1.30, 0.90], font_size=9)

    add_section_heading(doc, f"{num}.2", "Maintenance cadence", level=2)
    add_table(doc, ["Activity", "Frequency", "Owner"], [
        ["Disclaimer text refresh (per IRDAI updates)", "On demand",        "Compliance"],
        ["DLT template registry refresh",                "Per template add", "Compliance + Eng"],
        ["BSP template approval refresh",                "Per template add", "Eng (P5+)"],
        ["Dependency security scan",                      "Quarterly",         "Engineering"],
        ["Lint-rule audit",                               "Quarterly",         "Compliance + Eng"],
        ["Localisation queue review",                     "Monthly",           "Brand + Eng"],
        ["Audit-trail integrity check",                   "Monthly",           "Internal Audit"],
        ["Operational metrics review",                    "Quarterly",         "Project Lead + Sponsor"],
    ], col_widths_inches=[3.20, 1.45, 1.85], font_size=9)

    add_section_heading(doc, f"{num}.3", "Incident response (high-level)", level=2)
    add_state_flow(doc, [
        ("DETECTED",   "E74C3C"),
        ("TRIAGED",    "E67E22"),
        ("CONTAINED",  "2E75B6"),
        ("RESOLVED",   "27AE60"),
        ("REVIEWED",   "1B3A5C"),
        ("CLOSED",     "27AE60"),
    ])
    for b in [
        "Detected — automated monitor or user report; ticket auto-created with severity.",
        "Triaged — Engineering on-call assigns severity, scope and owner within 1h for Sev-1.",
        "Contained — stop-the-bleed: pause active dispatch, freeze affected templates.",
        "Resolved — root cause fixed, fix deployed, regression test added.",
        "Reviewed — post-mortem within 5 business days; published to all approvers.",
        "Closed — preventative actions added to backlog and tracked to completion.",
    ]:
        add_bullet(doc, b)

def section_confidentiality_ip(doc, num):
    """Confidentiality, IP & Data Ownership."""
    add_section_heading(doc, num, "Confidentiality, IP & Data Ownership")
    add_para(doc,
        "Internal SOWs that span multiple teams and external vendors need "
        "explicit treatment of confidentiality, intellectual property and "
        "data ownership. The clauses below are TPE-internal but written "
        "to survive any vendor handover that follows.",
        space_after=6)

    add_section_heading(doc, f"{num}.1", "Classification", level=2)
    add_table(doc, ["Classification", "Applies to", "Handling"], [
        ["Internal / Confidential",  "This SOW, the Content Reference, the trigger catalogue, the lint rules",  "Internal-only; named distribution"],
        ["Restricted",                "Customer PII, audit logs, channel-provider credentials",                    "Need-to-know; least-privilege; never external"],
        ["Public",                    "Open-source dependency licences",                                            "Standard open-source compliance"],
    ], col_widths_inches=[1.80, 3.10, 1.60], font_size=9)

    add_section_heading(doc, f"{num}.2", "Intellectual property", level=2)
    add_table(doc, ["Asset", "Owner", "Rights"], [
        ["This Statement of Work + Content Reference",
         "TPE (Fairvalue Insuretech Pvt. Ltd.)",
         "Full ownership; no third-party claims"],
        ["Trigger catalogue + cadence rules",
         "TPE",
         "Domain knowledge captured; no licence required to reuse internally"],
        ["Template content (drafted copy)",
         "TPE",
         "All approved copy is a TPE work-product"],
        ["Compliance lint rules (codified)",
         "TPE",
         "Encoded regulatory knowledge; updates follow regulator changes"],
        ["Channel-provider integrations (P5)",
         "Provider IP for SDK; TPE for integration code",
         "Standard SaaS terms apply"],
        ["Open-source dependencies",
         "Respective projects",
         "Licence inventory maintained"],
    ], col_widths_inches=[2.30, 2.10, 2.10], font_size=9)

    add_section_heading(doc, f"{num}.3", "Data ownership", level=2)
    add_table(doc, ["Data", "Owner", "Custodian"], [
        ["Customer PII (PH, INV)",         "Data principal (the customer)",  "TPE — processor under DPDPA"],
        ["Policy data",                     "Data principal + insurer",       "TPE — joint custody"],
        ["Audit log of communication",      "TPE",                            "TPE Engineering"],
        ["Insurer master data",             "Respective insurer",             "TPE — under bilateral SLA"],
        ["Template content (post-approval)","TPE",                            "Compliance + Brand"],
    ], col_widths_inches=[2.40, 2.20, 1.90], font_size=9)

    add_section_heading(doc, f"{num}.4", "Confidentiality clauses", level=2)
    for b in [
        "All persons named in the Distribution List (Section 4) are bound by their existing employment / engagement confidentiality terms with TPE.",
        "External vendors engaged in P5 (channel providers, BSPs, SMS gateways) are bound by separate NDAs that incorporate the data classification above.",
        "Public discussion of TPE's communication strategy, trigger catalogue, lint rules or compliance posture requires Sponsor approval.",
        "Audit evidence may be shared with regulators on demand; sharing follows the legal hold-and-disclose process and not this document's distribution rules.",
        "On termination of the project (Section 41 sign-off + 12 months) and absent renewal, this document moves to the archived state in the document repository.",
    ]:
        add_bullet(doc, b)


# ─────────────────────────────────────────────────────────────────────────────
# Main — v2.0 (41 sections in 11 parts)
# ─────────────────────────────────────────────────────────────────────────────

def main():
    doc = setup_document()

    # ── Front matter ────────────────────────────────────────────────────────
    add_cover_page(doc)
    add_table_of_contents(doc)
    section_executive_briefing(doc)

    # ── Part I — Document Control & Executive (1–7) ─────────────────────────
    section_doc_info(doc, "1")
    section_doc_control(doc, "2")
    section_version_history(doc, "3")
    doc.add_page_break()
    section_distribution(doc, "4")
    section_executive_summary(doc, "5")
    doc.add_page_break()
    section_executive_dashboard(doc, "6")
    doc.add_page_break()
    section_business_case(doc, "7")
    doc.add_page_break()

    # ── Part II — Context & Scope (8–11) ────────────────────────────────────
    section_business_context(doc, "8")
    section_problem_current_state(doc, "9")
    doc.add_page_break()
    section_vision_objectives(doc, "10")
    section_scope(doc, "11")
    doc.add_page_break()

    # ── Part III — Stakeholders & Audience (12–13) ──────────────────────────
    section_stakeholder_register(doc, "12")
    doc.add_page_break()
    section_audience_arch(doc, "13")
    doc.add_page_break()

    # ── Part IV — Domain Model (14–15) ──────────────────────────────────────
    section_trigger_catalog(doc, "14")
    doc.add_page_break()
    section_channel_strategy(doc, "15")

    # ── Part V — Solution (16–18) ───────────────────────────────────────────
    section_template_management(doc, "16")
    doc.add_page_break()
    section_solution_arch(doc, "17")
    section_skill_catalog(doc, "18")
    doc.add_page_break()

    # ── Part VI — Compliance, Security & Quality (19–21) ────────────────────
    section_compliance(doc, "19")
    doc.add_page_break()
    section_security_privacy(doc, "20")
    doc.add_page_break()
    section_qa_testing(doc, "21")
    doc.add_page_break()

    # ── Part VII — Delivery (22–24) ─────────────────────────────────────────
    section_phases(doc, "22")
    section_deliverables(doc, "23")
    doc.add_page_break()
    section_acceptance(doc, "24")

    # ── Part VIII — Governance (25–29) ──────────────────────────────────────
    section_raci(doc, "25")
    section_comms_management(doc, "26")
    doc.add_page_break()
    section_change_management(doc, "27")
    doc.add_page_break()
    section_sla(doc, "28")
    doc.add_page_break()
    section_governance(doc, "29")

    # ── Part IX — Risk & Dependencies (30–32) ───────────────────────────────
    section_assumptions(doc, "30")
    doc.add_page_break()
    section_dependencies(doc, "31")
    section_risks(doc, "32")
    doc.add_page_break()

    # ── Part X — Resourcing & Readiness (33–37) ─────────────────────────────
    section_effort(doc, "33")
    section_resource_plan(doc, "34")
    doc.add_page_break()
    section_operational_readiness(doc, "35")
    doc.add_page_break()
    section_knowledge_transfer(doc, "36")
    doc.add_page_break()
    section_post_impl_support(doc, "37")
    doc.add_page_break()

    # ── Part XI — Legal & Reference (38–41) ─────────────────────────────────
    section_confidentiality_ip(doc, "38")
    doc.add_page_break()
    section_glossary(doc, "39")
    section_references(doc, "40")
    section_signoff(doc, "41")

    out = "/Users/aksatyam/TPE_DIRECTORY/TPE - Communication System/docs/TPE_Communication_System_SOW.docx"
    doc.save(out)
    print(f"Saved: {out}  (v{DOC_VERSION} — 41 sections, 11 parts)")

if __name__ == "__main__":
    main()
