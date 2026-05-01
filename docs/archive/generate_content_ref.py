"""
TPE Communication Content Reference — Generator
Companion document to TPE_Communication_System_SOW.docx (v1.5).

This document is the operational data bible for the communication system:
  - Complete 49-trigger matrix with statuses, channels, cadence, variables
  - Per-trigger detail tables with message intent and compliance class
  - Channel content specifications (format rules, limits, gates)
  - Variable token catalogue (35+ tokens with type, source, validation)
  - Sample template library (actual message copy, ready to author)
  - Compliance quick-reference per regime

Run: python3 generate_content_ref.py
Output: TPE_Communication_Content_Reference.docx
"""

from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────────────────────
# Brand palette (per ~/.claude/CLAUDE.md)
# ─────────────────────────────────────────────────────────────────────────────
NAVY        = RGBColor(0x1B, 0x3A, 0x5C)
BLUE        = RGBColor(0x2E, 0x75, 0xB6)
GREEN       = RGBColor(0x27, 0xAE, 0x60)
ORANGE      = RGBColor(0xE6, 0x7E, 0x22)
RED         = RGBColor(0xE7, 0x4C, 0x3C)
DARK        = RGBColor(0x2C, 0x3E, 0x50)
MEDIUM      = RGBColor(0x7F, 0x8C, 0x8D)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT       = "F0F4F8"
ALT_ROW     = "F8F9FA"
HEADER_BG   = "1B3A5C"
SUCCESS_BG  = "E8F5E9"
WARN_BG     = "FFF8E1"
DANGER_BG   = "FFEBEE"

DOC_TITLE   = "TPE Communication Content Reference"
DOC_VERSION = "1.0"
DOC_ID      = "TPE-COMMS-REF-2026-001"
TODAY       = date(2026, 4, 26).strftime("%d %b %Y")
AUTHOR      = "Ashish Kumar Satyam"
ORG_NAME    = "The Policy Exchange"
ORG_LEGAL   = "Fairvalue Insuretech Pvt. Ltd."
ORG_FULL    = f"{ORG_NAME} ({ORG_LEGAL})"

# Audience accent colors (consistent with SOW v1.5)
COLOR_PH    = "2E75B6"   # Blue
COLOR_INV   = "27AE60"   # Green
COLOR_INS   = "E67E22"   # Orange
COLOR_OPS   = "1B3A5C"   # Navy
COLOR_REG   = "E74C3C"   # Red

# ─────────────────────────────────────────────────────────────────────────────
# Helpers (XML-level cell-width persistence — cell.width alone does not stick)
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_width_xml(cell, width_inches):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(int(width_inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')

def set_table_grid_xml(table, widths_inches):
    tbl = table._tbl
    grid = tbl.find(qn('w:tblGrid'))
    if grid is None:
        grid = OxmlElement('w:tblGrid')
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths_inches:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(int(w * 1440)))
        grid.append(col)

def set_table_width_xml(table, total_inches):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(int(total_inches * 1440)))
    tblW.set(qn('w:type'), 'dxa')

def lock_table_widths(table, widths_inches):
    total = sum(widths_inches)
    set_table_grid_xml(table, widths_inches)
    set_table_width_xml(table, total)
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            if ci < len(widths_inches):
                set_cell_width_xml(cell, widths_inches[ci])

def set_cell_borders(cell, color="BFBFBF", size=4):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(size))
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def style_run(run, *, font="Arial", size=10, bold=False, italic=False, color=DARK):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rFonts.set(qn(attr), font)

def add_para(doc, text="", *, font="Arial", size=10, bold=False, italic=False,
             color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4, indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if text:
        r = p.add_run(text)
        style_run(r, font=font, size=size, bold=bold, italic=italic, color=color)
    return p

def add_bullet(doc, text, *, indent=0.5, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    style_run(r, size=size)
    return p

def add_section_heading(doc, num, title, level=1):
    if level == 1:
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        cell = tbl.rows[0].cells[0]
        set_cell_bg(cell, HEADER_BG)
        cell.width = Inches(7.0)
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
        cell.paragraphs[0].text = ""
        run = cell.paragraphs[0].add_run(f"  {num}. {title}")
        style_run(run, font="Arial", size=14, bold=True, color=WHITE)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
    elif level == 2:
        accent = "2E75B6"
        tbl = doc.add_table(rows=1, cols=2)
        tbl.autofit = False
        bar = tbl.rows[0].cells[0]
        bar.width = Inches(0.08)
        set_cell_bg(bar, accent)
        bar.paragraphs[0].text = ""
        heading = tbl.rows[0].cells[1]
        heading.width = Inches(6.5)
        heading.paragraphs[0].text = ""
        heading.paragraphs[0].paragraph_format.space_before = Pt(2)
        heading.paragraphs[0].paragraph_format.space_after = Pt(2)
        r = heading.paragraphs[0].add_run(f"  {num}  {title}")
        style_run(r, size=12, bold=True, color=NAVY)
        add_para(doc, "", space_after=2)
    else:
        add_para(doc, f"{num}  {title}", size=11, bold=True, italic=True,
                 color=BLUE, space_after=2)

def add_table(doc, headers, rows, *, col_widths_inches=None, header_bg=HEADER_BG,
              header_fg=WHITE, alt_rows=True, font_size=9, code_cols=None):
    code_cols = code_cols or set()
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = False
    if col_widths_inches:
        lock_table_widths(tbl, col_widths_inches)
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        set_cell_bg(c, header_bg)
        set_cell_borders(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.paragraphs[0].text = ""
        r = c.paragraphs[0].add_run(h)
        style_run(r, font="Arial", size=font_size, bold=True, color=header_fg)
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]
        bg = ALT_ROW if (alt_rows and ri % 2 == 0) else "FFFFFF"
        for ci, val in enumerate(row):
            c = tr.cells[ci]
            set_cell_bg(c, bg)
            set_cell_borders(c)
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            c.paragraphs[0].text = ""
            font = "Consolas" if ci in code_cols else "Arial"
            r = c.paragraphs[0].add_run(str(val))
            style_run(r, font=font, size=font_size, color=DARK)
            c.paragraphs[0].paragraph_format.space_after = Pt(0)
    return tbl

def add_color_band(doc, hex_color, height_pt=4):
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    c = tbl.rows[0].cells[0]
    set_cell_bg(c, hex_color)
    c.paragraphs[0].text = ""
    c.paragraphs[0].paragraph_format.space_before = Pt(0)
    c.paragraphs[0].paragraph_format.space_after = Pt(0)
    tbl.rows[0].height = Pt(height_pt)
    tbl.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

def add_code_block(doc, code, *, font_size=8.5, label=None):
    if label:
        add_para(doc, label, size=9, bold=True, color=NAVY, space_after=2)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    lock_table_widths(tbl, [6.50])
    c = tbl.rows[0].cells[0]
    set_cell_bg(c, "F4F4F4")
    set_cell_borders(c, color="CCCCCC")
    c.paragraphs[0].text = ""
    for i, line in enumerate(code.splitlines() or [""]):
        if i == 0:
            p = c.paragraphs[0]
        else:
            p = c.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line if line else " ")
        style_run(r, font="Consolas", size=font_size, color=DARK)
    add_para(doc, "", space_after=4)

def add_callout(doc, title, body, kind="blue"):
    color_map = {
        'blue':   (BLUE,   LIGHT),
        'green':  (GREEN,  SUCCESS_BG),
        'orange': (ORANGE, WARN_BG),
        'red':    (RED,    DANGER_BG),
        'navy':   (NAVY,   LIGHT),
    }
    fg, bg = color_map[kind]
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    lock_table_widths(tbl, [6.50])
    c = tbl.rows[0].cells[0]
    set_cell_bg(c, bg)
    set_cell_borders(c, color="BFBFBF")
    c.paragraphs[0].text = ""
    r = c.paragraphs[0].add_run(title)
    style_run(r, size=10, bold=True, color=fg)
    p = c.add_paragraph()
    r2 = p.add_run(body)
    style_run(r2, size=9.5, color=DARK)
    add_para(doc, "", space_after=2)

def add_trigger_header(doc, trigger_id, name, audience_color, audience_label,
                        priority="STANDARD", volume="—"):
    """Trigger card header bar: ID + name + audience label + priority + volume."""
    priority_bg = {"HIGH": "FFEBEE", "STANDARD": "F0F4F8", "LOW": "F0F4F8"}.get(priority, "F0F4F8")
    priority_fg = {"HIGH": "E74C3C", "STANDARD": "2E75B6", "LOW": "7F8C8D"}.get(priority, "2E75B6")
    tbl = doc.add_table(rows=1, cols=4)
    tbl.autofit = False
    lock_table_widths(tbl, [0.85, 3.20, 1.30, 1.15])
    # ID cell (audience-colored)
    id_cell = tbl.rows[0].cells[0]
    set_cell_bg(id_cell, audience_color)
    set_cell_borders(id_cell, color=audience_color)
    id_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    id_cell.paragraphs[0].text = ""
    id_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = id_cell.paragraphs[0].add_run(trigger_id)
    style_run(r, font="Consolas", size=12, bold=True, color=WHITE)
    # Name cell (audience-colored)
    name_cell = tbl.rows[0].cells[1]
    set_cell_bg(name_cell, audience_color)
    set_cell_borders(name_cell, color=audience_color)
    name_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    name_cell.paragraphs[0].text = ""
    r2 = name_cell.paragraphs[0].add_run(f"  {name}")
    style_run(r2, size=12, bold=True, color=WHITE)
    # Audience label
    aud_cell = tbl.rows[0].cells[2]
    set_cell_bg(aud_cell, "FAFBFC")
    set_cell_borders(aud_cell, color="DDDDDD")
    aud_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    aud_cell.paragraphs[0].text = ""
    aud_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = aud_cell.paragraphs[0].add_run(audience_label)
    style_run(r3, size=8.5, bold=True, color=RGBColor.from_string(audience_color))
    # Priority + volume cell
    prio_cell = tbl.rows[0].cells[3]
    set_cell_bg(prio_cell, priority_bg)
    set_cell_borders(prio_cell, color="DDDDDD")
    prio_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    prio_cell.paragraphs[0].text = ""
    prio_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = prio_cell.paragraphs[0].add_run(f"● {priority}")
    style_run(r4, size=8.5, bold=True, color=RGBColor.from_string(priority_fg))
    p2 = prio_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(0)
    r5 = p2.add_run(volume)
    style_run(r5, size=8, color=MEDIUM, italic=True)

def add_trigger_meta(doc, statuses, channels, cadence, variables, compliance):
    """Inline metadata row beneath the trigger header."""
    tbl = doc.add_table(rows=5, cols=2)
    tbl.autofit = False
    lock_table_widths(tbl, [1.30, 5.20])
    rows_data = [
        ("Statuses",    " → ".join(statuses)),
        ("Channels",    " · ".join(channels)),
        ("Cadence",     cadence),
        ("Variables",   ", ".join(f"{{{{{v}}}}}" for v in variables)),
        ("Compliance",  compliance),
    ]
    for ri, (label, val) in enumerate(rows_data):
        l_cell = tbl.rows[ri].cells[0]
        set_cell_bg(l_cell, "F0F4F8")
        set_cell_borders(l_cell, color="DDDDDD")
        l_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        l_cell.paragraphs[0].text = ""
        r = l_cell.paragraphs[0].add_run(label)
        style_run(r, size=9, bold=True, color=NAVY)
        l_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        v_cell = tbl.rows[ri].cells[1]
        set_cell_bg(v_cell, "FFFFFF")
        set_cell_borders(v_cell, color="DDDDDD")
        v_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        v_cell.paragraphs[0].text = ""
        font = "Consolas" if label == "Variables" else "Arial"
        r2 = v_cell.paragraphs[0].add_run(val)
        style_run(r2, font=font, size=9, color=DARK)
        v_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    add_para(doc, "", space_after=4)

def add_sample_message(doc, channel, status_or_timing, body):
    """Render a sample message in a code-block style with channel/timing label."""
    label = f"SAMPLE  ·  {channel.upper()}  ·  {status_or_timing}"
    add_para(doc, label, size=8.5, bold=True, color=BLUE, space_after=2)
    add_code_block(doc, body, font_size=8.5)

# ─────────────────────────────────────────────────────────────────────────────
# Page setup, header, footer
# ─────────────────────────────────────────────────────────────────────────────

def setup_document():
    doc = Document()
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.4)
        section.footer_distance = Inches(0.4)
        h = section.header.paragraphs[0]
        h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rh = h.add_run(f"{DOC_TITLE}  |  v{DOC_VERSION}")
        style_run(rh, size=8, color=MEDIUM, italic=True)
        f = section.footer.paragraphs[0]
        f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rf = f.add_run(f"{DOC_ID}  |  Internal / Confidential  |  Page ")
        style_run(rf, size=8, color=MEDIUM)
        run = f.add_run()
        fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText'); instr.text = "PAGE"
        fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_end)
        style_run(run, size=8, color=MEDIUM, bold=True)
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.font.color.rgb = DARK
    return doc

# ─────────────────────────────────────────────────────────────────────────────
# Cover page
# ─────────────────────────────────────────────────────────────────────────────

def add_cover_page(doc):
    add_color_band(doc, HEADER_BG, height_pt=10)
    add_para(doc, "", space_after=18)
    add_para(doc, ORG_FULL.upper(), size=10, color=MEDIUM, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "CONTENT REFERENCE", size=22, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_para(doc, "TPE Communication System", size=28, bold=True, color=DARK,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(doc, "The operational data bible — every trigger, every channel, "
                  "every sample. Companion to the SOW.",
             size=11, italic=True, color=MEDIUM,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    # Snapshot tiles
    tbl = doc.add_table(rows=1, cols=4)
    tbl.autofit = False
    lock_table_widths(tbl, [1.625, 1.625, 1.625, 1.625])
    tiles = [
        ("49",    "Triggers documented",  COLOR_PH),
        ("38",    "Sample message drafts", COLOR_INV),
        ("37",    "Variables catalogued",  COLOR_INS),
        ("6",     "Compliance regimes",    COLOR_REG),
    ]
    for i, (n, lbl, col) in enumerate(tiles):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, "FAFBFC")
        set_cell_borders(c, color=col, size=18)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.paragraphs[0].text = ""
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.paragraphs[0].add_run(n)
        style_run(r, size=22, bold=True, color=RGBColor.from_string(col))
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(lbl)
        style_run(r2, size=8.5, color=DARK)
    add_para(doc, "", space_after=24)

    # Metadata block
    add_para(doc, "Document metadata", size=10, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    meta_tbl = doc.add_table(rows=6, cols=2)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_tbl.autofit = False
    lock_table_widths(meta_tbl, [2.10, 3.10])
    meta = [
        ("Document ID",    DOC_ID),
        ("Version",        DOC_VERSION),
        ("Date",           TODAY),
        ("Prepared by",    AUTHOR),
        ("Companion to",   "TPE-COMMS-SOW-2026-001 (SOW v1.5)"),
        ("Classification", "Internal / Confidential"),
    ]
    for ri, (k, v) in enumerate(meta):
        bg = ALT_ROW if ri % 2 == 0 else "FFFFFF"
        c0 = meta_tbl.rows[ri].cells[0]
        c1 = meta_tbl.rows[ri].cells[1]
        for c in (c0, c1):
            set_cell_bg(c, bg)
            set_cell_borders(c, color="DDDDDD")
            c.paragraphs[0].text = ""
        r0 = c0.paragraphs[0].add_run(k)
        style_run(r0, size=9.5, bold=True, color=NAVY)
        r1 = c1.paragraphs[0].add_run(v)
        style_run(r1, size=9.5, color=DARK)
    add_para(doc, "", space_after=18)
    add_color_band(doc, COLOR_PH, height_pt=4)
    add_para(doc, "Read this document with the SOW open in a second window. "
                  "The SOW defines what we are building; this document defines "
                  "what is in it.", size=9, italic=True, color=MEDIUM,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# Section 1 — How to use this document
# ─────────────────────────────────────────────────────────────────────────────

def section_how_to_use(doc):
    add_section_heading(doc, "1", "How to Use This Document")
    add_para(doc,
        "This Content Reference is the deliverable that lives alongside the "
        "Statement of Work. The SOW commits the team to building a versioned "
        "content library; this document is the substance of that library, "
        "expressed in the same form a template author will see when they sit "
        "down to draft.", space_after=8)

    add_para(doc, "Document map", size=11, bold=True, color=NAVY, space_after=2)
    add_table(doc, ["Section", "Use it to…", "Audience"], [
        ["2  Quick Matrix",
         "See every trigger at a glance — ID, audience, channel mix, volume, compliance class.",
         "Anyone needing a 30-second overview"],
        ["3  Channel Specs",
         "Look up the format rules for Email / SMS / WhatsApp / RCS / Slack / Email-attachment.",
         "Template authors, channel renderers"],
        ["4  Variable Catalogue",
         "Find the canonical name, type, source and validation rule for every personalisation token.",
         "Template authors, data integrators"],
        ["5  Policyholder Triggers",
         "Per-trigger detail card for all 18 PH triggers with status × channel timing matrix.",
         "PH-comms author, CX Head"],
        ["6  Investor Triggers",
         "Per-trigger detail card for all 16 INV triggers; INV-08 Premium Due gets full multi-touch detail.",
         "INV-comms author, CX Head"],
        ["7  Insurer / Internal / Regulatory",
         "Detail cards for the 15 B2B / internal / statutory triggers.",
         "Operations, Compliance"],
        ["8  Sample Template Library",
         "Drafted message copy for the 15 most-used templates, ready to lift into YAML.",
         "Template authors"],
        ["9  Compliance Quick-Reference",
         "Per-regime checklist (DLT, IRDAI, WhatsApp BSP, DPDPA, FATCA/KYC, TDS) of what every template must clear.",
         "Compliance reviewer"],
    ], col_widths_inches=[1.55, 3.30, 1.65], font_size=9)

    add_para(doc, "", space_after=6)
    add_callout(doc, "How this complements the SOW",
        "The SOW is read top-to-bottom by leadership for sign-off. This Content "
        "Reference is consulted laterally by authors and reviewers during "
        "delivery — you jump to the trigger or template you are working on. "
        "Both documents share the same trigger IDs, audience colours and "
        "compliance vocabulary so a reader can move between them without "
        "translation.", kind="navy")

    add_para(doc, "Conventions", size=11, bold=True, color=NAVY, space_after=2)
    for b in [
        "Trigger IDs are immutable (PH-03 means KYC forever, even if the wording changes).",
        "Audience colours: Policyholder = Blue, Investor = Green, Insurer = Orange, Internal = Navy, Regulator = Red.",
        "Variable names use snake_case and are wrapped {{like_this}} when shown inline.",
        "Sample messages are illustrative drafts, not approved copy. Approval gate: Section 9.",
        "Volume estimates are steady-state monthly figures based on the v1 scope (49 triggers, 5 audiences, ~120 templates).",
    ]:
        add_bullet(doc, b)
    doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# Section 2 — Complete trigger matrix (all 49)
# ─────────────────────────────────────────────────────────────────────────────

# Master data — drives both the matrix and the per-trigger cards
TRIGGERS = [
    # ── POLICYHOLDER ────────────────────────────────────────────────────────
    {"id":"PH-01","aud":"PH","name":"Lead Captured",
     "statuses":["New","Contacted","Qualified","Disqualified"],
     "channels":["Email","WhatsApp"],
     "cadence":"T+0 welcome → T+1 follow-up → T+3 value reminder → T+7 last nudge → close",
     "variables":["ph_name","source_channel","surrender_value_estimate","cta_link"],
     "compliance":"DPDPA opt-out; IRDAI solicitation disclaimer in promo content",
     "volume":"~500/mo","priority":"STANDARD"},
    {"id":"PH-02","aud":"PH","name":"Registration",
     "statuses":["Started","OTP Sent","Verified","Abandoned"],
     "channels":["SMS","WhatsApp","Email"],
     "cadence":"OTP: T+0; resend at T+10min on request. Abandoned: T+1h, T+24h, T+72h, then close.",
     "variables":["ph_name","otp","otp_expiry_min","reg_link"],
     "compliance":"DLT-Authentication for OTP; DLT-Service for nudges; opt-out on promo only",
     "volume":"~500/mo","priority":"HIGH"},
    {"id":"PH-03","aud":"PH","name":"KYC",
     "statuses":["Initiated","In-progress","Done","Failed","Pending"],
     "channels":["WhatsApp","Email","SMS"],
     "cadence":"Pending: weekly × 4, monthly × 4, then close. Done: immediate confirmation.",
     "variables":["ph_name","kyc_link","kyc_status","attempt_no","deadline_date","support_number"],
     "compliance":"DLT-Service; DPDPA opt-out; tone must remain non-coercive across all 8 attempts",
     "volume":"~2,000/mo","priority":"HIGH"},
    {"id":"PH-04","aud":"PH","name":"Policy Doc Upload",
     "statuses":["Uploaded","Reviewing","Verified","Rejected"],
     "channels":["Email","WhatsApp"],
     "cadence":"Per-state immediate; Reviewing: status update at T+24h if not resolved",
     "variables":["ph_name","policy_no","reject_reason","resubmit_link"],
     "compliance":"DLT-Service",
     "volume":"~400/mo","priority":"STANDARD"},
    {"id":"PH-05","aud":"PH","name":"Policy Verification",
     "statuses":["In-progress","Verified","Discrepancy","Insurer non-response"],
     "channels":["WhatsApp","Email"],
     "cadence":"Status update every 48h; escalation to Insurer non-response after 7d",
     "variables":["ph_name","policy_no","insurer_name","verification_status","next_step"],
     "compliance":"DLT-Service",
     "volume":"~300/mo","priority":"STANDARD"},
    {"id":"PH-06","aud":"PH","name":"Surrender Value Quote",
     "statuses":["Generated","Sent","Viewed","Expired (7d)","Counter-offer"],
     "channels":["Email","WhatsApp"],
     "cadence":"T+0 sent; T+2 if not viewed; T+5 urgency; T-1 (24h pre-expiry) final reminder",
     "variables":["ph_name","policy_no","insurer_name","surrender_value","quote_validity_date","accept_link","counter_offer_link"],
     "compliance":"IRDAI: must include 'estimate, subject to insurer confirmation' caveat",
     "volume":"~300/mo","priority":"HIGH"},
    {"id":"PH-07","aud":"PH","name":"Offer Decision",
     "statuses":["Accepted","Rejected","Negotiating","Silent"],
     "channels":["Email","WhatsApp"],
     "cadence":"Silent: T+3, T+7 nudge then close",
     "variables":["ph_name","policy_no","offer_amount","decision_link"],
     "compliance":"DLT-Service",
     "volume":"~250/mo","priority":"STANDARD"},
    {"id":"PH-08","aud":"PH","name":"Assignment Paperwork",
     "statuses":["Form Generated","e-Sign Initiated","e-Sign Done","Physical Required","NOC Pending"],
     "channels":["Email","WhatsApp","SMS"],
     "cadence":"e-Sign: T+0 link; T+1 reminder; T+3 last reminder before re-issue. NOC: per insurer SLA.",
     "variables":["ph_name","esign_link","esign_expiry","policy_no","insurer_name","noc_status"],
     "compliance":"IT Act 2000 e-Sign disclosures; IRDAI assignment intimation",
     "volume":"~250/mo","priority":"HIGH"},
    {"id":"PH-09","aud":"PH","name":"Investor Matched",
     "statuses":["Found","Confirmed"],
     "channels":["WhatsApp","Email"],
     "cadence":"Immediate on match; confirmation within 24h",
     "variables":["ph_name","policy_no","investor_id_masked","disbursement_eta"],
     "compliance":"DPDPA: investor identity is PII; share masked ID only",
     "volume":"~250/mo","priority":"STANDARD"},
    {"id":"PH-10","aud":"PH","name":"Disbursement",
     "statuses":["Initiated","In-process","Success","Failed"],
     "channels":["SMS","WhatsApp","Email"],
     "cadence":"State-driven: each transition triggers immediate notification",
     "variables":["ph_name","disbursement_amount","utr_number","bank_name","bank_last4","expected_credit_date","failure_reason"],
     "compliance":"RBI guidelines on payment confirmations; DLT-Service",
     "volume":"~250/mo","priority":"HIGH"},
    {"id":"PH-11","aud":"PH","name":"Post-Assignment Welcome",
     "statuses":["Recorded","Welcome"],
     "channels":["Email","WhatsApp"],
     "cadence":"T+1d welcome; quarterly newsletter thereafter",
     "variables":["ph_name","policy_no","tpe_dashboard_link"],
     "compliance":"DPDPA opt-out for newsletter",
     "volume":"~250/mo + ongoing","priority":"STANDARD"},
    {"id":"PH-12","aud":"PH","name":"Loan Application",
     "statuses":["Started","Docs Pending","Submitted","Under Review"],
     "channels":["WhatsApp","Email","SMS"],
     "cadence":"Docs Pending: T+1, T+3, T+7 nudge",
     "variables":["ph_name","loan_application_id","required_docs_list","upload_link"],
     "compliance":"DLT-Service",
     "volume":"~150/mo","priority":"STANDARD"},
    {"id":"PH-13","aud":"PH","name":"Loan Approval",
     "statuses":["Approved","Conditional","Rejected"],
     "channels":["Email","WhatsApp","SMS"],
     "cadence":"Immediate on decision",
     "variables":["ph_name","loan_amount","interest_rate","tenure_months","emi_amount","accept_by_date"],
     "compliance":"RBI lending fair-practice disclosures; IRDAI for loan-against-policy",
     "volume":"~120/mo","priority":"HIGH"},
    {"id":"PH-14","aud":"PH","name":"Loan Disbursement",
     "statuses":["Initiated","Disbursed","Failed"],
     "channels":["SMS","WhatsApp"],
     "cadence":"Immediate",
     "variables":["ph_name","loan_amount","utr_number","bank_last4","first_emi_date"],
     "compliance":"RBI; DLT-Service",
     "volume":"~120/mo","priority":"HIGH"},
    {"id":"PH-15","aud":"PH","name":"Loan EMI",
     "statuses":["T-7","T-3","T-day","Day+1","Day+7","Day+15","Day+30"],
     "channels":["SMS","WhatsApp","Email"],
     "cadence":"7-stage cadence per EMI cycle; voice escalation at Day+15",
     "variables":["ph_name","emi_amount","due_date","loan_number","payment_link","overdue_amount","late_charge"],
     "compliance":"RBI; DLT-Service for reminders, DLT-Promo not applicable",
     "volume":"~840/mo (7×120)","priority":"HIGH"},
    {"id":"PH-16","aud":"PH","name":"Loan Closure",
     "statuses":["Pre-close Quote","Foreclosed","Paid"],
     "channels":["Email","WhatsApp"],
     "cadence":"Per event",
     "variables":["ph_name","loan_number","foreclosure_amount","closure_certificate_link"],
     "compliance":"RBI; certificate retention 8 years",
     "volume":"~30/mo","priority":"STANDARD"},
    {"id":"PH-17","aud":"PH","name":"Re-engagement",
     "statuses":["30d Silent","90d Silent","6mo Silent","1y Silent"],
     "channels":["Email","WhatsApp"],
     "cadence":"Decay drip: progressively softer; suppress after 1y if no response",
     "variables":["ph_name","last_interaction_date","new_offer_link"],
     "compliance":"DPDPA opt-out mandatory; suppress on opt-out",
     "volume":"~600/mo","priority":"LOW"},
    {"id":"PH-18","aud":"PH","name":"Referral Program",
     "statuses":["Earned","Credited","Expiring"],
     "channels":["WhatsApp","Email"],
     "cadence":"Per event; Expiring at T-30d before reward expiry",
     "variables":["ph_name","reward_amount","reward_expiry_date","referee_name"],
     "compliance":"IRDAI: referral programs require explicit T&C disclosure",
     "volume":"~100/mo","priority":"LOW"},

    # ── INVESTOR ────────────────────────────────────────────────────────────
    {"id":"INV-01","aud":"INV","name":"Lead Captured",
     "statuses":["New","Webinar Reg","Webinar Attended","Brochure DL"],
     "channels":["Email","WhatsApp"],
     "cadence":"T+0, T+1, T+3, T+7, T+14 drip nurture",
     "variables":["inv_name","source_event","invest_calculator_link"],
     "compliance":"IRDAI 'subject of solicitation'; SEBI-style 'returns subject to market risks' style note",
     "volume":"~800/mo","priority":"STANDARD"},
    {"id":"INV-02","aud":"INV","name":"Registration",
     "statuses":["Started","Verified","KYC Pending","Active"],
     "channels":["Email","SMS","WhatsApp"],
     "cadence":"OTP T+0; KYC nudges per PH-03 cadence",
     "variables":["inv_name","otp","otp_expiry_min","reg_link","kyc_link"],
     "compliance":"DLT-Authentication; FATCA self-declaration prompt",
     "volume":"~400/mo","priority":"HIGH"},
    {"id":"INV-03","aud":"INV","name":"KYC + Risk Profile",
     "statuses":["CKYC","FATCA","Risk Profile","Bank Verified"],
     "channels":["Email","WhatsApp"],
     "cadence":"Sequential per step; reminders at T+24h, T+72h",
     "variables":["inv_name","kyc_link","fatca_link","risk_questionnaire_link","bank_verify_status"],
     "compliance":"FATCA mandatory; CKYC recommended; risk profile completion before any opportunity push",
     "volume":"~400/mo","priority":"HIGH"},
    {"id":"INV-04","aud":"INV","name":"Investment Opportunity Push",
     "statuses":["New Match","Weekly Curated","Premium Yield Alert"],
     "channels":["Email","WhatsApp","RCS"],
     "cadence":"Weekly digest (Mon 10am IST) + on-demand alerts for high-yield matches",
     "variables":["inv_name","policy_count","min_yield_pct","max_yield_pct","top_policy_no","top_policy_value","explore_link"],
     "compliance":"IRDAI promo disclaimers; no specific-return promises; risk disclosure",
     "volume":"~3,200/mo","priority":"HIGH"},
    {"id":"INV-05","aud":"INV","name":"Investment Soft-Commit",
     "statuses":["Interest","Doc Review","Clarifications"],
     "channels":["Email","WhatsApp"],
     "cadence":"Immediate ack; T+24h status update",
     "variables":["inv_name","policy_no","commit_amount","reservation_expiry"],
     "compliance":"Reservation T&C disclosure; cooling-off period reference",
     "volume":"~600/mo","priority":"STANDARD"},
    {"id":"INV-06","aud":"INV","name":"Funds Collection",
     "statuses":["Pay Link Sent","Received","Failed"],
     "channels":["SMS","Email","WhatsApp"],
     "cadence":"T+0 link; T+24h, T+48h reminders if pending",
     "variables":["inv_name","policy_no","commit_amount","payment_link","payment_methods","payment_expiry"],
     "compliance":"RBI payment confirmation; PCI-DSS reference for stored payment data",
     "volume":"~500/mo","priority":"HIGH"},
    {"id":"INV-07","aud":"INV","name":"Investment Confirmed",
     "statuses":["Assignment in Progress","Confirmed","Certificate Issued"],
     "channels":["Email","WhatsApp"],
     "cadence":"Immediate on each state",
     "variables":["inv_name","policy_no","investment_amount","certificate_link","first_premium_due"],
     "compliance":"IRDAI assignment intimation; certificate retention",
     "volume":"~400/mo","priority":"HIGH"},
    {"id":"INV-08","aud":"INV","name":"Premium Due (10–25 yr stream)",
     "statuses":["T-30","T-15","T-7","T-3","T-day","Day+1","Day+7","Day+15"],
     "channels":["Email","SMS","WhatsApp"],
     "cadence":"8-touchpoint cadence per premium cycle; HIGHEST volume trigger in the system",
     "variables":["inv_name","policy_no","insurer_name","premium_amount","due_date","payment_link","grace_period_end","late_fee"],
     "compliance":"DLT-Service; auto-pay reference; lapse-warning escalation post Day+15",
     "volume":"~15,000/mo","priority":"HIGH"},
    {"id":"INV-09","aud":"INV","name":"Premium Payment Confirmation",
     "statuses":["Paid to Insurer","Receipt Available"],
     "channels":["Email","WhatsApp"],
     "cadence":"Immediate on each state",
     "variables":["inv_name","policy_no","premium_amount","insurer_name","paid_date","receipt_link"],
     "compliance":"IRDAI receipt format; retention 8 years",
     "volume":"~2,000/mo","priority":"HIGH"},
    {"id":"INV-10","aud":"INV","name":"Premium Default → Lapse Risk",
     "statuses":["Grace Started","Lapse Warning","Lapsed"],
     "channels":["Email","SMS","WhatsApp","Voice (P5)"],
     "cadence":"Aggressive escalation: WA + SMS + Email same day; voice call at Day+22",
     "variables":["inv_name","policy_no","insurer_name","grace_period_end","days_overdue","potential_loss","payment_link","support_number"],
     "compliance":"IRDAI lapse rules; DLT-Service; voice consent for outbound calls",
     "volume":"~300/mo","priority":"HIGH"},
    {"id":"INV-11","aud":"INV","name":"Maturity Approaching",
     "statuses":["T-12mo","T-6mo","T-3mo","T-1mo","T-7d"],
     "channels":["Email","WhatsApp"],
     "cadence":"Per-stage with reinvestment opportunity at T-3mo",
     "variables":["inv_name","policy_no","insurer_name","maturity_date","maturity_amount","reinvest_link"],
     "compliance":"IRDAI maturity intimation; tax implication note",
     "volume":"~600/mo","priority":"STANDARD"},
    {"id":"INV-12","aud":"INV","name":"Maturity Received",
     "statuses":["Insurer Paid TPE","Settled to Investor"],
     "channels":["SMS","Email","WhatsApp"],
     "cadence":"Immediate on each state",
     "variables":["inv_name","policy_no","maturity_amount","settlement_date","bank_last4","utr_number","tax_certificate_link"],
     "compliance":"RBI; TDS certificate generation; IRDAI maturity intimation",
     "volume":"~150/mo","priority":"HIGH"},
    {"id":"INV-13","aud":"INV","name":"Portfolio Statement",
     "statuses":["Monthly","Quarterly","Annual"],
     "channels":["Email"],
     "cadence":"Calendar: Monthly 1st, Quarterly Q1-Q4 first day, Annual on FY close",
     "variables":["inv_name","period","total_invested","portfolio_value","unrealised_gain","statement_link"],
     "compliance":"IRDAI portfolio disclosure; SEBI-style risk caveat if return projection shown",
     "volume":"~3,000/mo","priority":"STANDARD"},
    {"id":"INV-14","aud":"INV","name":"Tax Documents",
     "statuses":["TDS Q1","TDS Q2","TDS Q3","TDS Q4","80C/80D","Annual Statement"],
     "channels":["Email"],
     "cadence":"Statutory deadlines: TDS within 15d of quarter close; 80C by 30 Apr",
     "variables":["inv_name","fy","tds_amount","tds_certificate_link","pan_last4","section"],
     "compliance":"Income Tax Act; TDS Form 16A format; PAN validation mandatory",
     "volume":"~3,000/mo (peak Q4+Apr)","priority":"HIGH"},
    {"id":"INV-15","aud":"INV","name":"Re-investment Opportunity",
     "statuses":["Maturity-to-Reinvest","New High-Yield Match"],
     "channels":["Email","WhatsApp","RCS"],
     "cadence":"Triggered by INV-11 T-3mo or by INV-04 high-yield event",
     "variables":["inv_name","matured_amount","matched_policy_no","matched_yield_pct","commit_link"],
     "compliance":"IRDAI promo; risk disclosure; no rollover-by-default",
     "volume":"~400/mo","priority":"STANDARD"},
    {"id":"INV-16","aud":"INV","name":"Account Notices",
     "statuses":["Monthly Newsletter","T&C Change"],
     "channels":["Email"],
     "cadence":"Monthly 1st; T&C change with 30d notice as required by DPDPA",
     "variables":["inv_name","newsletter_link","tnc_change_summary","tnc_effective_date"],
     "compliance":"DPDPA: 30d notice for material T&C change; IRDAI for product T&C",
     "volume":"~3,000/mo","priority":"STANDARD"},

    # ── INSURANCE PARTNER (B2B) ────────────────────────────────────────────
    {"id":"INS-01","aud":"INS","name":"Assignment Intimation",
     "statuses":["Submitted","Acknowledged","Approved","Rejected"],
     "channels":["Email + Portal"],
     "cadence":"Per case; follow-up at T+3d if no acknowledgement",
     "variables":["insurer_contact","policy_no","ph_name","investor_id","assignment_form_link","noc_request"],
     "compliance":"IRDAI assignment regulations; insurer-specific SLA",
     "volume":"~250/mo","priority":"HIGH"},
    {"id":"INS-02","aud":"INS","name":"Premium Remittance + Allocation MIS",
     "statuses":["Paid","MIS Attached","Confirmed","Discrepancy"],
     "channels":["Email + Excel attachment"],
     "cadence":"Monthly batch — 11 emails (one per insurer) on payment cycle",
     "variables":["insurer_name","insurer_contact","report_period","total_premium","policy_count","mis_attachment_name","utr_number"],
     "compliance":"IRDAI remittance reporting; insurer reconciliation SLA; reuses existing MIS Pipeline",
     "volume":"11/mo (one per insurer)","priority":"HIGH"},
    {"id":"INS-03","aud":"INS","name":"MIS Reconciliation",
     "statuses":["Sent","Confirmed","Variance"],
     "channels":["Email + Excel attachment"],
     "cadence":"Monthly; variance escalation within 7d",
     "variables":["insurer_name","report_period","variance_amount","variance_policies","resolution_owner"],
     "compliance":"Bilateral SLA; audit log",
     "volume":"11/mo","priority":"STANDARD"},
    {"id":"INS-04","aud":"INS","name":"Query Resolution",
     "statuses":["Raised","Response","Closed"],
     "channels":["Email"],
     "cadence":"Per case; SLA 48h ack, 7d resolve",
     "variables":["insurer_contact","ticket_id","query_subject","query_age_days"],
     "compliance":"Bilateral SLA",
     "volume":"~80/mo","priority":"STANDARD"},
    {"id":"INS-05","aud":"INS","name":"Maturity / Claim Coordination",
     "statuses":["Initiated","Docs Submitted","Approved","Settled"],
     "channels":["Email"],
     "cadence":"Per case; follow-up every 7d",
     "variables":["insurer_contact","policy_no","claim_type","claim_amount","docs_pending"],
     "compliance":"IRDAI claim regulations; investor intimation parallel",
     "volume":"~150/mo","priority":"HIGH"},
    {"id":"INS-06","aud":"INS","name":"Regulatory / KYC Refresh",
     "statuses":["IRDAI Shared","KYC Refresh"],
     "channels":["Email"],
     "cadence":"Per IRDAI / insurer cadence",
     "variables":["insurer_contact","refresh_period","compliance_doc_link"],
     "compliance":"IRDAI; KYC refresh per RBI cycle",
     "volume":"~30/mo","priority":"STANDARD"},

    # ── INTERNAL OPS / SALES / AGENTS ───────────────────────────────────────
    {"id":"OPS-01","aud":"OPS","name":"New Lead Allocated",
     "statuses":["Created","Assigned"],
     "channels":["Slack/Teams","Email"],
     "cadence":"Real-time on assignment",
     "variables":["agent_handle","lead_id","ph_name","source_channel","priority","crm_link"],
     "compliance":"Internal information classification only",
     "volume":"~1,300/mo (PH+INV)","priority":"STANDARD"},
    {"id":"OPS-02","aud":"OPS","name":"Daily Standup Briefing",
     "statuses":["Pipeline Snapshot","Stuck Items","KPI Roll-up"],
     "channels":["Slack","Email"],
     "cadence":"Daily 9:00 IST",
     "variables":["date","team_handle","new_leads","stuck_count","kpi_table"],
     "compliance":"Internal",
     "volume":"~22/mo (workdays)","priority":"STANDARD"},
    {"id":"OPS-03","aud":"OPS","name":"High-Value Escalation",
     "statuses":["Deal ≥ ₹50L Stuck","VIP Query"],
     "channels":["Slack","SMS"],
     "cadence":"Real-time",
     "variables":["agent_handle","lead_id","deal_value","stuck_reason","escalation_owner"],
     "compliance":"Internal; SMS only to designated escalation contacts",
     "volume":"~30/mo","priority":"HIGH"},
    {"id":"OPS-04","aud":"OPS","name":"Premium Default Alert",
     "statuses":["Investor Non-pay","Lapse Risk"],
     "channels":["Slack","Email"],
     "cadence":"Real-time on detection; daily digest of all open defaults",
     "variables":["inv_id","policy_no","days_overdue","collection_owner","action_needed"],
     "compliance":"Internal; PII handling per DPDPA",
     "volume":"~300/mo","priority":"HIGH"},
    {"id":"OPS-05","aud":"OPS","name":"KYC Stuck",
     "statuses":[">7 days in KYC"],
     "channels":["Slack"],
     "cadence":"Daily digest of all >7d open KYC cases",
     "variables":["date","stuck_kyc_count","oldest_age_days","owner_handle"],
     "compliance":"Internal",
     "volume":"~22/mo (digest)","priority":"STANDARD"},
    {"id":"OPS-06","aud":"OPS","name":"Compliance Deadline",
     "statuses":["IRDAI Filing","GST","TDS"],
     "channels":["Email","Calendar"],
     "cadence":"Pre-deadline T-30, T-7, T-1 reminders",
     "variables":["deadline_label","due_date","owner","filing_link"],
     "compliance":"Internal; sourced from compliance calendar",
     "volume":"~12/mo","priority":"HIGH"},

    # ── REGULATORY / STATUTORY ──────────────────────────────────────────────
    {"id":"REG-01","aud":"REG","name":"IRDAI Filing — quarterly + annual",
     "statuses":["Quarterly","Annual"],
     "channels":["Portal upload + signed email"],
     "cadence":"Per regulator schedule (quarterly within 30d of close; annual within 90d of FY close)",
     "variables":["filing_period","filing_type","filing_doc_link","signing_authority"],
     "compliance":"IRDAI; signed by Compliance Officer",
     "volume":"5/yr","priority":"HIGH"},
    {"id":"REG-02","aud":"REG","name":"TDS Certificates to Investors",
     "statuses":["Q1","Q2","Q3","Q4"],
     "channels":["Email"],
     "cadence":"Quarterly within 15d of quarter close",
     "variables":["inv_name","fy_quarter","tds_amount","section","tds_certificate_link"],
     "compliance":"Income Tax Act; Form 16A format",
     "volume":"~3,000/quarter","priority":"HIGH"},
    {"id":"REG-03","aud":"REG","name":"GST Invoices",
     "statuses":["Issued"],
     "channels":["Email"],
     "cadence":"Per transaction; B2B and B2C",
     "variables":["counterparty","invoice_no","gst_amount","hsn_code","invoice_link"],
     "compliance":"GST Act; invoice retention 8 years",
     "volume":"~250/mo","priority":"STANDARD"},
]

AUDIENCE_META = {
    "PH":  ("POLICYHOLDER",   COLOR_PH),
    "INV": ("INVESTOR",       COLOR_INV),
    "INS": ("INSURER (B2B)",  COLOR_INS),
    "OPS": ("INTERNAL OPS",   COLOR_OPS),
    "REG": ("REGULATOR",      COLOR_REG),
}

def section_quick_matrix(doc):
    add_section_heading(doc, "2", "Complete Trigger Matrix — All 49")
    add_para(doc,
        "Single source-of-truth view: every trigger, audience, channel mix, "
        "cadence headline, volume estimate and priority. The document expands "
        "each row into a detail card in Sections 5–7; this matrix is for "
        "lateral lookup. Sort by volume for capacity planning, by priority "
        "for review queue ordering, by audience for skill ownership.",
        space_after=8)
    rows = []
    for t in TRIGGERS:
        aud_label, _ = AUDIENCE_META[t["aud"]]
        rows.append([
            t["id"],
            aud_label.split()[0],
            t["name"],
            " · ".join(t["channels"][:3]),  # truncate to 3 channels
            t["cadence"][:60] + ("…" if len(t["cadence"]) > 60 else ""),
            t["volume"],
            t["priority"],
        ])
    add_table(doc,
        ["ID", "Aud", "Trigger", "Channels", "Cadence (headline)", "Volume", "Pri"],
        rows,
        col_widths_inches=[0.55, 0.50, 1.55, 1.30, 1.85, 0.55, 0.20],
        font_size=7.5, code_cols={0})
    add_para(doc, "", space_after=4)
    add_callout(doc, "Reading the Volume column",
        "Volume is steady-state monthly. The headline figure aggregates all "
        "messages a trigger sends across its statuses. INV-08 Premium Due is "
        "the highest-volume trigger in the system — eight touchpoints per "
        "premium cycle × ~1,800 active assigned policies = ~15,000 monthly "
        "messages. Capacity planning, channel-provider pricing and template-"
        "review SLAs all flow from this column.",
        kind="navy")
    doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# Section 3 — Channel content specifications
# ─────────────────────────────────────────────────────────────────────────────

def section_channel_specs(doc):
    add_section_heading(doc, "3", "Channel Content Specifications")
    add_para(doc,
        "Each channel imposes hard format constraints. A template that violates "
        "them either fails to render, fails to deliver, or fails compliance. "
        "These specs are the input to the channel renderer (Pillar 3 in the "
        "SOW architecture) and the input to the lint engine.",
        space_after=8)

    add_section_heading(doc, "3.1", "Email", level=2)
    add_table(doc, ["Aspect", "Specification"], [
        ["Subject line",       "≤ 78 chars; prefix with [TPE] for transactional, no prefix for marketing"],
        ["Pre-header",         "≤ 100 chars; first line of body acts as preview text"],
        ["Body format",        "MJML or responsive HTML; plain-text alternate mandatory"],
        ["Image-to-text ratio","Minimum 60% text by area to avoid spam filters"],
        ["Attachment limit",   "10 MB total; PDF preferred; XLSX for INS-02 MIS only"],
        ["Sender",             "noreply@thepolicyexchange.com (transactional) / hello@thepolicyexchange.com (marketing)"],
        ["Reply-to",           "support@thepolicyexchange.com (transactional) / care@thepolicyexchange.com (marketing)"],
        ["Required footer",    "Physical address, IRDAI Reg. No., Grievance link, Unsubscribe link (DPDPA-compliant)"],
        ["Tracking",           "UTM params on all CTAs; pixel for opens (suppress for INS B2B)"],
        ["IRDAI disclaimer",   "Long form for promotional; short form for service messages"],
    ], col_widths_inches=[1.50, 5.00], font_size=9)

    add_section_heading(doc, "3.2", "SMS (DLT-regulated)", level=2)
    add_table(doc, ["Aspect", "Specification"], [
        ["Char limit",         "160 chars for single-segment; 153/segment for multi-part (max 2 segments preferred)"],
        ["Encoding",           "GSM-7 only (no emoji, no curly quotes, no — em-dash); Unicode segments cost 70/segment"],
        ["DLT Header",         "6-char alpha (e.g., TPECRE for Customer Care, TPEINV for Investments)"],
        ["DLT Template ID",    "12-digit numeric registered with operator before send"],
        ["PE ID",              "Principal Entity ID, registered once per operator"],
        ["Category",           "Authentication (OTP) / Service (transactional) / Promotional — must match content"],
        ["Time window",        "Promotional: 9:00–21:00 IST; Service & Auth: 24×7"],
        ["Opt-out",            "Promo: 'STOP to 1909' or short URL; Service & Auth: not required"],
        ["Variable count",     "DLT registration locks variable count and order — cannot vary per send"],
        ["Sender for INV",     "TPEINV (Investments) — not TPECRE (Customer Care)"],
    ], col_widths_inches=[1.50, 5.00], font_size=9)

    add_section_heading(doc, "3.3", "WhatsApp (Meta BSP)", level=2)
    add_table(doc, ["Aspect", "Specification"], [
        ["Categories",         "Marketing / Utility / Authentication — chosen at template approval, locks behaviour"],
        ["Body limit",         "1024 chars; header (text/image/video) + footer + buttons all separate"],
        ["Variables",          "Numbered placeholders {{1}}, {{2}}; positional, no name binding at render"],
        ["Header media",       "Image ≤ 5 MB, Video ≤ 16 MB, Document ≤ 100 MB; 4:3 or 16:9"],
        ["Buttons",            "Up to 3 quick-reply OR up to 2 CTA (URL/phone) — not both in same template"],
        ["24h session window", "Free-text replies allowed within 24h of customer message; outside window = template only"],
        ["Approval timeline",  "Meta typically 24–72h; rejection reasons returned via BSP API"],
        ["Sender display",     "WABA-verified business name; green tick for verified accounts"],
        ["Opt-in proof",       "BSP must store opt-in evidence; revoked = sender block"],
        ["Throughput",         "Per-WABA limit (Tier 1 = 1k unique 24h, Tier 4 = unlimited); plan accordingly"],
    ], col_widths_inches=[1.50, 5.00], font_size=9)

    add_section_heading(doc, "3.4", "RCS (Carrier Rich Messaging)", level=2)
    add_table(doc, ["Aspect", "Specification"], [
        ["Body limit",         "2500 chars per message"],
        ["Card",               "Single rich card OR carousel (2–10 cards); image + title + description + CTAs"],
        ["Suggested replies",  "Up to 11 quick-reply chips per message"],
        ["Media",              "Image ≤ 1 MB recommended for fast carrier delivery"],
        ["Verified sender",    "Brand verification through Google's RCS Business Messaging onboarding"],
        ["Fallback",           "Define SMS fallback for handsets without RCS support — same DLT rules apply"],
        ["DLT regime",         "Same as SMS in India — Header, Template ID, Entity ID required"],
        ["Receipt events",     "Delivered / Read / Suggestion-clicked all return events"],
    ], col_widths_inches=[1.50, 5.00], font_size=9)

    add_section_heading(doc, "3.5", "Slack / Teams (internal)", level=2)
    add_table(doc, ["Aspect", "Specification"], [
        ["Format",             "Block Kit (Slack) / Adaptive Card (Teams); markdown also acceptable"],
        ["Routing",            "Channel-based for digests, DM for personal alerts (OPS-03 escalation)"],
        ["Attachments",        "Inline summary + link out to TPE dashboard; never attach PII files"],
        ["@mentions",          "Use sparingly; reserve for action-required, not digest"],
        ["Threading",          "Replies in thread for context retention; new top-level only for new event"],
        ["Bot identity",       "TPE System Bot — single identity across all OPS triggers"],
    ], col_widths_inches=[1.50, 5.00], font_size=9)

    add_section_heading(doc, "3.6", "Email + Excel attachment (B2B remittance)", level=2)
    add_table(doc, ["Aspect", "Specification"], [
        ["Use case",           "Insurer-facing: INS-02 (Premium Remittance), INS-03 (Reconciliation)"],
        ["Filename pattern",   "TPE_<InsurerCode>_PremiumMIS_<YYYYMM>.xlsx"],
        ["Sheet name",         "'Raw' (case-insensitive); some insurers expect 'RAW' or 'Raw Data'"],
        ["Schema",             "Per-insurer fixed columns; produced by existing MIS Pipeline"],
        ["Subject line",       "[TPE] Premium MIS — <Insurer Name> — <Mon YYYY>"],
        ["Body",               "Short formal cover note; UTR reference; primary contact + escalation contact"],
        ["Encryption",         "Password-protected XLSX; password shared via separate channel (per insurer SLA)"],
    ], col_widths_inches=[1.50, 5.00], font_size=9)
    doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# Section 4 — Variable token catalogue
# ─────────────────────────────────────────────────────────────────────────────

VARIABLES = [
    # Identity
    ("ph_name",                "string",   "TPE CRM",            "1–60 chars; no special chars",  "\"Customer\""),
    ("ph_phone",               "string",   "TPE CRM",            "Indian mobile (+91 / 10-digit)", "—"),
    ("ph_email",               "string",   "TPE CRM",            "Email format",                  "—"),
    ("ph_policy_no",           "string",   "TPE CRM",            "Insurer-specific format",       "—"),
    ("inv_name",               "string",   "TPE CRM",            "1–60 chars",                    "\"Investor\""),
    ("inv_id",                 "string",   "TPE CRM",            "INV-XXXXXXXX",                  "—"),
    ("inv_phone",              "string",   "TPE CRM",            "Indian mobile",                 "—"),
    ("inv_email",              "string",   "TPE CRM",            "Email format",                  "—"),
    ("agent_handle",           "string",   "TPE Slack/Teams",    "@handle",                       "—"),
    # Policy data
    ("policy_no",              "string",   "Insurer master",     "Per-insurer pattern",           "—"),
    ("policy_type",            "enum",     "Insurer master",     "ULIP / Traditional / Term",     "—"),
    ("insurer_name",           "string",   "TPE master",         "From 11-insurer registry",      "—"),
    ("insurer_contact",        "string",   "TPE master",         "Per-insurer ops contact",       "—"),
    ("premium_amount",         "decimal",  "Insurer data",       "₹ format with thousands sep",  "—"),
    ("premium_frequency",      "enum",     "Insurer data",       "Annual / Semi / Quarterly / Monthly", "—"),
    # Time
    ("due_date",               "date",     "TPE calendar",       "DD MMM YYYY",                   "—"),
    ("grace_period_end",       "date",     "Insurer data",       "DD MMM YYYY",                   "—"),
    ("quote_validity_date",    "date",     "TPE CRM",            "DD MMM YYYY (typically T+7)",   "—"),
    ("maturity_date",          "date",     "Insurer data",       "DD MMM YYYY",                   "—"),
    ("settlement_date",        "date",     "TPE finance",        "DD MMM YYYY",                   "—"),
    ("deadline_date",          "date",     "TPE config",         "DD MMM YYYY",                   "—"),
    ("expected_credit_date",   "date",     "Bank cycle",         "DD MMM YYYY",                   "—"),
    # Money
    ("surrender_value",        "decimal",  "TPE valuation",      "₹X,XX,XXX format",              "—"),
    ("disbursement_amount",    "decimal",  "TPE finance",        "₹X,XX,XXX",                     "—"),
    ("commit_amount",          "decimal",  "TPE CRM",            "₹X,XX,XXX",                     "—"),
    ("maturity_amount",        "decimal",  "Insurer data",       "₹X,XX,XXX",                     "—"),
    ("emi_amount",             "decimal",  "TPE loan",           "₹X,XX,XXX",                     "—"),
    ("loan_amount",            "decimal",  "TPE loan",           "₹X,XX,XXX",                     "—"),
    ("tds_amount",             "decimal",  "TPE finance",        "₹X,XX,XXX",                     "—"),
    # Links / Tokens
    ("kyc_link",               "URL",      "TPE KYC provider",   "HTTPS, time-limited",           "—"),
    ("payment_link",           "URL",      "TPE payment GW",     "HTTPS, time-limited",           "—"),
    ("esign_link",             "URL",      "e-Sign provider",    "HTTPS, single-use, ≤ 7d expiry","—"),
    ("otp",                    "string",   "TPE OTP service",    "6 digits, single-use",          "—"),
    ("otp_expiry_min",         "integer",  "TPE config",         "Default 10",                    "10"),
    ("utr_number",             "string",   "Bank",               "Alphanumeric ≤ 22 chars",       "—"),
    ("bank_last4",             "string",   "TPE CRM",            "4 digits, masked",              "XXXX"),
    ("certificate_link",       "URL",      "TPE document store", "HTTPS, signed URL",             "—"),
    # Compliance / Identity
    ("irdai_reg_no",           "string",   "TPE config",         "IRDAI registration number",     "—"),
    ("dlt_template_id",        "string",   "DLT registry",       "12-digit numeric",              "—"),
    ("support_number",         "string",   "TPE config",         "1800-XXX-XXXX",                 "1800-TPE-CARE"),
    ("grievance_link",         "URL",      "TPE config",         "Static URL",                    "—"),
    ("unsubscribe_link",       "URL",      "TPE CRM",            "DPDPA-compliant, per-recipient","—"),
]

def section_variable_catalogue(doc):
    add_section_heading(doc, "4", "Variable Token Catalogue (37 tokens)")
    add_para(doc,
        "Variables are the personalisation points in templates. They are "
        "case-sensitive snake_case identifiers, wrapped {{like_this}} in "
        "template content. Every variable has a single canonical source — "
        "duplication across templates is a lint failure. Fallbacks render "
        "when the source returns null; they are the only safe way to ship "
        "a template that may render before all data is available.",
        space_after=8)
    rows = [(name, typ, src, valid, fb) for name, typ, src, valid, fb in VARIABLES]
    add_table(doc,
        ["Variable", "Type", "Source", "Validation", "Fallback"],
        rows,
        col_widths_inches=[1.55, 0.70, 1.20, 1.95, 1.10],
        font_size=8.5, code_cols={0})
    add_para(doc, "", space_after=4)
    add_callout(doc, "Why fallbacks matter more than you think",
        "A template that renders 'Hello ,' (missing ph_name) is a brand "
        "incident. A template that renders 'Hello Customer,' (fallback fired) "
        "is acceptable. The lint engine flags any required variable without a "
        "fallback. Fallbacks also de-risk the channel renderer: when a data "
        "integration regresses, customers still receive coherent messages "
        "while the integration is fixed.", kind="green")
    doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# Helpers for trigger-detail sections
# ─────────────────────────────────────────────────────────────────────────────

def render_trigger_card(doc, trigger):
    aud_label, color = AUDIENCE_META[trigger["aud"]]
    add_trigger_header(doc, trigger["id"], trigger["name"],
                       audience_color=color, audience_label=aud_label,
                       priority=trigger["priority"], volume=trigger["volume"])
    add_trigger_meta(doc,
                     statuses=trigger["statuses"],
                     channels=trigger["channels"],
                     cadence=trigger["cadence"],
                     variables=trigger["variables"],
                     compliance=trigger["compliance"])

def render_trigger_card_with_samples(doc, trigger, samples):
    """trigger card + a list of sample messages (channel, label, body)."""
    render_trigger_card(doc, trigger)
    for s in samples:
        add_sample_message(doc, s["channel"], s["label"], s["body"])
    add_para(doc, "", space_after=6)

# ─────────────────────────────────────────────────────────────────────────────
# Sample message library — drafts for the 15 highest-impact triggers
# ─────────────────────────────────────────────────────────────────────────────

SAMPLES_BY_TRIGGER = {
    "PH-02": [
        {"channel": "SMS", "label": "OTP",
         "body": (
            "{{otp}} is your TPE registration OTP. Valid for {{otp_expiry_min}} "
            "minutes. Do not share with anyone. TPE will never ask for your OTP. "
            "-TPECRE"
         )},
        {"channel": "WhatsApp", "label": "Abandoned (T+24h)",
         "body": (
            "Hi {{ph_name}},\n\n"
            "You started registering on The Policy Exchange yesterday but didn't "
            "complete it. Picking up where you left off takes ~2 minutes:\n\n"
            "{{reg_link}}\n\n"
            "Need help? Reply HELP. To stop these messages, reply STOP."
         )},
    ],
    "PH-03": [
        {"channel": "WhatsApp", "label": "Pending — Week 1",
         "body": (
            "Hi {{ph_name}},\n\n"
            "Your KYC is pending on The Policy Exchange. It takes 2 minutes to "
            "complete:\n\n"
            "{{kyc_link}}\n\n"
            "Until KYC is complete, we cannot proceed with your policy "
            "valuation.\n\nReply HELP for assistance, STOP to opt out."
         )},
        {"channel": "WhatsApp", "label": "Pending — Week 4 (final weekly nudge)",
         "body": (
            "Hi {{ph_name}},\n\n"
            "This is our final weekly reminder — your KYC has been pending for "
            "4 weeks. We will switch to monthly reminders after this. Your "
            "application stays open, but no further progress is possible without "
            "KYC.\n\n"
            "Complete in 2 minutes: {{kyc_link}}\n\n"
            "Need help? Reply HELP. To opt out: STOP."
         )},
        {"channel": "Email", "label": "Pending — Month 4 (closure notice)",
         "body": (
            "Subject: Your TPE application — pending KYC for over 4 months\n\n"
            "Dear {{ph_name}},\n\n"
            "We have been trying to reach you for over 4 months to complete the "
            "KYC step on your TPE application. Without KYC, we are unable to "
            "proceed with your policy valuation or surrender request.\n\n"
            "If we do not receive your KYC by {{deadline_date}}, we will close "
            "your application. You can re-apply at any time.\n\n"
            "Complete KYC now: {{kyc_link}}\n"
            "Or call us: {{support_number}}\n\n"
            "Warm regards,\n"
            "TPE Customer Care\n\n"
            "Insurance is the subject matter of solicitation. IRDAI Reg. No. {{irdai_reg_no}}.\n"
            "To opt out of all communications: {{unsubscribe_link}}"
         )},
    ],
    "PH-06": [
        {"channel": "Email", "label": "Quote Sent (T+0)",
         "body": (
            "Subject: Your policy valuation from TPE — ₹{{surrender_value}} | "
            "Valid until {{quote_validity_date}}\n\n"
            "Dear {{ph_name}},\n\n"
            "Thank you for choosing The Policy Exchange. Based on the policy "
            "details you shared:\n\n"
            "  Policy: {{policy_no}}\n"
            "  Insurer: {{insurer_name}}\n"
            "  Estimated surrender value: ₹{{surrender_value}}\n"
            "  Quote valid until: {{quote_validity_date}}\n\n"
            "This is an estimate, subject to insurer confirmation at the time "
            "of assignment. Your next step:\n\n"
            "  ✓ Accept this quote: {{accept_link}}\n"
            "  ✗ Discuss alternatives (counter-offer): {{counter_offer_link}}\n\n"
            "Quotes typically expire after 7 days. We will send one reminder "
            "before expiry.\n\n"
            "Warm regards,\n"
            "TPE Customer Care\n"
            "{{support_number}}\n\n"
            "Insurance is the subject matter of solicitation. IRDAI Reg. No. {{irdai_reg_no}}.\n"
            "Grievance redressal: {{grievance_link}}\n"
            "To opt out: {{unsubscribe_link}}"
         )},
    ],
    "PH-08": [
        {"channel": "WhatsApp", "label": "e-Sign Initiated (T+0)",
         "body": (
            "Hi {{ph_name}},\n\n"
            "Your policy assignment paperwork is ready for e-signing.\n\n"
            "📋 Policy: {{policy_no}}\n"
            "🏢 Insurer: {{insurer_name}}\n"
            "📅 Link expires: {{esign_expiry}}\n\n"
            "👉 Sign here: {{esign_link}}\n\n"
            "The link is single-use and expires in 7 days. Need help? Reply HELP."
         )},
    ],
    "PH-10": [
        {"channel": "SMS", "label": "Disbursement Success",
         "body": (
            "Rs.{{disbursement_amount}} from TPE credited to your "
            "{{bank_name}} a/c ending {{bank_last4}}. UTR {{utr_number}}. "
            "For queries call {{support_number}}. -TPECRE"
         )},
        {"channel": "WhatsApp", "label": "Disbursement Failed",
         "body": (
            "Hi {{ph_name}},\n\n"
            "Your disbursement of ₹{{disbursement_amount}} could not be "
            "completed.\n\n"
            "Reason: {{failure_reason}}\n\n"
            "Our team will reach out within 24 hours to resolve this. You can "
            "also call us: {{support_number}}.\n\n"
            "We're sorry for the inconvenience."
         )},
    ],
    "PH-15": [
        {"channel": "SMS", "label": "EMI Due Today (T-day)",
         "body": (
            "Loan EMI Rs.{{emi_amount}} for {{loan_number}} due TODAY. "
            "Pay now: {{payment_link}} to avoid late charges. -TPECRE"
         )},
        {"channel": "WhatsApp", "label": "EMI Overdue (Day+15 escalation)",
         "body": (
            "Hi {{ph_name}},\n\n"
            "⚠️ Your loan EMI of ₹{{emi_amount}} is now 15 days overdue.\n\n"
            "Loan: {{loan_number}}\n"
            "Outstanding: ₹{{overdue_amount}} (incl. late charge ₹{{late_charge}})\n\n"
            "Pay immediately to avoid further charges and credit bureau "
            "reporting:\n{{payment_link}}\n\n"
            "Need to discuss? Call {{support_number}}."
         )},
    ],
    "INV-04": [
        {"channel": "Email", "label": "Weekly Curated Digest",
         "body": (
            "Subject: This week's policy picks — {{policy_count}} new opportunities, "
            "yields {{min_yield_pct}}–{{max_yield_pct}}%\n\n"
            "Dear {{inv_name}},\n\n"
            "Here are the top assigned-policy opportunities matching your risk "
            "profile this week:\n\n"
            "  ★ Top pick: Policy {{top_policy_no}}\n"
            "    Investment: ₹{{top_policy_value}}\n"
            "    Yield range: {{min_yield_pct}}–{{max_yield_pct}}%\n\n"
            "Explore all {{policy_count}} matches: {{explore_link}}\n\n"
            "Yields are projections based on premium-payment continuity to "
            "policy maturity. Returns are subject to insurer performance and "
            "are not guaranteed. Read the full risk disclosure at the link "
            "above.\n\n"
            "Warm regards,\n"
            "TPE Investments\n\n"
            "Insurance is the subject matter of solicitation. IRDAI Reg. No. {{irdai_reg_no}}.\n"
            "To opt out of opportunity alerts: {{unsubscribe_link}}"
         )},
    ],
    "INV-08": [
        {"channel": "SMS", "label": "Premium Due (T-7)",
         "body": (
            "{{inv_name}}: Premium Rs.{{premium_amount}} for policy "
            "{{policy_no}} ({{insurer_name}}) due {{due_date}}. "
            "Pay: {{payment_link}}. -TPEINV"
         )},
        {"channel": "WhatsApp", "label": "Premium Due (T-day)",
         "body": (
            "Hi {{inv_name}},\n\n"
            "Your premium payment is due TODAY.\n\n"
            "📋 Policy: {{policy_no}}\n"
            "🏢 Insurer: {{insurer_name}}\n"
            "💰 Amount: ₹{{premium_amount}}\n"
            "📅 Due date: {{due_date}}\n\n"
            "⚡ Pay now to avoid grace-period charges:\n"
            "{{payment_link}}\n\n"
            "Auto-pay set up? You can ignore this. For assistance, "
            "call {{support_number}}.\n\n"
            "Reply STOP to opt out."
         )},
        {"channel": "Email", "label": "Premium Due (T-3)",
         "body": (
            "Subject: Premium due in 3 days — ₹{{premium_amount}} for "
            "{{policy_no}}\n\n"
            "Dear {{inv_name}},\n\n"
            "Your next premium payment is due on {{due_date}}.\n\n"
            "  Policy:        {{policy_no}}\n"
            "  Insurer:       {{insurer_name}}\n"
            "  Amount due:    ₹{{premium_amount}}\n"
            "  Due date:      {{due_date}}\n"
            "  Grace period:  Until {{grace_period_end}} (then late fee ₹{{late_fee}})\n\n"
            "Pay now: {{payment_link}}\n\n"
            "If you have an auto-pay mandate active, no action is required. To "
            "set up auto-pay for future premiums, reply to this email.\n\n"
            "Warm regards,\n"
            "TPE Investments\n"
            "{{support_number}}\n\n"
            "Insurance is the subject matter of solicitation. IRDAI Reg. No. {{irdai_reg_no}}.\n"
            "Grievance redressal: {{grievance_link}}\n"
            "To opt out: {{unsubscribe_link}}"
         )},
    ],
    "INV-10": [
        {"channel": "WhatsApp", "label": "Lapse Warning (Grace started)",
         "body": (
            "⚠️ URGENT: {{inv_name}}\n\n"
            "Your premium for policy {{policy_no}} is overdue and the grace "
            "period has started.\n\n"
            "Days overdue: {{days_overdue}}\n"
            "Grace period ends: {{grace_period_end}}\n"
            "Potential loss if policy lapses: ₹{{potential_loss}}\n\n"
            "Pay immediately to keep your policy active:\n"
            "{{payment_link}}\n\n"
            "Need to discuss? Call {{support_number}} now. Our team is "
            "available {{support_hours}}."
         )},
    ],
    "INV-12": [
        {"channel": "Email", "label": "Maturity Settled to Investor",
         "body": (
            "Subject: 🎉 Policy {{policy_no}} has matured — ₹{{maturity_amount}} "
            "credited\n\n"
            "Dear {{inv_name}},\n\n"
            "Congratulations — your assigned policy {{policy_no}} ({{insurer_name}}) "
            "has matured. The proceeds have been settled to you.\n\n"
            "  Maturity amount:   ₹{{maturity_amount}}\n"
            "  Settlement date:   {{settlement_date}}\n"
            "  Credited to:       Bank a/c ending {{bank_last4}}\n"
            "  UTR:               {{utr_number}}\n\n"
            "Your TDS certificate for this maturity is ready: "
            "{{tax_certificate_link}}\n\n"
            "Considering re-investment? We'll send a curated set of new "
            "opportunities matching your profile separately.\n\n"
            "Warm regards,\n"
            "TPE Investments\n\n"
            "Insurance is the subject matter of solicitation. IRDAI Reg. No. {{irdai_reg_no}}."
         )},
    ],
    "INV-14": [
        {"channel": "Email", "label": "TDS Certificate (Quarterly)",
         "body": (
            "Subject: TDS Certificate — {{fy_quarter}} — TPE\n\n"
            "Dear {{inv_name}},\n\n"
            "Please find your TDS certificate for {{fy_quarter}}.\n\n"
            "  TDS deducted: ₹{{tds_amount}}\n"
            "  Section:      {{section}}\n"
            "  PAN:          XXXXX{{pan_last4}}\n"
            "  Certificate:  {{tds_certificate_link}}\n\n"
            "This is generated in Form 16A format and is valid for tax filing.\n\n"
            "Warm regards,\n"
            "TPE Compliance"
         )},
    ],
    "INS-02": [
        {"channel": "Email", "label": "Premium Allocation MIS — Monthly Batch",
         "body": (
            "Subject: [TPE] Premium Allocation MIS — {{insurer_name}} — "
            "{{report_period}}\n\n"
            "Dear {{insurer_contact}},\n\n"
            "Please find attached the premium allocation MIS for "
            "{{report_period}}.\n\n"
            "  Total premium remitted:  ₹{{total_premium}}\n"
            "  Policy count:            {{policy_count}}\n"
            "  UTR reference:           {{utr_number}}\n"
            "  Attachment:              {{mis_attachment_name}}\n\n"
            "Kindly reconcile and confirm receipt within 5 working days. For "
            "any variances, please use ticket subject 'INS-02 Variance — "
            "{{report_period}}' so we can route it correctly.\n\n"
            "Best regards,\n"
            "TPE Operations Team\n"
            "Fairvalue Insuretech Pvt. Ltd.\n"
            "ops@thepolicyexchange.com  |  {{support_number}}"
         )},
    ],
    "OPS-04": [
        {"channel": "Slack", "label": "Premium Default Alert (real-time)",
         "body": (
            "🚨 *Premium Default Alert*\n\n"
            "Investor: `{{inv_id}}`\n"
            "Policy: `{{policy_no}}`\n"
            "Days overdue: *{{days_overdue}}*\n"
            "Action needed: {{action_needed}}\n\n"
            "Owner: <@{{collection_owner}}>\n\n"
            "Open in CRM: <{{crm_link}}|View case>"
         )},
    ],
}

# ─────────────────────────────────────────────────────────────────────────────
# Sections 5–7: Per-audience trigger detail
# ─────────────────────────────────────────────────────────────────────────────

def render_audience_section(doc, sec_num, sec_title, audience_code,
                             audience_intro):
    add_section_heading(doc, sec_num, sec_title)
    add_para(doc, audience_intro, space_after=8)
    triggers = [t for t in TRIGGERS if t["aud"] == audience_code]
    for i, t in enumerate(triggers):
        samples = SAMPLES_BY_TRIGGER.get(t["id"])
        if samples:
            render_trigger_card_with_samples(doc, t, samples)
        else:
            render_trigger_card(doc, t)
            add_para(doc, "", space_after=4)
        # Page break every ~3 triggers to avoid orphan rows
        if (i + 1) % 3 == 0 and i < len(triggers) - 1:
            doc.add_page_break()
    doc.add_page_break()

def section_ph_triggers(doc):
    render_audience_section(doc, "5",
        "Policyholder Triggers (PH-01 to PH-18)",
        "PH",
        "Eighteen triggers serving the sell-side of the marketplace. Voice "
        "is empathetic, simple, vernacular-friendly. WhatsApp dominates "
        "channel mix (Indian retail). Six triggers carry sample message "
        "drafts (PH-02 OTP, PH-03 KYC, PH-06 Surrender Value Quote, PH-08 "
        "Assignment Paperwork, PH-10 Disbursement, PH-15 Loan EMI) — these "
        "are the highest-volume or highest-sensitivity touchpoints.")

def section_inv_triggers(doc):
    render_audience_section(doc, "6",
        "Investor Triggers (INV-01 to INV-16)",
        "INV",
        "Sixteen triggers for the buy-side, dominated by INV-08 Premium Due "
        "(eight touchpoints per cycle, ~15,000 messages/month at steady "
        "state). Voice is professional, data-led, returns-framed. Email "
        "leads channel mix; WhatsApp and RCS for time-critical alerts; SMS "
        "reserved for OTP and due-day urgency. Six triggers carry sample "
        "message drafts (INV-04 Opportunity Push, INV-08 Premium Due ×3 "
        "channels, INV-10 Lapse Warning, INV-12 Maturity, INV-14 TDS).")

def section_b2b_internal_reg(doc):
    add_section_heading(doc, "7",
        "Insurance Partner, Internal & Regulatory Triggers (15)")
    add_para(doc,
        "These three audiences share a common trait: low message volume per "
        "trigger but very high consequence per message. INS-02 Premium "
        "Remittance is monthly but each email carries an MIS that reconciles "
        "thousands of policies. OPS-03 escalations are rare but VIP-bound. "
        "REG-01 IRDAI filings are five per year and define our regulatory "
        "standing. Lower volume, higher review rigour.",
        space_after=8)

    # INS
    add_section_heading(doc, "7.1", "Insurance Partner — B2B (6)", level=2)
    for t in [t for t in TRIGGERS if t["aud"] == "INS"]:
        samples = SAMPLES_BY_TRIGGER.get(t["id"])
        if samples:
            render_trigger_card_with_samples(doc, t, samples)
        else:
            render_trigger_card(doc, t)
            add_para(doc, "", space_after=4)
    doc.add_page_break()

    # OPS
    add_section_heading(doc, "7.2", "Internal Ops / Sales / Agents (6)", level=2)
    for t in [t for t in TRIGGERS if t["aud"] == "OPS"]:
        samples = SAMPLES_BY_TRIGGER.get(t["id"])
        if samples:
            render_trigger_card_with_samples(doc, t, samples)
        else:
            render_trigger_card(doc, t)
            add_para(doc, "", space_after=4)
    doc.add_page_break()

    # REG
    add_section_heading(doc, "7.3", "Regulatory / Statutory (3)", level=2)
    for t in [t for t in TRIGGERS if t["aud"] == "REG"]:
        render_trigger_card(doc, t)
        add_para(doc, "", space_after=4)
    doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# Section 8 — Sample Template Library (illustrative YAML)
# ─────────────────────────────────────────────────────────────────────────────

def section_template_library(doc):
    add_section_heading(doc, "8", "Sample Template Library (YAML)")
    add_para(doc,
        "Three illustrative templates in their canonical YAML form. These "
        "are the actual file shape that lives at "
        "templates/<audience>/<channel>/<id>.yaml in the repository. The "
        "schema is described in SOW Section 11.3; the values here show the "
        "shape filled with realistic content.",
        space_after=8)

    add_section_heading(doc, "8.1", "TPL-OTP-SMS-001 — shared OTP template", level=2)
    add_code_block(doc, """template_id: TPL-OTP-SMS-001
name: "OTP SMS — shared (Registration / Login / e-Sign)"
version: 1.0.0
status: active
audience: shared
language: en
channel: sms
category: authentication                  # DLT category
compliance_class: transactional
sender_identity: TPECRE                   # 6-char DLT header
brand_voice: factual_minimal
intent: "Deliver a one-time passcode within DLT Authentication template constraints."
variables:
  - name: otp
    required: true
    type: string
    validation: "6 digits"
  - name: otp_expiry_min
    required: true
    type: integer
    fallback: 10
disclaimers: []                           # Authentication category — none required
footer: []                                # No opt-out for Authentication
content:
  body: |
    {{otp}} is your TPE OTP. Valid for {{otp_expiry_min}} minutes.
    Do not share with anyone. TPE will never ask for your OTP. -TPECRE
character_count_estimate: 142             # Within single-segment 160 limit
parent_template: null
variants:
  - TPL-OTP-SMS-001-hi                    # Hindi variant
used_by_triggers: [PH-02, INV-02, PH-08]
metadata:
  created_by: ashish.satyam
  created_at: 2026-04-26
  approved_by: compliance.lead
  reuse_count: 3
  delivery_estimate_per_month: 1400
""")

    add_section_heading(doc, "8.2",
        "TPL-INV-PREMIUM-DUE-WA-001 — premium due on day-of, WhatsApp", level=2)
    add_code_block(doc, """template_id: TPL-INV-PREMIUM-DUE-WA-001
name: "Premium Due (T-day) — WhatsApp"
version: 1.3.0
status: active
audience: investor
language: en
channel: whatsapp
category: utility                         # WA BSP category
compliance_class: service
sender_identity: TPE_INV_WABA
brand_voice: investor_professional
intent: "Notify investor of premium due today; provide one-tap pay link."
variables:
  - name: inv_name
    required: true
    fallback: "Investor"
  - name: policy_no
    required: true
  - name: insurer_name
    required: true
  - name: premium_amount
    required: true
    type: decimal
    format: "INR"
  - name: due_date
    required: true
    type: date
    format: "DD MMM YYYY"
  - name: payment_link
    required: true
    type: short_url
  - name: support_number
    required: true
    fallback: "1800-TPE-INVE"
disclaimers: []
footer: [opt_out_short]                   # Service messages still carry opt-out
content:
  body: |
    Hi {{inv_name}},

    Your premium payment is due TODAY.

    📋 Policy: {{policy_no}}
    🏢 Insurer: {{insurer_name}}
    💰 Amount: ₹{{premium_amount}}
    📅 Due date: {{due_date}}

    ⚡ Pay now to avoid grace-period charges:
    {{payment_link}}

    Auto-pay set up? You can ignore this. For assistance, call {{support_number}}.

    Reply STOP to opt out.
character_count_estimate: 412             # Within 1024 WA limit
buttons:
  - type: url
    label: "Pay Now"
    url: "{{payment_link}}"
parent_template: null
variants:
  - TPL-INV-PREMIUM-DUE-WA-001-hi
used_by_triggers: [INV-08]
metadata:
  created_by: ashish.satyam
  created_at: 2026-04-26
  approved_by: compliance.lead
  brand_review_by: cx.head
  meta_template_id: hsm:tpe:inv_premium_due_001
  reuse_count: 1
  delivery_estimate_per_month: 1800       # Day-of-due touchpoint only
""")

    add_section_heading(doc, "8.3",
        "TPL-INS-PREMIUM-REMITTANCE-EMAIL-001 — monthly insurer remittance", level=2)
    add_code_block(doc, """template_id: TPL-INS-PREMIUM-REMITTANCE-EMAIL-001
name: "Monthly Premium Allocation MIS — Insurer Email"
version: 1.1.0
status: active
audience: insurer
language: en
channel: email_attachment
category: b2b_transactional
compliance_class: transactional
sender_identity: TPE_OPS                  # ops@thepolicyexchange.com
brand_voice: formal_b2b
intent: "Deliver monthly premium remittance MIS with full reconciliation context."
variables:
  - name: insurer_name
    required: true
  - name: insurer_contact
    required: true
    fallback: "Operations Team"
  - name: report_period
    required: true
    type: string
    format: "MMM YYYY"
  - name: total_premium
    required: true
    type: decimal
    format: "INR"
  - name: policy_count
    required: true
    type: integer
  - name: utr_number
    required: true
  - name: mis_attachment_name
    required: true
    type: filename
    pattern: "TPE_<insurer_code>_PremiumMIS_<YYYYMM>.xlsx"
  - name: support_number
    required: true
disclaimers: []
footer: [b2b_signature]
content:
  subject: "[TPE] Premium Allocation MIS — {{insurer_name}} — {{report_period}}"
  body: |
    Dear {{insurer_contact}},

    Please find attached the premium allocation MIS for {{report_period}}.

      Total premium remitted:  ₹{{total_premium}}
      Policy count:            {{policy_count}}
      UTR reference:           {{utr_number}}
      Attachment:              {{mis_attachment_name}}

    Kindly reconcile and confirm receipt within 5 working days. For any
    variances, please use ticket subject 'INS-02 Variance — {{report_period}}'
    so we can route it correctly.

    Best regards,
    TPE Operations Team
    Fairvalue Insuretech Pvt. Ltd.
    ops@thepolicyexchange.com  |  {{support_number}}
attachments:
  - name: "{{mis_attachment_name}}"
    source: mis_pipeline                  # produced by existing MIS Pipeline
    encryption: password_protected_xlsx
    password_delivery: per_insurer_sla
parent_template: null
variants: []
used_by_triggers: [INS-02]
metadata:
  created_by: ashish.satyam
  created_at: 2026-04-26
  approved_by: compliance.lead
  reuse_count: 1
  delivery_estimate_per_month: 11         # One per insurer
""")

    add_para(doc, "", space_after=4)
    add_callout(doc, "Why YAML, not a database",
        "Templates are read-mostly, audit-critical, and benefit from Git "
        "history (every edit gets an author, timestamp, diff and reviewable "
        "PR). A database adds operational surface area (backups, migrations, "
        "PII handling) without paying for itself in v1. We can move to a "
        "database in v2 if and only if write-throughput becomes a bottleneck "
        "— which is unlikely until template count exceeds ~1,000.",
        kind="navy")
    doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# Section 9 — Compliance quick-reference
# ─────────────────────────────────────────────────────────────────────────────

def section_compliance_ref(doc):
    add_section_heading(doc, "9", "Compliance Quick-Reference")
    add_para(doc,
        "One page per regime. Each entry is the question a reviewer asks "
        "before transitioning a template from review → approved. The lint "
        "engine encodes these as machine-checkable rules where possible; "
        "the rest are manual approver checklists. A template that fails "
        "any required check cannot be approved.",
        space_after=8)

    add_section_heading(doc, "9.1", "TRAI / DLT (SMS, RCS)", level=2)
    add_table(doc, ["Check", "Applies to", "Required?"], [
        ["DLT Header registered (6-char alpha)",         "All SMS & RCS",        "Yes"],
        ["DLT Template ID present (12-digit numeric)",   "All SMS & RCS",        "Yes"],
        ["Principal Entity (PE) ID present",              "All SMS & RCS",        "Yes"],
        ["Category matches content (Auth/Service/Promo)", "All SMS & RCS",        "Yes"],
        ["Promo time-window: 9:00–21:00 IST",             "Promotional only",     "Yes"],
        ["GSM-7 character set only (or Unicode flag set)","All SMS",              "Yes"],
        ["Char count within 160 single / 153 multi",      "All SMS",              "Yes — lint enforced"],
        ["Variable count and order matches DLT template", "All SMS & RCS",        "Yes"],
        ["Promo opt-out instruction present",             "Promotional only",     "Yes"],
    ], col_widths_inches=[3.50, 2.00, 1.00], font_size=9)

    add_section_heading(doc, "9.2", "IRDAI", level=2)
    add_table(doc, ["Check", "Applies to", "Required?"], [
        ["'Insurance is the subject matter of solicitation' present",
                                                          "All promotional", "Yes"],
        ["IRDAI Reg. No. visible (footer or body)",        "All promotional", "Yes"],
        ["Returns disclaimer ('subject to insurer performance, not guaranteed')",
                                                          "Yield-mentioning", "Yes"],
        ["Grievance redressal contact present",            "All promotional", "Yes"],
        ["Surrender value framed as estimate, not final",  "PH-06 quote",     "Yes"],
        ["Assignment intimation language IRDAI-compliant", "PH-08, INS-01",   "Yes"],
        ["Lapse warning aligns with IRDAI policy lapse rules",
                                                          "INV-10",          "Yes"],
        ["Maturity intimation includes IRDAI required fields",
                                                          "INV-11, INV-12",  "Yes"],
        ["Loan-against-policy: explicit T&C reference",    "PH-12, PH-13",    "Yes"],
    ], col_widths_inches=[3.50, 2.00, 1.00], font_size=9)

    add_section_heading(doc, "9.3", "WhatsApp BSP (Meta)", level=2)
    add_table(doc, ["Check", "Applies to", "Required?"], [
        ["Meta-approved template ID present (hsm:...)",   "All WhatsApp",     "Yes"],
        ["Category matches content (Marketing/Utility/Auth)",
                                                          "All WhatsApp",     "Yes"],
        ["Variable count and order matches Meta template","All WhatsApp",     "Yes"],
        ["Body within 1024 chars",                         "All WhatsApp",     "Yes — lint enforced"],
        ["Header media within size limits (image 5MB, video 16MB, doc 100MB)",
                                                          "Templates with header media", "Yes"],
        ["Buttons: ≤ 3 quick-reply OR ≤ 2 CTA (not both)","Button-bearing templates", "Yes"],
        ["24h session window respected for free-text",    "Free-text replies","Yes"],
        ["Opt-in evidence stored with BSP",                "All WhatsApp",     "Yes — operational"],
        ["Opt-out instruction in promo + service",         "Marketing + Utility", "Yes"],
    ], col_widths_inches=[3.50, 2.00, 1.00], font_size=9)

    add_section_heading(doc, "9.4", "DPDPA (Digital Personal Data Protection Act)", level=2)
    add_table(doc, ["Check", "Applies to", "Required?"], [
        ["Consent reference for marketing categories",     "All Marketing",   "Yes"],
        ["Opt-out instruction in every promotional message",
                                                          "All Promotional", "Yes"],
        ["Unsubscribe link is per-recipient (not generic)","Email promo",     "Yes"],
        ["Material T&C change: 30-day prior notice",       "INV-16, account notices", "Yes"],
        ["PII masking in audit logs (PAN, bank, OTP)",     "All — operational","Yes"],
        ["Cross-border transfer disclosure (if applicable)","Foreign data flows", "Conditional"],
        ["Grievance redressal contact for data concerns",   "All",             "Yes"],
    ], col_widths_inches=[3.50, 2.00, 1.00], font_size=9)

    add_section_heading(doc, "9.5", "FATCA / KYC", level=2)
    add_table(doc, ["Check", "Applies to", "Required?"], [
        ["FATCA self-declaration prompt in INV onboarding","INV-02, INV-03",  "Yes"],
        ["CKYC reference where applicable",                 "INV-03, PH-03",   "Recommended"],
        ["Risk profile completion prompt before opportunity push",
                                                          "INV-04",          "Yes"],
        ["KYC refresh trigger per RBI cycle",               "INS-06",          "Yes"],
        ["PAN format validation in template variables",     "All financial",  "Yes — lint enforced"],
    ], col_widths_inches=[3.50, 2.00, 1.00], font_size=9)

    add_section_heading(doc, "9.6", "Income Tax (TDS) + GST", level=2)
    add_table(doc, ["Check", "Applies to", "Required?"], [
        ["TDS section reference (e.g., 194DA)",             "INV-14, REG-02", "Yes"],
        ["TDS rate matches FY",                              "INV-14, REG-02", "Yes — derived, not hard-coded"],
        ["FY references machine-derived",                    "All tax docs",   "Yes — lint enforced"],
        ["Form 16A format for TDS certificates",            "INV-14, REG-02", "Yes"],
        ["GST invoice format (HSN, GSTIN, place of supply)","REG-03",         "Yes"],
        ["Invoice retention 8 years",                        "REG-03 — operational", "Yes"],
    ], col_widths_inches=[3.50, 2.00, 1.00], font_size=9)

    add_section_heading(doc, "9.7", "Internal brand-voice", level=2)
    add_table(doc, ["Check", "Applies to", "Required?"], [
        ["No urgency-marketing terms ('act now', 'last chance') in PH",
                                                          "PH templates",     "Yes"],
        ["No specific-return promises ('₹X guaranteed', 'X% assured')",
                                                          "INV templates",    "Yes — lint enforced"],
        ["Tone matches audience profile (empathy/professional/formal)",
                                                          "All",              "Yes — manual"],
        ["Vernacular variants for PH templates",            "All PH",          "Yes (Hindi v1; regional v2)"],
        ["No PII in subject lines",                         "Email",           "Yes — lint enforced"],
        ["No emoji in formal channels (B2B, Regulator)",    "INS, REG",        "Yes — lint enforced"],
    ], col_widths_inches=[3.50, 2.00, 1.00], font_size=9)

    add_para(doc, "", space_after=4)
    add_callout(doc, "Approval gate, not retrospective audit",
        "Every check above is enforced (or surfaced for manual approval) at "
        "the review → approved transition. Retrospective discovery of a miss "
        "after dispatch is, in practice, irreversible — the regulator's clock "
        "has started. The override path exists for genuine emergencies; "
        "overrides are logged with actor + reason and reviewed weekly.",
        kind="orange")

# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

def main():
    doc = setup_document()
    add_cover_page(doc)
    section_how_to_use(doc)
    section_quick_matrix(doc)
    section_channel_specs(doc)
    section_variable_catalogue(doc)
    section_ph_triggers(doc)
    section_inv_triggers(doc)
    section_b2b_internal_reg(doc)
    section_template_library(doc)
    section_compliance_ref(doc)
    out = ("/Users/aksatyam/TPE_DIRECTORY/TPE - Communication System/docs/"
           "TPE_Communication_Content_Reference.docx")
    doc.save(out)
    print(f"Saved: {out}")

if __name__ == "__main__":
    main()
