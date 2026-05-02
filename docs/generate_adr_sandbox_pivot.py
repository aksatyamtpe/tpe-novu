"""
TPE Communication System — ADR-001: Sandbox-3.15 Promoted to Primary (v1.0)

Audience: Engineering team, architecture review board, future maintainers
Purpose: Architecture Decision Record for the 2026-05-01 pivot.

ADR documents the WHY of the architectural change so future engineers don't
relitigate, don't accidentally undo it, and have a clear reversal path if
circumstances change.

Run: python3 generate_adr_sandbox_pivot.py
Output: TPE_ADR-001_Sandbox_3.15_Primary_v1.0.docx
"""

from datetime import date
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.shared import Pt

from _doc_helpers import (
    NAVY, BLUE, GREEN, ORANGE, RED, DARK, MEDIUM, WHITE,
    LIGHT, ALT_ROW, HEADER_BG, SUCCESS_BG, WARN_BG, DANGER_BG,
    DEFAULT_AUTHOR, DEFAULT_FOR, DEFAULT_ORG_LEGAL, DEFAULT_ORG_FULL,
    set_cell_bg, set_cell_borders, lock_table_widths,
    style_run, add_para, add_bullet, add_section_heading, add_table,
    add_status_badge_table, add_color_band, add_callout, add_metric_tiles,
    add_code_block, setup_document, add_footer_block,
)

DOC_TITLE_FULL  = "ADR-001: Sandbox-3.15 Promoted to Primary"
DOC_TITLE_SHORT = "ADR-001 Sandbox Pivot"
DOC_VERSION     = "1.0"
DOC_ID          = "TPE-COMMS-ADR-2026-001"
TODAY           = date(2026, 5, 1).strftime("%d %b %Y")
DECISION_DATE   = date(2026, 5, 1).strftime("%d %b %Y")
OUTPUT_FILE     = "TPE_ADR-001_Sandbox_3.15_Primary_v1.0.docx"


def add_cover_page(doc):
    add_color_band(doc, HEADER_BG, height_pt=10)
    add_para(doc, "", space_after=18)
    add_para(doc, DEFAULT_ORG_FULL.upper(), size=10, color=MEDIUM, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "ARCHITECTURE DECISION RECORD", size=20, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "ADR-001", size=14, bold=True, color=BLUE,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_para(doc, "Sandbox-3.15 Promoted to Primary", size=24, bold=True, color=DARK,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(doc, "Why we pivoted from custom TPE Admin to the Novu CE 3.15 dashboard",
             size=11, italic=True, color=MEDIUM,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    # ADR status banner
    tbl = doc.add_table(rows=1, cols=3)
    tbl.autofit = False
    lock_table_widths(tbl, [2.166, 2.166, 2.168])
    tiles = [
        ("ACCEPTED",         "Status",            "27AE60"),
        (DECISION_DATE,      "Decision Date",     "1B3A5C"),
        ("Engineering Lead", "Decision Authority", "2E75B6"),
    ]
    for i, (val, lbl, col) in enumerate(tiles):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, "FAFBFC")
        set_cell_borders(c, color=col, size=18)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.paragraphs[0].text = ""
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        from docx.shared import RGBColor as RGB
        r = c.paragraphs[0].add_run(val)
        style_run(r, size=14, bold=True, color=RGB.from_string(col))
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(lbl)
        style_run(r2, size=8.5, color=DARK)
    add_para(doc, "", space_after=18)

    add_para(doc, "Document metadata", size=10, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    meta_tbl = doc.add_table(rows=8, cols=2)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_tbl.autofit = False
    lock_table_widths(meta_tbl, [2.10, 3.10])
    meta = [
        ("Document ID",     DOC_ID),
        ("Version",         DOC_VERSION),
        ("Decision date",   DECISION_DATE),
        ("Authors",         DEFAULT_AUTHOR),
        ("For",             DEFAULT_ORG_FULL),
        ("Audience",        "Engineering team, architecture review board, future maintainers"),
        ("Status",          "● ACCEPTED — implementation complete"),
        ("Supersedes",      "(none — first ADR for this project)"),
    ]
    for i, (k, v) in enumerate(meta):
        kc = meta_tbl.rows[i].cells[0]
        vc = meta_tbl.rows[i].cells[1]
        set_cell_bg(kc, ALT_ROW if i % 2 == 0 else "FFFFFF")
        set_cell_bg(vc, ALT_ROW if i % 2 == 0 else "FFFFFF")
        set_cell_borders(kc); set_cell_borders(vc)
        kc.paragraphs[0].text = ""; vc.paragraphs[0].text = ""
        kr = kc.paragraphs[0].add_run(k)
        style_run(kr, size=9, bold=True, color=NAVY)
        vr = vc.paragraphs[0].add_run(v)
        style_run(vr, size=9, color=DARK)
    add_para(doc, "", space_after=12)
    add_para(doc, "INTERNAL / CONFIDENTIAL — DO NOT REDISTRIBUTE",
             size=9, bold=True, color=RED,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    doc.add_page_break()


def s1_doc_info(doc):
    add_section_heading(doc, "1.", "Document Information", level=1)
    add_table(doc,
        ["Attribute", "Value"],
        [
            ["ADR ID",            "ADR-001"],
            ["Title",             "Sandbox-3.15 Promoted to Primary"],
            ["Status",            "ACCEPTED (implementation complete)"],
            ["Decision date",     DECISION_DATE],
            ["Authors",           f"{DEFAULT_AUTHOR}, {DEFAULT_ORG_FULL}"],
            ["For",               DEFAULT_ORG_FULL],
            ["Audience",          "Engineering, architecture review, future maintainers"],
            ["Approvers",         "Engineering Lead (decided); informed: CTO, Compliance Lead, Operations Lead"],
            ["Supersedes",        "(none)"],
            ["Future review",     "Re-evaluate in 6 months OR if 4+ workflows need ad-hoc operator authoring"],
        ],
        col_widths_inches=[2.00, 4.50])


def s2_context(doc):
    add_section_heading(doc, "2.", "Context — what was the situation before?", level=1)
    add_para(doc,
        "Through April 2026, the TPE Communication System ran on a single live Novu CE 2.3.0 stack at "
        "VPS port 80. Operator-facing functionality (channel allowlist, campaign builder, scheduled "
        "trigger UI, run history) was built into a custom Next.js application called \"TPE Admin\" "
        "served at /admin/. The 17 Charter §4.3 lifecycle workflows were already authored as Bridge "
        "code-first and dispatching via custom step.custom + lib/dispatch.ts to MSG91 (SMS) + ICPaaS "
        "(WhatsApp).",
        space_after=8)
    add_para(doc,
        "On 30 April 2026, a parallel Novu CE 3.15.0 sandbox stack was deployed at ports 8080 / 8081 "
        "/ 8083 specifically to evaluate compatibility — could the existing Bridge code-first "
        "workflows + dispatch architecture forward-port from 2.3.0 to 3.15.0 without source changes? "
        "(Answer: yes, with strict 3.x payload validation gotchas documented separately in "
        "memory/novu_3x_payload_gotchas.md.)",
        space_after=8)
    add_para(doc,
        "On 1 May 2026, before the team fully evaluated the 3.15 sandbox, leadership directed: "
        "\"we have to use the novu FE to manage all TPE Admin FE will be removed.\" Then refined to: "
        "\"Use this — TPE Sandbox 3.15. Only to manage all notifications only.\" That triggered "
        "this ADR.",
        space_after=4)
    add_callout(doc, "Trigger statement from leadership",
        '"We have to use the novu FE to manage all TPE Admin FE will be removed. Use this — TPE '
        'Sandbox 3.15. Only to manage all notifications only."  '
        "— Project lead, 1 May 2026 (verbatim from chat directive)",
        kind="navy")


def s3_problem(doc):
    add_section_heading(doc, "3.", "Problem statement", level=1)
    add_para(doc,
        "Five interlocking problems converged into the decision below:",
        space_after=4)
    add_bullet(doc,
        "P1 — Two operator surfaces (TPE Admin + Novu Dashboard) created confusion; "
        "operators didn't know which to use for what.")
    add_bullet(doc,
        "P2 — TPE Admin's compliance middleware lint was duplicating Novu's workflow validation, "
        "causing maintenance burden when Novu CE was upgraded.")
    add_bullet(doc,
        "P3 — Studio-UI workflow authoring (visible in Novu Dashboard) had no compliance gate; "
        "operators could create workflows that bypassed DLT template enforcement.")
    add_bullet(doc,
        "P4 — Live 2.3.0 was being evaluated for 3.15 upgrade; running TPE Admin v1 against 3.15 "
        "would require re-validating every custom admin endpoint.")
    add_bullet(doc,
        "P5 — Charter §E1 mandates code-first workflow authoring with PR review. TPE Admin was a "
        "custom app that had drifted slightly toward operator self-service authoring patterns.")


def s4_options(doc):
    add_section_heading(doc, "4.", "Considered options", level=1)
    add_section_heading(doc, "4.1", "Option A — Keep TPE Admin (status quo)", level=2)
    add_para(doc, "Description:", bold=True, color=NAVY, space_after=2)
    add_para(doc,
        "Keep the custom Next.js TPE Admin at /admin/ as the canonical operator surface. "
        "Maintain it across Novu version upgrades. Do not promote sandbox 3.15.",
        space_after=4)
    add_para(doc, "Pros:", bold=True, color=GREEN, space_after=2)
    add_bullet(doc, "Existing operator workflows continue uninterrupted")
    add_bullet(doc, "Custom features (channel allowlist UI, campaign builder, run history) remain accessible")
    add_bullet(doc, "No retraining of operations team")
    add_para(doc, "Cons:", bold=True, color=RED, space_after=2)
    add_bullet(doc, "Maintenance cost of custom Next.js app on top of Novu CE upgrades")
    add_bullet(doc, "Two operator surfaces to test on every change")
    add_bullet(doc, "Doesn't address Charter §E1 alignment risk (P5)")
    add_para(doc, "Verdict: REJECTED — leadership directive required different path", italic=True, color=DARK, space_after=8)

    add_section_heading(doc, "4.2", "Option B — Full Novu fork with native MSG91 provider", level=2)
    add_para(doc, "Description:", bold=True, color=NAVY, space_after=2)
    add_para(doc,
        "Fork Novu's open-source repo, add MSG91 to the SMS catalog (8 files across 4 packages), "
        "build custom Docker images, replace ghcr.io upstream images. Operators can now use the "
        "native Novu Integration Store to configure MSG91, and Studio-authored workflows can fire "
        "MSG91 SMS.",
        space_after=4)
    add_para(doc, "Pros:", bold=True, color=GREEN, space_after=2)
    add_bullet(doc, "MSG91 visible in Integration Store (Studio workflows can use it)")
    add_bullet(doc, "Operators can author ad-hoc workflows via Studio UI")
    add_bullet(doc, "Full feature parity with native Novu providers")
    add_para(doc, "Cons:", bold=True, color=RED, space_after=2)
    add_bullet(doc, "Fork maintenance burden — every Novu version (3.16, 3.17, 4.0) requires re-merge")
    add_bullet(doc, "Custom Docker image build + storage pipeline required")
    add_bullet(doc, "Initial effort 3-5 days + ongoing maintenance")
    add_bullet(doc, "DLT template_id + vars passing requires bridgeProviderData abstraction (still leaks knowledge to operators)")
    add_bullet(doc, "Doesn't address Charter §E1 (still allows Studio workflows that bypass compliance lint)")
    add_para(doc, "Verdict: REJECTED — fork burden + Charter §E1 alignment doesn't improve", italic=True, color=DARK, space_after=8)

    add_section_heading(doc, "4.3", "Option C — Sandbox 3.15 as primary, Bridge for SMS+WA, code-first authoring", level=2)
    add_para(doc, "Description:", bold=True, color=NAVY, space_after=2)
    add_para(doc,
        "Promote the existing Novu CE 3.15 sandbox dashboard to canonical operator surface. Stop "
        "(preserve) TPE Admin containers — code preserved for rollback. Bridge dispatch architecture "
        "stays exactly the same — MSG91 SMS + ICPaaS WhatsApp continue to flow through "
        "lib/dispatch.ts. Code-first authoring rule is formalized: Studio UI workflow creation "
        "becomes forbidden, drift detection added.",
        space_after=4)
    add_para(doc, "Pros:", bold=True, color=GREEN, space_after=2)
    add_bullet(doc, "Single operator surface (the standard Novu CE dashboard at :8080)")
    add_bullet(doc, "Zero fork burden — track upstream Novu CE releases unchanged")
    add_bullet(doc, "Charter §E1 alignment: code-first is the only authoring path → PR review enforced")
    add_bullet(doc, "Existing dispatch architecture preserved → zero risk to MSG91 + ICPaaS production paths")
    add_bullet(doc, "Reversible — TPE Admin containers stopped, not removed; can be revived via skill")
    add_bullet(doc, "Implementation effort minimal (~1 hour: stop containers, sync workflows to sandbox, smoke test)")
    add_para(doc, "Cons:", bold=True, color=RED, space_after=2)
    add_bullet(doc, "Operations team loses TPE Admin's bespoke features (channel allowlist UI, campaign builder)")
    add_bullet(doc, "Operators cannot author ad-hoc workflows — every new trigger requires engineering")
    add_bullet(doc, "Channel allowlist becomes Mongo-direct edit (no UI) until alternative built")
    add_para(doc, "Verdict: ACCEPTED", italic=True, bold=True, color=GREEN, space_after=4)


def s5_decision(doc):
    add_section_heading(doc, "5.", "Decision", level=1)
    add_para(doc,
        "Adopt Option C — Sandbox 3.15 as primary, Bridge for SMS+WA, code-first authoring rule.",
        size=12, bold=True, color=NAVY, space_after=8)
    add_para(doc, "Specifically:", space_after=4)
    add_bullet(doc, "Operator UI: Novu CE 3.15 dashboard at http://103.138.96.180:8080")
    add_bullet(doc, "API: Novu CE 3.15 at http://103.138.96.180:8081")
    add_bullet(doc, "Bridge: next-bridge container with REAL MSG91 + ICPaaS creds (cloned from live novu-bridge)")
    add_bullet(doc, "Workflow authoring: code-first only, Studio UI authoring forbidden")
    add_bullet(doc, "TPE Admin: docker-stopped (preserved); rollback skill exists")
    add_bullet(doc, "Live 2.3.0 stack: kept running for redundancy, de-emphasised")
    add_para(doc, "", space_after=4)


def s6_tradeoffs(doc):
    add_section_heading(doc, "6.", "Trade-offs accepted", level=1)
    add_para(doc, "By choosing Option C we explicitly accept the following trade-offs:", space_after=6)
    add_table(doc,
        ["Trade-off accepted", "Why we accept it"],
        [
            ["Operations team loses TPE Admin's custom features (channel allowlist UI, campaign builder)",
             "Custom features are not Charter-required; standard Novu UI provides 80%+ of operator needs; rollback path exists if 20% becomes blocking"],
            ["Operators cannot author ad-hoc workflows from the Dashboard (Studio UI banned)",
             "Charter §E1 mandates PR-reviewed code-first authoring; operator self-service authoring would bypass compliance middleware lint, DLT enforcement, and audit-row format requirements"],
            ["Channel allowlist becomes Mongo-direct (no UI)",
             "Allowlist is engineering-managed config; operators rarely change it; UI can be rebuilt in a future Novu CE plugin if needed"],
            ["Two parallel WhatsApp paths exist (Bridge → ICPaaS, native whatsapp-business → Meta direct)",
             "User configured native whatsapp-business in Integration Store separately; doesn't conflict with Bridge path; future cleanup task"],
            ["Single-VPS deployment (Track B) continues for sandbox primary",
             "Pre-production stage; Track A migration planned for Phase 4 (Oct 2026) anyway"],
        ],
        col_widths_inches=[2.80, 3.70], font_size=8.5)


def s7_implementation(doc):
    add_section_heading(doc, "7.", "Implementation summary (executed 2026-05-01)", level=1)
    add_para(doc, "The pivot was executed in approximately 1 hour with the following sequence:", space_after=4)
    add_table(doc,
        ["#", "Action", "Outcome"],
        [
            ["1", "Stop tpe-admin + tpe-admin-worker containers (preserved, not removed)",
             "Both containers stopped; data + code intact for rollback"],
            ["2", "Drop placeholder-cred next-bridge container",
             "Removed compat-test artifact"],
            ["3", "Clone live novu-bridge env vars (MSG91, ICPaaS, NODE_ENV)",
             "Captured 27 env-var values"],
            ["4", "Recreate next-bridge with real creds + sandbox-specific NOVU_SECRET_KEY + MONGO_URL",
             "Container ready in 219ms; discovery returns 19 workflows / 79 steps"],
            ["5", "Run npx novu sync against sandbox API",
             "21 workflows registered (17 Charter + welcome-onboarding + tpe-multichannel-test + 2 stale Studio Welcome dupes)"],
            ["6", "Smoke-trigger ph-02-registration with crypto.randomUUID()",
             "Bridge dispatch: MSG91 v5/Flow HTTP 200 (status=sent); ICPaaS HTTP 200 (wamid issued); Email skipped (SES deferred); In-app sent"],
            ["7", "Document the pivot in memory/sandbox_315_primary.md",
             "9 KB memory file with rollback path documented"],
            ["8", "Establish code-first rule in memory/architecture_code_first_only.md",
             "Architectural rule with drift-check commands"],
            ["9", "Create skill novu-ce-pivot-rollback for reversal",
             "Operational runbook for restoring TPE Admin if needed"],
        ],
        col_widths_inches=[0.40, 3.30, 2.80], font_size=8.5)


def s8_verification(doc):
    add_section_heading(doc, "8.", "Verification — PH-02 smoke test results", level=1)
    add_para(doc, "End-to-end verification at 2026-05-01:", space_after=4)
    add_code_block(doc, """
trigger payload  : { otp: "482910", ph_id: "PH-SANDBOX-001",
                     testMode: true, triggerInstanceId: <crypto.randomUUID()> }
subscriber       : test-icpaas-sandbox-1 (phone +919465185365)
Novu trigger     : status 201, transactionId txn_69f458d8shgs5i2ntcjh

Bridge dispatch:
  SMS      : MSG91 v5/flow → HTTP 200, status=sent,
             template 6409f80fd6fc05100b559b13
  WhatsApp : ICPaaS dynamictemp2 → HTTP 200, wamid issued
             (wamid.HBgMOTE5NDY1MTg1MzY1FQIAERgSM0Q2ODU0MjExQkI5RkU3REUyAA==)
  Email    : skipped (SesNotConfigured: SES creds not provisioned — expected)
  In-app   : sent via Novu native (messages collection has the entry)

Audit emission verified — all four channels produced audit rows with
transactionId, triggerInstanceId, channel, status, ts.
""")
    add_section_heading(doc, "8.1", "Acceptance criteria", level=2)
    add_status_badge_table(doc, [
        ("17 Charter workflows registered on sandbox 3.15",  "Confirmed via Mongo notificationtemplates count",  "green"),
        ("MSG91 SMS dispatches successfully",                "HTTP 200 from MSG91 v5/Flow API",  "green"),
        ("ICPaaS WhatsApp dispatches successfully",          "HTTP 200 + wamid issued",  "green"),
        ("Email skips with descriptive reason",              "SesNotConfigured surfaced in audit row",  "green"),
        ("In-app notification creates a messages doc",       "Confirmed via messages collection lookup",  "green"),
        ("TPE Admin containers safely stopped",              "Both containers state = Exited (0)",  "green"),
        ("Rollback path verified",                           "novu-ce-pivot-rollback skill written + tested commands",  "green"),
    ])


def s9_reversal(doc):
    add_section_heading(doc, "9.", "Reversal plan", level=1)
    add_para(doc,
        "If this decision proves wrong, reversal is straightforward and documented:",
        space_after=4)
    add_bullet(doc, "Invoke the novu-ce-pivot-rollback skill (.claude/skills/novu-ce-pivot-rollback/SKILL.md)")
    add_bullet(doc, "Restart tpe-admin + tpe-admin-worker containers (data + code preserved)")
    add_bullet(doc, "Re-seed tpe_channel_gating Mongo doc if it was wiped (recipe in skill)")
    add_bullet(doc, "Decide what to do with sandbox 3.15 stack (keep / stop / decommission)")
    add_bullet(doc, "Notify operators of the URL change back to /admin")
    add_bullet(doc, "Supersede this ADR with a follow-up ADR documenting the reversal")
    add_para(doc, "", space_after=4)
    add_callout(doc, "Reversibility cost",
        "Estimated reversal effort: ~1 hour. Same as the original pivot since the work is "
        "essentially mirrored. The pivot was deliberately designed to be reversible.",
        kind="navy")


def s10_review_trigger(doc):
    add_section_heading(doc, "10.", "Future review triggers", level=1)
    add_para(doc, "Re-evaluate this decision (open ADR-002) if any of the following occur:", space_after=4)
    add_bullet(doc, "Operators report 4+ workflows that they need to author ad-hoc without engineering involvement")
    add_bullet(doc, "Novu CE 3.15 → 3.16+ upgrade introduces breaking changes in the Bridge protocol")
    add_bullet(doc, "MSG91 contract changes (vendor switch, deprecation) requires SMS provider migration")
    add_bullet(doc, "Charter §E1 requirement for code-first authoring is relaxed by Compliance Lead")
    add_bullet(doc, "An operations-driven need for in-dashboard campaign builder emerges (i.e. Phase 1A-1C resurrection)")
    add_bullet(doc, "Track A (production ECS Fargate) migration begins — re-evaluate which surface to expose")
    add_para(doc, "", space_after=4)
    add_para(doc, "Default review cadence: every 6 months OR on any of the above triggers.", italic=True, color=MEDIUM)


def s11_references(doc):
    add_section_heading(doc, "11.", "References", level=1)
    add_table(doc,
        ["Resource", "Path / URL"],
        [
            ["Project memory: pivot context",  "memory/sandbox_315_primary.md"],
            ["Project memory: code-first rule", "memory/architecture_code_first_only.md"],
            ["Project memory: 3.x payload gotchas", "memory/novu_3x_payload_gotchas.md"],
            ["Project memory: workflow wipe history", "memory/workflow_wipe_2026-05-01.md"],
            ["Skill: rollback this pivot",    ".claude/skills/novu-ce-pivot-rollback/SKILL.md"],
            ["Skill: recreate Bridge",         ".claude/skills/novu-ce-sandbox-bridge-recreate/SKILL.md"],
            ["Operator Manual",                "docs/TPE_Communication_System_Operator_Manual_v1.0.docx"],
            ["Project Status",                 "docs/TPE_Communication_System_Project_Status_v2.0.docx"],
            ["Charter §E1 mandate",            "docs/TPE_Communication_System_Charter_v1.0.docx (§E1)"],
            ["GitHub repo",                    "https://github.com/aksatyamtpe/tpe-novu"],
        ],
        col_widths_inches=[2.50, 4.00], font_size=9)


def main():
    doc = setup_document(DOC_TITLE_SHORT, DOC_ID, DOC_VERSION)
    add_cover_page(doc)
    s1_doc_info(doc)
    s2_context(doc)
    s3_problem(doc)
    s4_options(doc)
    s5_decision(doc)
    s6_tradeoffs(doc)
    s7_implementation(doc)
    s8_verification(doc)
    s9_reversal(doc)
    s10_review_trigger(doc)
    s11_references(doc)
    add_footer_block(doc, DEFAULT_AUTHOR, DEFAULT_ORG_FULL)
    doc.save(OUTPUT_FILE)
    print(f"✓ Generated {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
