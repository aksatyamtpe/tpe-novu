"""
TPE Communication System — Sandbox 3.15 Operator Manual Generator (v1.0)

Audience: Operations team using the Novu CE 3.15 Dashboard at port 8080
Purpose: Step-by-step reference for triggering and managing notifications post-pivot.

Captures the 2026-05-01 architectural pivot in operator-actionable terms:
- Dashboard at 103.138.96.180:8080 is now primary
- TPE Admin custom UI sunsetted (containers stopped, preserved)
- Bridge dispatch architecture for MSG91 SMS + ICPaaS WhatsApp
- 17 Charter §4.3 workflows operational

Run: python3 generate_operator_manual.py
Output: TPE_Communication_System_Operator_Manual_v1.0.docx
"""

from datetime import date
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.shared import Pt

from _doc_helpers import (
    NAVY, BLUE, GREEN, ORANGE, RED, DARK, MEDIUM, WHITE,
    LIGHT, ALT_ROW, HEADER_BG, SUCCESS_BG, WARN_BG, DANGER_BG,
    DEFAULT_AUTHOR, DEFAULT_BY, DEFAULT_FOR, DEFAULT_ORG_LEGAL, DEFAULT_ORG_FULL,
    set_cell_bg, set_cell_borders, lock_table_widths,
    style_run, add_para, add_bullet, add_section_heading, add_table,
    add_status_badge_table, add_color_band, add_callout, add_metric_tiles,
    add_code_block, setup_document, add_footer_block,
)

DOC_TITLE_FULL  = "TPE Communication System — Sandbox 3.15 Operator Manual"
DOC_TITLE_SHORT = "Operator Manual"
DOC_VERSION     = "1.0"
DOC_ID          = "TPE-COMMS-OPS-MANUAL-2026-001"
TODAY           = date(2026, 5, 1).strftime("%d %b %Y")
OUTPUT_FILE     = "TPE_Communication_System_Operator_Manual_v1.0.docx"


def add_cover_page(doc):
    add_color_band(doc, HEADER_BG, height_pt=10)
    add_para(doc, "", space_after=18)
    add_para(doc, DEFAULT_ORG_FULL.upper(), size=10, color=MEDIUM, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "OPERATOR MANUAL", size=22, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_para(doc, "TPE Communication System", size=26, bold=True, color=DARK,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(doc, "Sandbox 3.15 Dashboard — daily reference for the operations team",
             size=11, italic=True, color=MEDIUM,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    # Status snapshot tiles
    add_metric_tiles(doc, [
        ("17",  "Workflows live",     "1B3A5C"),
        ("4",   "Channels supported", "2E75B6"),
        ("2",   "Providers active",   "27AE60"),
        ("32",  "Workflows pending",  "E67E22"),
    ])
    add_para(doc, "", space_after=18)

    # Metadata table
    add_para(doc, "Document metadata", size=10, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    meta_tbl = doc.add_table(rows=8, cols=2)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_tbl.autofit = False
    lock_table_widths(meta_tbl, [2.10, 3.10])
    meta = [
        ("Document ID",     DOC_ID),
        ("Version",         DOC_VERSION),
        ("Date issued",     TODAY),
        ("Prepared by",     f"{DEFAULT_AUTHOR}  ·  {DEFAULT_BY}"),
        ("Prepared for",    DEFAULT_FOR),
        ("Audience",        "TPE Operations Team"),
        ("Classification",  "Internal / Confidential"),
        ("Document family", "Companion to Charter v1.0 + End User Guide v1.0"),
    ]
    for i, (k, v) in enumerate(meta):
        kc = meta_tbl.rows[i].cells[0]
        vc = meta_tbl.rows[i].cells[1]
        set_cell_bg(kc, ALT_ROW if i % 2 == 0 else "FFFFFF")
        set_cell_bg(vc, ALT_ROW if i % 2 == 0 else "FFFFFF")
        set_cell_borders(kc); set_cell_borders(vc)
        kc.paragraphs[0].text = ""; vc.paragraphs[0].text = ""
        kr = kc.paragraphs[0].add_run(k)
        style_run(kr, size=9, bold=True, color=NAVY)
        vr = vc.paragraphs[0].add_run(v)
        style_run(vr, size=9, color=DARK)
    add_para(doc, "", space_after=12)

    # Classification
    add_para(doc, "INTERNAL / CONFIDENTIAL — DO NOT REDISTRIBUTE",
             size=9, bold=True, color=RED,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    doc.add_page_break()


def add_section_1_doc_info(doc):
    add_section_heading(doc, "1.", "Document Information", level=1)

    add_para(doc, "Purpose", size=11, bold=True, color=NAVY)
    add_para(doc,
        "This manual is the day-to-day operating reference for the TPE Operations team using "
        "the Novu Community Edition 3.15.0 Dashboard at http://103.138.96.180:8080. It covers "
        "every operator-facing function — managing subscribers, triggering workflows, monitoring "
        "activity, and troubleshooting failed dispatches.",
        space_after=8)

    add_para(doc, "Scope", size=11, bold=True, color=NAVY)
    add_para(doc,
        "Operator workflows only. Engineering operations (workflow authoring, provider configuration, "
        "infrastructure) are out of scope and live in the engineering runbooks (deployment/Novu-Combined-"
        "Deployment-Guide.docx) and the Operational Skills catalogue at .claude/skills/INDEX.md.",
        space_after=8)

    add_para(doc, "Companion documents", size=11, bold=True, color=NAVY)
    add_table(doc,
        ["Document", "Purpose"],
        [
            ["TPE_Communication_System_Charter_v1.0.docx",
             "Project charter — strategic scope, budget, governance"],
            ["TPE_Communication_System_Charter_Addendum_v1.1.docx",
             "Tier 1+2 gap-fill: SLAs, security framework, glossary"],
            ["TPE_Communication_System_Trigger_Conditions_v2.0.docx",
             "Per-trigger conditions table (49 lifecycle triggers)"],
            ["TPE_Communication_System_End_User_Guide_v1.0.docx",
             "End-user-facing guide (recipients of notifications)"],
            ["TPE_Communication_System_Operator_Manual_v1.0.docx",
             "THIS DOCUMENT — operations team daily reference"],
            ["deployment/Novu-Combined-Deployment-Guide.docx",
             "Engineering deployment + ops runbook"],
        ],
        col_widths_inches=[3.20, 3.30])
    add_para(doc, "", space_after=8)

    add_para(doc, "Conventions used in this manual", size=11, bold=True, color=NAVY)
    add_bullet(doc, "● Green badge = component active and healthy")
    add_bullet(doc, "● Orange badge = degraded, manual intervention possible")
    add_bullet(doc, "● Red badge = down or blocked")
    add_bullet(doc, "Code in Consolas font is meant to be copy-pasted verbatim")
    add_bullet(doc, "Section §X.Y references mean Section X.Y of THIS document unless prefixed otherwise")


def add_section_2_audience(doc):
    add_section_heading(doc, "2.", "Audience Profile", level=1)

    add_para(doc,
        "This manual assumes the reader is a member of the Operations team with the following profile:",
        space_after=6)

    add_table(doc,
        ["Skill area", "Expected level", "Why it matters"],
        [
            ["Web browser usage", "Intermediate",
             "Dashboard is browser-based; you'll click through forms and tables"],
            ["JSON syntax", "Basic — read + edit",
             "Workflow payloads are JSON; you'll edit them in the trigger modal"],
            ["TPE business domain", "Intermediate to advanced",
             "You'll know what 'PH-02 OTP', 'INV-08 Premium Due', 'EMI cadence' mean"],
            ["Command line / SSH", "Not required for daily ops",
             "Bridge audit log access is helpful but Engineering can pull it for you"],
            ["Coding (TypeScript/Python)", "Not required",
             "Workflow authoring is engineering's job, not operations'"],
        ],
        col_widths_inches=[1.50, 1.50, 3.50])
    add_para(doc, "", space_after=8)

    add_callout(doc, "If you came from the legacy TPE Admin custom UI",
        "Some features that lived at /admin (channel allowlist, campaign builder, scheduled triggers) "
        "are sunsetted as of 2026-05-01. Their workflows are documented in §11 Common Operator Scenarios "
        "with the equivalent action in the Novu Dashboard. The legacy UI containers are stopped but "
        "preserved — Engineering can revive them via the novu-ce-pivot-rollback skill if needed.",
        kind="orange")


def add_section_3_architecture(doc):
    add_section_heading(doc, "3.", "Architecture in One Page", level=1)

    add_para(doc,
        "Understanding where you fit in the data flow makes everything else easier. Here's the path "
        "of a single notification from your trigger click to the customer's phone:",
        space_after=6)

    flow_code = """
   YOU (operations team)
    │
    ▼  click "Trigger Workflow" or send curl
┌──────────────────────────────────────────┐
│  Novu Dashboard 3.15 — port 8080         │  ← lists workflows from Mongo
│  http://103.138.96.180:8080              │
└────────────┬─────────────────────────────┘
             │ POST /v1/events/trigger { name, to, payload }
             ▼
┌──────────────────────────────────────────┐
│  Novu API + Worker (next-api / next-worker)
│  • Validates the workflow exists         │
│  • Queues a Notification job             │
│  • Worker pulls + calls Bridge           │
└────────────┬─────────────────────────────┘
             │ HTTP /api/novu (Bridge protocol)
             ▼
┌──────────────────────────────────────────┐
│  Bridge (next-bridge) — OUR CODE         │
│  bridge/workflows/<trigger>.ts            │
│   ├─ payloadSchema check (zod)           │
│   ├─ step.inApp(...)         → Inbox     │
│   ├─ step.custom('sms-otp', dispatch)    │
│   │     → MSG91 v5/Flow API              │
│   └─ step.custom('wa-otp', dispatch)     │
│         → ICPaaS WhatsApp Cloud API      │
└────────────┬─────────────────────────────┘
             │ HTTPS to providers
             ▼
   📨 MSG91     →  SMS to phone
   💬 ICPaaS    →  WhatsApp to phone
   ✉ AWS SES   →  Email (deferred)
   🔔 Novu     →  In-app entry on dashboard
"""
    add_code_block(doc, flow_code)
    add_para(doc, "", space_after=4)

    add_section_heading(doc, "3.1", "What runs where", level=2)
    add_table(doc,
        ["Component", "Container", "URL", "What it does"],
        [
            ["Dashboard UI", "next-dashboard",
             "http://103.138.96.180:8080", "What you click in"],
            ["API", "next-api",
             ":8081 (loopback :3000)", "Receives /v1/events/trigger calls"],
            ["Worker", "next-worker",
             "internal", "Picks up jobs from queue, calls Bridge"],
            ["Bridge", "next-bridge",
             ":4001 (loopback)", "Runs OUR code-first workflows + dispatches to providers"],
            ["Mongo", "next-mongodb",
             "internal", "Stores workflows, subscribers, jobs, activity"],
            ["Redis", "next-redis",
             "internal", "Job queue"],
        ],
        col_widths_inches=[1.20, 1.30, 1.80, 2.20], font_size=8.5)
    add_para(doc, "", space_after=4)

    add_callout(doc, "Why our SMS doesn't go through Novu's catalog",
        "MSG91 isn't in Novu CE 3.15's built-in SMS catalog (38 providers, none MSG91). "
        "Our 17 Charter workflows use a custom Bridge dispatch layer that calls MSG91's v5 Flow API "
        "directly with our DLT-registered template IDs. This is required for Indian SMS compliance — "
        "TRAI/DLT regulations need template_id + variable substitution, which generic SMS providers "
        "can't supply. WhatsApp goes the same way (ICPaaS, not Novu's native WhatsApp Business).",
        kind="blue")


def add_section_4_quickstart(doc):
    add_section_heading(doc, "4.", "Quick Start — Your First 5 Minutes", level=1)

    add_para(doc, "If this is your first time, do these five things in order:", space_after=6)

    add_section_heading(doc, "4.1", "Open the dashboard + sign in", level=2)
    add_bullet(doc, "Open http://103.138.96.180:8080 in your browser")
    add_bullet(doc, "Email: sandbox@tpe-test.local")
    add_bullet(doc, "Password: SandboxStage2026!")
    add_bullet(doc, "On first login, accept any prompts to enter the Development environment")
    add_para(doc, "", space_after=6)

    add_section_heading(doc, "4.2", "Verify the 17 workflows are visible", level=2)
    add_bullet(doc, "Click 'Workflows' in the left sidebar")
    add_bullet(doc, "You should see at least 19 entries: 17 Charter §4.3 workflows + welcome-onboarding + tpe-multichannel-test")
    add_bullet(doc, "Each shows a 'code-first' badge (read-only — they're managed in source)")
    add_para(doc, "", space_after=4)
    add_callout(doc, "If you see fewer than 17 workflows",
        "The Bridge sync may not have run. Contact Engineering — they'll run "
        "the novu-ce-bridge-sync skill against the sandbox API.",
        kind="orange")

    add_section_heading(doc, "4.3", "Verify your test subscriber exists", level=2)
    add_bullet(doc, "Click 'Subscribers' in the left sidebar")
    add_bullet(doc, "Look for a subscriber with ID 'test-icpaas-sandbox-1'")
    add_bullet(doc, "If missing, see §5.1 below to create one")

    add_section_heading(doc, "4.4", "Fire your first test trigger (PH-02 OTP)", level=2)
    add_bullet(doc, "Workflows tab → click 'ph-02-registration'")
    add_bullet(doc, "Top-right: click 'Trigger Workflow'")
    add_bullet(doc, "Subscriber ID: test-icpaas-sandbox-1")
    add_bullet(doc, "Payload (paste this JSON):")
    add_code_block(doc, '{\n  "otp": "482910",\n  "ph_id": "PH-TEST-001",\n  "testMode": true\n}')
    add_bullet(doc, "Click 'Run Workflow'")
    add_bullet(doc, "You'll get a transactionId like 'txn_69f...'")

    add_section_heading(doc, "4.5", "Verify dispatch happened", level=2)
    add_bullet(doc, "Click 'Activity Feed' in the left sidebar")
    add_bullet(doc, "Find your transactionId — click it")
    add_bullet(doc, "Each step (in_app / sms / whatsapp / email) should show a status")
    add_bullet(doc, "SMS + WhatsApp should land on the phone number associated with the subscriber")


def add_section_5_subscribers(doc):
    add_section_heading(doc, "5.", "Subscribers Tab", level=1)

    add_para(doc,
        "Subscribers are the 'who' of every notification. A subscriber represents one customer "
        "(or test account) and holds their contact info — phone, email, locale, custom data. "
        "Every trigger references a subscriberId to know who to notify.",
        space_after=8)

    add_section_heading(doc, "5.1", "Add a new subscriber", level=2)
    add_bullet(doc, "Subscribers tab → click 'Add Subscriber' (top-right)")
    add_bullet(doc, "Fill the form:")
    add_table(doc,
        ["Field", "Required?", "Example", "Notes"],
        [
            ["Subscriber ID", "Yes (unique)", "test-icpaas-sandbox-1",
             "Your app's user ID. Cannot be changed later."],
            ["First Name", "Recommended", "Sandbox",
             "Used in template substitution: {{firstName}}"],
            ["Last Name", "Optional", "Tester", ""],
            ["Email", "Yes if email channel", "no-reply@thepolicyexchange.com",
             "Where email goes"],
            ["Phone", "Yes if SMS or WhatsApp", "+919465185365",
             "E.164 format with country code; same field for SMS AND WA"],
            ["Avatar URL", "Optional", "https://...",
             "Shown in In-app messages"],
            ["Locale", "Recommended", "en  or  hi",
             "Picks the template language at dispatch time"],
        ],
        col_widths_inches=[1.30, 1.20, 1.80, 2.20], font_size=8.5)
    add_para(doc, "", space_after=4)
    add_bullet(doc, "Click 'Save'")
    add_bullet(doc, "Subscriber appears in the list immediately; can now be referenced in triggers")
    add_para(doc, "", space_after=4)

    add_section_heading(doc, "5.2", "Edit a subscriber", level=2)
    add_bullet(doc, "Click any subscriber in the list")
    add_bullet(doc, "Profile drawer opens — click the pencil icon next to any field")
    add_bullet(doc, "Change phone / email / locale / custom data (data field is free-form JSON)")
    add_bullet(doc, "Save")

    add_section_heading(doc, "5.3", "Common subscriber data shape for TPE", level=2)
    add_para(doc, "Custom 'data' JSON field — recommended schema for TPE subscribers:", space_after=4)
    add_code_block(doc, """
{
  "ph_id": "PH-12345",                    // policyholder ID
  "ph_state": "active",                   // or: lead | contacted | qualified | etc.
  "policy_number": "20356001",
  "investor_id": "INV-9876",              // if subscriber is an investor
  "investor_state": "soft_committed",
  "preferred_channels": ["sms", "whatsapp"],
  "dlt_consent": true,
  "language": "en"
}
""")
    add_callout(doc, "Compliance reminder",
        "Subscriber phone numbers are PII subject to DPDPA. Never enter test data with real "
        "customer phone numbers from production systems. Use synthetic test numbers or your own.",
        kind="red")

    add_section_heading(doc, "5.4", "Delete a subscriber", level=2)
    add_bullet(doc, "Subscriber profile → 'Delete subscriber' button (bottom)")
    add_bullet(doc, "Confirm twice (this is irreversible)")
    add_bullet(doc, "Activity history is preserved per Charter §4.8 audit retention")


def add_section_6_workflows(doc):
    add_section_heading(doc, "6.", "Workflows Tab", level=1)

    add_para(doc,
        "Workflows are the notification programs themselves — they define what channels fire, "
        "in what order, with what content. Under our code-first architecture, you cannot edit "
        "workflow steps from the dashboard. You CAN view them, trigger them, and inspect their schemas.",
        space_after=8)

    add_section_heading(doc, "6.1", "List view — what you see", level=2)
    add_table(doc,
        ["Column", "What it shows"],
        [
            ["Name / Trigger ID", "The workflow identifier (e.g. ph-02-registration)"],
            ["Channels", "Which channels this workflow uses (SMS / WhatsApp / Email / In-app)"],
            ["Updated", "When the workflow was last synced from source code"],
            ["Status", "Active / Inactive — whether triggers fire it"],
            ["Read-only badge", "Confirms it's code-first (cannot be edited in UI)"],
        ],
        col_widths_inches=[1.80, 4.70])
    add_para(doc, "", space_after=4)

    add_section_heading(doc, "6.2", "View workflow details", level=2)
    add_bullet(doc, "Click any workflow row")
    add_bullet(doc, "Detail page shows the step graph (visual flowchart)")
    add_bullet(doc, "Each step shows its type (inApp / SMS / WhatsApp / Email / Delay)")
    add_bullet(doc, "Click a step to see its content templates and conditions")
    add_bullet(doc, "Payload schema panel shows required + optional fields")

    add_section_heading(doc, "6.3", "Trigger a workflow", level=2)
    add_para(doc, "Two-step process:", space_after=4)
    add_bullet(doc, "1. Click 'Trigger Workflow' button (top-right of detail page)")
    add_bullet(doc, "2. Fill the trigger form")
    add_para(doc, "Trigger form fields:", space_after=4)
    add_table(doc,
        ["Field", "Type", "Notes"],
        [
            ["Subscriber ID", "Required string", "Must match an existing subscriber's ID"],
            ["Payload", "Required JSON object", "Workflow-specific; see §10 for per-workflow schema"],
            ["Override", "Optional JSON object", "Per-channel content override; rarely needed"],
        ],
        col_widths_inches=[1.30, 1.30, 3.90])
    add_para(doc, "", space_after=4)
    add_bullet(doc, "Click 'Run Workflow' — you get a transactionId immediately")
    add_bullet(doc, "Workflow execution happens asynchronously (typically completes in 5-15 seconds)")

    add_section_heading(doc, "6.4", "Why some workflows show 'invalid payload' on trigger", level=2)
    add_para(doc,
        "If your trigger fails with an 'invalid payload' error, it's because Bridge's strict "
        "zod schema validation rejected what you sent. In Novu 3.15 (vs older 2.3.0), validation "
        "is strict — wrong field names, missing required fields, or malformed UUIDs all 400.",
        space_after=4)
    add_para(doc, "Common gotchas (from observed sandbox failures):", space_after=4)
    add_table(doc,
        ["Mistake", "Error you see", "Fix"],
        [
            ["Sent 'otpCode' instead of 'otp'", "/otp Required",
             "Match exact field name from §10 reference"],
            ["Padded fake UUID", "Invalid uuid",
             "Use a real RFC-4122 UUID (engineering can generate)"],
            ["Missing required field", "/<field> Required",
             "Check §10 for what's required vs optional"],
            ["Sent string for number field", "Expected number, received string",
             "Cast to number before sending"],
        ],
        col_widths_inches=[2.00, 1.80, 2.70], font_size=8.5)


def add_section_7_topics(doc):
    add_section_heading(doc, "7.", "Topics Tab", level=1)

    add_para(doc,
        "Topics let you trigger one workflow to many subscribers at once. Useful for broadcast "
        "scenarios — service announcements, regulatory notices, mass campaigns. Without topics, "
        "you'd loop through subscribers and call /v1/events/trigger N times; with topics, one call "
        "fans out automatically.",
        space_after=8)

    add_section_heading(doc, "7.1", "Create a topic", level=2)
    add_bullet(doc, "Topics tab → 'Create Topic' button")
    add_bullet(doc, "Topic Key: lowercase-with-hyphens (this is the routing identifier)")
    add_bullet(doc, "Topic Name: human-readable description")
    add_bullet(doc, "Save → topic created with zero subscribers")

    add_section_heading(doc, "7.2", "Add subscribers to a topic", level=2)
    add_bullet(doc, "Click into the topic")
    add_bullet(doc, "'Add Subscribers' → select from list (multi-select)")
    add_bullet(doc, "Save → membership updated")

    add_section_heading(doc, "7.3", "Trigger to a topic", level=2)
    add_para(doc, "From the dashboard:", space_after=4)
    add_bullet(doc, "Workflows → pick workflow → Trigger")
    add_bullet(doc, "In the 'To' field, switch from 'Subscriber' to 'Topic'")
    add_bullet(doc, "Select the topic key")
    add_bullet(doc, "Payload: same as for a single subscriber — same payload goes to ALL members")

    add_section_heading(doc, "7.4", "Suggested topics for TPE", level=2)
    add_table(doc,
        ["Topic Key", "Use case"],
        [
            ["all-policyholders-en", "All English-speaking policyholders"],
            ["all-policyholders-hi", "All Hindi-speaking policyholders"],
            ["all-investors", "All investor accounts"],
            ["maturing-this-month", "Policies maturing in current month (rebuild monthly)"],
            ["lapsed-policies", "Lapsed accounts targeted for re-engagement (PH-17)"],
            ["test-subscribers", "Internal QA accounts — fire here before any real broadcast"],
        ],
        col_widths_inches=[2.30, 4.20])
    add_para(doc, "", space_after=4)

    add_callout(doc, "Topic vs Subscriber list — when to use which",
        "Use a Topic when the recipient list is stable (e.g. 'all investors'). Build a "
        "subscriber-list trigger loop instead when the list is dynamic (e.g. 'investors with "
        "policy maturing in the next 30 days' — recompute the list each run from your business DB).",
        kind="navy")


def add_section_8_integrations(doc):
    add_section_heading(doc, "8.", "Integrations Store", level=1)

    add_para(doc,
        "Integrations are the providers that actually send notifications — SMS gateways, email services, "
        "WhatsApp APIs, push services. The dashboard's Integrations Store lists what's configured and lets "
        "you add more.",
        space_after=8)

    add_section_heading(doc, "8.1", "Currently configured integrations (sandbox 3.15)", level=2)
    add_table(doc,
        ["Provider", "Channel", "Status", "Notes"],
        [
            ["Novu Inbox", "in_app", "● Active",
             "Default — auto-configured on every Novu install"],
            ["WhatsApp Business", "chat", "● Active",
             "Meta Cloud API direct (parallel to ICPaaS Bridge path)"],
            ["MSG91", "sms", "(not in catalog)",
             "Goes through Bridge dispatch — see §3 architecture"],
            ["ICPaaS", "whatsapp", "(not in catalog)",
             "Goes through Bridge dispatch — see §3 architecture"],
            ["AWS SES", "email", "(deferred)",
             "Not yet configured — emails currently skip"],
            ["FCM / APNs", "push", "(not yet)",
             "Push channel not in scope for current phase"],
        ],
        col_widths_inches=[1.40, 1.10, 1.30, 2.70], font_size=8.5)
    add_para(doc, "", space_after=4)

    add_callout(doc, "Don't try to add MSG91 from the dashboard",
        "MSG91 is not in Novu's catalog (any version). Attempting to add it via 'Add a Provider' "
        "won't show MSG91 as an option. Our SMS dispatch path is engineered through Bridge specifically "
        "because of this. If MSG91 needs reconfiguration, that's an engineering task — see "
        "novu-ce-providers-config skill.",
        kind="orange")

    add_section_heading(doc, "8.2", "Add a new catalog integration (e.g. SES email later)", level=2)
    add_bullet(doc, "Integrations Store → 'Add a Provider'")
    add_bullet(doc, "Filter by channel (SMS / Email / Push / etc.)")
    add_bullet(doc, "Pick from list — Novu CE 3.15 ships 38 SMS, 10+ email, 5 push providers")
    add_bullet(doc, "Fill credential fields per the provider's documentation")
    add_bullet(doc, "Save → toggle 'Active' to enable")
    add_bullet(doc, "Test from the integration's 'Test' button before relying on it for production")


def add_section_9_activity(doc):
    add_section_heading(doc, "9.", "Activity Feed — Post-Trigger Inspection", level=1)

    add_para(doc,
        "Activity Feed is your audit + debug surface. Every trigger ever fired produces an entry "
        "with full execution trace — which steps ran, what each one returned, what was sent.",
        space_after=8)

    add_section_heading(doc, "9.1", "Find a specific trigger", level=2)
    add_bullet(doc, "Activity Feed tab")
    add_bullet(doc, "Search bar: paste a transactionId (e.g. 'txn_69f458d8shgs5i2ntcjh')")
    add_bullet(doc, "Or filter: by workflow name, subscriber, status, date range")

    add_section_heading(doc, "9.2", "Read an activity entry", level=2)
    add_para(doc, "Each entry shows:", space_after=4)
    add_table(doc,
        ["Field", "Meaning"],
        [
            ["transactionId", "Unique trigger ID (txn_<base32>)"],
            ["Workflow", "Which workflow was fired"],
            ["Subscriber", "Recipient subscriberId"],
            ["Triggered at", "Timestamp of /v1/events/trigger call"],
            ["Status", "queued / running / completed / errored"],
            ["Steps", "Per-step status (in_app: sent, sms: sent, etc.)"],
            ["Payload", "What was sent in the trigger"],
            ["Error", "If errored, the failure detail"],
        ],
        col_widths_inches=[1.50, 5.00])
    add_para(doc, "", space_after=4)

    add_section_heading(doc, "9.3", "Channel-step status meanings", level=2)
    add_status_badge_table(doc, [
        ("Pending",    "Step queued, not yet picked up by worker", "navy"),
        ("Running",    "Step is executing (rare to see; usually transient)", "blue"),
        ("Completed",  "Step finished without error", "green"),
        ("Skipped",    "Step intentionally bypassed (channel allowlist, missing data)", "orange"),
        ("Failed",     "Step errored — Bridge returned 4xx/5xx, provider rejected, etc.", "red"),
    ])
    add_para(doc, "", space_after=4)

    add_section_heading(doc, "9.4", "Trigger says 'Completed' but customer didn't receive — what now?", level=2)
    add_para(doc,
        "The activity feed only shows what NOVU saw. If a step says 'Completed' but the customer "
        "didn't receive the message, the failure happened downstream of Novu (at the provider). "
        "You need the Bridge audit log:",
        space_after=4)
    add_code_block(doc,
        "ssh -i ~/.ssh/id_novu_vps -p 7576 root@103.138.96.180\n"
        "docker logs next-bridge --since 10m | grep <transactionId>")
    add_para(doc,
        "If you don't have SSH access, escalate to Engineering with the transactionId. "
        "They'll pull the audit log and tell you whether MSG91 / ICPaaS / SES rejected the send.",
        space_after=4)


def add_section_10_payload_reference(doc):
    add_section_heading(doc, "10.", "17 Charter Workflows — Payload Reference", level=1)

    add_para(doc,
        "Each workflow has a payload schema enforced at Bridge. Below is the per-workflow reference "
        "of required and optional fields, plus a sample payload you can paste into the trigger form.",
        space_after=8)

    add_callout(doc, "All workflows share the tpeBasePayload optional fields",
        "Every workflow accepts (optional): insurerId, triggerInstanceId (UUID), forceLocale ('en'|'hi'). "
        "These don't need to be in the per-workflow tables below — they always work.",
        kind="navy")
    add_para(doc, "", space_after=4)

    workflows = [
        ("ph-02-registration", "Policyholder OTP (registration)",
         [("otp", "string, length 6", "Required"),
          ("ph_id", "string", "Optional"),
          ("testMode", "boolean", "Optional"),
          ("abandonedAfterMinutes", "number", "Optional")],
         '{"otp":"482910","ph_id":"PH-001","testMode":true}'),
        ("ph-08-assignment-paperwork", "E-Sign nudge for policy assignment",
         [("ph_id", "string", "Required"),
          ("documentUrl", "string (URL)", "Required"),
          ("expiresAt", "ISO date string", "Optional")],
         '{"ph_id":"PH-001","documentUrl":"https://docs.tpe/x.pdf"}'),
        ("ph-09-investor-matched", "PH notified that an investor matched their policy",
         [("ph_id", "string", "Required"),
          ("investor_id", "string", "Required"),
          ("matchedAt", "ISO date", "Optional")],
         '{"ph_id":"PH-001","investor_id":"INV-9876"}'),
        ("ph-11-post-assignment-welcome", "Onboarding after investor accepts",
         [("ph_id", "string", "Required"),
          ("investor_id", "string", "Required")],
         '{"ph_id":"PH-001","investor_id":"INV-9876"}'),
        ("ph-12-loan-application", "Loan application acknowledgement",
         [("loan_application_id", "string", "Required"),
          ("ph_id", "string", "Required")],
         '{"loan_application_id":"LAP-001","ph_id":"PH-001"}'),
        ("ph-13-loan-approval", "Loan approval decision",
         [("loan_application_id", "string", "Required"),
          ("approvedAmount", "number", "Required"),
          ("ph_id", "string", "Required")],
         '{"loan_application_id":"LAP-001","approvedAmount":50000,"ph_id":"PH-001"}'),
        ("ph-14-loan-disbursement", "Loan funds disbursed",
         [("loan_application_id", "string", "Required"),
          ("disbursedAmount", "number", "Required"),
          ("ph_id", "string", "Required"),
          ("disbursedAt", "ISO date", "Optional")],
         '{"loan_application_id":"LAP-001","disbursedAmount":50000,"ph_id":"PH-001"}'),
        ("ph-15-loan-emi-reminder", "EMI reminder cadence (channel-matrix-by-stage)",
         [("loan_application_id", "string", "Required"),
          ("emiAmount", "number", "Required"),
          ("dueDate", "ISO date", "Required"),
          ("stage", "enum: due | 3d | 24h | 0h | post", "Required"),
          ("ph_id", "string", "Required")],
         '{"loan_application_id":"LAP-001","emiAmount":5500,"dueDate":"2026-05-15","stage":"3d","ph_id":"PH-001"}'),
        ("ph-16-loan-closure", "Loan closed — no more EMIs",
         [("loan_application_id", "string", "Required"),
          ("ph_id", "string", "Required"),
          ("closedAt", "ISO date", "Optional")],
         '{"loan_application_id":"LAP-001","ph_id":"PH-001"}'),
        ("ph-17-re-engagement", "Reactivation for dormant accounts",
         [("ph_id", "string", "Required"),
          ("dormancyDays", "number", "Required"),
          ("lastEvent", "string", "Optional")],
         '{"ph_id":"PH-001","dormancyDays":90,"lastEvent":"login"}'),
        ("inv-02-registration", "Investor OTP (registration)",
         [("otp", "string, length 6", "Required"),
          ("investor_id", "string", "Optional"),
          ("testMode", "boolean", "Optional")],
         '{"otp":"123456","investor_id":"INV-9876","testMode":true}'),
        ("inv-05-investment-soft-commit", "Investment soft-commit captured",
         [("investor_id", "string", "Required"),
          ("policy_number", "string", "Required"),
          ("amount", "number", "Required")],
         '{"investor_id":"INV-9876","policy_number":"20356001","amount":100000}'),
        ("inv-07-investment-confirmed", "Investment confirmed + receipt",
         [("investor_id", "string", "Required"),
          ("policy_number", "string", "Required"),
          ("amount", "number", "Required"),
          ("receiptUrl", "string (URL)", "Optional")],
         '{"investor_id":"INV-9876","policy_number":"20356001","amount":100000}'),
        ("inv-08-premium-due", "FLAGSHIP — Premium-due cadence (highest volume)",
         [("investor_id", "string", "Required"),
          ("policy_number", "string", "Required"),
          ("premiumAmount", "number", "Required"),
          ("dueDate", "ISO date", "Required"),
          ("stage", "enum: 30d | 7d | 24h | 0h | overdue", "Required")],
         '{"investor_id":"INV-9876","policy_number":"20356001","premiumAmount":12500,"dueDate":"2026-06-01","stage":"7d"}'),
        ("inv-09-premium-payment-confirmation", "Payment receipt",
         [("investor_id", "string", "Required"),
          ("policy_number", "string", "Required"),
          ("amount", "number", "Required"),
          ("paidAt", "ISO date", "Optional")],
         '{"investor_id":"INV-9876","policy_number":"20356001","amount":12500}'),
        ("inv-11-maturity-reminder", "Maturity-approaching cadence",
         [("investor_id", "string", "Required"),
          ("policy_number", "string", "Required"),
          ("maturityDate", "ISO date", "Required"),
          ("stage", "enum: 90d | 30d | 7d | due | post", "Required")],
         '{"investor_id":"INV-9876","policy_number":"20356001","maturityDate":"2026-08-01","stage":"30d"}'),
        ("inv-12-maturity-received", "Maturity payout receipt",
         [("investor_id", "string", "Required"),
          ("policy_number", "string", "Required"),
          ("payoutAmount", "number", "Required"),
          ("paidAt", "ISO date", "Optional")],
         '{"investor_id":"INV-9876","policy_number":"20356001","payoutAmount":250000}'),
    ]

    for trig, title, fields, sample in workflows:
        add_section_heading(doc, "", trig, level=3)
        add_para(doc, title, italic=True, color=MEDIUM, size=9.5, space_after=4)
        add_table(doc,
            ["Field", "Type", "Required?"],
            fields,
            col_widths_inches=[1.80, 2.50, 2.20], font_size=8.5)
        add_para(doc, "Sample payload:", size=9, bold=True, color=NAVY, space_after=2)
        add_code_block(doc, sample)
        add_para(doc, "", space_after=2)


def add_section_11_scenarios(doc):
    add_section_heading(doc, "11.", "Common Operator Scenarios", level=1)

    add_para(doc,
        "Real-world tasks — what to click for each.",
        space_after=8)

    add_section_heading(doc, "11.1", "Send a one-off OTP to a new policyholder", level=2)
    add_bullet(doc, "Subscribers tab → Add Subscriber (subscriberId, phone with country code)")
    add_bullet(doc, "Workflows tab → ph-02-registration → Trigger Workflow")
    add_bullet(doc, "Subscriber ID + payload {\"otp\":\"<6 digits>\"} → Run")
    add_bullet(doc, "Activity Feed → confirm transactionId completed; phone receives SMS+WA")

    add_section_heading(doc, "11.2", "Fire premium-due reminders to a list of investors", level=2)
    add_bullet(doc, "Topics tab → Create topic 'premium-due-this-week'")
    add_bullet(doc, "Add subscribers (your investor list)")
    add_bullet(doc, "Workflows tab → inv-08-premium-due → Trigger Workflow")
    add_bullet(doc, "To: Topic premium-due-this-week")
    add_bullet(doc, "Payload (per-subscriber values won't vary in this approach — for that, see §11.3)")
    add_para(doc, "", space_after=4)
    add_callout(doc, "When per-subscriber payload differs",
        "Topics broadcast the SAME payload to all members. If each investor needs different "
        "premiumAmount/dueDate, you can't use a Topic — you need a loop. That requires Engineering "
        "to write a one-off script (or revive the campaign builder if rolled back).",
        kind="orange")

    add_section_heading(doc, "11.3", "Verify a customer received an OTP they're claiming didn't arrive", level=2)
    add_bullet(doc, "Activity Feed → search by subscriber ID OR by phone number (data filter)")
    add_bullet(doc, "Find the most recent ph-02-registration trigger for that subscriber")
    add_bullet(doc, "Click into it → check SMS step status")
    add_bullet(doc, "If 'Completed' → message was accepted by MSG91; verify with Engineering's Bridge log")
    add_bullet(doc, "If 'Failed' → click the step → see the failure reason")
    add_bullet(doc, "If no recent trigger at all → the upstream system (registration service) didn't fire it")

    add_section_heading(doc, "11.4", "A workflow keeps failing — how to diagnose", level=2)
    add_table(doc,
        ["Symptom", "Most likely cause", "Action"],
        [
            ["All steps say 'Failed: Bridge execution failed'",
             "Payload schema mismatch (3.x strict validation)",
             "Click step → read raw error → fix payload field name/type"],
            ["SMS step says 'Skipped: channel disabled'",
             "tpe_channel_gating doc has SMS off",
             "Engineering re-seeds Mongo doc"],
            ["SMS step 'Completed' but no SMS arrives",
             "MSG91 accepted but DLT template rejected by carrier",
             "Engineering pulls Bridge audit log + MSG91 portal status"],
            ["WhatsApp step 'Failed: templateName required'",
             "ICPAAS_TEMPLATE_<TRIGGER>_WHATSAPP env not set",
             "Engineering fixes Bridge .env; recreate next-bridge"],
            ["Workflow not found at all",
             "Sync didn't run after a deploy",
             "Engineering runs novu-ce-bridge-sync"],
        ],
        col_widths_inches=[2.40, 2.20, 1.90], font_size=8)


def add_section_12_troubleshooting(doc):
    add_section_heading(doc, "12.", "Troubleshooting", level=1)

    add_para(doc,
        "When something's wrong, work down this list in order. The fixes get progressively more "
        "engineering-heavy — start at the top and only escalate when the easier options don't apply.",
        space_after=8)

    add_section_heading(doc, "12.1", "Operator-fixable issues", level=2)
    add_table(doc,
        ["Issue", "Quick check", "Fix"],
        [
            ["Can't log in to dashboard",
             "Try the credentials in §4.1 verbatim",
             "Caps lock off; if still failing, escalate to Engineering for password reset"],
            ["Workflow not in the list",
             "Filter is set or wrong env",
             "Clear filters; check top-right env selector says 'Development'"],
            ["Subscriber not findable",
             "Search box may be substring-strict",
             "Try partial ID; or filter by 'Has phone'"],
            ["Trigger button greyed out",
             "Workflow may be Inactive",
             "Toggle Active in workflow settings (still no edit to steps possible)"],
            ["Trigger 400: Subscriber not found",
             "subscriberId typed wrong",
             "Copy from Subscribers tab to avoid typos"],
        ],
        col_widths_inches=[1.90, 2.00, 2.60], font_size=8.5)

    add_section_heading(doc, "12.2", "Engineering-required issues", level=2)
    add_para(doc, "Escalate any of the following to Engineering with the transactionId and screenshot:",
             space_after=4)
    add_bullet(doc, "All steps fail with 'Bridge execution failed' across multiple workflows (Bridge container down)")
    add_bullet(doc, "Sandbox dashboard returns 500/502 on basic page loads (API or worker down)")
    add_bullet(doc, "MSG91 SMS step shows 'Completed' but customer never receives (carrier issue or DLT rejection)")
    add_bullet(doc, "ICPaaS WhatsApp step 'Failed: templateName required' (env-var bug)")
    add_bullet(doc, "Activity Feed empty for triggers that you know fired (Mongo or worker issue)")
    add_bullet(doc, "Need to add a new workflow (always engineering — code-first rule)")
    add_bullet(doc, "Need a new MSG91 DLT template registered (compliance + engineering process)")

    add_section_heading(doc, "12.3", "Escalation contacts", level=2)
    add_table(doc,
        ["Issue area", "Owner", "Channel"],
        [
            ["Dashboard / triggers", "Engineering Lead", "TBD per team"],
            ["Compliance / DLT / DPDPA", "Compliance Lead", "TBD per team"],
            ["MSG91 + ICPaaS provider issues", "Engineering + Vendor SLA", "TBD per team"],
            ["Customer copy / language", "CX Lead", "TBD per team"],
            ["Infrastructure / VPS", "Infrastructure Lead", "TBD per team"],
        ],
        col_widths_inches=[2.20, 2.00, 2.30])


def add_section_13_daily_checklist(doc):
    add_section_heading(doc, "13.", "Daily Operations Checklist", level=1)

    add_para(doc,
        "What an operator should verify each morning before the day's notification work begins.",
        space_after=8)

    add_section_heading(doc, "13.1", "Morning health check (5 min)", level=2)
    add_bullet(doc, "Open dashboard at http://103.138.96.180:8080 — should load within 3 sec")
    add_bullet(doc, "Click 'Workflows' — confirm 17+ workflows visible")
    add_bullet(doc, "Click 'Activity Feed' — confirm yesterday's last triggers all show 'Completed'")
    add_bullet(doc, "Trigger a test ph-02-registration to test-icpaas-sandbox-1 → confirm SMS + WA arrive")

    add_section_heading(doc, "13.2", "End-of-day verification (3 min)", level=2)
    add_bullet(doc, "Activity Feed → filter today → count 'Completed' vs 'Failed'")
    add_bullet(doc, "Failed > 5% of total → escalate to Engineering")
    add_bullet(doc, "Note any new transactionIds in your daily log")

    add_section_heading(doc, "13.3", "Weekly review (15 min, Mondays)", level=2)
    add_bullet(doc, "Review previous week's failure rate per workflow")
    add_bullet(doc, "Cross-check with MSG91 portal delivery reports for SMS")
    add_bullet(doc, "Cross-check ICPaaS dashboard for WhatsApp delivery (read receipts)")
    add_bullet(doc, "Surface anomalies to Engineering + CX standup")


def add_section_14_glossary(doc):
    add_section_heading(doc, "14.", "Glossary", level=1)

    glossary = [
        ("Bridge",
         "Custom Next.js HTTP server we wrote. Holds the 17 workflow definitions in TypeScript "
         "code and dispatches to MSG91/ICPaaS. Novu's worker calls Bridge over HTTP; Bridge calls providers."),
        ("Charter §4.3",
         "Section of the TPE Communication System Charter that lists the 49 lifecycle triggers. "
         "17 of them are currently live."),
        ("Channel allowlist",
         "Operator-controlled gate (Mongo `tpe_channel_gating` doc) that can disable specific channels "
         "(SMS, WA, email, in-app) at runtime. UI for managing this was sunsetted with TPE Admin."),
        ("Code-first authoring",
         "Architectural rule: workflows are TypeScript files in source control, NOT created via the "
         "Studio UI. Enforced via PR review + compliance middleware lint."),
        ("DLT",
         "Distributed Ledger Technology registry maintained by TRAI for India SMS. Every business "
         "SMS template must be DLT-registered with template_id; sending raw text is non-compliant."),
        ("DPDPA",
         "Digital Personal Data Protection Act 2023 (India). Why TPE data must stay India-resident."),
        ("ICPaaS",
         "WhatsApp aggregator that proxies Meta WhatsApp Business Cloud API. TPE uses their template "
         "messaging endpoint (POST .../v23.0/<phone-id>/messages)."),
        ("In-app",
         "Notifications shown inside the app (mobile/web Inbox widget), not delivered to phone/email. "
         "Powered by Novu's native Inbox provider."),
        ("MSG91",
         "Indian SMS gateway. TPE uses v5/Flow API with DLT-registered templates. Sender ID: TPolEx."),
        ("Novu CE",
         "Novu Community Edition — the open-source self-hosted version. We run 3.15.0 in sandbox at :8080 "
         "(primary) and 2.3.0 in live at :80 (legacy)."),
        ("Pivot",
         "The 2026-05-01 architectural decision to make sandbox 3.15 the primary operator surface, "
         "sunset the TPE Admin custom UI, and keep all Bridge dispatch (MSG91/ICPaaS) intact."),
        ("Studio UI",
         "Novu's in-dashboard 'Create Workflow' feature. FORBIDDEN for use under our code-first rule "
         "because it bypasses Git, compliance lint, and DLT enforcement."),
        ("Subscriber",
         "Novu's term for a recipient of notifications. One subscriber = one customer/test account "
         "with associated phone/email/locale/custom data."),
        ("Topic",
         "Group of subscribers, identified by a topic key. Triggering to a topic fans out to all members."),
        ("transactionId",
         "Unique ID for a trigger event. Format: txn_<base32> in 3.x (was raw UUID in 2.3.0). "
         "Use to trace a notification through the Activity Feed and Bridge audit log."),
        ("Workflow",
         "A notification program. Defines which channels fire, in what order, with what content. "
         "Each workflow has a unique trigger ID like 'ph-02-registration'."),
    ]
    for term, defn in glossary:
        add_para(doc, term, size=10, bold=True, color=NAVY, space_after=2)
        add_para(doc, defn, size=9.5, color=DARK, space_after=6)


def add_section_15_quickref(doc):
    add_section_heading(doc, "15.", "Quick Reference Card", level=1)

    add_para(doc,
        "Tear-out single-page summary. Print this and keep it next to your monitor.",
        size=10, italic=True, color=MEDIUM, space_after=8)

    add_section_heading(doc, "15.1", "URLs + credentials", level=2)
    add_table(doc,
        ["Surface", "URL", "Login"],
        [
            ["Primary Dashboard", "http://103.138.96.180:8080",
             "sandbox@tpe-test.local / SandboxStage2026!"],
            ["API endpoint",      "http://103.138.96.180:8081/v1/events/trigger",
             "ApiKey 54ad9d4dd50398489413be350237bd88"],
            ["Live Dashboard (legacy)", "http://103.138.96.180/", "(same creds)"],
        ],
        col_widths_inches=[1.50, 2.50, 2.50], font_size=8)

    add_section_heading(doc, "15.2", "Test subscriber", level=2)
    add_table(doc,
        ["Field", "Value"],
        [
            ["Subscriber ID", "test-icpaas-sandbox-1"],
            ["Phone (E.164)", "+919465185365"],
            ["Email", "no-reply@thepolicyexchange.com"],
        ],
        col_widths_inches=[1.50, 5.00])

    add_section_heading(doc, "15.3", "Most-used workflows + minimum payload", level=2)
    add_table(doc,
        ["Trigger ID", "Minimum payload"],
        [
            ["ph-02-registration",            '{"otp":"482910"}'],
            ["inv-02-registration",           '{"otp":"123456"}'],
            ["inv-08-premium-due",            '{"investor_id":"X","policy_number":"Y","premiumAmount":N,"dueDate":"...","stage":"7d"}'],
            ["ph-15-loan-emi-reminder",       '{"loan_application_id":"X","emiAmount":N,"dueDate":"...","stage":"3d","ph_id":"X"}'],
            ["inv-11-maturity-reminder",      '{"investor_id":"X","policy_number":"Y","maturityDate":"...","stage":"30d"}'],
            ["ph-17-re-engagement",           '{"ph_id":"X","dormancyDays":N}'],
        ],
        col_widths_inches=[2.00, 4.50], font_size=8.5, code_cols={1})

    add_section_heading(doc, "15.4", "Top 5 troubleshooting steps", level=2)
    add_bullet(doc, "1. Search Activity Feed by transactionId — get full trace")
    add_bullet(doc, "2. Click into the failed step — read the raw error message")
    add_bullet(doc, "3. Check that subscriber has phone (for SMS+WA) and email")
    add_bullet(doc, "4. Verify payload field names match §10 reference exactly")
    add_bullet(doc, "5. If still stuck after 5 min, escalate to Engineering with txn ID + screenshot")


def main():
    doc = setup_document(DOC_TITLE_SHORT, DOC_ID, DOC_VERSION)
    add_cover_page(doc)
    add_section_1_doc_info(doc)
    add_section_2_audience(doc)
    add_section_3_architecture(doc)
    add_section_4_quickstart(doc)
    add_section_5_subscribers(doc)
    add_section_6_workflows(doc)
    add_section_7_topics(doc)
    add_section_8_integrations(doc)
    add_section_9_activity(doc)
    add_section_10_payload_reference(doc)
    add_section_11_scenarios(doc)
    add_section_12_troubleshooting(doc)
    add_section_13_daily_checklist(doc)
    add_section_14_glossary(doc)
    add_section_15_quickref(doc)
    add_footer_block(doc, DEFAULT_AUTHOR, DEFAULT_BY, DEFAULT_FOR)
    doc.save(OUTPUT_FILE)
    print(f"✓ Generated {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
