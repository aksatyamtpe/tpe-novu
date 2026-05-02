"""
TPE enterprise-grade DOCX helpers — shared module.

Used by generators created 2026-05-01:
- generate_operator_manual.py        (Operator Manual)
- generate_project_status.py         (Master Status, 18 sections)
- generate_adr_sandbox_pivot.py      (ADR for Sandbox-3.15 pivot)

Brand colors + helpers match the patterns in:
- generate_charter_addendum.py
- generate_end_user_guide.py

Following CLAUDE.md global enterprise document standards.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────────────────────
# Brand palette (per CLAUDE.md global standards)
# ─────────────────────────────────────────────────────────────────────────────
NAVY    = RGBColor(0x1B, 0x3A, 0x5C)
BLUE    = RGBColor(0x2E, 0x75, 0xB6)
GREEN   = RGBColor(0x27, 0xAE, 0x60)
ORANGE  = RGBColor(0xE6, 0x7E, 0x22)
RED     = RGBColor(0xE7, 0x4C, 0x3C)
DARK    = RGBColor(0x2C, 0x3E, 0x50)
MEDIUM  = RGBColor(0x7F, 0x8C, 0x8D)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT       = "F0F4F8"
ALT_ROW     = "F8F9FA"
HEADER_BG   = "1B3A5C"
SUCCESS_BG  = "E8F5E9"
WARN_BG     = "FFF8E1"
DANGER_BG   = "FFEBEE"

# Default attribution per CLAUDE.md
DEFAULT_AUTHOR     = "Ashish Kumar Satyam"
DEFAULT_BY         = "TechDigital WishTree"
DEFAULT_FOR        = "The Policy Exchange"
DEFAULT_ORG_LEGAL  = "Fairvalue Insuretech Pvt. Ltd."
DEFAULT_ORG_FULL   = f"{DEFAULT_FOR} ({DEFAULT_ORG_LEGAL})"


# ─────────────────────────────────────────────────────────────────────────────
# Cell + table XML helpers
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_width_xml(cell, width_inches):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(int(width_inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')


def set_table_grid_xml(table, widths_inches):
    tbl = table._tbl
    grid = tbl.find(qn('w:tblGrid'))
    if grid is None:
        grid = OxmlElement('w:tblGrid')
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths_inches:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(int(w * 1440)))
        grid.append(col)


def set_table_width_xml(table, total_inches):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(int(total_inches * 1440)))
    tblW.set(qn('w:type'), 'dxa')


def lock_table_widths(table, widths_inches):
    total = sum(widths_inches)
    set_table_grid_xml(table, widths_inches)
    set_table_width_xml(table, total)
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            if ci < len(widths_inches):
                set_cell_width_xml(cell, widths_inches[ci])


def set_cell_borders(cell, color="BFBFBF", size=4):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(size))
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


# ─────────────────────────────────────────────────────────────────────────────
# Text + paragraph helpers
# ─────────────────────────────────────────────────────────────────────────────

def style_run(run, *, font="Arial", size=10, bold=False, italic=False, color=DARK):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rFonts.set(qn(attr), font)


def add_para(doc, text="", *, font="Arial", size=10, bold=False, italic=False,
             color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4, indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if text:
        r = p.add_run(text)
        style_run(r, font=font, size=size, bold=bold, italic=italic, color=color)
    return p


def add_bullet(doc, text, *, indent=0.5, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    style_run(r, size=size)
    return p


def add_section_heading(doc, num, title, level=1):
    """Plain typographic heading hierarchy (no left bars, no fills)."""
    if level == 1:
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{num}  {title}")
        style_run(run, font="Arial", size=14, bold=True, color=DARK)
        rule = doc.add_table(rows=1, cols=1)
        rule.autofit = False
        lock_table_widths(rule, [6.50])
        rc = rule.rows[0].cells[0]
        set_cell_bg(rc, "2C3E50")
        rc.paragraphs[0].text = ""
        rc.paragraphs[0].paragraph_format.space_before = Pt(0)
        rc.paragraphs[0].paragraph_format.space_after = Pt(0)
        rule.rows[0].height = Pt(1)
        rule.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
    elif level == 2:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{num}  {title}")
        style_run(r, size=12, bold=True, color=NAVY)
    else:
        add_para(doc, f"{num}  {title}", size=11, bold=True, italic=True,
                 color=BLUE, space_after=2)


def add_table(doc, headers, rows, *, col_widths_inches=None,
              alt_rows=True, font_size=9, code_cols=None, header_bg=None):
    """Headers: bold dark text on white; heavy bottom border separating header.
    Data rows: alternating background. Code columns use Consolas."""
    code_cols = code_cols or set()
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = False
    if col_widths_inches:
        lock_table_widths(tbl, col_widths_inches)
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        set_cell_bg(c, header_bg or "FFFFFF")
        set_cell_borders(c, color="2C3E50", size=8)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.paragraphs[0].text = ""
        r = c.paragraphs[0].add_run(h)
        style_run(r, font="Arial", size=font_size, bold=True,
                  color=WHITE if header_bg else DARK)
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]
        bg = ALT_ROW if (alt_rows and ri % 2 == 0) else "FFFFFF"
        for ci, val in enumerate(row):
            c = tr.cells[ci]
            set_cell_bg(c, bg)
            set_cell_borders(c)
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            c.paragraphs[0].text = ""
            font = "Consolas" if ci in code_cols else "Arial"
            r = c.paragraphs[0].add_run(str(val))
            style_run(r, font=font, size=font_size, color=DARK)
            c.paragraphs[0].paragraph_format.space_after = Pt(0)
    return tbl


def add_status_badge_table(doc, items):
    """Status badge table: items = list of (label, status, color_kind)
    color_kind one of 'green','orange','red','navy','blue'."""
    color_map = {
        'green':  (GREEN, SUCCESS_BG),
        'orange': (ORANGE, WARN_BG),
        'red':    (RED, DANGER_BG),
        'navy':   (NAVY, LIGHT),
        'blue':   (BLUE, LIGHT),
    }
    tbl = doc.add_table(rows=len(items), cols=2)
    tbl.autofit = False
    lock_table_widths(tbl, [3.0, 3.5])
    for i, (label, status, kind) in enumerate(items):
        fg, bg = color_map[kind]
        lc = tbl.rows[i].cells[0]
        set_cell_bg(lc, ALT_ROW if i % 2 == 0 else "FFFFFF")
        set_cell_borders(lc)
        lc.paragraphs[0].text = ""
        lc.paragraphs[0].add_run(label)
        style_run(lc.paragraphs[0].runs[0], size=9, bold=True, color=DARK)
        sc = tbl.rows[i].cells[1]
        set_cell_bg(sc, bg)
        set_cell_borders(sc)
        sc.paragraphs[0].text = ""
        r = sc.paragraphs[0].add_run(f"●  {status}")
        style_run(r, size=9, bold=True, color=fg)


def add_color_band(doc, hex_color, height_pt=4):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    lock_table_widths(tbl, [6.50])
    c = tbl.rows[0].cells[0]
    set_cell_bg(c, hex_color)
    c.paragraphs[0].text = ""
    c.paragraphs[0].paragraph_format.space_before = Pt(0)
    c.paragraphs[0].paragraph_format.space_after = Pt(0)
    tbl.rows[0].height = Pt(height_pt)
    tbl.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY


def add_callout(doc, title, body, kind="blue"):
    color_map = {
        'blue':   (BLUE,   LIGHT),
        'green':  (GREEN,  SUCCESS_BG),
        'orange': (ORANGE, WARN_BG),
        'red':    (RED,    DANGER_BG),
        'navy':   (NAVY,   LIGHT),
    }
    fg, bg = color_map[kind]
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    lock_table_widths(tbl, [6.50])
    c = tbl.rows[0].cells[0]
    set_cell_bg(c, bg)
    set_cell_borders(c, color="BFBFBF")
    c.paragraphs[0].text = ""
    r = c.paragraphs[0].add_run(title)
    style_run(r, size=10, bold=True, color=fg)
    p = c.add_paragraph()
    r2 = p.add_run(body)
    style_run(r2, size=9.5, color=DARK)
    add_para(doc, "", space_after=2)


def add_metric_tiles(doc, tiles):
    """4-tile horizontal metric card row.
    tiles = list of (number/value, label, hex_color_string).
    Up to 4 tiles per call."""
    n = len(tiles)
    widths = [6.50 / n] * n
    tbl = doc.add_table(rows=1, cols=n)
    tbl.autofit = False
    lock_table_widths(tbl, widths)
    for i, (val, label, col) in enumerate(tiles):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, "FAFBFC")
        set_cell_borders(c, color=col, size=18)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.paragraphs[0].text = ""
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.paragraphs[0].add_run(str(val))
        style_run(r, size=22, bold=True, color=RGBColor.from_string(col))
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(label)
        style_run(r2, size=8.5, color=DARK)


def add_code_block(doc, code, *, font_size=8.5):
    """Monospace block with light gray background — for command examples, JSON, etc."""
    lines = code.strip("\n").split("\n")
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    lock_table_widths(tbl, [6.50])
    c = tbl.rows[0].cells[0]
    set_cell_bg(c, "F5F7FA")
    set_cell_borders(c, color="DDE2E8")
    c.paragraphs[0].text = ""
    for i, line in enumerate(lines):
        if i == 0:
            p = c.paragraphs[0]
        else:
            p = c.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line if line else " ")
        style_run(r, font="Consolas", size=font_size, color=DARK)
    add_para(doc, "", space_after=2)


# ─────────────────────────────────────────────────────────────────────────────
# Page setup + cover + footer block
# ─────────────────────────────────────────────────────────────────────────────

def setup_document(doc_title_short, doc_id, version):
    """Configure US Letter, 1in margins, header (right-aligned title|version),
    footer (centered doc-id | confidentiality | page number)."""
    doc = Document()
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.4)
        section.footer_distance = Inches(0.4)
        h = section.header.paragraphs[0]
        h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rh = h.add_run(f"{doc_title_short}  |  v{version}")
        style_run(rh, size=8, color=MEDIUM, italic=True)
        f = section.footer.paragraphs[0]
        f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rf = f.add_run(f"{doc_id}  |  Internal & Confidential  |  Page ")
        style_run(rf, size=8, color=MEDIUM)
        run = f.add_run()
        fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText'); instr.text = "PAGE"
        fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_end)
        style_run(run, size=8, color=MEDIUM, bold=True)
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.font.color.rgb = DARK
    return doc


def add_footer_block(doc, prepared_by, by_org, for_org):
    """Navy footer block at the very end of the document with attribution."""
    add_para(doc, "", space_after=12)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    lock_table_widths(tbl, [6.50])
    c = tbl.rows[0].cells[0]
    set_cell_bg(c, HEADER_BG)
    set_cell_borders(c, color="0D1B2A", size=8)
    c.paragraphs[0].text = ""
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Prepared by")
    style_run(r, size=9, bold=True, color=WHITE)
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = p2.add_run(prepared_by)
    style_run(r2, size=11, bold=True, color=WHITE)
    p3 = c.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r3 = p3.add_run(f"{by_org}  ·  for {for_org}")
    style_run(r3, size=9, color=WHITE)
    p4 = c.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p4.paragraph_format.space_after = Pt(8)
    r4 = p4.add_run("Internal / Confidential")
    style_run(r4, size=8, italic=True, color=WHITE)
